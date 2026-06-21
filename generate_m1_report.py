"""Module 1 Comprehensive Report Generator.

Builds a DOCX report covering all 9 Module-1 deliverables:
  1.  Automatic CO Generation
  2.  Intelligent Question Bank Creation
  3.  Smart Question Paper (auto-generated sample)
  4.  Bloom's Taxonomy Compliance Check
  5.  CO Coverage Verification
  6.  Difficulty Level Analysis
  7.  Scenario & Case Study Analysis
  8.  CO-wise Marks Distribution
  9.  OBE Compliance Report
  10. Question Quality vs Programme Outcomes (PO)
  11. Question Quality vs Sustainable Development Goals (SDG)
"""

import re
import json
import math
import datetime
import textwrap

try:
    import report_charts as _rc
except Exception:
    _rc = None

# ── Bloom's action-verb mapping ───────────────────────────────────────────────
_BLOOM_LEVELS = {
    "remember":   ["define", "list", "recall", "state", "identify", "name", "recognise", "label"],
    "understand": ["explain", "describe", "summarise", "interpret", "classify", "compare", "illustrate"],
    "apply":      ["solve", "compute", "use", "implement", "demonstrate", "calculate", "execute", "apply"],
    "analyse":    ["analyse", "differentiate", "examine", "distinguish", "break down", "investigate"],
    "evaluate":   ["evaluate", "justify", "assess", "critique", "argue", "appraise", "defend", "judge"],
    "create":     ["design", "develop", "construct", "generate", "formulate", "plan", "produce", "build"],
}

_SCENARIO_KEYWORDS = [
    "scenario", "case study", "consider", "given that", "in a system",
    "suppose", "assume", "a company", "an organization", "discuss the",
    "analyse the following", "with reference to", "in context of",
    "real-world", "industry", "application of",
]

_DIRECT_VERBS = [
    r"^define\b", r"^list\b", r"^state\b", r"^what is\b",
    r"^name\b", r"^write\b", r"^give\b",
]


def _detect_bloom(text: str) -> str:
    t = text.lower()
    for level, verbs in reversed(list(_BLOOM_LEVELS.items())):
        if any(t.startswith(v) or f" {v} " in t for v in verbs):
            return level.capitalize()
    return "Remember"   # default fallback


def _is_scenario(text: str) -> bool:
    t = text.lower()
    return any(kw in t for kw in _SCENARIO_KEYWORDS)


def _is_direct(text: str) -> bool:
    t = text.lower().strip()
    return any(re.match(pat, t) for pat in _DIRECT_VERBS)


# ── Learning analytics constants ─────────────────────────────────────────────

_LA_LEVEL_NAMES = {
    'L1': 'Remember', 'L2': 'Understand', 'L3': 'Apply',
    'L4': 'Analyze',  'L5': 'Evaluate',   'L6': 'Create',
}
_LA_LEVEL_ORDER = ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']
_LA_KDIMS_ORDER = ['Factual', 'Conceptual', 'Procedural', 'Meta-Cognitive']


def _compute_bloom_analytics(cos: list, taxonomy_grid: dict) -> dict:
    """Derive cognitive-level learning analytics from the RBT taxonomy grid."""
    level_co_map = {lv: [] for lv in _LA_LEVEL_ORDER}
    kdim_co_map  = {kd: [] for kd in _LA_KDIMS_ORDER}
    co_to_level  = {}

    if taxonomy_grid:
        for kd, levels in taxonomy_grid.items():
            kd_clean = kd if kd in _LA_KDIMS_ORDER else 'Factual'
            for lv, co_list in levels.items():
                if lv not in _LA_LEVEL_ORDER:
                    continue
                for co in co_list:
                    if co not in co_to_level:
                        co_to_level[co] = lv
                    if co not in level_co_map[lv]:
                        level_co_map[lv].append(co)
                    if co not in kdim_co_map[kd_clean]:
                        kdim_co_map[kd_clean].append(co)
    else:
        _lv_map = {
            'Remember': 'L1', 'Understand': 'L2', 'Apply': 'L3',
            'Analyse': 'L4', 'Evaluate': 'L5', 'Create': 'L6',
        }
        for c in cos:
            co_key = f"CO{c['num']}"
            lv = _lv_map.get(_detect_bloom(c["statement"]), 'L1')
            co_to_level[co_key] = lv
            if co_key not in level_co_map[lv]:
                level_co_map[lv].append(co_key)

    total_mapped = max(sum(len(v) for v in level_co_map.values()), 1)
    lots_cos = sum(len(level_co_map[lv]) for lv in ['L1', 'L2', 'L3'])
    hots_cos = sum(len(level_co_map[lv]) for lv in ['L4', 'L5', 'L6'])
    weighted_sum = sum((i + 1) * len(level_co_map[lv])
                       for i, lv in enumerate(_LA_LEVEL_ORDER))
    progression_score = round(weighted_sum / total_mapped, 1)

    lots_pct = round(lots_cos / total_mapped * 100, 1)
    hots_pct = round(hots_cos / total_mapped * 100, 1)

    # Cognitive breadth, dominant level and empty levels
    level_counts   = {lv: len(level_co_map[lv]) for lv in _LA_LEVEL_ORDER}
    covered_levels = [lv for lv in _LA_LEVEL_ORDER if level_counts[lv] > 0]
    empty_levels   = [lv for lv in _LA_LEVEL_ORDER if level_counts[lv] == 0]
    breadth        = len(covered_levels)
    dominant_level = (max(_LA_LEVEL_ORDER, key=lambda lv: level_counts[lv])
                      if any(level_counts.values()) else None)

    # Cognitive balance index (normalised Shannon entropy, 0-100) — measures how
    # evenly COs are spread across the six Bloom levels (100 = perfectly even).
    probs = [level_counts[lv] / total_mapped for lv in _LA_LEVEL_ORDER if level_counts[lv] > 0]
    if len(probs) > 1:
        entropy = -sum(p * math.log(p, 2) for p in probs)
        balance_index = round(entropy / math.log(len(_LA_LEVEL_ORDER), 2) * 100)
    else:
        balance_index = 0

    # Knowledge-dimension breadth
    kdim_covered = [kd for kd in _LA_KDIMS_ORDER if kdim_co_map.get(kd)]
    kdim_empty   = [kd for kd in _LA_KDIMS_ORDER if not kdim_co_map.get(kd)]

    return {
        'level_co_map': level_co_map,
        'kdim_co_map':  kdim_co_map,
        'co_to_level':  co_to_level,
        'total_mapped': total_mapped,
        'lots_cos':     lots_cos,
        'hots_cos':     hots_cos,
        'lots_pct':     lots_pct,
        'hots_pct':     hots_pct,
        'progression_score': progression_score,
        'level_counts':   level_counts,
        'covered_levels': covered_levels,
        'empty_levels':   empty_levels,
        'breadth':        breadth,
        'dominant_level': dominant_level,
        'balance_index':  balance_index,
        'kdim_covered':   kdim_covered,
        'kdim_empty':     kdim_empty,
        'hots_target_met': hots_pct >= 40,
    }


def _la_alignment_checks(la: dict) -> list:
    """Return OBE alignment indicators as (criterion, result, status) rows.

    The qualitative status is derived internally; the numeric cut-offs used to
    decide it are intentionally not exposed in the report.
    """
    hots_pct = la['hots_pct']
    ps       = la['progression_score']
    bi       = la['balance_index']
    kd_cov   = len(la['kdim_covered'])
    return [
        ("HOTS Coverage (L4-L6)",       f"{hots_pct}%",
         "Met" if hots_pct >= 40 else "Not Met"),
        ("Bloom's Progression Score",   f"{ps} / 6.0",
         "Met" if ps >= 2.5 else "Not Met"),
        ("Cognitive Breadth",           f"{la['breadth']} of 6",
         "Met" if la['breadth'] >= 3 else "Not Met"),
        ("Cognitive Balance Index",     f"{bi} / 100",
         "Met" if bi >= 50 else "Not Met"),
        ("Knowledge-Dimension Breadth", f"{kd_cov} of 4",
         "Met" if kd_cov >= 2 else "Not Met"),
    ]


def _la_observations(la: dict) -> list:
    """Return a list of plain-text interpretive observations and recommendations."""
    obs = []
    hots_pct = la['hots_pct']
    lots_pct = la['lots_pct']

    if hots_pct >= 40:
        obs.append(
            f"Strong higher-order focus: {hots_pct}% of COs target Analyze, Evaluate "
            f"or Create, indicating healthy higher-order thinking coverage for OBE."
        )
    elif hots_pct >= 25:
        obs.append(
            f"Moderate higher-order focus: {hots_pct}% HOTS coverage. Reframe one or two "
            f"lower-order COs toward Analyze/Evaluate/Create to strengthen higher-order "
            f"thinking."
        )
    else:
        obs.append(
            f"Low higher-order focus: only {hots_pct}% of COs reach HOTS (L4-L6); the "
            f"course is weighted toward lower-order skills ({lots_pct}% LOTS). Add or "
            f"rewrite COs to strengthen Analyze, Evaluate and Create."
        )

    if la['dominant_level'] and la['total_mapped']:
        dl  = la['dominant_level']
        idx = _LA_LEVEL_ORDER.index(dl) + 1
        n   = la['level_counts'][dl]
        obs.append(
            f"The cognitive profile peaks at L{idx} - {_LA_LEVEL_NAMES[dl]} "
            f"({n} CO(s)), the course's dominant thinking level."
        )

    if la['empty_levels']:
        names = ", ".join(_LA_LEVEL_NAMES[lv] for lv in la['empty_levels'])
        obs.append(
            f"No COs were mapped to: {names}. Coverage gaps at these levels limit the "
            f"cognitive range and should be reviewed against the syllabus depth."
        )
    else:
        obs.append(
            "All six Bloom's levels are represented across the COs - an ideal cognitive spread."
        )

    bi = la['balance_index']
    if bi >= 70:
        obs.append(
            f"Cognitive balance is well distributed (balance index {bi}/100); COs are "
            f"spread evenly rather than clustering at a single level."
        )
    elif bi >= 40:
        obs.append(
            f"Cognitive distribution is moderately balanced (balance index {bi}/100); "
            f"some levels carry noticeably more COs than others."
        )
    else:
        obs.append(
            f"COs cluster around a few levels (balance index {bi}/100); broadening the "
            f"spread across more Bloom's levels would improve assessment depth."
        )

    # Note: not every knowledge dimension needs to be covered for a given course, so
    # uncovered dimensions are reported descriptively elsewhere rather than flagged
    # as a gap here.
    return obs


def _learning_analytics_intro(title, code, la) -> str:
    """Course-specific Learning Analytics introduction.

    Replaces the previous generic boilerplate with a sentence that cites this
    course's actual cognitive profile (levels used, dominant level, LOTS/HOTS split
    and knowledge-dimension breadth) so the section reads as analysis of THIS course.
    """
    total   = la.get('total_mapped', 0)
    breadth = la.get('breadth', 0)
    dom     = la.get('dominant_level')
    dom_name = _LA_LEVEL_NAMES.get(dom, '-') if dom else '-'
    hots    = la.get('hots_pct', 0)
    lots    = la.get('lots_pct', 0)
    kd_cov  = len(la.get('kdim_covered', []))
    return (
        f"For {title} ({code}), this section profiles the cognitive design of the "
        f"{total} mapped Course Outcome(s) using the Revised Bloom's Taxonomy (RBT). "
        f"These COs span {breadth} of the 6 cognitive levels and peak at "
        f"{dom_name}, with a {lots}% Lower-Order (LOTS) to {hots}% Higher-Order (HOTS) "
        f"split across {kd_cov} of the 4 knowledge dimensions. The metrics below detail "
        f"this course's cognitive-level distribution, LOTS/HOTS balance, knowledge-"
        f"dimension spread, cognitive balance, and OBE alignment - showing whether "
        f"{code} builds from recall toward analysis, evaluation and creation, and where "
        f"cognitive gaps remain for OBE accreditation."
    )


# ── CO placement rationale (why each CO sits in its RBT cell) ─────────────────

_KDIM_RATIONALE = {
    'Factual':        "recall of terminology, specific facts and the basic elements of the subject",
    'Conceptual':     "interrelationships among elements - classifications, principles, theories and models",
    'Procedural':     "methods, techniques, algorithms and how to carry out a process",
    'Meta-Cognitive': "strategic and self-reflective knowledge about one's own learning and approach",
}

_LEVEL_RATIONALE = {
    'L1': "recall or recognise information",
    'L2': "explain ideas and interpret meaning",
    'L3': "apply knowledge to solve problems in new situations",
    'L4': "break content into parts and examine how they relate",
    'L5': "judge, justify or critique against defined criteria",
    'L6': "design, formulate or produce original work",
}

_LEVEL_VERB_KEY = {
    'L1': 'remember', 'L2': 'understand', 'L3': 'apply',
    'L4': 'analyse',  'L5': 'evaluate',   'L6': 'create',
}

# Broad set of measurable action verbs used to judge CO Formulation. Includes the
# verbs the CO generator itself is restricted to (both -ize/-ise spellings) so a
# correctly generated CO is never falsely flagged as poorly formulated.
_MEASURABLE_VERBS = {v for verbs in _BLOOM_LEVELS.values() for v in verbs} | {
    'recognize', 'recognise', 'recall', 'interpret', 'exemplify', 'classify',
    'summarize', 'summarise', 'infer', 'compare', 'explain', 'execute', 'implement',
    'differentiate', 'organize', 'organise', 'attribute', 'check', 'critique',
    'generate', 'plan', 'produce', 'analyze', 'analyse', 'evaluate', 'create',
    'apply', 'design', 'develop', 'construct', 'formulate', 'illustrate',
    'demonstrate', 'calculate', 'compute', 'solve', 'describe', 'define', 'list',
    'identify', 'estimate', 'derive', 'model', 'assess',
}


def _qqi_label(score) -> str:
    """Map a 0-100 quality score to a qualitative rating.

    Three-band scale (shown alongside the numeric score, not instead of it):
        Good    : > 75
        Average : 60-75
        Poor    : < 60
    """
    try:
        s = float(score)
    except (TypeError, ValueError):
        return "—"
    if s > 75:
        return "Good"
    if s >= 60:
        return "Average"
    return "Poor"


def _qqi_score_label(score) -> str:
    """Render a 0-100 quality score as 'NN/100 (Band)' so both the number and the
    qualitative band are always shown together."""
    label = _qqi_label(score)
    try:
        return f"{round(float(score))}/100 ({label})"
    except (TypeError, ValueError):
        return label


def _obe_band(score) -> str:
    """Map a 0-100 OBE compliance score to a qualitative band only (no number).

    Reuses the three-band thresholds of _qqi_label, but the lowest band is
    labelled 'Needs Improvement' rather than 'Poor'.
    """
    band = _qqi_label(score)
    return "Needs Improvement" if band == "Poor" else band


def _co_assessment_count(key, analytics, co_tally, is_lab) -> int:
    """Per-CO item count for the CO-coverage table.

    For a lab course the assessment items are experiments (held in
    analytics['co_marks']); for a written course they are the 2/5/10-mark
    exam questions in co_tally. Using the lab source keeps this table in step
    with the lab-aware coverage summary instead of always reporting 0.
    """
    if is_lab:
        return analytics.get("co_marks", {}).get(key, 0)
    t = co_tally.get(key, {})
    return t.get(2, 0) + t.get(5, 0) + t.get(10, 0)


# ── Accreditation improvement guidance ───────────────────────────────────────

# Statuses that are considered healthy and need no improvement prompt.
_OK_METRIC_STATUSES = {"excellent", "strong", "good", "ready"}

# Keyword groups -> (regeneration target, concrete improvement advice) for a weak
# metric. 'target' decides what an in-app "Apply" rebuilds: the COs or the QB.
# Order matters: more specific keyword groups are listed first (first match wins).
_IMPROVEMENT_TIPS = [
    (("formulation", "co statement", "statement quality"), "cos",
     "Rewrite the CO statements using strong, measurable action verbs with a clear "
     "condition and standard, so every outcome is specific and assessable."),
    (("higher-order", "higher order", "hots"), "cos",
     "Elevate selected Course Outcomes to higher-order Bloom levels (Analyze, Evaluate, "
     "Create) using stronger action verbs where the syllabus supports deeper cognition."),
    (("knowledge dimension", "knowledge depth"), "cos",
     "Spread the COs across the knowledge dimensions (Factual, Conceptual, Procedural, "
     "Meta-Cognitive) so the course is not concentrated in a single dimension."),
    (("marks distribution", "marks balance", "marks"), "qbank",
     "Even out the marks allocated to each CO so no outcome is over- or under-weighted; "
     "add questions for the under-weighted COs."),
    (("coverage",), "qbank",
     "Add questions to every CO - especially any that are uncovered - so each outcome "
     "carries a balanced mix of 2-, 5- and 10-mark items."),
    (("bloom", "alignment", "cognitive"), "qbank",
     "Match each question's demand to its CO's stated Bloom level and add more "
     "higher-order (Analyze/Evaluate/Create) questions where those levels are thin."),
    (("difficulty", "balance"), "qbank",
     "Rebalance the easy/medium/hard mix so the paper is neither too easy nor too hard, "
     "aiming for a graded spread of marks across every CO."),
    (("scenario", "case", "integration", "application", "real-world"), "qbank",
     "Add more scenario, case-study and application-based questions that place concepts "
     "in real-world or industry contexts."),
    (("clarity", "assessment", "rubric"), "qbank",
     "Sharpen question wording - state the task, the expected output and the marks "
     "clearly - and remove ambiguity so students know exactly what is required."),
    (("industry", "relevance", "employab", "sdg"), "qbank",
     "Link outcomes and questions to current industry practice and tools, and map them "
     "to relevant programme outcomes and SDGs."),
    (("quality", "qqi", "variety"), "qbank",
     "Increase the number and variety of questions per CO and balance the mark "
     "distribution to lift overall question quality."),
]

_FALLBACK_TARGET = "qbank"
_FALLBACK_TIP = ("Review this area against OBE expectations and apply targeted "
                 "improvements, drawing on the recommendations and action items above.")


def _metric_needs_work(status) -> bool:
    return str(status).strip().lower() not in _OK_METRIC_STATUSES


def _lookup_improvement(metric_name: str):
    """Return (target, tip) for a metric name via keyword match."""
    m = str(metric_name).lower()
    for keys, target, tip in _IMPROVEMENT_TIPS:
        if any(k in m for k in keys):
            return target, tip
    return _FALLBACK_TARGET, _FALLBACK_TIP


def _improvement_tip(metric_name: str) -> str:
    return _lookup_improvement(metric_name)[1]


def _weak_metric_tips(key_metrics: list) -> list:
    """Return (metric, status, how-to-improve) rows for every metric not yet 'Good'."""
    rows = []
    for mtr in key_metrics or []:
        status = mtr.get("status", "")
        if _metric_needs_work(status):
            rows.append((mtr.get("metric", ""), status, _improvement_tip(mtr.get("metric", ""))))
    return rows


def _evenness_score(values: list) -> int:
    """0-100 score for how evenly a set of non-negative values is distributed
    (100 = perfectly even). Used for marks/difficulty balance metrics."""
    vals = [v for v in values]
    tot  = sum(vals)
    n    = len(vals)
    if not tot or n <= 1:
        return 0
    ideal = 1 / n
    devs  = sum(abs(v / tot - ideal) for v in vals)
    max_dev = 2 * (1 - ideal)   # all mass in one bucket
    return round(100 * (1 - devs / max_dev)) if max_dev else 100


def build_key_metrics(analytics: dict, la: dict = None, cos: list = None) -> list:
    """Deterministic key metrics spanning the WHOLE Module 1 report, derived from the
    rule-based analytics (not the AI's free-form output).

    Fixed metric names keep the dashboard and the in-app improvement panel STABLE
    across rebuilds, so applying a fix measurably moves the right metric and the weak
    list converges. Every name matches a keyword in _IMPROVEMENT_TIPS (no generic
    fallback). Coverage by report section:
      CO Coverage (S5), Question Quality (S8), Difficulty Balance (S6),
      Scenario Integration (S7), Higher-Order Thinking + Knowledge Dimension
      Coverage (S1/S4 Learning Analytics), CO Marks Balance (S9),
      CO Formulation (S1).
    """
    metrics = []

    cov = analytics.get("coverage_pct", 0) or 0
    metrics.append({"metric": "CO Coverage", "score": round(cov),
                    "status": _qqi_label(cov)})

    qqi = analytics.get("overall_qqi", 0) or 0
    metrics.append({"metric": "Question Quality", "score": round(qqi),
                    "status": _qqi_label(qqi)})

    # Higher-order thinking and knowledge-dimension spread (Learning Analytics)
    if la:
        hots = la.get("hots_pct", 0) or 0
        metrics.append({"metric": "Higher-Order Thinking", "score": round(hots),
                        "status": _qqi_label(hots)})
        kd_cov   = len(la.get("kdim_covered", []))
        kd_score = round(kd_cov / len(_LA_KDIMS_ORDER) * 100)
        metrics.append({"metric": "Knowledge Dimension Coverage", "score": kd_score,
                        "status": _qqi_label(kd_score)})

    # CO Formulation: share of COs whose statement uses a recognised, measurable verb
    if cos:
        good = sum(
            1 for c in cos
            if any(re.search(r'\b' + re.escape(v) + r'\b', (c.get("statement") or "").lower())
                   for v in _MEASURABLE_VERBS)
        )
        cf = round(good / len(cos) * 100)
        metrics.append({"metric": "CO Formulation", "score": cf,
                        "status": _qqi_label(cf)})

    if not analytics.get("is_lab"):
        # Difficulty balance: how evenly exam questions split across 2/5/10-mark tiers.
        bal = _evenness_score([analytics.get("total_2", 0),
                               analytics.get("total_5", 0),
                               analytics.get("total_10", 0)])
        metrics.append({"metric": "Difficulty Balance", "score": bal,
                        "status": _qqi_label(bal)})

        # Scenario integration: scaled so ~40% scenario questions reaches full score.
        sc_pct   = analytics.get("scenario_pct", 0) or 0
        sc_score = min(100, round(sc_pct / 40 * 100)) if sc_pct else 0
        metrics.append({"metric": "Scenario Integration", "score": sc_score,
                        "status": _qqi_label(sc_score)})

        # CO Marks Balance: how evenly QB marks are spread across the COs (Section 9).
        co_marks = analytics.get("co_marks", {}) or {}
        if co_marks:
            mb = _evenness_score(list(co_marks.values()))
            metrics.append({"metric": "CO Marks Balance", "score": mb,
                            "status": _qqi_label(mb)})

    return metrics


def weak_metric_actions(key_metrics: list) -> list:
    """Like _weak_metric_tips, but each item also carries the regeneration target
    ('cos' or 'qbank') so the in-app preview can offer an actionable Apply button."""
    actions = []
    for mtr in key_metrics or []:
        status = mtr.get("status", "")
        if _metric_needs_work(status):
            metric = mtr.get("metric", "")
            target, tip = _lookup_improvement(metric)
            actions.append({"metric": metric, "status": status,
                            "tip": tip, "target": target})
    return actions


def _co_placement_rationale(cos: list, grid: dict) -> list:
    """Explain why each CO sits in its RBT cell.

    Returns rows of (CO, "Dimension / Ln-Level", rationale) reconstructed from the
    action verb in the CO statement (cognitive level) and the knowledge type
    (knowledge dimension) recorded in the taxonomy grid. The rationale names the
    Bloom's action verb, the cognitive-level name and the knowledge-dimension name
    explicitly, and suggests level-appropriate verbs to refine the CO without
    changing the number of COs.
    """
    if not grid:
        return []
    stmt_by_co = {f"CO{c['num']}": c.get('statement', '') for c in cos}
    rows = []
    seen = set()
    for kd in _LA_KDIMS_ORDER:
        levels = grid.get(kd, {})
        for lv in _LA_LEVEL_ORDER:
            for co in levels.get(lv, []):
                if co in seen:
                    continue
                seen.add(co)
                stmt  = stmt_by_co.get(co, "")
                t     = stmt.lower()
                verbs = _BLOOM_LEVELS.get(_LEVEL_VERB_KEY[lv], [])
                matched = next(
                    (v for v in verbs if re.search(r'\b' + re.escape(v) + r'\b', t)),
                    None,
                )
                idx     = _LA_LEVEL_ORDER.index(lv) + 1
                lname   = _LA_LEVEL_NAMES[lv]          # e.g. "Understand"
                suggest = ", ".join(verbs[:3]) if verbs else ""
                if matched:
                    verb_part = (f'The Bloom\'s action verb "{matched}" signals the '
                                 f'{lname} cognitive level (L{idx}), where students '
                                 f'{_LEVEL_RATIONALE[lv]}')
                else:
                    verb_part = (f'The outcome\'s cognitive demand maps to the {lname} '
                                 f'level (L{idx}), where students {_LEVEL_RATIONALE[lv]}')
                knowledge_part = (f'It engages the {kd} knowledge dimension - '
                                  f'{_KDIM_RATIONALE.get(kd, "subject knowledge")}')
                suggest_part = (f' Suggested {lname}-level verbs to keep this CO measurable: '
                                f'{suggest}.' if suggest else "")
                rationale = f"{verb_part}. {knowledge_part}.{suggest_part}"
                rows.append((co, f"{kd} / L{idx}-{lname}", rationale))
    return rows


# ── Quantitative analytics (no AI needed) ────────────────────────────────────

def _compute_lab_analytics(cos: list, lab_tally: dict) -> dict:
    """Compute analytics for a lab course using experiment counts per CO."""
    co_nums   = {str(c["num"]) for c in cos}
    covered   = {k for k in lab_tally if lab_tally[k].get("exp", 0) > 0}
    uncovered = co_nums - covered
    coverage_pct = round(len(covered) / len(co_nums) * 100, 1) if co_nums else 0

    total_exp = sum(v.get("exp", 0) for v in lab_tally.values())

    co_marks  = {k: v.get("exp", 0) for k, v in lab_tally.items()}
    total_marks = sum(co_marks.values()) or 1

    qqi_scores = {}
    for c in cos:
        key   = str(c["num"])
        n_exp = lab_tally.get(key, {}).get("exp", 0)
        adequacy = min(100, n_exp / max(1, 3) * 100)   # expect ~3 experiments per CO
        qqi_scores[f"CO{key}"] = round(adequacy)
    overall_qqi = round(sum(qqi_scores.values()) / max(len(qqi_scores), 1))

    return {
        "covered": sorted(covered),
        "uncovered": sorted(uncovered),
        "coverage_pct": coverage_pct,
        "total_2": 0, "total_5": 0, "total_10": 0,
        "total_assign": 0, "total_quiz": 0,
        "total_exam": total_exp,
        "total_exp": total_exp,
        "diff_dist": {"Lab Experiments": total_exp},
        "co_marks": co_marks,
        "total_marks": total_marks,
        "scenario_count": 0, "direct_count": 0,
        "unclassified": 0, "total_analyzed": 0, "scenario_pct": 0,
        "qqi_scores": qqi_scores,
        "overall_qqi": overall_qqi,
        "is_lab": True,
    }


def compute_analytics(cos: list, co_tally: dict, raw_qb_text: str = "",
                      is_lab: bool = False, lab_tally: dict = None) -> dict:
    """Return all rule-based metrics derived from COs and QB tally."""
    if is_lab and lab_tally:
        return _compute_lab_analytics(cos, lab_tally)

    co_nums = {str(c["num"]) for c in cos}

    # --- CO Coverage ---
    covered = {k for k in co_tally if co_tally[k].get(2, 0) + co_tally[k].get(5, 0) + co_tally[k].get(10, 0) > 0}
    uncovered = co_nums - covered
    coverage_pct = round(len(covered) / len(co_nums) * 100, 1) if co_nums else 0

    # --- Difficulty distribution ---
    total_2  = sum(v.get(2,  0) for v in co_tally.values())
    total_5  = sum(v.get(5,  0) for v in co_tally.values())
    total_10 = sum(v.get(10, 0) for v in co_tally.values())
    total_assign = sum(v.get("assign", 0) for v in co_tally.values())
    total_quiz   = sum(v.get("quiz",   0) for v in co_tally.values())
    total_exam = total_2 + total_5 + total_10

    # Difficulty labels: 2-mark→Easy, 5-mark→Medium, 10-mark→Hard
    diff_dist = {
        "Easy (2-mark)":   total_2,
        "Medium (5-mark)": total_5,
        "Hard (10-mark)":  total_10,
    }

    # --- CO-wise marks distribution ---
    co_marks = {}
    for k, v in co_tally.items():
        co_marks[k] = v.get(2, 0) * 2 + v.get(5, 0) * 5 + v.get(10, 0) * 10
    total_marks = sum(co_marks.values()) or 1

    # --- Scenario analysis (from raw QB text) ---
    # Count ALL Q\d+. lines (exam + quiz) so unclassified stays non-negative
    scenario_count = 0
    direct_count   = 0
    total_analyzed = 0
    if raw_qb_text:
        for line in raw_qb_text.split("\n"):
            s = line.strip()
            if re.match(r"Q\d+\.", s):
                total_analyzed += 1
                text = re.sub(r"Q\d+\.\s*", "", s)
                text = re.sub(r"\[CO\d+\]", "", text).strip()
                if _is_scenario(text):
                    scenario_count += 1
                elif _is_direct(text):
                    direct_count += 1
    unclassified = max(0, total_analyzed - scenario_count - direct_count)
    scenario_pct = round(scenario_count / total_analyzed * 100, 1) if total_analyzed else 0

    # --- Rule-based QQI per CO ---
    qqi_scores = {}
    for c in cos:
        key = str(c["num"])
        t   = co_tally.get(key, {})
        n2, n5, n10 = t.get(2, 0), t.get(5, 0), t.get(10, 0)
        total = n2 + n5 + n10
        # Score components (0–100 each)
        adequacy    = min(100, total / max(1, 12 / max(len(cos), 1)) * 100)  # expect ~12 qs total / n_cos
        variety     = min(100, sum(1 for x in [n2, n5, n10] if x > 0) / 3 * 100)
        mark_balance = 100 - abs(n2 - n5) * 5 - abs(n5 - n10) * 5
        mark_balance = max(0, mark_balance)
        qqi_scores[f"CO{key}"] = round((adequacy * 0.4 + variety * 0.35 + mark_balance * 0.25))
    overall_qqi = round(sum(qqi_scores.values()) / max(len(qqi_scores), 1))

    return {
        "covered": sorted(covered),
        "uncovered": sorted(uncovered),
        "coverage_pct": coverage_pct,
        "total_2": total_2,
        "total_5": total_5,
        "total_10": total_10,
        "total_assign": total_assign,
        "total_quiz": total_quiz,
        "total_exam": total_exam,
        "diff_dist": diff_dist,
        "co_marks": co_marks,
        "total_marks": total_marks,
        "scenario_count": scenario_count,
        "direct_count": direct_count,
        "unclassified": unclassified,
        "total_analyzed": total_analyzed,
        "scenario_pct": scenario_pct,
        "qqi_scores": qqi_scores,
        "overall_qqi": overall_qqi,
    }


# ── Question quality vs POs and SDGs (Bloom sub-element heuristic) ────────────
#
# Module 1 has no CO-PO / CO-SDG mapping in session, so these two analyses are
# self-contained: each question inherits its CO's Revised-Bloom cognitive level
# (from the taxonomy grid) and its CO statement is scanned for the Bloom keyword
# sub-element. POs are scored by whether the question set reaches the cognitive
# level each Programme Outcome demands; SDGs by theme keywords in the CO
# statements weighted by higher-order (Apply+) question share.

_LEVEL_NUM = {lv: i + 1 for i, lv in enumerate(_LA_LEVEL_ORDER)}

# Bloom keyword sub-elements per RBT cognitive level (the cognitive processes
# that sit under each of the six levels — shown to the reader as the "keywords").
_BLOOM_SUBELEMENTS = {
    'L1': ['Recognising', 'Recalling', 'Defining', 'Listing', 'Naming', 'Stating', 'Identifying'],
    'L2': ['Interpreting', 'Exemplifying', 'Classifying', 'Summarising', 'Inferring', 'Comparing', 'Explaining'],
    'L3': ['Executing', 'Implementing', 'Solving', 'Using', 'Demonstrating', 'Calculating', 'Applying'],
    'L4': ['Differentiating', 'Organising', 'Attributing', 'Analysing', 'Examining', 'Distinguishing'],
    'L5': ['Checking', 'Critiquing', 'Evaluating', 'Justifying', 'Assessing', 'Judging', 'Defending'],
    'L6': ['Generating', 'Planning', 'Producing', 'Designing', 'Developing', 'Constructing', 'Formulating'],
}

# All 12 NBA Programme Outcomes, each handled by how a question bank can assess it:
#   'cognitive' : scored by the minimum Revised-Bloom level the PO demands.
#   'thematic'  : scored by content-theme keywords in the COs (society / environment
#                 / ethics) - assessable by questions but content- not level-driven.
#   'process'   : behavioural outcomes (teamwork, communication, project management)
#                 that a written/lab question bank cannot assess; reported as assessed
#                 through projects, labs, seminars and internships instead.
# (Descriptors mirror the NBA PO set in generate_po_mapping.py.)
_PO_DEFINITIONS = [
    ('PO1',  'Engineering Knowledge',           'cognitive', 'L3'),  # Apply
    ('PO2',  'Problem Analysis',                'cognitive', 'L4'),  # Analyze
    ('PO3',  'Design/Development of Solutions', 'cognitive', 'L6'),  # Create
    ('PO4',  'Conduct Investigations',          'cognitive', 'L4'),  # Analyze
    ('PO5',  'Modern Tool Usage',               'cognitive', 'L3'),  # Apply
    ('PO6',  'The Engineer and Society',        'thematic',
     ['societ', 'social', 'community', 'public', 'cultural', 'legal', 'welfare', 'human']),
    ('PO7',  'Environment and Sustainability',  'thematic',
     ['environment', 'sustainab', 'green', 'energy', 'climate', 'emission', 'ecolog', 'pollution', 'waste', 'renewable']),
    ('PO8',  'Ethics',                          'thematic',
     ['ethic', 'moral', 'integrity', 'plagiar', 'responsib', 'norms', 'professional conduct']),
    ('PO9',  'Individual and Team Work',        'process',   None),
    ('PO10', 'Communication',                   'process',   None),
    ('PO11', 'Project Management and Finance',  'process',   None),
    ('PO12', 'Life-long Learning',              'cognitive', 'L2'),  # Understand
]

# SDG themes detected from CO-statement keywords (lower-cased substring match).
# Keyword sets are deliberately discipline-agnostic: foundational science, maths
# and computing courses contribute to the SDGs through data, modelling, analysis
# and decision-making, so SDG 8/9 capture quantitative and analytical themes, not
# only physical engineering ones. A CO may match several SDGs.
_SDG_THEMES = [
    ('SDG 3',  'Good Health & Well-being',           ['health', 'safety', 'medical', 'hazard', 'risk', 'disease', 'clinical', 'patient', 'biomed', 'epidemi']),
    ('SDG 4',  'Quality Education',                   ['learn', 'educat', 'skill', 'train', 'pedagog', 'literac', 'curriculum', 'teaching', 'foundation']),
    ('SDG 6',  'Clean Water & Sanitation',           ['water', 'sanitation', 'wastewater', 'hydr', 'irrigation']),
    ('SDG 7',  'Affordable & Clean Energy',          ['energy', 'renewable', 'solar', 'power', 'electric', 'grid', 'fuel', 'battery', 'thermal']),
    ('SDG 8',  'Decent Work & Economic Growth',      ['econom', 'finance', 'financial', 'cost', 'productiv', 'business', 'market', 'employ', 'entrepreneur', 'bank', 'invest', 'forecast']),
    ('SDG 9',  'Industry, Innovation & Infrastructure', ['design', 'develop', 'innovat', 'manufactur', 'industr', 'infrastructure', 'construct', 'automat',
                                                         'data', 'statistic', 'probabilit', 'random', 'distribution', 'regress', 'model', 'algorithm',
                                                         'comput', 'analyt', 'analys', 'optimis', 'optimiz', 'simulat', 'network', 'graph', 'digital',
                                                         'software', 'machine', 'technolog', 'research', 'signal', 'control', 'logic']),
    ('SDG 11', 'Sustainable Cities & Communities',   ['urban', 'city', 'transport', 'building', 'housing', 'community', 'traffic', 'smart city']),
    ('SDG 12', 'Responsible Consumption & Production', ['waste', 'recycl', 'efficien', 'material', 'resource', 'optimis', 'optimiz', 'reuse', 'lifecycle']),
    ('SDG 13', 'Climate Action',                     ['climate', 'emission', 'carbon', 'environment', 'green', 'pollution', 'sustainab', 'ecolog']),
]


def _co_subelement(statement: str, level: str) -> str:
    """Return the Bloom keyword sub-element evidenced by a CO statement.

    Prefers the actual action verb found in the statement (rendered as its
    gerund/keyword form); falls back to the level's representative sub-element.
    """
    t = (statement or '').lower()
    verbs = _BLOOM_LEVELS.get(_LEVEL_VERB_KEY.get(level, ''), [])
    for v in verbs:
        if re.search(r'\b' + re.escape(v) + r'\b', t):
            # Render as a keyword sub-element (e.g. "design" -> "Designing").
            base = v[:-1] + 'ing' if v.endswith('e') else v + 'ing'
            return base.capitalize()
    subs = _BLOOM_SUBELEMENTS.get(level, [])
    return subs[0] if subs else '-'


def _co_question_counts(cos: list, co_tally: dict, analytics: dict, is_lab: bool) -> dict:
    """Questions (or lab experiments) attributed to each CO -> {'CO1': n, ...}."""
    qn = {}
    for c in cos:
        key = str(c['num'])
        if is_lab:
            qn[f"CO{c['num']}"] = analytics.get('co_marks', {}).get(key, 0)
        else:
            t = co_tally.get(key, {})
            qn[f"CO{c['num']}"] = t.get(2, 0) + t.get(5, 0) + t.get(10, 0)
    return qn


def _po_sdg_rating(pct: float) -> str:
    """Coverage-share rating for a single PO/SDG (gentler bands than _qqi_label,
    since few question sets put 60%+ of items at any one higher-order level)."""
    if pct >= 40:
        return "Good"
    if pct >= 15:
        return "Average"
    return "Poor"


def compute_po_quality(cos: list, co_tally: dict, analytics: dict,
                       taxonomy_grid: dict, is_lab: bool = False) -> tuple:
    """Score question quality against all 12 Programme Outcomes.

    Returns (rows, summary). Each row is a dict with the PO, the basis used to
    assess it, the item count/share, the Bloom sub-elements that carry it, and a
    rating. Cognitive and thematic POs are assessable by the question bank; process
    POs are reported as assessed through other components (not scored).
    """
    co_lvl = _compute_bloom_analytics(cos, taxonomy_grid)['co_to_level']
    qn = _co_question_counts(cos, co_tally, analytics, is_lab)
    total_q = sum(qn.values()) or 1

    rows = []
    addressed = 0
    assessable = 0
    for po, name, mode, spec in _PO_DEFINITIONS:
        if mode == 'process':
            rows.append({
                'po': po, 'name': name, 'mode': mode,
                'basis': 'Project / Lab / Seminar',
                'count_str': '-', 'pct': None, 'pct_str': '-',
                'verbs': '-', 'rating': 'Indirect',
            })
            continue

        assessable += 1
        n = 0
        verbs = []
        if mode == 'cognitive':
            exp_n = _LEVEL_NUM[spec]
            basis = f"{spec}-{_LA_LEVEL_NAMES[spec]}"
            for c in cos:
                key = f"CO{c['num']}"
                lv = co_lvl.get(key)
                if lv and _LEVEL_NUM[lv] >= exp_n and qn.get(key, 0) > 0:
                    n += qn[key]
                    sv = _co_subelement(c['statement'], lv)
                    if sv and sv not in verbs:
                        verbs.append(sv)
        else:  # thematic
            basis = 'Content theme'
            for c in cos:
                t = (c['statement'] or '').lower()
                if any(k in t for k in spec) and qn.get(f"CO{c['num']}", 0) > 0:
                    n += qn[f"CO{c['num']}"]
                    lv = co_lvl.get(f"CO{c['num']}")
                    if lv:
                        sv = _co_subelement(c['statement'], lv)
                        if sv and sv not in verbs:
                            verbs.append(sv)

        pct = round(n / total_q * 100)
        if n > 0:
            addressed += 1
        rows.append({
            'po': po, 'name': name, 'mode': mode,
            'basis': basis,
            'count_str': str(n), 'pct': pct, 'pct_str': f"{pct}%",
            'verbs': ', '.join(verbs) or '-',
            'rating': _po_sdg_rating(pct) if n > 0 else 'Poor',
        })
    summary = {
        'addressed': addressed, 'assessable': assessable, 'total': len(_PO_DEFINITIONS),
        'process': [r['po'] for r in rows if r['mode'] == 'process'],
        'total_q': total_q,
    }
    return rows, summary


def compute_sdg_quality(cos: list, co_tally: dict, analytics: dict,
                        taxonomy_grid: dict, is_lab: bool = False) -> tuple:
    """Score question quality against the SDGs the COs touch.

    Returns (rows, summary). Only SDGs with at least one matching CO are listed.
    Quality reflects the higher-order (Apply+) share of the SDG-relevant questions,
    since meaningful SDG competency needs application, analysis or design - not recall.
    """
    co_lvl = _compute_bloom_analytics(cos, taxonomy_grid)['co_to_level']
    qn = _co_question_counts(cos, co_tally, analytics, is_lab)
    total_q = sum(qn.values()) or 1

    rows = []
    for sdg, name, keys in _SDG_THEMES:
        matched = []
        q = 0
        hots = 0
        verbs = []
        for c in cos:
            t = (c['statement'] or '').lower()
            if any(k in t for k in keys):
                key = f"CO{c['num']}"
                matched.append(key)
                q += qn.get(key, 0)
                lv = co_lvl.get(key)
                if lv and _LEVEL_NUM[lv] >= 3:
                    hots += qn.get(key, 0)
                if lv:
                    sv = _co_subelement(c['statement'], lv)
                    if sv and sv not in verbs:
                        verbs.append(sv)
        if not matched:
            continue
        hots_pct = round(hots / q * 100) if q else 0
        if q == 0:
            rating = "Poor"
        elif hots_pct >= 50:
            rating = "Good"
        elif hots_pct >= 25:
            rating = "Average"
        else:
            rating = "Poor"
        rows.append({
            'sdg': sdg, 'name': name,
            'cos': ', '.join(matched),
            'count': q, 'pct': round(q / total_q * 100),
            'hots_pct': hots_pct,
            'verbs': ', '.join(verbs) or '-',
            'rating': rating,
        })
    summary = {'covered': len(rows), 'total': len(_SDG_THEMES), 'total_q': total_q}
    return rows, summary


def _po_quality_intro(title: str, code: str, summary: dict) -> str:
    proc = ", ".join(summary.get('process', [])) or "none"
    return (
        f"This section assesses how the {summary['total_q']} assessment item(s) for "
        f"{title} ({code}) exercise all {summary['total']} Programme Outcomes (POs). "
        f"Cognitive POs (PO1-PO5, PO12) are scored by whether questions reach the Bloom "
        f"level the PO demands, evidenced by the keyword sub-elements (e.g. Analysing, "
        f"Designing) in the mapped COs; thematic POs (PO6 Society, PO7 Environment, PO8 "
        f"Ethics) are scored by the relevant content in the COs. Process POs ({proc}) - "
        f"teamwork, communication and project management - are assessed through projects, "
        f"labs and seminars, not the written question bank. Of the "
        f"{summary['assessable']} question-assessable POs, {summary['addressed']} are "
        f"exercised by the current question set."
    )


def _sdg_quality_intro(title: str, code: str, summary: dict) -> str:
    return (
        f"This section assesses how the question set for {title} ({code}) supports the "
        f"UN Sustainable Development Goals (SDGs). SDG themes are detected from the CO "
        f"statements and weighted by the higher-order (Apply and above) share of their "
        f"questions, since SDG competency requires applying, analysing and designing - "
        f"not recall. {summary['covered']} SDG theme(s) are touched by the current COs."
    )


# ── Build sample question paper text ─────────────────────────────────────────

def build_sample_paper_text(qb_data: dict, code: str, title: str) -> list:
    """Return a list of (indent_level, text) tuples for the sample exam paper."""
    from generate_qpaper import _flat_list

    lines = []
    def add(text="", bold=False, indent=0):
        lines.append((indent, text, bold))

    add(f"{code} — {title}", bold=True)
    add("INTERNAL ASSESSMENT EXAMINATION (SAMPLE)", bold=True)
    add(f"Duration: 3 Hours   |   Maximum Marks: 50")
    add()
    add("Instructions:", bold=True)
    add("• Part A: Answer all questions                    (5 × 2 = 10 marks)", indent=1)
    add("• Part B: Answer any 2 out of 3 questions         (2 × 5 = 10 marks)", indent=1)
    add("• Part C: Answer any 1 out of 2 questions         (1 × 10 = 10 marks)", indent=1)
    add()
    add("-" * 70)
    add()
    add("PART A — Short Answer Questions (2 marks each)", bold=True)
    add()

    sel_2 = (_flat_list(qb_data["units"], "exam", 2))[:5]
    for i, (unit, q) in enumerate(sel_2, 1):
        co_tag = f"  [{q.get('co', '')}]" if q.get("co") else ""
        add(f"Q{i}.  {q.get('text', '')} {co_tag}", indent=1)
    add()
    add("-" * 70)
    add()
    add("PART B — Medium Answer Questions (5 marks each)", bold=True)
    add()
    sel_5 = (_flat_list(qb_data["units"], "exam", 5))[:3]
    for i, (unit, q) in enumerate(sel_5, 1):
        co_tag = f"  [{q.get('co', '')}]" if q.get("co") else ""
        add(f"Q{i}.  {q.get('text', '')} {co_tag}", indent=1)
    add()
    add("-" * 70)
    add()
    add("PART C — Long Answer Questions (10 marks each)", bold=True)
    add()
    sel_10 = (_flat_list(qb_data["units"], "exam", 10))[:2]
    for i, (unit, q) in enumerate(sel_10, 1):
        co_tag = f"  [{q.get('co', '')}]" if q.get("co") else ""
        add(f"Q{i}.  {q.get('text', '')} {co_tag}", indent=1)

    return lines


def _build_co_pct_rows(cos: list, co_tally: dict) -> list:
    """Return rows: (co, 2mk_str, 5mk_str, 10mk_str, assign_str, quiz_str, total_str).
    Each count cell shows 'N (X%)' where X% is share of that question type across all COs."""
    total_2    = sum(v.get(2,       0) for v in co_tally.values())
    total_5    = sum(v.get(5,       0) for v in co_tally.values())
    total_10   = sum(v.get(10,      0) for v in co_tally.values())
    total_asgn = sum(v.get("assign",0) for v in co_tally.values())
    total_quiz = sum(v.get("quiz",  0) for v in co_tally.values())

    def _fmt(n, tot):
        return f"{n} ({round(n/tot*100)}%)" if tot else str(n)

    rows = []
    for c in cos:
        key = str(c["num"])
        v   = co_tally.get(key, {})
        n2, n5, n10 = v.get(2,0), v.get(5,0), v.get(10,0)
        na, nq      = v.get("assign",0), v.get("quiz",0)
        rows.append((
            f"CO{key}",
            _fmt(n2,  total_2),
            _fmt(n5,  total_5),
            _fmt(n10, total_10),
            _fmt(na,  total_asgn),
            _fmt(nq,  total_quiz),
            str(n2 + n5 + n10 + na + nq),
        ))
    return rows


def _build_co_lab_pct_rows(cos: list, analytics: dict) -> list:
    """Return rows: (co, experiments, share%) for lab courses."""
    co_marks  = analytics.get("co_marks", {})
    total_exp = analytics.get("total_exp", analytics.get("total_exam", 0))
    rows = []
    for c in cos:
        key   = str(c["num"])
        n_exp = co_marks.get(key, 0)
        pct   = round(n_exp / total_exp * 100) if total_exp else 0
        rows.append((f"CO{key}", str(n_exp), f"{pct}%"))
    return rows


def build_sample_lab_paper_text(lab_data: dict, code: str, title: str) -> list:
    """Return a sample lab practical exam as (indent, text, bold) tuples."""
    from generate_qpaper import _flat_lab_list

    lines = []
    def add(text="", bold=False, indent=0):
        lines.append((indent, text, bold))

    add(f"{code} -- {title}", bold=True)
    add("LAB PRACTICAL EXAMINATION (SAMPLE)", bold=True)
    add("Duration: 3 Hours   |   Maximum Marks: 50")
    add()
    add("Instructions:", bold=True)
    add("* Perform the experiment and record your observations.", indent=1)
    add("* Write the aim, procedure, observations, and result in your lab record.", indent=1)
    add("* A brief viva will be conducted at the end of the session.", indent=1)
    add()
    add("-" * 70)
    add()
    add("PART A -- Practical Experiments", bold=True)
    add("(Attempt any ONE experiment from the list below)", indent=1)
    add()

    flat = _flat_lab_list(lab_data["units"])
    for i, (unit, exp) in enumerate(flat[:5], 1):
        co_tag = f"  [{exp.get('co', '')}]" if exp.get("co") else ""
        exp_title = exp.get("title") or exp.get("aim") or ""
        add(f"Exp {i}.  {exp_title}{co_tag}", indent=1)
    add()

    add("-" * 70)
    add()
    add("PART B -- Viva Voce Questions", bold=True)
    add("(Answer any FIVE questions — 2 marks each)", indent=1)
    add()

    viva_q = []
    for _, exp in flat:
        vivas = exp.get("viva") or []
        for vq in (vivas if isinstance(vivas, list) else []):
            text = vq.get("question", str(vq)) if isinstance(vq, dict) else str(vq)
            co   = exp.get("co", "")
            viva_q.append((text, co))
        if len(viva_q) >= 10:
            break

    if viva_q:
        for i, (text, co) in enumerate(viva_q[:10], 1):
            co_tag = f"  [{co}]" if co else ""
            add(f"Q{i}.  {text}{co_tag}", indent=1)
    else:
        add("(Viva questions will be asked based on the experiment performed.)", indent=1)

    return lines


# ── AI qualitative analysis call ─────────────────────────────────────────────

def get_ai_analysis(client, cos: list, co_tally: dict, bloom_summary: str,
                    code: str, title: str, analytics: dict) -> dict:
    """Single Claude call returning JSON for all qualitative report sections."""
    co_lines = "\n".join(f"  CO{c['num']}: {c['statement']}" for c in cos)
    tally_lines = "\n".join(
        f"  CO{k}: {v.get(2,0)}×2-mark, {v.get(5,0)}×5-mark, {v.get(10,0)}×10-mark"
        for k, v in co_tally.items()
    )

    prompt = f"""You are an OBE accreditation expert for Indian engineering colleges.

Course: {title} ({code})
Course Outcomes:
{co_lines}

Bloom's Summary: {bloom_summary or "Not available"}

Question Bank Coverage per CO:
{tally_lines}

CO Coverage: {analytics["coverage_pct"]}% ({len(analytics["covered"])} of {len(analytics["covered"]) + len(analytics["uncovered"])} COs covered)
Total Questions: {analytics["total_2"]}×2-mark + {analytics["total_5"]}×5-mark + {analytics["total_10"]}×10-mark = {analytics["total_exam"]} exam questions
Scenario/Case-study questions detected: {analytics["scenario_count"]} ({analytics["scenario_pct"]}%)

Provide a structured analysis in this EXACT JSON format (no markdown, no extra text):
{{
  "blooms_compliance": [
    {{"co": "CO1", "stated_level": "Apply", "qb_level": "Apply/Analyse mix", "rating": "Good", "remark": "..."}}
  ],
  "scenario_analysis": {{
    "assessment": "...(2-3 sentences on scenario coverage)",
    "recommendation": "...(1-2 sentences)"
  }},
  "obe_compliance": {{
    "co_statement_quality": "Excellent|Good|Satisfactory|Needs Improvement",
    "bloom_alignment": "Excellent|Good|Satisfactory|Needs Improvement",
    "co_coverage_status": "Excellent|Good|Satisfactory|Needs Improvement",
    "overall_obe_score": 85,
    "compliance_level": "Fully Compliant|Substantially Compliant|Partially Compliant",
    "strengths": ["...", "..."],
    "gaps": ["...", "..."],
    "recommendations": ["...", "..."]
  }}
}}"""

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=3500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        # Strip markdown fences
        raw = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
        # Find the outermost JSON object bounds robustly
        start = raw.find('{')
        end   = raw.rfind('}')
        if start >= 0 and end > start:
            raw = raw[start:end + 1]
        return json.loads(raw)
    except Exception:
        return {}


# ── DOCX report builder ───────────────────────────────────────────────────────

def build_docx(
    cos: list,
    co_text: str,
    qb_data,            # parsed qb_data dict or None
    co_tally: dict,
    bloom_summary: str,
    analytics: dict,
    ai: dict,
    sample_paper,       # list of (indent, text, bold) tuples or []
    code: str,
    title: str,
    semester,
    output_path: str,
    is_lab: bool = False,
    taxonomy_grid=None,
):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    # ── Helpers ──────────────────────────────────────────────────────────────
    INDIGO   = RGBColor(0x31, 0x2E, 0x81)
    TEAL     = RGBColor(0x06, 0x4E, 0x5B)
    GREEN    = RGBColor(0x06, 0x65, 0x28)
    AMBER    = RGBColor(0x78, 0x35, 0x00)
    RED_C    = RGBColor(0x7F, 0x1D, 0x1D)
    GRAY     = RGBColor(0x37, 0x41, 0x51)
    LIGHT_BG = RGBColor(0xEE, 0xEE, 0xFF)

    STATUS_COLOR = {
        "Excellent": GREEN, "Strong": GREEN, "Good": GREEN, "Ready": GREEN,
        "Satisfactory": AMBER, "Adequate": AMBER, "Average": AMBER, "Mostly Ready": AMBER,
        "Needs Improvement": RED_C, "Needs Work": RED_C, "Poor": RED_C,
        "Fully Compliant": GREEN, "Substantially Compliant": AMBER,
        "Partially Compliant": RED_C,
    }

    def _add_heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color and p.runs:
            p.runs[0].font.color.rgb = color
        return p

    def _add_para(text="", bold=False, italic=False, color=None, size=10, space_after=4, indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(0.3 * indent)
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def _add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            tc = hdr[i]._tc
            tcp = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "DDDDF0")
            tcp.append(shd)
        for row_data in rows:
            row = t.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)
                for run in row[i].paragraphs[0].runs:
                    run.font.size = Pt(9)
        if col_widths:
            for row in t.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    def _insert_chart(buf, width=5.2):
        if buf is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(_fname, width=Inches(width))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        except Exception:
            pass
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    def _add_bar_chart(label_count_pairs, total, label_header="Category", unit_label=""):
        """Visual bar chart table using colored ■□ characters in Courier New."""
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate([label_header, "Visual Distribution", "Count"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True; run.font.size = Pt(9)
            tc = hdr[i]._tc
            tcp = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DDDDF0")
            tcp.append(shd)
        for label, count in label_count_pairs:
            pct = round(count / total * 100) if total else 0
            filled_n = max(0, min(20, pct // 5))
            row = t.add_row().cells
            row[0].text = str(label)
            for run in row[0].paragraphs[0].runs:
                run.font.size = Pt(9)
            # Bar column: two colored runs
            p = row[1].paragraphs[0]
            if filled_n:
                r1 = p.add_run("■" * filled_n)
                r1.font.name = "Courier New"; r1.font.size = Pt(9)
                r1.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)   # indigo
            if filled_n < 20:
                r2 = p.add_run("□" * (20 - filled_n))
                r2.font.name = "Courier New"; r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0xC4, 0xB5, 0xFD)   # light lavender
            pct_run = p.add_run(f"  {pct}%")
            pct_run.font.size = Pt(8.5)
            pct_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            count_str = f"{count}{unit_label}"
            row[2].text = count_str
            for run in row[2].paragraphs[0].runs:
                run.font.size = Pt(9)
        for row in t.rows:
            row.cells[0].width = Inches(2.1)
            row.cells[1].width = Inches(3.4)
            row.cells[2].width = Inches(0.8)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _badge(p, text, color):
        run = p.add_run(f"  {text}  ")
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = color

    def _section_divider(num, label):
        _add_heading(f"Section {num}: {label}", level=1, color=INDIGO)

    # ── Cover page ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("MyOBE — Module 1 Comprehensive Report")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = INDIGO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AI-Powered Question Paper System")
    r2.italic = True; r2.font.size = Pt(12); r2.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    meta = [
        ("Course Code",  code),
        ("Course Title", title),
        ("Semester",     str(semester) if semester else "—"),
        ("Generated on", datetime.date.today().strftime("%d %B %Y")),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        mt.rows[i].cells[0].text = k
        mt.rows[i].cells[1].text = v
        for run in mt.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(10)
        for run in mt.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
    for row in mt.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.1)

    doc.add_page_break()

    # ── Executive Summary ────────────────────────────────────────────────────
    _add_heading("Executive Summary", level=1, color=INDIGO)
    n_cos   = len(cos)
    n_total = analytics["total_exam"]
    cov     = analytics["coverage_pct"]

    if is_lab:
        qb_desc = (
            f"The Lab Question Bank comprises {n_total} experiment(s) "
            f"covering {cov}% of the defined COs. "
        )
    else:
        qb_desc = (
            f"The Intelligent Question Bank comprises {n_total} examination questions "
            f"(2-mark, 5-mark, and 10-mark) covering {cov}% of the defined COs. "
        )
    _add_para(
        f"This report presents the complete Module 1 output for {title} ({code}). "
        f"A total of {n_cos} Course Outcomes were generated and mapped to Bloom's taxonomy. "
        + qb_desc,
        size=10
    )

    key_stats = [
        ("Course Outcomes Generated", str(n_cos)),
        ("Total Exam Questions",      str(n_total)),
        ("Assignments in QB",         str(analytics["total_assign"])),
        ("Quiz Questions in QB",      str(analytics["total_quiz"])),
        ("CO Coverage",               f"{cov}%"),
    ]
    _add_table(["Metric", "Value"], key_stats, col_widths=[3.5, 2.8])

    doc.add_page_break()

    # ── Section 1: Automatic CO Generation ──────────────────────────────────
    _section_divider(1, "Automatic CO Generation")
    _add_heading("Course Outcome Statements", level=2, color=TEAL)

    co_rows = []
    for c in cos:
        stmt = c["statement"]
        level = _detect_bloom(stmt)
        co_rows.append((f"CO{c['num']}", stmt, level))

    _add_table(["CO", "Statement", "Bloom's Level"], co_rows,
               col_widths=[0.6, 4.5, 1.2])

    # Revised Bloom's Taxonomy (RBT) Table — full 4×6 knowledge-dimension matrix
    _add_heading("Revised Bloom's Taxonomy (RBT) Table", level=2, color=TEAL)

    _RBT_LNAMES  = ['REMEMBER', 'UNDERSTAND', 'APPLY', 'ANALYZE', 'EVALUATE', 'CREATE']
    _RBT_LVERBS  = [
        'Recognizing, Recalling',
        'Interpreting, Exemplifying, Classifying, Summarizing, Inferring, Comparing, Explaining',
        'Executing, Implementing',
        'Differentiating, Organizing, Attributing',
        'Checking, Critiquing',
        'Generating, Planning, Producing',
    ]
    _RBT_KDIMS = [
        ("Factual",        "FACTUAL\nKNOWLEDGE"),
        ("Conceptual",     "CONCEPTUAL\nKNOWLEDGE"),
        ("Procedural",     "PROCEDURAL\nKNOWLEDGE"),
        ("Meta-Cognitive", "METACOGNITIVE\nKNOWLEDGE"),
    ]
    _RBT_LEVELS = ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']

    if taxonomy_grid:
        _raw_grid = taxonomy_grid
    else:
        from generate_cos import build_taxonomy_grid as _btg
        _raw_grid = _btg(co_text)
    rbt_grid = {kd: {lv: _raw_grid.get(kd, {}).get(lv, []) for lv in _RBT_LEVELS}
                for kd, _ in _RBT_KDIMS}

    rbt_tbl = doc.add_table(rows=6, cols=7)  # row 0: names, row 1: verbs, rows 2-5: data
    rbt_tbl.style = 'Table Grid'

    def _shd(cell, fill_hex):
        tc  = cell._tc
        tcp = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcp.append(shd)

    # Header row 0 — level numbers and names (bold)
    hdr0 = rbt_tbl.rows[0].cells
    p0   = hdr0[0].paragraphs[0]
    r0   = p0.add_run('KNOWLEDGE\nDIMENSION')
    r0.bold = True; r0.font.size = Pt(8)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0); p0.paragraph_format.space_after = Pt(0)
    _shd(hdr0[0], 'C8C8C8')
    for i, name in enumerate(_RBT_LNAMES, 1):
        p  = hdr0[i].paragraphs[0]
        r  = p.add_run(f'{i}. {name}')
        r.bold = True; r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        _shd(hdr0[i], 'C8C8C8')

    # Header row 1 — verbs (italic, smaller)
    hdr1 = rbt_tbl.rows[1].cells
    _shd(hdr1[0], 'C8C8C8')
    for i, verbs in enumerate(_RBT_LVERBS, 1):
        p  = hdr1[i].paragraphs[0]
        r  = p.add_run(verbs)
        r.italic = True; r.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        _shd(hdr1[i], 'C8C8C8')

    # Data rows (indices 2-5)
    for row_idx, (kd, label) in enumerate(_RBT_KDIMS, 2):
        row = rbt_tbl.rows[row_idx].cells
        pr  = row[0].paragraphs[0]
        rr  = pr.add_run(label); rr.bold = True; rr.font.size = Pt(8)
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col_idx, lv in enumerate(_RBT_LEVELS, 1):
            cell_text = ', '.join(rbt_grid[kd].get(lv, []))
            row[col_idx].text = cell_text
            cp = row[col_idx].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── CO Placement Rationale ───────────────────────────────────────────────
    _rationale_rows = _co_placement_rationale(cos, rbt_grid)
    if _rationale_rows:
        _add_heading("CO Placement Rationale", level=2, color=TEAL)
        _add_para(
            "Why each Course Outcome is mapped to its cell (knowledge dimension x "
            "cognitive level) in the RBT table above. Each rationale names the Bloom's "
            "action verb, its cognitive-level name (Remember-Create) and the "
            "knowledge-dimension name (Factual, Conceptual, Procedural, Meta-Cognitive), "
            "and suggests level-appropriate verbs to strengthen the CO while keeping the "
            "same number of Course Outcomes.",
            italic=True, size=9
        )
        _add_table(
            ["CO", "RBT Cell", "Why placed here"],
            _rationale_rows,
            col_widths=[0.55, 1.75, 4.0]
        )
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Learning Analytics (RBT-derived) ─────────────────────────────────────
    _add_heading("Learning Analytics", level=2, color=TEAL)
    _la = _compute_bloom_analytics(cos, rbt_grid)
    _la_total = _la['total_mapped']
    _add_para(_learning_analytics_intro(title, code, _la), italic=True, size=9)

    _add_heading("A.  Cognitive Level Distribution", level=3, color=TEAL)
    _add_bar_chart(
        [(f"L{i+1} - {_LA_LEVEL_NAMES[lv]}", len(_la['level_co_map'][lv]))
         for i, lv in enumerate(_LA_LEVEL_ORDER)],
        _la_total, label_header="Bloom's Level"
    )
    _level_detail_rows = []
    for _li, _lv in enumerate(_LA_LEVEL_ORDER, 1):
        _co_list = _la['level_co_map'][_lv]
        _pct = round(len(_co_list) / _la_total * 100, 1)
        _level_detail_rows.append((
            f"L{_li} - {_LA_LEVEL_NAMES[_lv]}",
            ', '.join(_co_list) or "—",
            str(len(_co_list)),
            f"{_pct}%",
            "HOTS" if _li >= 4 else "LOTS",
        ))
    _add_table(
        ["Bloom's Level", "COs", "Count", "%", "Type"],
        _level_detail_rows,
        col_widths=[1.5, 2.5, 0.5, 0.6, 0.65]
    )

    _add_heading("B.  Higher vs Lower Order Thinking (HOTS / LOTS)", level=3, color=TEAL)
    _lots_pct = round(_la['lots_cos'] / _la_total * 100, 1)
    _hots_pct = round(_la['hots_cos'] / _la_total * 100, 1)
    _add_table(
        ["Category", "Levels", "CO Count", "%"],
        [
            ("LOTS (Lower Order)", "L1-L3: Remember, Understand, Apply",
             str(_la['lots_cos']), f"{_lots_pct}%"),
            ("HOTS (Higher Order)", "L4-L6: Analyze, Evaluate, Create",
             str(_la['hots_cos']), f"{_hots_pct}%"),
        ],
        col_widths=[1.3, 2.8, 0.85, 0.8]
    )
    _ps = _la['progression_score']
    _ps_label = "Strong" if _ps >= 4.0 else "Moderate" if _ps >= 2.5 else "Low"
    _add_para(
        f"Bloom's Progression Score: {_ps} / 6.0 ({_ps_label} cognitive ambition). "
        "A healthy share of higher-order thinking (L4-L6) supports OBE expectations.",
        size=9, italic=True
    )

    if any(_la['kdim_co_map'].values()):
        _add_heading("C.  Knowledge Dimension Coverage", level=3, color=TEAL)
        _add_para(
            "Distribution of COs across the four RBT knowledge dimensions, showing how "
            "the course balances factual, conceptual, procedural and metacognitive learning.",
            italic=True, size=9
        )
        _kdim_rows = []
        for _kd in _LA_KDIMS_ORDER:
            _kd_cos = _la['kdim_co_map'].get(_kd, [])
            _kd_pct = round(len(_kd_cos) / _la_total * 100, 1)
            _kdim_rows.append((_kd, ', '.join(_kd_cos) or "—", str(len(_kd_cos)), f"{_kd_pct}%"))
        _add_table(
            ["Knowledge Dimension", "COs", "Count", "%"],
            _kdim_rows,
            col_widths=[1.6, 3.0, 0.55, 0.6]
        )
        _add_bar_chart(
            [(_kd, len(_la['kdim_co_map'].get(_kd, []))) for _kd in _LA_KDIMS_ORDER],
            _la_total, label_header="Knowledge Dimension"
        )

    # D. Cognitive Balance & Diversity
    _add_heading("D.  Cognitive Balance & Diversity", level=3, color=TEAL)
    _dom = _la['dominant_level']
    _dom_label = (f"L{_LA_LEVEL_ORDER.index(_dom) + 1} - {_LA_LEVEL_NAMES[_dom]}"
                  if _dom else "—")
    _empty_label = (", ".join(_LA_LEVEL_NAMES[lv] for lv in _la['empty_levels'])
                    if _la['empty_levels'] else "None — all levels covered")
    _bal = _la['balance_index']
    _bal_word = "Well balanced" if _bal >= 70 else "Moderately balanced" if _bal >= 40 else "Concentrated"
    _add_table(
        ["Indicator", "Value"],
        [
            ("Cognitive Breadth (levels used)", f"{_la['breadth']} of 6"),
            ("Dominant Cognitive Level",        _dom_label),
            ("Uncovered Levels",                _empty_label),
            ("Cognitive Balance Index",         f"{_bal} / 100 ({_bal_word})"),
            ("Bloom's Progression Score",       f"{_ps} / 6.0 ({_ps_label})"),
        ],
        col_widths=[3.0, 3.3]
    )

    # E. OBE Alignment Indicators
    _add_heading("E.  OBE Alignment Indicators", level=3, color=TEAL)
    _align_rows = _la_alignment_checks(_la)
    _at = _add_table(
        ["Criterion", "Result", "Status"],
        _align_rows,
        col_widths=[3.0, 1.6, 1.4]
    )
    # Colour the Status column by Met / Not Met
    for _ri, (_, _, _status) in enumerate(_align_rows, 1):
        _cell = _at.rows[_ri].cells[2]
        for _run in _cell.paragraphs[0].runs:
            _run.bold = True
            _run.font.color.rgb = GREEN if _status == "Met" else RED_C
    _met = sum(1 for r in _align_rows if r[2] == "Met")
    _add_para(
        f"Indicators met: {_met} of {len(_align_rows)}.",
        bold=True, size=9.5
    )

    # F. Analysis and Learning Path
    _add_heading("F.  Analysis and Learning Path", level=3, color=TEAL)
    for _ob in _la_observations(_la):
        _add_para(f"• {_ob}", size=9.5, indent=1)

    doc.add_page_break()

    # ── Section 2: Question Bank Creation ────────────────────────────────────
    _section_divider(2, "Lab Question Bank Creation" if is_lab else "Intelligent Question Bank Creation")

    if is_lab:
        total_exp = analytics.get("total_exp", analytics["total_exam"])
        qb_stats  = [
            ("Lab Experiments", str(total_exp)),
            ("Total Lab Work",  str(total_exp)),
        ]
        _add_table(["Category", "Count"], qb_stats, col_widths=[3.5, 2.8])
        _add_heading("CO-wise Experiment Distribution", level=2, color=TEAL)
        _add_para("Count and share (%) of experiments allocated to each CO.", italic=True, size=9)
        lab_pct_rows = _build_co_lab_pct_rows(cos, analytics)
        if lab_pct_rows:
            _add_table(["CO", "Experiments", "Share (%)"],
                       lab_pct_rows, col_widths=[0.8, 1.2, 1.0])
    else:
        qb_stats = [
            ("2-Mark Questions",    str(analytics["total_2"])),
            ("5-Mark Questions",    str(analytics["total_5"])),
            ("10-Mark Questions",   str(analytics["total_10"])),
            ("Assignment Tasks",    str(analytics["total_assign"])),
            ("Quiz Questions",      str(analytics["total_quiz"])),
            ("Total Exam Questions", str(analytics["total_exam"])),
        ]
        _add_table(["Category", "Count"], qb_stats, col_widths=[3.5, 2.8])
        _add_heading("CO-wise Question Distribution", level=2, color=TEAL)
        _add_para("Count and share (%) of each question type allocated to each CO.", italic=True, size=9)
        co_pct_rows = _build_co_pct_rows(cos, co_tally)
        if co_pct_rows:
            _add_table(
                ["CO", "2-Mk (%)", "5-Mk (%)", "10-Mk (%)", "Assign (%)", "Quiz (%)", "Total"],
                co_pct_rows,
                col_widths=[0.55, 0.85, 0.85, 0.85, 0.82, 0.73, 0.6]
            )

    doc.add_page_break()

    # ── Section 3: Sample Paper ───────────────────────────────────────────────
    _section_divider(3, "Lab Practical Examination (Sample)" if is_lab
                     else "Smart Question Paper / Assignment Generation")

    if sample_paper:
        for indent, text, bold in sample_paper:
            if not text:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            if indent:
                p.paragraph_format.left_indent = Inches(0.35 * indent)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(9.5)
    else:
        _add_para("Sample paper could not be generated — Question Bank is empty.", italic=True)

    doc.add_page_break()

    # ── Section 4: Bloom's Taxonomy Compliance Check ─────────────────────────
    _section_divider(4, "Bloom's Taxonomy Compliance Check")

    blooms_rows = []
    for entry in ai.get("blooms_compliance", []):
        rating = entry.get("rating", "—")
        blooms_rows.append((
            entry.get("co", ""),
            entry.get("stated_level", ""),
            entry.get("qb_level", ""),
            rating,
            entry.get("remark", ""),
        ))

    if blooms_rows:
        _add_table(["CO", "Stated Level", "QB Level Coverage", "Rating", "Remark"],
                   blooms_rows, col_widths=[0.6, 1.0, 1.3, 0.9, 2.5])
    else:
        _add_para("Bloom's compliance data not available.", italic=True)

    doc.add_page_break()

    # ── Section 5: CO Coverage Verification ─────────────────────────────────
    _section_divider(5, "CO Coverage Verification")

    cov_set = set(analytics.get("covered", []))
    item_header = "Experiments" if is_lab else "Questions"
    cov_rows = []
    for c in cos:
        key = str(c["num"])
        n   = _co_assessment_count(key, analytics, co_tally, is_lab)
        status = "Covered" if key in cov_set else "Not Covered"
        cov_rows.append((f"CO{key}", c["statement"], str(n), status))

    _add_table(["CO", "Statement", item_header, "Status"],
               cov_rows, col_widths=[0.6, 3.7, 0.9, 1.1])

    _add_para(f"Overall Coverage: {cov}%  |  Covered: {len(analytics['covered'])}  |  "
              f"Not Covered: {len(analytics['uncovered'])}", bold=True, size=9.5)

    doc.add_page_break()

    # ── Section 6: Difficulty Level Analysis ────────────────────────────────
    _section_divider(6, "Difficulty Level Analysis")

    diff_rows = []
    total_q = analytics["total_exam"] or 1
    for label, count in analytics["diff_dist"].items():
        pct = round(count / total_q * 100, 1)
        diff_rows.append((label, str(count), f"{pct}%"))

    _add_table(["Difficulty Level", "Question Count", "Percentage"],
               diff_rows, col_widths=[2.5, 1.8, 1.5])

    _add_heading("Distribution Bar", level=2, color=TEAL)
    _add_bar_chart(list(analytics["diff_dist"].items()), total_q,
                   label_header="Difficulty Level", unit_label=" questions")
    if _rc:
        _insert_chart(_rc.difficulty_pie_chart(analytics["diff_dist"]))

    doc.add_page_break()

    # ── Section 7: Scenario & Case Study Analysis ────────────────────────────
    _section_divider(7, "Scenario & Case Study Analysis")

    sc = analytics["scenario_count"]
    dc = analytics["direct_count"]
    uc = analytics.get("unclassified", 0)
    ta = analytics.get("total_analyzed", sc + dc + uc) or 1

    scenario_stats = [
        ("Scenario / Case Study questions", str(sc)),
        ("Direct / Factual questions",      str(dc)),
        ("Other / Unclassified",            str(uc)),
        ("Total Questions Analysed",        str(ta)),
        ("Scenario Coverage %",             f"{analytics['scenario_pct']}%"),
    ]
    _add_table(["Category", "Count"], scenario_stats, col_widths=[3.5, 2.8])

    sc_ai = ai.get("scenario_analysis", {})
    if sc_ai.get("assessment"):
        _add_heading("Assessment", level=2, color=TEAL)
        _add_para(sc_ai["assessment"], size=9.5)
    if sc_ai.get("recommendation"):
        _add_heading("Recommendation", level=2, color=TEAL)
        _add_para(sc_ai["recommendation"], size=9.5)

    doc.add_page_break()

    # ── Section 8: CO-wise Marks Distribution ───────────────────────────────
    _section_divider(8, "CO-wise Marks Distribution")

    marks_rows = []
    for c in cos:
        key   = str(c["num"])
        marks = analytics["co_marks"].get(key, 0)
        pct   = round(marks / analytics["total_marks"] * 100, 1)
        marks_rows.append((f"CO{key}", str(marks), f"{pct}%"))

    _add_table(["CO", "Total Marks in QB", "Share (%)"],
               marks_rows, col_widths=[0.9, 2.2, 1.5])

    _add_heading("Distribution Bar", level=2, color=TEAL)
    marks_pairs = [(f"CO{c['num']}", analytics["co_marks"].get(str(c["num"]), 0)) for c in cos]
    _add_bar_chart(marks_pairs, analytics["total_marks"],
                   label_header="CO", unit_label=" marks")
    if _rc:
        _insert_chart(_rc.co_marks_chart(cos, analytics["co_marks"], analytics["total_marks"]))

    doc.add_page_break()

    # ── Section 9: OBE Compliance Report ──────────────────────
    _section_divider(9, "OBE Compliance Report")

    obe = ai.get("obe_compliance", {})
    obe_rows = [
        ("CO Statement Quality",  obe.get("co_statement_quality", "—")),
        ("Bloom's Alignment",     obe.get("bloom_alignment", "—")),
        ("CO Coverage Status",    obe.get("co_coverage_status", "—")),
        ("Overall OBE Score",     _obe_band(obe.get('overall_obe_score'))),
        ("Compliance Level",      obe.get("compliance_level", "—")),
    ]
    _add_table(["Parameter", "Status"], obe_rows, col_widths=[3.0, 3.3])

    if obe.get("strengths"):
        _add_heading("Strengths", level=2, color=GREEN)
        for s in obe["strengths"]:
            _add_para(f"• {s}", size=9.5, indent=1)
    if obe.get("gaps"):
        _add_heading("Compliance Gaps", level=2, color=AMBER)
        for s in obe["gaps"]:
            _add_para(f"• {s}", size=9.5, indent=1)
    if obe.get("recommendations"):
        _add_heading("Recommendations", level=2, color=TEAL)
        for s in obe["recommendations"]:
            _add_para(f"• {s}", size=9.5, indent=1)

    doc.add_page_break()

    # ── Section 10: Question Quality vs Programme Outcomes (PO) ───────────────
    _section_divider(10, "Question Quality vs Programme Outcomes (PO)")
    _po_rows, _po_sum = compute_po_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    _add_para(_po_quality_intro(title, code, _po_sum), italic=True, size=9)

    _po_tbl = _add_table(
        ["PO", "Programme Outcome", "Basis", "Items", "%", "Bloom Sub-elements", "Quality"],
        [(r['po'], r['name'], r['basis'], r['count_str'], r['pct_str'], r['verbs'], r['rating'])
         for r in _po_rows],
        col_widths=[0.5, 1.6, 1.0, 0.5, 0.5, 1.6, 0.7]
    )
    for _ri, _r in enumerate(_po_rows, 1):
        _cell = _po_tbl.rows[_ri].cells[6]
        for _run in _cell.paragraphs[0].runs:
            _run.bold = True
            _run.font.color.rgb = STATUS_COLOR.get(_r['rating'], GRAY)
    _add_bar_chart(
        [(r['po'], r['pct']) for r in _po_rows if r['pct'] is not None], 100,
        label_header="PO", unit_label="%"
    )
    _add_para(
        f"Question-assessable POs exercised: {_po_sum['addressed']} of "
        f"{_po_sum['assessable']}. Cognitive POs (PO2 Analysis, PO3 Design) need "
        "Analyse/Create-level questions; thematic POs (PO6-PO8) need societal, "
        "environmental or ethical content - add such questions where a PO is rated "
        f"Poor. PO9-PO11 ({', '.join(_po_sum['process'])}) are assessed via projects, "
        "labs and seminars, not the question bank.",
        size=9, italic=True
    )

    doc.add_page_break()

    # ── Section 11: Question Quality vs Sustainable Development Goals (SDG) ────
    _section_divider(11, "Question Quality vs Sustainable Development Goals (SDG)")
    _sdg_rows, _sdg_sum = compute_sdg_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    _add_para(_sdg_quality_intro(title, code, _sdg_sum), italic=True, size=9)

    if _sdg_rows:
        _sdg_tbl = _add_table(
            ["SDG", "Goal", "COs", "Items", "HOTS %", "Bloom Sub-elements", "Quality"],
            [(r['sdg'], r['name'], r['cos'], str(r['count']), f"{r['hots_pct']}%", r['verbs'], r['rating'])
             for r in _sdg_rows],
            col_widths=[0.6, 1.5, 1.0, 0.5, 0.6, 1.5, 0.7]
        )
        for _ri, _r in enumerate(_sdg_rows, 1):
            _cell = _sdg_tbl.rows[_ri].cells[6]
            for _run in _cell.paragraphs[0].runs:
                _run.bold = True
                _run.font.color.rgb = STATUS_COLOR.get(_r['rating'], GRAY)
        _add_bar_chart(
            [(r['sdg'], r['hots_pct']) for r in _sdg_rows], 100,
            label_header="SDG", unit_label="% HOTS"
        )
        _add_para(
            "HOTS % is the share of each SDG's questions at Apply level or above. "
            "Raise it by reframing recall-level questions on these themes toward "
            "application, analysis or design.",
            size=9, italic=True
        )
    else:
        _add_para(
            "No SDG themes were detected in the current Course Outcomes. Introduce COs "
            "that connect the subject to sustainability, energy, environment, health, "
            "industry or community contexts to strengthen SDG alignment.",
            size=9.5
        )

    doc.add_page_break()

    _add_para()
    _add_para("─" * 80, size=8, color=GRAY)
    _add_para(
        f"Report generated by MyOBE on {datetime.date.today().strftime('%d %B %Y')}. "
        "This report is intended for internal academic and accreditation purposes.",
        italic=True, size=8, color=GRAY
    )

    doc.save(output_path)


# ── Plain-text report builder ─────────────────────────────────────────────────

def build_txt(
    cos, co_text, qb_data, co_tally, bloom_summary, analytics, ai,
    sample_paper, code, title, semester, output_path, is_lab=False, taxonomy_grid=None,
):
    SEP  = "=" * 72
    SEP2 = "-" * 72
    date = datetime.date.today().strftime("%d %B %Y")

    def _sec(n, label):
        lines.append("")
        lines.append(SEP)
        lines.append(f"SECTION {n}: {label.upper()}")
        lines.append(SEP)

    def _sub(label):
        lines.append("")
        lines.append(label)
        lines.append(SEP2)

    def _tbl(headers, rows, widths=None):
        if not widths:
            widths = [max(len(h), max((len(str(r[i])) for r in rows), default=0))
                      for i, h in enumerate(headers)]
        fmt = "  ".join(f"{{:<{w}}}" for w in widths)
        lines.append(fmt.format(*headers))
        lines.append("  ".join("-" * w for w in widths))
        for row in rows:
            lines.append(fmt.format(*[str(v)[:w] for v, w in zip(row, widths)]))

    lines = []
    lines.append("MYOBE — MODULE 1 COMPREHENSIVE REPORT")
    lines.append("AI-Powered Question Paper System")
    lines.append(SEP)
    lines.append(f"Course Code  : {code}")
    lines.append(f"Course Title : {title}")
    lines.append(f"Semester     : {semester or '—'}")
    lines.append(f"Generated on : {date}")
    lines.append("")

    # Executive Summary
    lines.append(SEP)
    lines.append("EXECUTIVE SUMMARY")
    lines.append(SEP)
    lines.append(f"COs Generated          : {len(cos)}")
    lines.append(f"Total Exam Questions   : {analytics['total_exam']}")
    lines.append(f"Assignments in QB      : {analytics['total_assign']}")
    lines.append(f"Quiz Questions in QB   : {analytics['total_quiz']}")
    lines.append(f"CO Coverage            : {analytics['coverage_pct']}%")

    # Section 1: CO Generation
    _sec(1, "Automatic CO Generation")
    co_rows = [(f"CO{c['num']}", _detect_bloom(c["statement"]), c["statement"]) for c in cos]
    _tbl(["CO", "Bloom's", "Statement"], co_rows, [5, 10, 54])
    if bloom_summary:
        _sub("Bloom's Taxonomy Summary")
        lines.append(bloom_summary.strip())

    # Learning Analytics
    _sub("Learning Analytics")
    _la_txt = _compute_bloom_analytics(cos, taxonomy_grid)
    _la_txt_total = _la_txt['total_mapped']
    lines.append(textwrap.fill(_learning_analytics_intro(title, code, _la_txt), 78))
    lines.append("")

    _rationale_txt = _co_placement_rationale(cos, taxonomy_grid)
    if _rationale_txt:
        lines.append("CO Placement Rationale")
        lines.append("-" * 40)
        lines.append("Why each CO sits in its RBT cell (knowledge dimension x level).")
        lines.append("Each rationale names the Bloom's action verb, the cognitive-level")
        lines.append("name and the knowledge-dimension name, and suggests level-appropriate")
        lines.append("verbs to strengthen the CO while keeping the same number of COs:")
        for _co_r, _cell_r, _why_r in _rationale_txt:
            lines.append(f"{_co_r} [{_cell_r}]:")
            lines.append(f"  {_why_r}")
        lines.append("")

    lines.append("Cognitive Level Distribution")
    lines.append("-" * 40)
    _la_level_rows = []
    for _li2, _lv2 in enumerate(_LA_LEVEL_ORDER, 1):
        _co_list2 = _la_txt['level_co_map'][_lv2]
        _pct2 = round(len(_co_list2) / _la_txt_total * 100, 1)
        _cat2 = "HOTS" if _li2 >= 4 else "LOTS"
        _la_level_rows.append((
            f"L{_li2} - {_LA_LEVEL_NAMES[_lv2]}",
            ', '.join(_co_list2) or "—",
            str(len(_co_list2)),
            f"{_pct2}%",
            _cat2,
        ))
    _tbl(["Level", "COs", "N", "%", "Type"], _la_level_rows, [15, 22, 3, 6, 5])
    _lots_pct2 = round(_la_txt['lots_cos'] / _la_txt_total * 100, 1)
    _hots_pct2 = round(_la_txt['hots_cos'] / _la_txt_total * 100, 1)
    _ps2 = _la_txt['progression_score']
    _ps2_label = "Strong" if _ps2 >= 4.0 else "Moderate" if _ps2 >= 2.5 else "Low"
    lines.append("")
    lines.append(f"LOTS (L1-L3 Remember/Understand/Apply) : {_la_txt['lots_cos']} COs ({_lots_pct2}%)")
    lines.append(f"HOTS (L4-L6 Analyze/Evaluate/Create)   : {_la_txt['hots_cos']} COs ({_hots_pct2}%)")
    lines.append(f"Bloom's Progression Score               : {_ps2} / 6.0  ({_ps2_label})")
    lines.append("(A healthy share of higher-order thinking supports OBE expectations)")
    if any(_la_txt['kdim_co_map'].values()):
        lines.append("")
        lines.append("Knowledge Dimension Coverage")
        lines.append("-" * 40)
        _kdim_txt_rows = [
            (kd, ', '.join(_la_txt['kdim_co_map'].get(kd, [])) or "—",
             str(len(_la_txt['kdim_co_map'].get(kd, []))),
             f"{round(len(_la_txt['kdim_co_map'].get(kd, [])) / _la_txt_total * 100, 1)}%")
            for kd in _LA_KDIMS_ORDER
        ]
        _tbl(["Dimension", "COs", "N", "%"], _kdim_txt_rows, [17, 25, 5, 7])

    # Cognitive Balance & Diversity
    lines.append("")
    lines.append("Cognitive Balance & Diversity")
    lines.append("-" * 40)
    _dom_txt = _la_txt['dominant_level']
    _dom_txt_label = (f"L{_LA_LEVEL_ORDER.index(_dom_txt) + 1} - {_LA_LEVEL_NAMES[_dom_txt]}"
                      if _dom_txt else "—")
    _empty_txt = (", ".join(_LA_LEVEL_NAMES[lv] for lv in _la_txt['empty_levels'])
                  if _la_txt['empty_levels'] else "None - all levels covered")
    _bal_txt = _la_txt['balance_index']
    lines.append(f"Cognitive Breadth (levels used) : {_la_txt['breadth']} of 6")
    lines.append(f"Dominant Cognitive Level        : {_dom_txt_label}")
    lines.append(f"Uncovered Levels                : {_empty_txt}")
    lines.append(f"Cognitive Balance Index         : {_bal_txt} / 100")

    # OBE Alignment Indicators
    lines.append("")
    lines.append("OBE Alignment Indicators")
    lines.append("-" * 40)
    _align_txt = _la_alignment_checks(_la_txt)
    _tbl(["Criterion", "Result", "Status"], _align_txt, [30, 14, 8])
    _met_txt = sum(1 for r in _align_txt if r[2] == "Met")
    lines.append(f"Indicators met: {_met_txt} of {len(_align_txt)}")

    # Analysis and Learning Path
    lines.append("")
    lines.append("Analysis and Learning Path")
    lines.append("-" * 40)
    for _ob_txt in _la_observations(_la_txt):
        lines.append(f"- {_ob_txt}")

    # Section 2: Question Bank
    _sec(2, "Lab Question Bank Creation" if is_lab else "Intelligent Question Bank Creation")
    if is_lab:
        total_exp = analytics.get("total_exp", analytics["total_exam"])
        _tbl(["Category", "Count"], [
            ("Lab Experiments", total_exp),
            ("Total Lab Work",  total_exp),
        ], [22, 8])
        _sub("CO-wise Experiment Distribution (count and % share)")
        lab_pct_rows = _build_co_lab_pct_rows(cos, analytics)
        _tbl(["CO", "Experiments", "Share (%)"], lab_pct_rows, [5, 12, 10])
    else:
        _tbl(["Category", "Count"], [
            ("2-Mark Questions",    analytics["total_2"]),
            ("5-Mark Questions",    analytics["total_5"]),
            ("10-Mark Questions",   analytics["total_10"]),
            ("Assignment Tasks",    analytics["total_assign"]),
            ("Quiz Questions",      analytics["total_quiz"]),
            ("Total Exam Questions",analytics["total_exam"]),
        ], [22, 8])
        _sub("CO-wise Question Distribution (count and % share of each type)")
        co_pct_rows = _build_co_pct_rows(cos, co_tally)
        _tbl(["CO", "2-Mk(%)", "5-Mk(%)", "10-Mk(%)", "Asgn(%)", "Quiz(%)", "Total"],
             co_pct_rows, [5, 10, 10, 10, 10, 10, 7])

    # Section 3: Sample Paper
    _sec(3, "Lab Practical Examination (Sample)" if is_lab
         else "Smart Question Paper / Assignment Generation")
    if sample_paper:
        for indent, text, bold in sample_paper:
            prefix = "  " * indent
            lines.append(prefix + text)
    else:
        lines.append("Sample paper not available.")

    # Section 4: Bloom's Compliance
    _sec(4, "Bloom's Taxonomy Compliance Check")
    bc = ai.get("blooms_compliance", [])
    if bc:
        _tbl(["CO", "Stated Level", "QB Level", "Rating", "Remark"],
             [(e.get("co",""), e.get("stated_level",""), e.get("qb_level",""),
               e.get("rating",""), e.get("remark","")[:35]) for e in bc],
             [5, 13, 18, 11, 35])
    else:
        lines.append("AI analysis not available.")

    # Section 5: CO Coverage
    _sec(5, "CO Coverage Verification")
    cov_set = set(analytics.get("covered", []))
    item_header = "Experiments" if is_lab else "Questions"
    cov_rows = []
    for c in cos:
        key = str(c["num"])
        n   = _co_assessment_count(key, analytics, co_tally, is_lab)
        cov_rows.append((f"CO{key}", c["statement"], str(n),
                         "Covered" if key in cov_set else "NOT COVERED"))
    _tbl(["CO", "Statement", item_header, "Status"], cov_rows, [5, 55, 10, 12])
    lines.append(f"\nOverall Coverage: {analytics['coverage_pct']}%  |  "
                 f"Covered: {len(analytics['covered'])}  |  "
                 f"Not Covered: {len(analytics['uncovered'])}")

    # Section 6: Difficulty
    _sec(6, "Difficulty Level Analysis")
    total_q = analytics["total_exam"] or 1
    diff_rows = [(lbl, count, f"{round(count/total_q*100,1)}%")
                 for lbl, count in analytics["diff_dist"].items()]
    _tbl(["Level", "Count", "Pct"], diff_rows, [22, 7, 7])
    lines.append("")
    for lbl, count in analytics["diff_dist"].items():
        pct = round(count / total_q * 100)
        bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
        lines.append(f"  {lbl:<20}  {bar}  {count}")

    # Section 7: Scenario Analysis
    _sec(7, "Scenario & Case Study Analysis")
    ta = analytics.get("total_analyzed", analytics["scenario_count"] + analytics["direct_count"]) or 1
    lines.append(f"Scenario / Case Study questions : {analytics['scenario_count']}")
    lines.append(f"Direct / Factual questions      : {analytics['direct_count']}")
    lines.append(f"Other / Unclassified            : {analytics.get('unclassified', 0)}")
    lines.append(f"Total Questions Analysed        : {ta}")
    lines.append(f"Scenario Coverage               : {analytics['scenario_pct']}%")
    sc_ai = ai.get("scenario_analysis", {})
    if sc_ai.get("assessment"):
        lines.append("")
        lines.append("Assessment:")
        lines.append(sc_ai["assessment"])
    if sc_ai.get("recommendation"):
        lines.append("")
        lines.append("Recommendation:")
        lines.append(sc_ai["recommendation"])

    # Section 8: CO-wise Marks
    _sec(8, "CO-wise Marks Distribution")
    marks_rows = [(f"CO{c['num']}", analytics["co_marks"].get(str(c["num"]),0),
                   f"{round(analytics['co_marks'].get(str(c['num']),0)/analytics['total_marks']*100,1)}%")
                  for c in cos]
    _tbl(["CO", "Marks in QB", "Share"], marks_rows, [5, 12, 8])
    lines.append("")
    for c in cos:
        key   = str(c["num"])
        marks = analytics["co_marks"].get(key, 0)
        pct   = round(marks / analytics["total_marks"] * 100)
        bar   = "█" * (pct // 5) + "░" * (20 - pct // 5)
        lines.append(f"  CO{key:<4}  {bar}  {marks} marks ({pct}%)")

    # Section 9: OBE Compliance
    _sec(9, "OBE Compliance Report")
    obe = ai.get("obe_compliance", {})
    if obe:
        lines.append(f"CO Statement Quality  : {obe.get('co_statement_quality','—')}")
        lines.append(f"Bloom's Alignment     : {obe.get('bloom_alignment','—')}")
        lines.append(f"CO Coverage Status    : {obe.get('co_coverage_status','—')}")
        lines.append(f"Overall OBE Score     : {_obe_band(obe.get('overall_obe_score'))}")
        lines.append(f"Compliance Level      : {obe.get('compliance_level','—')}")
        for label, key in [("Strengths","strengths"),("Gaps","gaps"),("Recommendations","recommendations")]:
            items = obe.get(key, [])
            if items:
                lines.append(f"\n{label}:")
                for s in items:
                    lines.append(f"  • {s}")
    else:
        lines.append("OBE analysis not available.")

    # Section 10: Question Quality vs Programme Outcomes
    _sec(10, "Question Quality vs Programme Outcomes (PO)")
    po_rows, po_sum = compute_po_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    lines.append(textwrap.fill(_po_quality_intro(title, code, po_sum), 72))
    lines.append("")
    _tbl(["PO", "Programme Outcome", "Basis", "Items", "%", "Quality"],
         [(r['po'], r['name'], r['basis'], r['count_str'], r['pct_str'], r['rating'])
          for r in po_rows],
         [5, 26, 16, 6, 6, 9])
    lines.append("")
    for r in po_rows:
        if r['mode'] != 'process':
            lines.append(f"  {r['po']:<5} Bloom sub-elements: {r['verbs']}")
    lines.append("")
    lines.append(f"Question-assessable POs exercised: {po_sum['addressed']} of {po_sum['assessable']}")
    lines.append(f"Process POs (assessed via projects/labs/seminars): {', '.join(po_sum['process'])}")

    # Section 11: Question Quality vs Sustainable Development Goals
    _sec(11, "Question Quality vs Sustainable Development Goals (SDG)")
    sdg_rows, sdg_sum = compute_sdg_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    lines.append(textwrap.fill(_sdg_quality_intro(title, code, sdg_sum), 72))
    lines.append("")
    if sdg_rows:
        _tbl(["SDG", "Goal", "Items", "HOTS%", "Quality"],
             [(r['sdg'], r['name'], r['count'], f"{r['hots_pct']}%", r['rating'])
              for r in sdg_rows],
             [7, 34, 6, 7, 8])
        lines.append("")
        for r in sdg_rows:
            lines.append(f"  {r['sdg']:<7} COs: {r['cos']}  |  Sub-elements: {r['verbs']}")
    else:
        lines.append("No SDG themes were detected in the current Course Outcomes.")
        lines.append("Introduce COs connecting the subject to sustainability, energy,")
        lines.append("environment, health, industry or community contexts.")

    lines.append("")
    lines.append(SEP)
    lines.append(f"Report generated by MyOBE on {date}.")
    lines.append("This report is intended for internal academic and accreditation purposes.")
    lines.append(SEP)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ── PDF report builder ────────────────────────────────────────────────────────

def build_pdf(
    cos, co_text, qb_data, co_tally, bloom_summary, analytics, ai,
    sample_paper, code, title, semester, output_path, is_lab=False, taxonomy_grid=None,
):
    from fpdf import FPDF, XPos, YPos

    def _s(v):
        return str(v).replace("—", " - ").replace("–", " - ").replace(
            "’", "'").replace("‘", "'").replace(
            "“", '"').replace("”", '"').replace(
            "•", "*").replace("█", "#").replace(
            "░", ".").encode("latin-1", "replace").decode("latin-1")

    date = datetime.date.today().strftime("%d %B %Y")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(15, 15, 15)
    PAGE_W = 180  # usable width in mm (A4 - 2*15)

    # ── helpers ────────────────────────────────────────────────────────────────
    def _heading1(text):
        if pdf.get_y() > pdf.h - 40:
            pdf.add_page()
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_fill_color(220, 220, 240)
        pdf.cell(0, 9, _s(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True)
        pdf.ln(1)

    def _heading2(text):
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.cell(0, 7, _s(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(0.5)

    def _body(text, indent=0):
        pdf.set_font("Helvetica", "", 9.5)
        if indent:
            pdf.set_x(pdf.l_margin + indent * 4)
        pdf.multi_cell(0, 5.5, _s(str(text)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _kv(key, val):
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.cell(50, 6, _s(key + ":"))
        pdf.set_font("Helvetica", "", 9.5)
        pdf.cell(0, 6, _s(str(val)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _tbl(headers, rows, col_ws):
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_fill_color(210, 210, 235)
        for h, w in zip(headers, col_ws):
            pdf.cell(w, 7, _s(h), border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8.5)
        for row in rows:
            # check page break before each row
            if pdf.get_y() > pdf.h - 20:
                pdf.add_page()
            for val, w in zip(row, col_ws):
                pdf.cell(w, 6, _s(str(val))[:int(w*1.5)], border=1)
            pdf.ln()
        pdf.ln(2)

    def _bullets(items, indent=5):
        pdf.set_font("Helvetica", "", 9.5)
        for item in items:
            pdf.set_x(pdf.l_margin + indent)
            pdf.multi_cell(PAGE_W - indent, 5.5, _s("* " + str(item)),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _bar_row(label, pct, count):
        filled = int(pct / 5)
        bar = "#" * filled + "." * (20 - filled)
        pdf.set_font("Courier", "", 8.5)
        pdf.cell(0, 5, _s(f"  {label:<22} [{bar}] {count}"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    try:
        import report_charts as _rc
    except Exception:
        _rc = None

    def _pdf_chart(buf, w=155):
        if buf is None or _rc is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            x = pdf.l_margin + (PAGE_W - w) / 2
            pdf.image(_fname, x=x, w=w)
            pdf.ln(3)
        except Exception:
            pass
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    # ── Cover page ──────────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_fill_color(49, 46, 129)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 14, "MyOBE - Module 1 Comprehensive Report",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True, align="C")
    pdf.set_font("Helvetica", "I", 11)
    pdf.set_fill_color(99, 102, 241)
    pdf.cell(0, 9, "AI-Powered Question Paper System",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)
    for k, v in [("Course Code", code), ("Course Title", title),
                  ("Semester", str(semester) if semester else "-"),
                  ("Generated on", date)]:
        _kv(k, v)

    # Executive Summary
    pdf.ln(6)
    _heading1("EXECUTIVE SUMMARY")
    _tbl(
        ["Metric", "Value"],
        [
            ("COs Generated",          len(cos)),
            ("Total Exam Questions",    analytics["total_exam"]),
            ("Assignments in QB",       analytics["total_assign"]),
            ("Quiz Questions in QB",    analytics["total_quiz"]),
            ("CO Coverage",             f"{analytics['coverage_pct']}%"),
        ],
        [110, 60],
    )

    # Section 1
    pdf.add_page()
    _heading1("SECTION 1: AUTOMATIC CO GENERATION")
    # Section 1 CO table — use multi-line rows so full statement is visible
    _heading2("Course Outcome Statements")
    COL_CO1  = 14
    COL_BL1  = 26
    COL_STMT1 = PAGE_W - COL_CO1 - COL_BL1
    LH1, PAD1 = 5.0, 1.5
    chars_per_line1 = max(1, int(COL_STMT1 / 2.0))

    def _co_hdr():
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_fill_color(210, 210, 235)
        for lbl, w in [("CO", COL_CO1), ("Bloom's Level", COL_BL1), ("Statement", COL_STMT1)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()

    _co_hdr()
    pdf.set_font("Helvetica", "", 8.5)
    for c in cos:
        stmt_safe = _s(c["statement"])
        stmt_lines = [stmt_safe[i:i+chars_per_line1]
                      for i in range(0, max(1, len(stmt_safe)), chars_per_line1)]
        row_h = max(7, len(stmt_lines) * LH1 + PAD1 * 2)
        if pdf.get_y() + row_h > pdf.h - 18:
            pdf.add_page(); _co_hdr(); pdf.set_font("Helvetica", "", 8.5)
        y0 = pdf.get_y(); x0 = pdf.l_margin
        pdf.rect(x0, y0, COL_CO1, row_h)
        pdf.rect(x0 + COL_CO1, y0, COL_BL1, row_h)
        pdf.rect(x0 + COL_CO1 + COL_BL1, y0, COL_STMT1, row_h)
        pdf.set_xy(x0 + PAD1, y0 + (row_h - LH1) / 2)
        pdf.cell(COL_CO1 - PAD1, LH1, _s(f"CO{c['num']}"), border=0)
        pdf.set_xy(x0 + COL_CO1 + PAD1, y0 + (row_h - LH1) / 2)
        pdf.cell(COL_BL1 - PAD1, LH1, _s(_detect_bloom(c["statement"])), border=0)
        for li, line in enumerate(stmt_lines):
            pdf.set_xy(x0 + COL_CO1 + COL_BL1 + PAD1, y0 + PAD1 + li * LH1)
            pdf.cell(COL_STMT1 - PAD1 * 2, LH1, line, border=0)
        pdf.set_y(y0 + row_h)
    pdf.ln(2)
    if bloom_summary:
        _heading2("Bloom's Taxonomy Summary")
        _body(bloom_summary.strip())

    # RBT table
    _heading2("Revised Bloom's Taxonomy (RBT) Table")
    _PDF_KDIMS = [
        ("Factual",        "FACTUAL\nKNOWLEDGE"),
        ("Conceptual",     "CONCEPTUAL\nKNOWLEDGE"),
        ("Procedural",     "PROCEDURAL\nKNOWLEDGE"),
        ("Meta-Cognitive", "METACOGNITIVE\nKNOWLEDGE"),
    ]
    _PDF_LEVELS = ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']
    _PDF_LNAMES = ['REMEMBER', 'UNDERSTAND', 'APPLY', 'ANALYZE', 'EVALUATE', 'CREATE']
    _PDF_LVERBS = [
        'Recognizing, Recalling',
        'Interpreting, Exemplifying, Classifying, Summarizing, Inferring, Comparing, Explaining',
        'Executing, Implementing',
        'Differentiating, Organizing, Attributing',
        'Checking, Critiquing',
        'Generating, Planning, Producing',
    ]
    if taxonomy_grid:
        _rbt_pdf_raw = taxonomy_grid
    else:
        from generate_cos import build_taxonomy_grid as _btg2
        _rbt_pdf_raw = _btg2(co_text)
    rbt_pdf_grid = {kd: {lv: _rbt_pdf_raw.get(kd, {}).get(lv, []) for lv in _PDF_LEVELS}
                    for kd, _ in _PDF_KDIMS}

    _C0  = 32.0
    _CW  = (PAGE_W - _C0) / 6.0
    _FH  = 6.5
    _FD  = 7.5
    _LH  = 3.8
    _PAD = 1.0
    _FILL_RBT = (200, 200, 220)

    def _rbt_lines(text, col_w, fsize):
        cpl = max(1, int(col_w / (fsize * 0.45)))
        words = text.replace('\n', ' ').split()
        n, cur = 0, 0
        for w2 in words:
            if cur == 0:
                cur = len(w2); n = 1
            elif cur + 1 + len(w2) <= cpl:
                cur += 1 + len(w2)
            else:
                n += 1; cur = len(w2)
        return max(1, n)

    def _rbt_cell_h(texts, col_w, fsize):
        return max(_rbt_lines(t, col_w, fsize) for t in texts) * _LH + _PAD * 2

    def _draw_rbt_cell(x, y, w, h, text, bold=False, italic=False, fsize=7.0, fill=None):
        if fill:
            pdf.set_fill_color(*fill)
            pdf.rect(x, y, w, h, style='FD')
        else:
            pdf.rect(x, y, w, h)
        style = ('B' if bold else '') + ('I' if italic else '')
        pdf.set_font("Helvetica", style, fsize)
        pdf.set_xy(x + _PAD, y + _PAD)
        pdf.multi_cell(w - _PAD * 2, _LH, _s(text.replace('\n', ' ')), border=0, align='C')

    h0 = _rbt_cell_h(
        ['KNOWLEDGE DIMENSION'] + [f'{i+1}. {_PDF_LNAMES[i]}' for i in range(6)],
        _CW - _PAD * 2, _FH)
    h1 = _rbt_cell_h([''] + _PDF_LVERBS, _CW - _PAD * 2, _FH - 0.5)
    hd = _rbt_cell_h(['METACOGNITIVE KNOWLEDGE'], _C0 - _PAD * 2, _FD)

    if pdf.get_y() + h0 + h1 + hd * 4 > pdf.h - 18:
        pdf.add_page()

    _ry = pdf.get_y()
    _draw_rbt_cell(pdf.l_margin, _ry, _C0, h0, 'KNOWLEDGE\nDIMENSION', bold=True, fsize=_FH, fill=_FILL_RBT)
    for _i, _nm in enumerate(_PDF_LNAMES):
        _draw_rbt_cell(pdf.l_margin + _C0 + _i * _CW, _ry, _CW, h0,
                       f'{_i+1}. {_nm}', bold=True, fsize=_FH, fill=_FILL_RBT)
    _ry += h0

    _draw_rbt_cell(pdf.l_margin, _ry, _C0, h1, '', fsize=_FH - 0.5, fill=_FILL_RBT)
    for _i, _vb in enumerate(_PDF_LVERBS):
        _draw_rbt_cell(pdf.l_margin + _C0 + _i * _CW, _ry, _CW, h1,
                       _vb, italic=True, fsize=_FH - 0.5, fill=_FILL_RBT)
    _ry += h1

    for _kd, _kd_lbl in _PDF_KDIMS:
        _draw_rbt_cell(pdf.l_margin, _ry, _C0, hd, _kd_lbl, bold=True, fsize=_FD)
        for _j, _lv in enumerate(_PDF_LEVELS):
            _ct = ', '.join(rbt_pdf_grid[_kd].get(_lv, []))
            _draw_rbt_cell(pdf.l_margin + _C0 + _j * _CW, _ry, _CW, hd, _ct, fsize=_FD)
        _ry += hd
    pdf.set_y(_ry + 3)

    # CO Placement Rationale (PDF)
    _rationale_pdf = _co_placement_rationale(cos, rbt_pdf_grid)
    if _rationale_pdf:
        _heading2("CO Placement Rationale")
        _body(
            "Why each Course Outcome is mapped to its cell (knowledge dimension x "
            "cognitive level) in the RBT table above. Each rationale names the Bloom's "
            "action verb, its cognitive-level name (Remember-Create) and the "
            "knowledge-dimension name (Factual, Conceptual, Procedural, Meta-Cognitive), "
            "and suggests level-appropriate verbs to strengthen the CO while keeping the "
            "same number of Course Outcomes."
        )
        for _co_rp, _cell_rp, _why_rp in _rationale_pdf:
            pdf.set_font("Helvetica", "B", 9.5)
            pdf.cell(0, 5.5, _s(f"{_co_rp}  [{_cell_rp}]"),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            _body(_why_rp, indent=1)

    # Learning Analytics (PDF)
    _heading2("Learning Analytics")
    _la_pdf = _compute_bloom_analytics(cos, rbt_pdf_grid)
    _la_pdf_total = _la_pdf['total_mapped']
    _body(_learning_analytics_intro(title, code, _la_pdf))

    _heading2("Cognitive Level Distribution")
    _la_pdf_level_rows = []
    for _li3, _lv3 in enumerate(_LA_LEVEL_ORDER, 1):
        _co_list3 = _la_pdf['level_co_map'][_lv3]
        _pct3 = round(len(_co_list3) / _la_pdf_total * 100, 1)
        _cat3 = "HOTS" if _li3 >= 4 else "LOTS"
        _la_pdf_level_rows.append((
            f"L{_li3} - {_LA_LEVEL_NAMES[_lv3]}",
            ', '.join(_co_list3) or "-",
            str(len(_co_list3)),
            f"{_pct3}%",
            _cat3,
        ))
    _tbl(["Bloom's Level", "COs", "N", "%", "Type"],
         _la_pdf_level_rows, [45, 62, 10, 16, 17])

    _lots_pct3 = round(_la_pdf['lots_cos'] / _la_pdf_total * 100, 1)
    _hots_pct3 = round(_la_pdf['hots_cos'] / _la_pdf_total * 100, 1)
    _tbl(
        ["Category", "Levels", "Count", "%"],
        [
            ("LOTS (Lower Order)", "L1-L3: Remember, Understand, Apply",
             str(_la_pdf['lots_cos']), f"{_lots_pct3}%"),
            ("HOTS (Higher Order)", "L4-L6: Analyze, Evaluate, Create",
             str(_la_pdf['hots_cos']), f"{_hots_pct3}%"),
        ],
        [42, 80, 20, 18]
    )
    _ps3 = _la_pdf['progression_score']
    _ps3_label = "Strong" if _ps3 >= 4.0 else "Moderate" if _ps3 >= 2.5 else "Low"
    _body(
        f"Bloom's Progression Score: {_ps3} / 6.0 ({_ps3_label} cognitive ambition). "
        "A healthy share of higher-order thinking (L4-L6) supports OBE expectations."
    )
    if any(_la_pdf['kdim_co_map'].values()):
        _heading2("Knowledge Dimension Coverage")
        _kdim_pdf_rows = [
            (kd, ', '.join(_la_pdf['kdim_co_map'].get(kd, [])) or "-",
             str(len(_la_pdf['kdim_co_map'].get(kd, []))),
             f"{round(len(_la_pdf['kdim_co_map'].get(kd, [])) / _la_pdf_total * 100, 1)}%")
            for kd in _LA_KDIMS_ORDER
        ]
        _tbl(["Knowledge Dimension", "COs", "Count", "%"], _kdim_pdf_rows, [50, 82, 18, 20])

    # Cognitive Balance & Diversity (PDF)
    _heading2("Cognitive Balance & Diversity")
    _dom_pdf = _la_pdf['dominant_level']
    _dom_pdf_label = (f"L{_LA_LEVEL_ORDER.index(_dom_pdf) + 1} - {_LA_LEVEL_NAMES[_dom_pdf]}"
                      if _dom_pdf else "-")
    _empty_pdf = (", ".join(_LA_LEVEL_NAMES[lv] for lv in _la_pdf['empty_levels'])
                  if _la_pdf['empty_levels'] else "None - all levels covered")
    _bal_pdf = _la_pdf['balance_index']
    _bal_pdf_word = ("Well balanced" if _bal_pdf >= 70
                     else "Moderately balanced" if _bal_pdf >= 40 else "Concentrated")
    _tbl(
        ["Indicator", "Value"],
        [
            ("Cognitive Breadth (levels used)", f"{_la_pdf['breadth']} of 6"),
            ("Dominant Cognitive Level",        _dom_pdf_label),
            ("Uncovered Levels",                _empty_pdf),
            ("Cognitive Balance Index",         f"{_bal_pdf} / 100 ({_bal_pdf_word})"),
            ("Bloom's Progression Score",       f"{_ps3} / 6.0 ({_ps3_label})"),
        ],
        [60, 110]
    )

    # OBE Alignment Indicators (PDF)
    _heading2("OBE Alignment Indicators")
    _align_pdf = _la_alignment_checks(_la_pdf)
    _tbl(["Criterion", "Result", "Status"], _align_pdf, [100, 40, 30])
    _met_pdf = sum(1 for r in _align_pdf if r[2] == "Met")
    _body(f"Indicators met: {_met_pdf} of {len(_align_pdf)}.")

    # Analysis and Learning Path (PDF)
    _heading2("Analysis and Learning Path")
    for _ob_pdf in _la_observations(_la_pdf):
        _body(f"- {_ob_pdf}")

    # Section 2
    pdf.add_page()
    _heading1("SECTION 2: LAB QUESTION BANK CREATION" if is_lab
              else "SECTION 2: INTELLIGENT QUESTION BANK CREATION")
    if is_lab:
        total_exp = analytics.get("total_exp", analytics["total_exam"])
        _tbl(["Category", "Count"], [
            ("Lab Experiments", total_exp),
            ("Total Lab Work",  total_exp),
        ], [130, 40])
        _heading2("CO-wise Experiment Distribution")
        _body("Count and share (%) of experiments allocated to each CO.")
        lab_pct_rows = _build_co_lab_pct_rows(cos, analytics)
        _tbl(["CO", "Experiments", "Share (%)"], lab_pct_rows, [20, 40, 30])
    else:
        _tbl(["Category", "Count"], [
            ("2-Mark Questions",     analytics["total_2"]),
            ("5-Mark Questions",     analytics["total_5"]),
            ("10-Mark Questions",    analytics["total_10"]),
            ("Assignment Tasks",     analytics["total_assign"]),
            ("Quiz Questions",       analytics["total_quiz"]),
            ("Total Exam Questions", analytics["total_exam"]),
        ], [130, 40])
        _heading2("CO-wise Question Distribution")
        _body("Count and share (%) of each question type allocated to each CO.")
        co_pct_rows = _build_co_pct_rows(cos, co_tally)
        # Columns: CO(12) | 2-Mk%(26) | 5-Mk%(26) | 10-Mk%(26) | Assign%(26) | Quiz%(22) | Total(14) = 152mm
        _tbl(
            ["CO", "2-Mk (%)", "5-Mk (%)", "10-Mk (%)", "Assign (%)", "Quiz (%)", "Total"],
            co_pct_rows,
            [12, 26, 26, 26, 26, 22, 14]
        )

    # Section 3: Sample Paper
    pdf.add_page()
    _heading1("SECTION 3: LAB PRACTICAL EXAMINATION (SAMPLE)" if is_lab
              else "SECTION 3: SMART QUESTION PAPER (SAMPLE)")
    if sample_paper:
        for indent, text, bold in sample_paper:
            pdf.set_font("Helvetica", "B" if bold else "", 9)
            x_offset = indent * 5
            if x_offset:
                pdf.set_x(pdf.l_margin + x_offset)
            pdf.multi_cell(PAGE_W - x_offset, 5.5, _s(text),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        _body("Sample paper not available.")

    # Section 4: Bloom's Compliance — wider columns to avoid overflow
    pdf.add_page()
    _heading1("SECTION 4: BLOOM'S TAXONOMY COMPLIANCE CHECK")
    bc = ai.get("blooms_compliance", [])
    if bc:
        _tbl(["CO", "Stated Level", "QB Level Coverage", "Rating"],
             [(e.get("co",""), e.get("stated_level",""), e.get("qb_level",""),
               e.get("rating","")) for e in bc],
             [14, 38, 90, 38])
        for e in bc:
            if e.get("remark"):
                _body(f"{e.get('co','')}: {e['remark']}")
    else:
        _body("AI analysis not available.")

    # Section 5: CO Coverage — uses per-row multi_cell for full statement wrapping
    pdf.add_page()
    _heading1("SECTION 5: CO COVERAGE VERIFICATION")

    COL_CO   = 16
    COL_STMT = 105
    COL_Q    = 22
    COL_ST   = 37
    LH       = 5.5   # line height per text line
    PAD      = 1.5

    cov_set     = set(analytics.get("covered", []))
    item_header = "Experiments" if is_lab else "Questions"

    def _cov_header():
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_fill_color(210, 210, 235)
        for lbl, w in [("CO", COL_CO), ("Statement", COL_STMT),
                        (item_header, COL_Q), ("Status", COL_ST)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()

    _cov_header()
    pdf.set_font("Helvetica", "", 8.5)
    for c in cos:
        key  = str(c["num"])
        tot  = _co_assessment_count(key, analytics, co_tally, is_lab)
        stat = "Covered" if key in cov_set else "NOT COVERED"

        stmt_safe = _s(c["statement"])
        chars_per_line = max(1, int(COL_STMT / 2.1))
        stmt_lines = [stmt_safe[i:i+chars_per_line]
                      for i in range(0, max(1, len(stmt_safe)), chars_per_line)]
        row_h = max(7, len(stmt_lines) * LH + PAD * 2)

        if pdf.get_y() + row_h > pdf.h - 18:
            pdf.add_page(); _cov_header()
            pdf.set_font("Helvetica", "", 8.5)

        y0 = pdf.get_y(); x0 = pdf.l_margin

        pdf.rect(x0,              y0, COL_CO,   row_h)
        pdf.rect(x0 + COL_CO,     y0, COL_STMT, row_h)
        pdf.rect(x0 + COL_CO + COL_STMT,       y0, COL_Q, row_h)
        pdf.rect(x0 + COL_CO + COL_STMT + COL_Q, y0, COL_ST, row_h)

        # CO cell
        pdf.set_xy(x0 + PAD, y0 + (row_h - LH) / 2)
        pdf.cell(COL_CO - PAD, LH, _s(f"CO{key}"), border=0)

        # Statement cell — line by line
        for li, line in enumerate(stmt_lines):
            pdf.set_xy(x0 + COL_CO + PAD, y0 + PAD + li * LH)
            pdf.cell(COL_STMT - PAD * 2, LH, line, border=0)

        # Questions cell
        pdf.set_xy(x0 + COL_CO + COL_STMT + PAD, y0 + (row_h - LH) / 2)
        pdf.cell(COL_Q - PAD, LH, str(tot), border=0, align="C")

        # Status cell
        pdf.set_xy(x0 + COL_CO + COL_STMT + COL_Q + PAD, y0 + (row_h - LH) / 2)
        pdf.cell(COL_ST - PAD, LH, _s(stat), border=0)

        pdf.set_y(y0 + row_h)
    pdf.ln(2)
    _body(f"Overall Coverage: {analytics['coverage_pct']}%  |  "
          f"Covered: {len(analytics['covered'])}  |  Not Covered: {len(analytics['uncovered'])}")

    # Section 6: Difficulty
    pdf.add_page()
    _heading1("SECTION 6: DIFFICULTY LEVEL ANALYSIS")
    total_q = analytics["total_exam"] or 1
    _tbl(["Level", "Count", "Percentage"],
         [(lbl, cnt, f"{round(cnt/total_q*100,1)}%") for lbl, cnt in analytics["diff_dist"].items()],
         [60, 25, 25])
    _heading2("Distribution")
    for lbl, cnt in analytics["diff_dist"].items():
        _bar_row(lbl, round(cnt / total_q * 100), cnt)
    _pdf_chart(_rc.difficulty_pie_chart(analytics["diff_dist"]) if _rc else None)

    # Section 7: Scenario Analysis
    pdf.add_page()
    _heading1("SECTION 7: SCENARIO & CASE STUDY ANALYSIS")
    _tbl(["Category", "Count"], [
        ("Scenario / Case Study questions", analytics["scenario_count"]),
        ("Direct / Factual questions",      analytics["direct_count"]),
        ("Scenario Coverage %",             f"{analytics['scenario_pct']}%"),
    ], [120, 50])
    sc_ai = ai.get("scenario_analysis", {})
    if sc_ai.get("assessment"):
        _heading2("Assessment")
        _body(sc_ai["assessment"])
    if sc_ai.get("recommendation"):
        _heading2("Recommendation")
        _body(sc_ai["recommendation"])

    # Section 8: CO Marks
    pdf.add_page()
    _heading1("SECTION 8: CO-WISE MARKS DISTRIBUTION")
    marks_rows = [(f"CO{c['num']}",
                   analytics["co_marks"].get(str(c["num"]), 0),
                   f"{round(analytics['co_marks'].get(str(c['num']),0)/analytics['total_marks']*100,1)}%")
                  for c in cos]
    _tbl(["CO", "Marks in QB", "Share"], marks_rows, [20, 40, 30])
    _heading2("Distribution")
    for c in cos:
        key   = str(c["num"])
        marks = analytics["co_marks"].get(key, 0)
        pct   = round(marks / analytics["total_marks"] * 100)
        _bar_row(f"CO{key}", pct, f"{marks} marks ({pct}%)")
    _pdf_chart(_rc.co_marks_chart(cos, analytics["co_marks"], analytics["total_marks"]) if _rc else None)

    # Section 9: OBE Compliance
    pdf.add_page()
    _heading1("SECTION 9: OBE COMPLIANCE REPORT")
    obe = ai.get("obe_compliance", {})
    if obe:
        _tbl(["Parameter", "Status"], [
            ("CO Statement Quality",  obe.get("co_statement_quality", "-")),
            ("Bloom's Alignment",     obe.get("bloom_alignment", "-")),
            ("CO Coverage Status",    obe.get("co_coverage_status", "-")),
            ("Overall OBE Score",     _obe_band(obe.get('overall_obe_score'))),
            ("Compliance Level",      obe.get("compliance_level", "-")),
        ], [90, 80])
        if obe.get("strengths"):
            _heading2("Strengths")
            _bullets(obe["strengths"])
        if obe.get("gaps"):
            _heading2("Compliance Gaps")
            _bullets(obe["gaps"])
        if obe.get("recommendations"):
            _heading2("Recommendations")
            _bullets(obe["recommendations"])
    else:
        _body("OBE analysis not available.")

    # Section 10: Question Quality vs Programme Outcomes
    pdf.add_page()
    _heading1("SECTION 10: QUESTION QUALITY VS PROGRAMME OUTCOMES (PO)")
    po_rows, po_sum = compute_po_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    _body(_po_quality_intro(title, code, po_sum))
    pdf.ln(1)
    _tbl(["PO", "Programme Outcome", "Basis", "Items", "%", "Bloom Sub-elements", "Quality"],
         [(r['po'], r['name'], r['basis'], r['count_str'], r['pct_str'], r['verbs'], r['rating'])
          for r in po_rows],
         [14, 38, 26, 14, 14, 50, 24])
    _body(f"Question-assessable POs exercised: {po_sum['addressed']} of {po_sum['assessable']}. "
          "Cognitive POs (PO2 Analysis, PO3 Design) need Analyse/Create-level questions; "
          "thematic POs (PO6-PO8) need societal, environmental or ethical content - add such "
          f"questions where a PO is rated Poor. PO9-PO11 ({', '.join(po_sum['process'])}) are "
          "assessed via projects, labs and seminars, not the question bank.")

    # Section 11: Question Quality vs Sustainable Development Goals
    pdf.add_page()
    _heading1("SECTION 11: QUESTION QUALITY VS SUSTAINABLE DEVELOPMENT GOALS (SDG)")
    sdg_rows, sdg_sum = compute_sdg_quality(cos, co_tally, analytics, taxonomy_grid, is_lab)
    _body(_sdg_quality_intro(title, code, sdg_sum))
    pdf.ln(1)
    if sdg_rows:
        _tbl(["SDG", "Goal", "COs", "Items", "HOTS %", "Bloom Sub-elements", "Quality"],
             [(r['sdg'], r['name'], r['cos'], r['count'], f"{r['hots_pct']}%", r['verbs'], r['rating'])
              for r in sdg_rows],
             [16, 40, 24, 14, 18, 38, 20])
        _body("HOTS % is the share of each SDG's questions at Apply level or above. "
              "Raise it by reframing recall-level questions on these themes toward "
              "application, analysis or design.")
    else:
        _body("No SDG themes were detected in the current Course Outcomes. Introduce COs "
              "that connect the subject to sustainability, energy, environment, health, "
              "industry or community contexts to strengthen SDG alignment.")

    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, _s(f"Report generated by MyOBE on {date}. For internal academic and accreditation use."),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(output_path)
