"""
generate_qpaper.py  -  Question Paper Generator

Usage:
  python generate_qpaper.py <question_bank_file>

Parses a question bank produced by generate_qbank.py and assembles a
formatted question paper: Assignment, Quiz, Internal Assessment, or Final Exam.
Supports .txt / .pdf / .docx question bank files.
"""

import sys
import os
import re
from collections import OrderedDict
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

try:
    from generate_cos import extract as _extract_generic
except ImportError:
    print("Error: generate_cos.py must be in the same directory.")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Question bank loader  (preserves ## markers for docx)
# ---------------------------------------------------------------------------

def _load_qbank_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        from docx import Document
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                lines.append('')
                continue
            style = (p.style.name or '').lower()
            if 'title' in style:
                lines.append(f'# {text}')
            elif 'heading 1' in style:
                lines.append(f'## {text}')   # unit headers saved at level=1
            elif 'heading 2' in style:
                lines.append(f'### {text}')  # CO headers saved at level=2
            elif 'heading 3' in style:
                lines.append(f'### {text}')
            else:
                lines.append(text)
        return '\n'.join(lines)
    return _extract_generic(path)


# ---------------------------------------------------------------------------
# Question bank parser
# ---------------------------------------------------------------------------

_ASSIGN_HDR_RE = re.compile(r'^\[Assignment\b', re.IGNORECASE)
_QUIZ_HDR_RE   = re.compile(r'^\[Quiz\b',       re.IGNORECASE)
_MARK_HDR_RE   = re.compile(r'^\[(\d+)\s*Marks?\]')
# Match "## Unit 1: Title" (txt) OR bare "Unit 1: Title" (extracted from pdf/docx)
_UNIT_HDR_RE   = re.compile(r'^(?:##\s+)?(Unit\s+\d+)(.*?)$', re.IGNORECASE)
_SKIP_LINE_RE  = re.compile(
    r'^(#?\s*Question Bank|Exam questions|## CO Coverage|Questions per unit'
    r'|Course Outcomes:|CO Coverage Summary)',
    re.IGNORECASE
)


def _detect_qtype(text):
    t = text.lower()
    if t.startswith('true or false'):      return 'tf'
    if t.startswith('fill in the blank'):  return 'fib'
    return 'mcq'


def _flush(q, data, unit, section):
    if not q or not unit or not section or unit not in data['units']:
        return
    u = data['units'][unit]
    if section.startswith('exam_'):
        marks = int(section.split('_')[1])
        if marks in u['exam']:
            u['exam'][marks].append(q)
    elif section == 'assign':
        u['assign'].append(q)
    elif section == 'quiz':
        u['quiz'].append(q)


def parse_qbank(text):
    """Parse question bank text into structured data.

    Returns dict with keys: course_code, course_title, semester, units.
    units is an OrderedDict: label -> {title, exam:{2:[],5:[],10:[]}, assign:[], quiz:[]}.
    Each question is a dict with keys: text/title, co, unit, marks/qtype/options/context/body.
    """
    data = {
        'course_code': '', 'course_title': '', 'semester': None,
        'units': OrderedDict(),
    }
    current_unit = current_section = current_q = None
    in_preamble = False  # True between a unit header and the first section marker

    for raw in text.split('\n'):
        s = raw.strip()

        if _SKIP_LINE_RE.match(s):
            continue

        # blank line
        if not s:
            if current_section == 'assign' and current_q:
                current_q['body'].append('')
            continue

        # Course header: "Course: CODE — Title  [Semester N]"
        cm = re.match(r'Course:\s*(.+?)\s*[—\-]+\s*(.+?)(?:\s*\[Semester\s*(\d+)\])?\s*$', s)
        if cm:
            data['course_code']  = cm.group(1).strip()
            data['course_title'] = cm.group(2).strip()
            data['semester']     = cm.group(3) if cm.group(3) else None
            continue

        # Unit header: "## Unit 1: Title"
        um = _UNIT_HDR_RE.match(s)
        if um:
            _flush(current_q, data, current_unit, current_section)
            current_q = None
            label = um.group(1).strip()
            rest  = re.sub(r'^[:\-–—]\s*', '', um.group(2).strip())
            rest  = re.sub(r'\s*\|.*$', '', rest).strip()
            current_unit = label
            current_section = None
            in_preamble = True  # skip CO headers / metadata until first section marker
            if label not in data['units']:
                data['units'][label] = {
                    'title': rest,
                    'exam': {2: [], 5: [], 10: []},
                    'assign': [],
                    'quiz': [],
                }
            continue

        # Skip CO header lines and document metadata (markdown form)
        if s.startswith('###') or s.startswith('# '):
            continue

        # Section markers
        mm = _MARK_HDR_RE.match(s)
        if mm:
            _flush(current_q, data, current_unit, current_section)
            current_q = None
            current_section = f'exam_{int(mm.group(1))}'
            in_preamble = False
            continue

        if _ASSIGN_HDR_RE.match(s):
            _flush(current_q, data, current_unit, current_section)
            current_q = None
            current_section = 'assign'
            in_preamble = False
            continue

        if _QUIZ_HDR_RE.match(s):
            _flush(current_q, data, current_unit, current_section)
            current_q = None
            current_section = 'quiz'
            in_preamble = False
            continue

        # Skip preamble content (CO statements, metadata between unit header and first section)
        if in_preamble:
            continue

        # Question content
        if current_section and current_section.startswith('exam_'):
            marks = int(current_section.split('_')[1])
            qm = re.match(r'Q\d+\.\s*(.*)', s)
            if qm:
                _flush(current_q, data, current_unit, current_section)
                co_m = re.search(r'\[CO(\d+)\]', s)
                text_clean = re.sub(r'\s*\[CO\d+\]\s*$', '', qm.group(1)).strip()
                current_q = {
                    'text': text_clean, 'co': f"CO{co_m.group(1)}" if co_m else None,
                    'unit': current_unit, 'marks': marks, 'context': [],
                }
            elif current_q and not re.match(r'Q\d+\.', s):
                # Continuation of a wrapped Q-line — pick up CO tag and full text
                co_m = re.search(r'\[CO(\d+)\]', s)
                if co_m and current_q['co'] is None:
                    current_q['co'] = f"CO{co_m.group(1)}"
                extra = re.sub(r'\s*\[CO\d+\]\s*$', '', s).strip()
                if extra:
                    if marks == 10:
                        current_q['context'].append(extra)
                    else:
                        current_q['text'] += ' ' + extra

        elif current_section == 'assign':
            am = re.match(r'A\d+\.\s*(.*)', s)
            if am:
                _flush(current_q, data, current_unit, current_section)
                co_m = re.search(r'\[CO(\d+)\]', s)
                title_clean = re.sub(r'\s*\[CO\d+\]\s*$', '', am.group(1)).strip()
                current_q = {
                    'title': title_clean, 'co': f"CO{co_m.group(1)}" if co_m else None,
                    'unit': current_unit, 'body': [],
                }
            elif current_q:
                current_q['body'].append(s)

        elif current_section == 'quiz':
            qm = re.match(r'Q\d+\.\s*(.*)', s)
            if qm:
                _flush(current_q, data, current_unit, current_section)
                co_m = re.search(r'\[CO(\d+)\]', s)
                text_clean = re.sub(r'\s*\[CO\d+\]\s*$', '', qm.group(1)).strip()
                current_q = {
                    'text': text_clean, 'co': f"CO{co_m.group(1)}" if co_m else None,
                    'unit': current_unit, 'qtype': _detect_qtype(text_clean), 'options': [],
                }
            elif current_q and re.match(r'^\([a-d]\)', s):
                current_q['options'].append(s)
            elif current_q and not current_q['options']:
                # Continuation of a wrapped Q-line (PDF multi_cell splits long questions)
                co_m = re.search(r'\[CO(\d+)\]', s)
                if co_m and current_q['co'] is None:
                    current_q['co'] = f"CO{co_m.group(1)}"
                extra = re.sub(r'\s*\[CO\d+\]\s*$', '', s).strip()
                if extra:
                    current_q['text'] += ' ' + extra

    _flush(current_q, data, current_unit, current_section)
    return data


# ---------------------------------------------------------------------------
# Lab question bank parser
# ---------------------------------------------------------------------------

_LAB_EXP_HDR_RE  = re.compile(r'^\[Experiments?\b', re.IGNORECASE | re.MULTILINE)
_LAB_EXP_LINE_RE = re.compile(r'^Exp\s+(\d+)\.\s*(.*)', re.IGNORECASE)
_LAB_SEC_RE      = re.compile(
    r'^(Aim|Equipment(?:[^:]*)?|Tools?\s+Required|Materials?\s+Required'
    r'|Theory|Procedure|Expected\s+(?:Observations?|Output)'
    r'|Observations?/Output|Result|Viva\s+Questions?)\s*:',
    re.IGNORECASE
)


def is_lab_qbank(text):
    """Return True if *text* looks like a lab / practical question bank."""
    return bool(_LAB_EXP_HDR_RE.search(text))


def _flush_exp(exp, data, unit):
    if exp and unit and unit in data['units']:
        data['units'][unit]['experiments'].append(exp)


def parse_lab_qbank(text):
    """Parse a lab question bank into structured data.

    Returns dict: course_code, course_title, semester,
    units: OrderedDict(label -> {title, experiments: [exp_dict]}).
    exp_dict keys: num, title, co, unit,
                   aim, equipment, theory, procedure, observations, result, viva
    (each value is a list of str).
    """
    data = {
        'course_code': '', 'course_title': '', 'semester': None,
        'units': OrderedDict(),
    }
    current_unit = None
    current_exp  = None
    current_sec  = None

    for raw in text.split('\n'):
        s = raw.strip()
        if _SKIP_LINE_RE.match(s) or not s:
            continue

        cm = re.match(r'Course:\s*(.+?)\s*[—\-]+\s*(.+?)(?:\s*\[Semester\s*(\d+)\])?\s*$', s)
        if cm:
            data['course_code']  = cm.group(1).strip()
            data['course_title'] = cm.group(2).strip()
            data['semester']     = cm.group(3) if cm.group(3) else None
            continue

        um = _UNIT_HDR_RE.match(s)
        if um:
            _flush_exp(current_exp, data, current_unit)
            current_exp = current_sec = None
            label = um.group(1).strip()
            rest  = re.sub(r'^[:\-–—]\s*', '', um.group(2).strip())
            rest  = re.sub(r'\s*\|.*$', '', rest).strip()
            current_unit = label
            if label not in data['units']:
                data['units'][label] = {'title': rest, 'experiments': []}
            continue

        if s.startswith('###') or s.startswith('# ') or s.startswith('## '):
            continue
        if _LAB_EXP_HDR_RE.match(s):
            continue

        em = _LAB_EXP_LINE_RE.match(s)
        if em:
            _flush_exp(current_exp, data, current_unit)
            co_m        = re.search(r'\[CO(\d+)\]', s)
            title_clean = re.sub(r'\s*\[CO\d+\]\s*$', '', em.group(2)).strip()
            current_exp = {
                'num': int(em.group(1)), 'title': title_clean,
                'co':  f"CO{co_m.group(1)}" if co_m else None,
                'unit': current_unit,
                'aim': [], 'equipment': [], 'theory': [],
                'procedure': [], 'observations': [], 'result': [], 'viva': [],
            }
            current_sec = None
            continue

        sm = _LAB_SEC_RE.match(s)
        if sm and current_exp is not None:
            sec = sm.group(1).lower()
            if   'aim'       in sec:                                    current_sec = 'aim'
            elif any(k in sec for k in ('equipment', 'tool', 'material')): current_sec = 'equipment'
            elif 'theory'    in sec:                                    current_sec = 'theory'
            elif 'procedure' in sec:                                    current_sec = 'procedure'
            elif any(k in sec for k in ('observ', 'output')):          current_sec = 'observations'
            elif 'result'    in sec:                                    current_sec = 'result'
            elif 'viva'      in sec:                                    current_sec = 'viva'
            rest_text = s[sm.end():].strip()
            if rest_text and current_sec:
                current_exp[current_sec].append(rest_text)
            continue

        if current_exp is not None and current_sec:
            current_exp[current_sec].append(s)

    _flush_exp(current_exp, data, current_unit)
    return data


def _flat_lab_list(units):
    """Return flat [(unit_label, exp_dict)] for all experiments."""
    return [(unit, exp)
            for unit, udata in units.items()
            for exp in udata.get('experiments', [])]


# ---------------------------------------------------------------------------
# Listing and selection helpers
# ---------------------------------------------------------------------------

def _flat_list(units, section, marks=None):
    result = []
    for unit, udata in units.items():
        if section == 'exam':
            for q in udata['exam'].get(marks, []):
                result.append((unit, q))
        elif section == 'assign':
            for q in udata['assign']:
                result.append((unit, q))
        elif section == 'quiz':
            for q in udata['quiz']:
                result.append((unit, q))
    return result


def _display(items, section, marks=None):
    if not items:
        print("    (none available)")
        return
    kind = {
        'exam':   f"{marks}-Mark Exam Questions",
        'assign': "Assignment Tasks",
        'quiz':   "Quiz Questions",
    }[section]
    print(f"\n  {kind}  ({len(items)} available)")
    print(f"  {'─' * 58}")
    for i, (unit, q) in enumerate(items, 1):
        co = f"  [{q['co']}]" if q.get('co') else ""
        if section == 'assign':
            preview = q['title'][:68]
        elif section == 'quiz':
            tp = {'mcq': 'MCQ', 'tf': 'T/F', 'fib': 'FIB'}.get(q.get('qtype', 'mcq'), '')
            preview = f"({tp}) {q['text'][:62]}"
        else:
            preview = q['text'][:68]
        print(f"  {i:3}.  [{unit}]  {preview}{co}")


def _select(items, prompt=""):
    if not items:
        return []
    raw = input(f"  Select{prompt} (numbers/ranges/all/skip) [default: all]: ").strip().lower()
    if not raw or raw == 'all':
        return list(items)
    if raw == 'skip':
        return []
    sel = set()
    for part in raw.split(','):
        part = part.strip()
        if '-' in part:
            a, _, b = part.partition('-')
            try:
                sel.update(range(int(a) - 1, int(b)))
            except ValueError:
                pass
        else:
            try:
                sel.add(int(part) - 1)
            except ValueError:
                pass
    return [items[i] for i in sorted(sel) if 0 <= i < len(items)]


def _ask(prompt, default=""):
    raw = input(f"  {prompt}" + (f" [{default}]: " if default else ": ")).strip()
    return raw if raw else default


def _ask_int(prompt, default):
    raw = input(f"  {prompt} [default {default}]: ").strip()
    try:
        return max(0, int(raw))
    except ValueError:
        return default


# ---------------------------------------------------------------------------
# Paper builders  (return list of text blocks)
# ---------------------------------------------------------------------------

def _header_block(title, course_code, course_title, semester,
                  date="", duration="", max_marks="", instructions=None):
    sem = f"  [Semester {semester}]" if semester else ""
    details = "  |  ".join(filter(None, [
        f"Date: {date}"           if date       else None,
        f"Duration: {duration}"   if duration   else None,
        f"Max Marks: {max_marks}" if max_marks  else None,
    ]))
    lines = [f"# {title}", f"Course: {course_code} — {course_title}{sem}"]
    if details:
        lines.append(details)
    if instructions:
        lines += ["", "Instructions:"] + [f"  - {inst}" for inst in instructions]
    return "\n".join(lines) + "\n"


def build_assignment_paper(selected, data, cfg):
    blocks = [_header_block(
        cfg['title'], data['course_code'], data['course_title'], data['semester'],
        cfg.get('date', ''), max_marks=cfg.get('max_marks', ''),
        instructions=cfg.get('instructions', []),
    )]
    blocks.append("─" * 60 + "\n")
    for i, (unit, q) in enumerate(selected, 1):
        co = f"  [{q['co']}]" if q.get('co') and cfg.get('show_co') else ""
        body = [l for l in q.get('body', []) if l is not None]
        while body and not body[-1].strip():
            body.pop()
        lines = [f"Assignment {i}: {q['title']}{co}", f"[{unit}]", ""]
        lines += body
        blocks.append("\n".join(lines) + "\n\n")
    return blocks


def build_quiz_paper(selected, data, cfg):
    marks_each = cfg.get('marks_each', 1)
    blocks = [_header_block(
        cfg['title'], data['course_code'], data['course_title'], data['semester'],
        cfg.get('date', ''), cfg.get('duration', ''), cfg.get('max_marks', ''),
        instructions=cfg.get('instructions', []),
    )]
    blocks.append("─" * 60 + "\n")
    q_lines = []
    for i, (unit, q) in enumerate(selected, 1):
        co = f"  [{q['co']}]" if q.get('co') and cfg.get('show_co') else ""
        tp = {'mcq': 'MCQ', 'tf': 'True or False', 'fib': 'Fill in the Blank'}.get(
            q.get('qtype', 'mcq'), '')
        mark_tag = f"[{marks_each} Mark{'s' if marks_each != 1 else ''}]"
        q_lines.append(f"Q{i}. ({tp})  {q['text']}{co}\t{mark_tag}")
        for opt in q.get('options', []):
            q_lines.append(f"   {opt}")
        q_lines.append("")
    blocks.append("\n".join(q_lines))
    return blocks


def build_exam_paper(sel_2, sel_5, sel_10, data, cfg):
    blocks = [_header_block(
        cfg['title'], data['course_code'], data['course_title'], data['semester'],
        cfg.get('date', ''), cfg.get('duration', ''), cfg.get('max_marks', ''),
        instructions=cfg.get('instructions', []),
    )]

    def _part(selected, marks, part_cfg):
        if not selected:
            return None
        label      = part_cfg.get('label', f'PART {chr(64 + marks // 2)}')
        answer_any = part_cfg.get('answer_any') or len(selected)
        answer_any = min(answer_any, len(selected))
        total_m    = answer_any * marks
        kind       = {2: 'Short Answer', 5: 'Descriptive', 10: 'Long Answer / Case Study'}[marks]
        instr      = ('Answer all questions' if answer_any == len(selected)
                      else f'Answer any {answer_any} of {len(selected)}')
        lines = [
            f"\n{label} — {kind} Questions  [{marks} × {answer_any} = {total_m} Marks]",
            f"({instr})",
        ]
        for i, (unit, q) in enumerate(selected, 1):
            co = f"  [{q['co']}]" if q.get('co') and cfg.get('show_co') else ""
            lines.append(f"  {i}. {q['text']}{co}")
            for ctx in q.get('context', []):
                lines.append(f"     {ctx}")
        return "\n".join(lines) + "\n"

    parts = cfg.get('parts', {})
    for sel, marks in [(sel_2, 2), (sel_5, 5), (sel_10, 10)]:
        blk = _part(sel, marks, parts.get(marks, {}))
        if blk:
            blocks.append(blk)
    return blocks


# ---------------------------------------------------------------------------
# Lab paper builders
# ---------------------------------------------------------------------------

def build_lab_practical_paper(selected, data, cfg):
    """Build a practical examination paper from selected experiments."""
    blocks = [_header_block(
        cfg['title'], data['course_code'], data['course_title'], data['semester'],
        cfg.get('date', ''), cfg.get('duration', ''), cfg.get('max_marks', ''),
        instructions=cfg.get('instructions', []),
    )]
    blocks.append("─" * 60 + "\n")
    for i, (unit, exp) in enumerate(selected, 1):
        co = f"  [{exp['co']}]" if exp.get('co') and cfg.get('show_co') else ""
        lines = [f"Experiment {i}: {exp['title']}{co}", f"[{unit}]", ""]
        if exp.get('aim'):
            lines.append("Aim: " + " ".join(exp['aim']))
        if exp.get('equipment'):
            lines.append("Equipment/Tools Required: " + " ".join(exp['equipment']))
        if exp.get('theory'):
            lines += ["", "Theory:"] + exp['theory']
        if exp.get('procedure'):
            lines += ["", "Procedure:"] + exp['procedure']
        lines += ["", "Observations/Output:", "", "Result:", ""]
        blocks.append("\n".join(lines) + "\n\n")
    return blocks


def build_lab_viva_paper(selected, data, cfg):
    """Build a viva / oral examination paper listing viva questions from selected experiments."""
    blocks = [_header_block(
        cfg['title'], data['course_code'], data['course_title'], data['semester'],
        cfg.get('date', ''), cfg.get('duration', ''), cfg.get('max_marks', ''),
        instructions=cfg.get('instructions', []),
    )]
    blocks.append("─" * 60 + "\n")
    q_num = 1
    for unit, exp in selected:
        if not exp.get('viva'):
            continue
        co = f"  [{exp['co']}]" if exp.get('co') and cfg.get('show_co') else ""
        lines = [f"\n[{unit}] {exp['title']}{co}"]
        for vq in exp['viva']:
            # Strip leading labels like "V1.", "1.", "Q1." before renumbering
            vq_clean = re.sub(r'^[A-Za-z]?\d+\.\s*', '', vq).strip()
            if vq_clean:
                lines.append(f"Q{q_num}. {vq_clean}")
                q_num += 1
        blocks.append("\n".join(lines) + "\n")
    return blocks


# ---------------------------------------------------------------------------
# Save helpers
# ---------------------------------------------------------------------------

_UNICODE_SUBS = str.maketrans({
    # dashes
    '\u2014': '--', '\u2013': '-', '\u2015': '-',
    # quotes
    '\u2018': "'", '\u2019': "'", '\u201a': ',',
    '\u201c': '"', '\u201d': '"', '\u201e': '"',
    # ellipsis / bullets
    '\u2026': '...', '\u2022': '-', '\u2023': '-', '\u25cf': '-',
    # math operators
    '\u00d7': 'x', '\u00f7': '/', '\u00b1': '+/-', '\u00b0': 'deg',
    # arrows
    '\u2192': '->', '\u2190': '<-', '\u21d2': '=>',
    '\u21d0': '<=', '\u21d4': '<=>',
    # relations
    '\u2261': '=', '\u2260': '!=', '\u2264': '<=',
    '\u2265': '>=', '\u2248': '~=',
    # logic / set
    '\u2295': 'XOR', '\u2227': 'AND', '\u2228': 'OR', '\u00ac': 'NOT',
    '\u2208': 'in', '\u2209': 'not in',
    '\u2200': 'for all', '\u2203': 'exists',
    '\u221e': 'inf', '\u2211': 'sum', '\u221a': 'sqrt',
    # Greek lower
    '\u03b1': 'alpha', '\u03b2': 'beta', '\u03b3': 'gamma',
    '\u03b4': 'delta', '\u03b5': 'epsilon', '\u03b6': 'zeta',
    '\u03b7': 'eta', '\u03b8': 'theta', '\u03ba': 'kappa',
    '\u03bb': 'lambda', '\u03bc': 'mu', '\u03bd': 'nu',
    '\u03c0': 'pi', '\u03c1': 'rho', '\u03c3': 'sigma',
    '\u03c4': 'tau', '\u03c6': 'phi', '\u03c7': 'chi',
    '\u03c8': 'psi', '\u03c9': 'omega',
    # Greek upper
    '\u0394': 'Delta', '\u03a9': 'Omega',
    '\u03a3': 'Sigma', '\u03a0': 'Pi',
    # superscripts
    '\u00b2': '^2', '\u00b3': '^3', '\u00b9': '^1',
    '\u2070': '^0', '\u207f': '^n',
    # misc
    '\u00ae': '(R)', '\u00a9': '(C)', '\u2122': '(TM)',
    '\u2032': "'", '\u2033': '"',
    '\u00a0': ' ', '\u200b': '', '\ufeff': '',
})


def _pdf_safe(text):
    import unicodedata
    text = text.translate(_UNICODE_SUBS)
    result = []
    for ch in text:
        try:
            ch.encode('latin-1')
            result.append(ch)
        except UnicodeEncodeError:
            nfkd = unicodedata.normalize('NFKD', ch)
            ascii_form = nfkd.encode('ascii', errors='ignore').decode('ascii')
            result.append(ascii_form if ascii_form else ' ')
    return ''.join(result)


def _count_co_coverage(sel_2=None, sel_5=None, sel_10=None, sel_assign=None, sel_quiz=None):
    """Count selected questions per CO, keyed by tier."""
    tally = {}

    def _blank():
        return {2: 0, 5: 0, 10: 0, 'assign': 0, 'quiz': 0}

    def _add(co_str, tier):
        if not co_str:
            return
        m = re.match(r'CO(\d+)', co_str)
        if not m:
            return
        cn = m.group(1)
        if cn not in tally:
            tally[cn] = _blank()
        tally[cn][tier] += 1

    for _, q in (sel_2     or []): _add(q.get('co'), 2)
    for _, q in (sel_5     or []): _add(q.get('co'), 5)
    for _, q in (sel_10    or []): _add(q.get('co'), 10)
    for _, q in (sel_assign or []): _add(q.get('co'), 'assign')
    for _, q in (sel_quiz   or []): _add(q.get('co'), 'quiz')

    return dict(sorted(tally.items(), key=lambda kv: int(kv[0])))


def _co_summary_txt(tally):
    """Return a plain-text CO coverage summary table."""
    if tally is None:
        return ""
    if not tally:
        return "\n## CO Coverage\n\n  (No CO tags found in the source question bank)\n"
    rows = []
    tot2 = tot5 = tot10 = tota = totq = 0
    for cn, c in tally.items():
        rows.append((f"CO{cn}", c[2], c[5], c[10], c['assign'], c['quiz'],
                     c[2] + c[5] + c[10] + c['assign'] + c['quiz']))
        tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
        tota += c['assign']; totq += c['quiz']
    rows.append(("Total", tot2, tot5, tot10, tota, totq, tot2 + tot5 + tot10 + tota + totq))

    hdr = (f"  {'CO':<8} {'2-mark':>8} {'5-mark':>8} {'10-mark':>9}"
           f" {'Assign':>8} {'Quiz':>8} {'Total':>8}")
    sep = "  " + "-" * 8 + (" " + "-" * 8) * 3 + " " + "-" * 9 + " " + "-" * 8 + " " + "-" * 8
    lines = ["\n## CO Coverage\n", hdr, sep]
    for i, (co, m2, m5, m10, ma, mq, tot) in enumerate(rows):
        if i == len(rows) - 1:
            lines.append(sep)
        lines.append(f"  {co:<8} {m2:>8} {m5:>8} {m10:>9} {ma:>8} {mq:>8} {tot:>8}")
    return '\n'.join(lines) + '\n'


def _count_co_lab_coverage(selected):
    """Count experiments per CO from a flat [(unit, exp)] list.

    Returns {co_num_str: {'exp': int}, '_lab': True} so save functions can
    distinguish this tally from a regular 5-tier tally.
    """
    tally = {}
    for _, exp in selected:
        co = exp.get('co') or ''
        m = re.match(r'CO(\d+)', co)
        if m:
            cn = m.group(1)
            tally.setdefault(cn, {'exp': 0})
            tally[cn]['exp'] += 1
    result = dict(sorted(tally.items(), key=lambda kv: int(kv[0])))
    result['_lab'] = True   # sentinel: this is a lab tally
    return result


def _co_lab_summary_txt(tally):
    """Plain-text CO coverage table for lab papers (CO | Experiments | Total)."""
    rows = [(f"CO{cn}", c['exp']) for cn, c in tally.items() if cn != '_lab']
    if not rows:
        return "\n## CO Coverage\n\n  (No CO tags found in selected experiments)\n"
    total = sum(r[1] for r in rows)
    rows.append(("Total", total))
    hdr = f"  {'CO':<8} {'Experiments':>13} {'Total':>8}"
    sep = "  " + "-"*8 + " " + "-"*13 + " " + "-"*8
    lines = ["\n## CO Coverage\n", hdr, sep]
    for i, (co, exp) in enumerate(rows):
        if i == len(rows) - 1:
            lines.append(sep)
        lines.append(f"  {co:<8} {exp:>13} {exp:>8}")
    return '\n'.join(lines) + '\n'


def save_txt(blocks, output_path, co_tally=None):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(blocks))
        if co_tally is not None:
            if co_tally.get('_lab'):
                f.write(_co_lab_summary_txt(co_tally))
            else:
                f.write(_co_summary_txt(co_tally))


def save_docx(blocks, output_path, co_tally=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    import docx.oxml.ns as qns
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    def _shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qns.qn('w:val'),   'clear')
        shd.set(qns.qn('w:color'), 'auto')
        shd.set(qns.qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def _bold_para(text, color_rgb=None, size_pt=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        if size_pt:
            run.font.size = Pt(size_pt)
        return p

    def _add_question(num, text, co_shown):
        p = doc.add_paragraph()
        run_n = p.add_run(f"{num}.  ")
        run_n.bold = True
        co_m = re.search(r'\s*(\[CO\d+\])\s*(\[.+?\])?\s*$', text)
        if co_m:
            body = text[:co_m.start()].strip()
            tags = co_m.group(0).strip()
            p.add_run(body + "  ")
            co_run = p.add_run(tags)
            co_run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
            co_run.font.size = Pt(8)
        else:
            p.add_run(text)

    for block in blocks:
        for raw in block.split('\n'):
            line = raw.strip()
            if not line or set(line) <= {'─', '-', '='}:
                doc.add_paragraph()
                continue

            # Document title
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)

            # Course / date-duration line
            elif line.startswith('Course:') or re.match(r'^Date:', line) or re.match(r'^Duration:', line):
                _bold_para(line) if line.startswith('Course:') else doc.add_paragraph(line)

            # Instructions header
            elif line == 'Instructions:':
                _bold_para('Instructions:', size_pt=10)

            # Instruction bullet
            elif re.match(r'^-\s+', line):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])

            # Part header: "PART A — ..."
            elif re.match(r'^PART [A-Z]', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            # Answer instruction: "(Answer any...)"
            elif re.match(r'^\(Answer', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.size = Pt(9)

            # Numbered exam question: "  1. text"
            elif re.match(r'^\d+\.', line):
                num, _, rest = line.partition('.')
                _add_question(num.strip(), rest.strip(), True)

            # 10-mark context continuation: "     context"
            elif re.match(r'^\s{4,}', raw) and not re.match(r'^\s+\([a-d]\)', raw):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.add_run(line).italic = True

            # MCQ option: "(a) ..."
            elif re.match(r'^\([a-d]\)', line):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                p.add_run(line)

            # Quiz / viva question: "Q1. text"
            elif re.match(r'^Q\d+\.', line):
                qm = re.match(r'(Q\d+)\.\s*(.*)', line)
                if qm:
                    parts = qm.group(2).split('\t', 1)
                    body = parts[0].rstrip()
                    mark_label = parts[1].strip() if len(parts) > 1 else ''
                    p = doc.add_paragraph()
                    run_n = p.add_run(f"{qm.group(1)}.  ")
                    run_n.bold = True
                    if mark_label:
                        pPr = p._p.get_or_add_pPr()
                        tabs_el = OxmlElement('w:tabs')
                        tab_el = OxmlElement('w:tab')
                        tab_el.set(qns.qn('w:val'), 'right')
                        tab_el.set(qns.qn('w:pos'), '9072')
                        tabs_el.append(tab_el)
                        pPr.append(tabs_el)
                        p.add_run(body)
                        p.add_run('\t')
                        run_m = p.add_run(mark_label)
                        run_m.italic = True
                        run_m.font.size = Pt(9)
                        run_m.font.color.rgb = RGBColor(80, 80, 80)
                    else:
                        # Add body directly to the existing paragraph (not via _add_question
                        # which would open a second paragraph and produce a duplicate number line)
                        co_m = re.search(r'\s*(\[CO\d+\])\s*(\[.+?\])?\s*$', body)
                        if co_m:
                            p.add_run(body[:co_m.start()].strip() + "  ")
                            co_run = p.add_run(co_m.group(0).strip())
                            co_run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
                            co_run.font.size = Pt(8)
                        else:
                            p.add_run(body)

            # Experiment title: "Experiment N: title"
            elif re.match(r'^Experiment \d+:', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(75, 0, 130)  # indigo

            # Assignment title: "Assignment N: title"
            elif re.match(r'^Assignment \d+:', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xBF, 0x47, 0x00)

            # Unit reference: "[Unit N]"
            elif re.match(r'^\[Unit', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            # Section labels (regular + lab)
            elif re.match(r'^(Context|Tasks|Deliverables|Evaluation Criteria'
                          r'|Aim|Theory|Procedure|Equipment[^:]*'
                          r'|Observations?(?:/Output)?|Result):', line):
                _bold_para(line, size_pt=10)

            # Sub-task: "(i) ...", "(ii) ..."
            elif re.match(r'^\([ivx]+\)', line):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.add_run(line)

            else:
                doc.add_paragraph(line)

    if co_tally is not None:
        doc.add_heading("CO Coverage", level=1)
        real_rows = {cn: c for cn, c in co_tally.items() if cn != '_lab'}
        if not real_rows:
            p = doc.add_paragraph("No CO tags found in the source question bank.")
            p.runs[0].italic = True
            doc.save(output_path)
            return

        if co_tally.get('_lab'):
            # Lab tally: CO | Experiments | Total
            headers = ["CO", "Experiments", "Total"]
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade_cell(hdr_cells[i], "4B0082")
            total_exp = 0
            for cn, c in real_rows.items():
                rc = table.add_row().cells
                rc[0].text = f"CO{cn}"
                rc[1].text = str(c['exp'])
                rc[2].text = str(c['exp'])
                total_exp += c['exp']
            tot_cells = table.add_row().cells
            for cell in tot_cells:
                _shade_cell(cell, "EDE0F5")
            tot_cells[0].text = "Total"
            tot_cells[1].text = str(total_exp)
            tot_cells[2].text = str(total_exp)
            for cell in tot_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
        else:
            # Regular tally: CO | 2-mark | 5-mark | 10-mark | Assign | Quiz | Total
            headers = ["CO", "2-mark", "5-mark", "10-mark", "Assign", "Quiz", "Total"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade_cell(hdr_cells[i], "1F497D")
            tot2 = tot5 = tot10 = tota = totq = 0
            for cn, c in real_rows.items():
                row_cells = table.add_row().cells
                row_cells[0].text = f"CO{cn}"
                row_cells[1].text = str(c[2])
                row_cells[2].text = str(c[5])
                row_cells[3].text = str(c[10])
                row_cells[4].text = str(c['assign'])
                row_cells[5].text = str(c['quiz'])
                row_cells[6].text = str(c[2] + c[5] + c[10] + c['assign'] + c['quiz'])
                tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
                tota += c['assign']; totq += c['quiz']
            tot_cells = table.add_row().cells
            for cell in tot_cells:
                _shade_cell(cell, "DCE6F1")
            tot_cells[0].text = "Total"
            tot_cells[1].text = str(tot2)
            tot_cells[2].text = str(tot5)
            tot_cells[3].text = str(tot10)
            tot_cells[4].text = str(tota)
            tot_cells[5].text = str(totq)
            tot_cells[6].text = str(tot2 + tot5 + tot10 + tota + totq)
            for cell in tot_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.save(output_path)


def save_pdf(blocks, output_path, co_tally=None):
    from fpdf import FPDF, XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    def _h1(text):
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_fill_color(31, 73, 125)
        pdf.set_text_color(255, 255, 255)
        pdf.multi_cell(0, 10, text, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def _course_line(text):
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _detail_line(text):
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _instructions_label():
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, "Instructions:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _instruction_bullet(text):
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(6, 5, "-")
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _part_hdr(text):
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_fill_color(220, 230, 241)
        pdf.multi_cell(0, 8, text, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _answer_instr(text):
        pdf.set_font("Helvetica", "I", 9)
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _exam_q(num, text):
        co_m = re.search(r'\s*(\[CO\d+\].*?)$', text)
        body = text[:co_m.start()].strip() if co_m else text
        tags = co_m.group(1).strip() if co_m else ""
        label = f"{num}.  "
        pdf.set_font("Helvetica", "B", 10)
        lw = pdf.get_string_width(label)
        pdf.cell(lw, 6, label)
        pdf.set_font("Helvetica", "", 10)
        full = (body + "  " + tags) if tags else body
        pdf.multi_cell(0, 6, full, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _context_line(text):
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(10, 5, "")
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _mcq_option(text):
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(12, 5, "")
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _assign_title(text):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(191, 71, 0)
        pdf.multi_cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    def _unit_ref(text):
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    def _section_label(text):
        pdf.set_font("Helvetica", "B", 10)
        pdf.multi_cell(0, 6, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _sub_task(text):
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(8, 5, "")
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _quiz_q(num_label, rest_text):
        mark_m = re.search(r'\t(\[[\d]+ Marks?\])\s*$', rest_text, re.IGNORECASE)
        if mark_m:
            mark_label = mark_m.group(1)
            body = rest_text[:mark_m.start()].strip()
        else:
            mark_label = ''
            body = rest_text
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin
        num_str = f"{num_label}.  "
        pdf.set_font("Helvetica", "B", 10)
        num_w = pdf.get_string_width(num_str)
        if mark_label:
            pdf.set_font("Helvetica", "I", 9)
            mark_w = pdf.get_string_width(mark_label) + 2
        else:
            mark_w = 0
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(num_w, 6, num_str)
        pdf.set_font("Helvetica", "", 10)
        text_w = usable_w - num_w - mark_w
        y0 = pdf.get_y()
        pdf.multi_cell(text_w, 6, body, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        y1 = pdf.get_y()
        if mark_label:
            pdf.set_xy(pdf.l_margin + usable_w - mark_w, y0)
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(mark_w, 6, mark_label, align='R', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            if pdf.get_y() < y1:
                pdf.set_y(y1)
        pdf.ln(1)

    for block in blocks:
        for raw in block.split('\n'):
            line = _pdf_safe(raw.strip())

            if not line or set(line) <= {'-', '=', chr(0x2500)}:
                pdf.ln(2)
                continue

            if line.startswith('# '):
                _h1(line[2:])
            elif line.startswith('Course:'):
                _course_line(line)
            elif re.match(r'^(Date:|Duration:|Max Marks:)', line):
                _detail_line(line)
            elif line == 'Instructions:':
                _instructions_label()
            elif re.match(r'^-\s+', line):
                _instruction_bullet(line[2:])
            elif re.match(r'^PART [A-Z]', line):
                _part_hdr(line)
            elif re.match(r'^\(Answer', line):
                _answer_instr(line)
            elif re.match(r'^\d+\.', line):
                num, _, rest = line.partition('.')
                _exam_q(num.strip(), rest.strip())
            elif re.match(r'^\s{4,}', raw) and not re.match(r'^\s+\([a-d]\)', raw):
                _context_line(line)
            elif re.match(r'^\([a-d]\)', line):
                _mcq_option(line)
            elif re.match(r'^Q\d+\.', line):
                qm = re.match(r'(Q\d+)\.\s*(.*)', line)
                if qm:
                    _quiz_q(qm.group(1), qm.group(2))
            elif re.match(r'^Experiment \d+:', line):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(75, 0, 130)
                pdf.multi_cell(0, 7, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_text_color(0, 0, 0)
            elif re.match(r'^Assignment \d+:', line):
                _assign_title(line)
            elif re.match(r'^\[Unit', line):
                _unit_ref(line)
            elif re.match(r'^(Context|Tasks|Deliverables|Evaluation Criteria'
                          r'|Aim|Theory|Procedure|Equipment[^:]*'
                          r'|Observations?(?:/Output)?|Result):', line):
                _section_label(line)
            elif re.match(r'^\([ivx]+\)', line):
                _sub_task(line)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(0, 5, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if co_tally is not None:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(220, 230, 241)
        pdf.multi_cell(0, 8, "CO Coverage", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        real_rows = {cn: c for cn, c in co_tally.items() if cn != '_lab'}
        if not real_rows:
            pdf.set_font("Helvetica", "I", 10)
            pdf.multi_cell(0, 6, "No CO tags found in the source question bank.",
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.output(output_path)
            return

        if co_tally.get('_lab'):
            col_widths = [22, 80, 24]
            headers    = ["CO", "Experiments", "Total"]

            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(75, 0, 130)
            pdf.set_text_color(255, 255, 255)
            for w, h in zip(col_widths, headers):
                pdf.cell(w, 8, h, border=1, fill=True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)

            tot_exp = 0
            for idx, (cn, c) in enumerate(real_rows.items()):
                fill = idx % 2 == 0
                pdf.set_fill_color(240, 235, 255) if fill else pdf.set_fill_color(255, 255, 255)
                pdf.set_font("Helvetica", "", 9)
                exp_count = c['exp']
                vals = [f"CO{cn}", str(exp_count), str(exp_count)]
                for w, v in zip(col_widths, vals):
                    pdf.cell(w, 7, v, border=1, fill=fill)
                pdf.ln()
                tot_exp += exp_count

            pdf.set_fill_color(220, 210, 240)
            pdf.set_font("Helvetica", "B", 9)
            for w, v in zip(col_widths, ["Total", str(tot_exp), str(tot_exp)]):
                pdf.cell(w, 8, v, border=1, fill=True)
            pdf.ln()

        else:
            col_widths = [22, 24, 24, 28, 24, 24, 24]
            headers    = ["CO", "2-mark", "5-mark", "10-mark", "Assign", "Quiz", "Total"]

            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(31, 73, 125)
            pdf.set_text_color(255, 255, 255)
            for w, h in zip(col_widths, headers):
                pdf.cell(w, 8, h, border=1, fill=True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)

            tot2 = tot5 = tot10 = tota = totq = 0
            for idx, (cn, c) in enumerate(real_rows.items()):
                fill = idx % 2 == 0
                pdf.set_fill_color(240, 245, 255) if fill else pdf.set_fill_color(255, 255, 255)
                pdf.set_font("Helvetica", "", 9)
                vals = [f"CO{cn}", str(c[2]), str(c[5]), str(c[10]),
                        str(c['assign']), str(c['quiz']),
                        str(c[2] + c[5] + c[10] + c['assign'] + c['quiz'])]
                for w, v in zip(col_widths, vals):
                    pdf.cell(w, 7, v, border=1, fill=fill)
                pdf.ln()
                tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
                tota += c['assign']; totq += c['quiz']

            pdf.set_fill_color(220, 230, 241)
            pdf.set_font("Helvetica", "B", 9)
            for w, v in zip(col_widths, ["Total", str(tot2), str(tot5), str(tot10),
                                          str(tota), str(totq),
                                          str(tot2 + tot5 + tot10 + tota + totq)]):
                pdf.cell(w, 8, v, border=1, fill=True)
            pdf.ln()

    pdf.output(output_path)


# ---------------------------------------------------------------------------
# Syllabus → Question Bank generation (delegates to generate_qbank)
# ---------------------------------------------------------------------------

def _generate_from_syllabus(syllabus_path, co_path=None):
    """Generate a question bank interactively from a syllabus.

    Imports generate_qbank and reuses its generation pipeline.
    Returns (qbank_text_str, saved_qbank_path) or (None, None) on error.
    The returned text string uses the txt/markdown format with ## markers
    so parse_qbank() can consume it directly without re-reading from disk.
    """
    try:
        import generate_qbank as qb
        import anthropic
    except ImportError as exc:
        print(f"Error: {exc}.  generate_qbank.py must be in the same directory.")
        return None, None

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("Error: ANTHROPIC_API_KEY is not set in the environment.")
        return None, None

    # ---- Load syllabus ----
    syllabus_text = _extract_generic(syllabus_path)
    courses = qb.split_into_courses(syllabus_text)
    if not courses:
        print("Error: No course codes detected in the syllabus.")
        return None, None

    # ---- Auto-detect course from CO filename, or ask ----
    auto_code = None
    if co_path:
        m = re.search(r'[_\-]([A-Z0-9]{4,20})_COs\.', os.path.basename(co_path), re.IGNORECASE)
        if m and m.group(1).upper() in courses:
            auto_code = m.group(1).upper()

    print("\n" + "=" * 60)
    print("  Available Courses")
    print("=" * 60)
    for code, info in courses.items():
        sem = f"  [Sem {info['semester']}]" if info['semester'] else ""
        print(f"  {code:<20} {info['title'][:45]}{sem}")
    print("=" * 60)

    if len(courses) == 1:
        matched = list(courses.keys())[0]
        print(f"\nAuto-selected: {matched}")
    elif auto_code:
        matched = auto_code
        print(f"\nAuto-detected course: {matched}")
    else:
        raw = input("\nEnter Course Code: ").strip().upper()
        matched = raw if raw in courses else next((c for c in courses if raw in c), None)
        if not matched:
            print(f"Course code not found.")
            return None, None

    info      = courses[matched]
    sem_label = f"  [Semester {info['semester']}]" if info['semester'] else ""
    print(f"\nSelected: {matched} — {info['title']}{sem_label}")

    client = anthropic.Anthropic(api_key=api_key, timeout=120.0)

    # ---- Load or generate COs ----
    if co_path:
        print(f"\nLoading CO file: {co_path}")
        co_ext = os.path.splitext(co_path)[1].lower()
        if co_ext == '.pdf':
            cos = qb.load_cos_from_pdf(co_path, matched)
        else:
            raw_co  = _extract_generic(co_path)
            co_text = qb._extract_co_section(raw_co, matched)
            cos     = qb.parse_cos_from_text(co_text)
        print(f"  Loaded from: {os.path.basename(co_path)}")
    else:
        NUM_COS = 15
        print(f"\nNo CO file provided — generating {NUM_COS} COs for {matched}...")
        print("=" * 60)
        co_text = qb.generate_cos_for_course(client, matched, info['title'], info['text'], num_cos=NUM_COS)
        print("\nEnter CO numbers to keep (e.g. 1,3,5  |  1-5  |  all) [default: all]:")
        raw_sel  = input("Selection: ").strip()
        selected = qb.parse_co_selection(raw_sel, NUM_COS)
        co_text  = qb.filter_cos(co_text, selected)
        cos      = qb.parse_cos_from_text(co_text)

    if not cos:
        print("Error: No COs could be parsed.")
        return None, None

    print(f"\nFound {len(cos)} Course Outcome(s):")
    for co in cos:
        print(f"  CO{co['num']}: {co['statement'][:70]}")

    # ---- Build unit plan ----
    unit_plan  = qb.build_unit_plan(client, cos, info['text'])
    num_units  = len(unit_plan)
    print(f"\nOrganised into {num_units} unit(s):")
    for ulabel, udata in unit_plan.items():
        co_nums = ', '.join(f"CO{c['num']}" for c in udata['cos'])
        print(f"  {ulabel} — {udata['title'] or ''}  [{co_nums or 'no COs mapped'}]")

    # ---- Question counts ----
    print("\n" + "=" * 60)
    print("Question Distribution per Unit  (press Enter for defaults)")
    print("=" * 60)
    n2       = max(1, _ask_int("2-mark  questions per unit", 4))
    n5       = max(1, _ask_int("5-mark  questions per unit", 4))
    n10      = max(1, _ask_int("10-mark questions per unit", 4))
    n_assign = _ask_int("Assignment tasks  per unit (0 to skip)", 2)
    n_quiz   = _ask_int("Quiz questions    per unit (0 to skip)", 5)
    total_per_unit = n2 + n5 + n10

    # ---- Question bank output format ----
    print("\nSave question bank as:")
    print("  1. txt")
    print("  2. docx  (Word)")
    print("  3. pdf")
    fmt_raw = input("Enter 1, 2, or 3 [default: 1]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_raw, "txt")

    # ---- Generate ----
    print(f"\nGenerating question bank...")
    print("=" * 60)

    all_blocks = []
    extras = []
    if n_assign: extras.append(f"{n_assign} assignment task(s)")
    if n_quiz:   extras.append(f"{n_quiz} quiz question(s)")
    extra_line = ("  |  " + " + ".join(extras) + " per unit\n") if extras else "\n"
    all_blocks.append(
        f"# Question Bank\n"
        f"Course: {matched} — {info['title']}{sem_label}\n"
        f"Exam questions per unit: {n2} x 2-mark + {n5} x 5-mark + {n10} x 10-mark"
        f" = {total_per_unit} per unit  |  {total_per_unit * num_units} total\n"
        + extra_line
    )

    for idx, (unit_label, udata) in enumerate(unit_plan.items(), 1):
        unit_title = udata['title']
        unit_cos   = udata['cos']
        full_unit  = f"{unit_label}: {unit_title}" if unit_title else unit_label
        co_summary = "  |  ".join(f"CO{c['num']}" for c in unit_cos) or "no COs mapped"

        header = f"\n## {full_unit}\n"
        if unit_cos:
            header += f"### Course Outcomes: {co_summary}\n"
            for co in unit_cos:
                header += f"### CO{co['num']}: {co['statement']}\n"
        else:
            header += "### Questions based on syllabus content\n"
        all_blocks.append(header)

        print(f"\n[{idx}/{num_units}] {full_unit}  ({co_summary})")
        print("-" * 60)

        qtext = qb.generate_questions_for_unit(
            client, unit_label, unit_title, unit_cos,
            matched, info['title'], info['text'],
            n2=n2, n5=n5, n10=n10,
        )
        qtext = qb._normalize_qblock(qtext, unit="")
        if qtext.strip():
            all_blocks.append(qtext + "\n")
        else:
            print(f"  WARNING: no exam questions generated for {full_unit}.")

        if n_assign:
            print(f"\n  Generating {n_assign} assignment task(s)...")
            print("-" * 60)
            atext = qb.generate_assignments_for_unit(
                client, unit_label, unit_title, unit_cos,
                matched, info['title'], info['text'], n_assign=n_assign,
            )
            atext = qb._normalize_assignment_block(atext)
            if atext.strip():
                all_blocks.append(atext + "\n")
            else:
                print(f"  WARNING: no assignments generated for {full_unit}.")

        if n_quiz:
            print(f"\n  Generating {n_quiz} quiz question(s)...")
            print("-" * 60)
            qztext = qb.generate_quiz_for_unit(
                client, unit_label, unit_title, unit_cos,
                matched, info['title'], info['text'], n_quiz=n_quiz,
            )
            qztext = qb._normalize_quiz_block(qztext)
            if qztext.strip():
                all_blocks.append(qztext + "\n")
            else:
                print(f"  WARNING: no quiz questions generated for {full_unit}.")

    # ---- Save question bank ----
    base_path  = os.path.splitext(syllabus_path)[0]
    qbank_path = f"{base_path}_{matched}_QuestionBank.{fmt}"
    co_tally   = qb.count_co_questions(all_blocks)

    if fmt == "txt":
        qb.save_txt(all_blocks, qbank_path, co_tally=co_tally)
    elif fmt == "docx":
        qb.save_docx(all_blocks, qbank_path, co_tally=co_tally)
    elif fmt == "pdf":
        qb.save_pdf(all_blocks, qbank_path, co_tally=co_tally)

    print(f"\nQuestion bank saved to: {qbank_path}")

    # Return the markdown text directly so parse_qbank() can consume it
    # without re-reading the saved file (avoids PDF-wrapping issues on re-load)
    return "\n".join(all_blocks), qbank_path


# ---------------------------------------------------------------------------
# Lab paper interactive flow (called from main when lab QB detected)
# ---------------------------------------------------------------------------

def _run_lab_paper_flow(data, qb_path):
    sem = f"  [Semester {data['semester']}]" if data['semester'] else ""
    print(f"\nCourse: {data['course_code']} — {data['course_title']}{sem}")
    flat_all = _flat_lab_list(data['units'])
    print(f"\nAvailable experiments:")
    print(f"  {len(data['units'])} unit(s)  |  {len(flat_all)} experiment(s)")
    print(f"  {'─' * 58}")
    for i, (unit, exp) in enumerate(flat_all, 1):
        co = f"  [{exp['co']}]" if exp.get('co') else ""
        print(f"  {i:3}.  [{unit}]  {exp['title'][:65]}{co}")

    print("\n" + "=" * 60)
    print("  Select Paper Type")
    print("=" * 60)
    print("  1. Practical Exam  (experiment + blank observation/result sheet)")
    print("  2. Viva / Oral Exam  (viva questions only)")
    ptype_raw = input("\nEnter 1-2: ").strip()
    ptype = {"1": "practical", "2": "viva"}.get(ptype_raw)
    if not ptype:
        print("Invalid choice. Exiting.")
        sys.exit(1)

    show_co = input("\n  Show CO tags on the paper? (y/n) [n]: ").strip().lower() == 'y'
    print()

    # Select experiments
    raw_sel = input(f"  Select experiments (numbers/ranges/all) [default: all]: ").strip().lower()
    if not raw_sel or raw_sel == 'all':
        selected = list(flat_all)
    else:
        idxs = set()
        for part in raw_sel.split(','):
            part = part.strip()
            if '-' in part:
                a, _, b = part.partition('-')
                try:
                    idxs.update(range(int(a) - 1, int(b)))
                except ValueError:
                    pass
            else:
                try:
                    idxs.add(int(part) - 1)
                except ValueError:
                    pass
        selected = [flat_all[i] for i in sorted(idxs) if 0 <= i < len(flat_all)]

    if not selected:
        print("No experiments selected. Exiting.")
        sys.exit(0)

    if ptype == 'practical':
        cfg = {
            'title':    _ask("Paper title", "Practical Examination"),
            'date':     _ask("Date (leave blank to fill later)", ""),
            'duration': _ask("Duration", "3 Hours"),
            'max_marks': _ask("Max marks", ""),
            'show_co':  show_co,
            'instructions': [
                "Perform the given experiment and record your observations.",
                "Write the aim, procedure, observations, and result in your lab record.",
                "A brief viva will be conducted at the end of the session.",
            ],
        }
        blocks = build_lab_practical_paper(selected, data, cfg)
        label  = "Practical"
    else:
        cfg = {
            'title':    _ask("Paper title", "Viva Voce Examination"),
            'date':     _ask("Date (leave blank to fill later)", ""),
            'duration': _ask("Duration", "30 Minutes"),
            'max_marks': _ask("Max marks", ""),
            'show_co':  show_co,
            'instructions': ["Answer all questions clearly and concisely."],
        }
        blocks = build_lab_viva_paper(selected, data, cfg)
        label  = "Viva"

    # Output format
    print("\nOutput format:")
    print("  1. txt\n  2. docx  (Word)\n  3. pdf")
    fmt_input = input("Enter 1, 2, or 3 [default: 1]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_input, "txt")

    base        = os.path.splitext(qb_path)[0]
    output_path = f"{base}_{label}_Paper.{fmt}"
    if fmt == 'txt':
        save_txt(blocks, output_path)
    elif fmt == 'docx':
        save_docx(blocks, output_path)
    elif fmt == 'pdf':
        save_pdf(blocks, output_path)
    print(f"\nSaved to: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_qpaper.py <question_bank_file|syllabus_file> [co_file]")
        sys.exit(1)

    qb_path  = sys.argv[1]
    co_path  = sys.argv[2] if len(sys.argv) > 2 else None
    if not os.path.exists(qb_path):
        print(f"Error: File not found: {qb_path}")
        sys.exit(1)
    if co_path and not os.path.exists(co_path):
        print(f"Error: CO file not found: {co_path}")
        sys.exit(1)

    print(f"\nLoading: {qb_path}")
    raw_text = _load_qbank_text(qb_path)
    if not raw_text.strip():
        print("Error: No text could be extracted from the file.")
        sys.exit(1)

    # Lab question banks use a different structure — detect first
    if is_lab_qbank(raw_text):
        data = parse_lab_qbank(raw_text)
        if not data['units']:
            print("Error: Could not parse lab question bank structure.")
            sys.exit(1)
        _run_lab_paper_flow(data, qb_path)
        return

    data = parse_qbank(raw_text)

    if not data['units']:
        # Try treating the file as a syllabus and generating a question bank
        print("\nNo question bank detected — treating as syllabus.")
        print("Generating question bank first...\n")
        raw_text, qb_path = _generate_from_syllabus(qb_path, co_path)
        if not raw_text:
            sys.exit(1)
        data = parse_qbank(raw_text)
        if not data['units']:
            print("Error: Generated question bank could not be parsed.")
            sys.exit(1)
        print()

    # Summary
    sem = f"  [Semester {data['semester']}]" if data['semester'] else ""
    print(f"\nCourse: {data['course_code']} — {data['course_title']}{sem}")
    total = {
        2:  sum(len(u['exam'][2])  for u in data['units'].values()),
        5:  sum(len(u['exam'][5])  for u in data['units'].values()),
        10: sum(len(u['exam'][10]) for u in data['units'].values()),
        'a': sum(len(u['assign'])  for u in data['units'].values()),
        'q': sum(len(u['quiz'])    for u in data['units'].values()),
    }
    print(f"\nAvailable questions:")
    print(f"  {len(data['units'])} unit(s)")
    print(f"  2-mark:     {total[2]}")
    print(f"  5-mark:     {total[5]}")
    print(f"  10-mark:    {total[10]}")
    print(f"  Assignment: {total['a']}")
    print(f"  Quiz:       {total['q']}")

    # Paper type
    print("\n" + "=" * 60)
    print("  Select Paper Type")
    print("=" * 60)
    print("  1. Assignment")
    print("  2. Quiz")
    print("  3. Internal Assessment")
    print("  4. Final Exam")
    ptype_raw = input("\nEnter 1-4: ").strip()
    ptype = {"1": "assign", "2": "quiz", "3": "internal", "4": "final"}.get(ptype_raw)
    if not ptype:
        print("Invalid choice. Exiting.")
        sys.exit(1)

    show_co = input("\n  Show CO tags on the paper? (y/n) [n]: ").strip().lower() == 'y'
    print()

    # ---- Assignment ----
    if ptype == 'assign':
        cfg = {
            'title':        _ask("Paper title", "Assignment 1"),
            'date':         _ask("Submission deadline (leave blank to fill later)", ""),
            'max_marks':    _ask("Max marks (leave blank to fill later)", ""),
            'instructions': [
                "Complete all tasks independently.",
                "Cite all references. Plagiarism will result in a zero grade.",
            ],
            'show_co': show_co,
        }
        print()
        items = _flat_list(data['units'], 'assign')
        _display(items, 'assign')
        selected = _select(items, " assignments to include")
        if not selected:
            print("No assignments selected. Exiting.")
            sys.exit(0)
        blocks = build_assignment_paper(selected, data, cfg)

    # ---- Quiz ----
    elif ptype == 'quiz':
        marks_each = _ask_int("Marks per question", 1)
        cfg = {
            'title':        _ask("Paper title", "Quiz 1"),
            'date':         _ask("Date (leave blank to fill later)", ""),
            'duration':     _ask("Duration", "20 minutes"),
            'marks_each':   marks_each,
            'instructions': [
                "For MCQ questions, choose the single best answer.",
                "For Fill in the Blank, write the exact missing term.",
                "No negative marking.",
            ],
            'show_co': show_co,
        }
        print()
        items = _flat_list(data['units'], 'quiz')
        _display(items, 'quiz')
        selected = _select(items, " quiz questions to include")
        if not selected:
            print("No quiz questions selected. Exiting.")
            sys.exit(0)
        cfg['max_marks'] = str(len(selected) * marks_each)
        blocks = build_quiz_paper(selected, data, cfg)

    # ---- Internal / Final ----
    else:
        defaults = {
            'internal': {'title': 'Internal Assessment Test - I', 'duration': '1.5 Hours', 'max_marks': '30'},
            'final':    {'title': 'Semester End Examination',     'duration': '3 Hours',   'max_marks': '100'},
        }[ptype]

        cfg = {
            'title':     _ask("Paper title",  defaults['title']),
            'date':      _ask("Exam date (leave blank to fill later)", ""),
            'duration':  _ask("Duration",     defaults['duration']),
            'max_marks': _ask("Max marks",    defaults['max_marks']),
            'show_co':   show_co,
            'parts':     {},
            'instructions': [],
        }

        part_defaults = {2: 'PART A', 5: 'PART B', 10: 'PART C'}
        sel_2 = sel_5 = sel_10 = []
        instructions = []

        for marks in (2, 5, 10):
            avail = _flat_list(data['units'], 'exam', marks)
            if not avail:
                continue
            print(f"\n{'='*60}")
            print(f"  {marks}-Mark Questions for the Paper")
            print(f"{'='*60}")
            _display(avail, 'exam', marks)
            sel = _select(avail, f" {marks}-mark questions to include")
            if not sel:
                continue

            n = len(sel)
            label = _ask(f"Part label for {marks}-mark section", part_defaults[marks])
            if n == 1:
                answer_any = 1
            else:
                answer_any = _ask_int(
                    f"'Answer any X of {n}' — enter X (or {n} for answer-all)", n
                )
                answer_any = min(max(1, answer_any), n)

            cfg['parts'][marks] = {'label': label, 'answer_any': answer_any if answer_any < n else None}
            instr_text = ("Answer all questions" if answer_any == n
                          else f"Answer any {answer_any} of {n}")
            instructions.append(f"{label}: {instr_text} ({marks} marks each)")

            if marks == 2:   sel_2  = sel
            elif marks == 5: sel_5  = sel
            else:            sel_10 = sel

        if not sel_2 and not sel_5 and not sel_10:
            print("No questions selected. Exiting.")
            sys.exit(0)

        cfg['instructions'] = instructions
        blocks = build_exam_paper(sel_2, sel_5, sel_10, data, cfg)

    # Output format
    print("\nOutput format:")
    print("  1. txt")
    print("  2. docx  (Word)")
    print("  3. pdf")
    fmt_input = input("Enter 1, 2, or 3 [default: 1]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_input, "txt")
    if fmt_input and fmt_input not in ("1", "2", "3"):
        print("Invalid choice — using txt.")

    suffix = {'assign': 'Assignment', 'quiz': 'Quiz', 'internal': 'Internal', 'final': 'Final'}[ptype]
    base        = os.path.splitext(qb_path)[0]
    output_path = f"{base}_{suffix}_Paper.{fmt}"

    if ptype == 'assign':
        co_tally = _count_co_coverage(sel_assign=selected)
    elif ptype == 'quiz':
        co_tally = _count_co_coverage(sel_quiz=selected)
    else:
        co_tally = _count_co_coverage(sel_2=sel_2, sel_5=sel_5, sel_10=sel_10)

    if fmt == 'txt':
        save_txt(blocks, output_path, co_tally=co_tally)
    elif fmt == 'docx':
        save_docx(blocks, output_path, co_tally=co_tally)
    elif fmt == 'pdf':
        save_pdf(blocks, output_path, co_tally=co_tally)

    print(f"\nSaved to: {output_path}")

    print("\nCO Coverage:")
    if not co_tally:
        print("  (No CO tags found in the source question bank — regenerate the question bank to include them)")
    else:
        print(f"  {'CO':<8} {'2-mark':>8} {'5-mark':>8} {'10-mark':>9} {'Assign':>8} {'Quiz':>8} {'Total':>8}")
        print(f"  {'-'*8} {'-'*8} {'-'*8} {'-'*9} {'-'*8} {'-'*8} {'-'*8}")
        tot2 = tot5 = tot10 = tota = totq = 0
        for cn, c in co_tally.items():
            row_tot = c[2] + c[5] + c[10] + c['assign'] + c['quiz']
            print(f"  {'CO'+cn:<8} {c[2]:>8} {c[5]:>8} {c[10]:>9} {c['assign']:>8} {c['quiz']:>8} {row_tot:>8}")
            tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
            tota += c['assign']; totq += c['quiz']
        print(f"  {'-'*8} {'-'*8} {'-'*8} {'-'*9} {'-'*8} {'-'*8} {'-'*8}")
        print(f"  {'Total':<8} {tot2:>8} {tot5:>8} {tot10:>9} {tota:>8} {totq:>8} {tot2+tot5+tot10+tota+totq:>8}")


if __name__ == "__main__":
    main()
