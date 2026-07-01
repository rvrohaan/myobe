"""Module 4 Comprehensive Assessment Report Generator.

Builds a DOCX/PDF/TXT report covering all Module-4 deliverables:
  1.  Course Overview
  2.  CO-PO-PSO Mapping Table
  3.  CO Strength Summary (avg PO score per CO)
  4.  CO Attainment Analysis
  5.  PO Attainment Summary
  6.  Attainment Level Distribution
  7.  CO to SDG Contribution Analysis
  8.  PO to SDG Contribution Analysis
  9.  Composite SDG Index
 10.  OBE Assessment Compliance Checklist
 11.  Action Taken Report
"""

import re
import textwrap
import datetime

try:
    import report_charts as _rc
except Exception:
    _rc = None


# ── Section explanations ──────────────────────────────────────────────────────
# Plain-language notes that interpret each section's results for the reader.
# They explain WHY the numbers look the way they do without exposing any of the
# underlying scoring formulas.

def build_explanations(a, ai=None):
    """Return a dict of section-key -> explanation paragraph(s).

    Each note is data-aware where it helps, but never reveals how a value is
    computed - only what it means and why it turned out that way.
    """
    mean = a.get("mean_att_pct", 0)
    ld   = a.get("level_dist", {}) or {}
    high = ld.get(3, 0) + ld.get(2, 0)
    low  = ld.get(1, 0) + ld.get(0, 0)
    comp = a.get("sdgpo_composite", 0)
    n_sdg = len(a.get("sdgs_covered", []) or [])

    comp_band = ("a strong overall alignment" if comp >= 70 else
                 "a moderate alignment"        if comp >= 50 else
                 "a developing alignment")
    dist_read = ("most outcomes were attained at a high standard"
                 if high >= low else
                 "a sizeable share of outcomes still sit at the lower levels")
    mean_read = ("the cohort cleared the bulk of the intended outcomes"
                 if mean >= 60 else
                 "several outcomes fell short of the expected bar")

    t3 = (a.get("thresholds", {}) or {}).get("t3", 50)

    def _names(items, n=3):
        items = [str(x) for x in items if str(x).strip()]
        if not items:
            return ""
        head = ", ".join(items[:n])
        return head + (f" and {len(items) - n} more" if len(items) > n else "")

    # CO performance extremes
    _co_sorted = sorted([r for r in a.get("co_results", []) if r.get("co")],
                        key=lambda r: r.get("pct", 0))
    top_co    = _co_sorted[-1] if _co_sorted else None
    bottom_co = _co_sorted[0]  if _co_sorted else None
    weak_cos  = [r.get("co", "") for r in _co_sorted if (r.get("pct") or 0) < t3]

    # Broadest-reaching CO (by average PO support)
    _str_sorted = sorted((a.get("co_strength", {}) or {}).items(),
                         key=lambda kv: kv[1], reverse=True)
    strong_co = _str_sorted[0] if _str_sorted else None

    # PO column averages (from the mapping)
    _po_sorted = sorted((a.get("po_avg", {}) or {}).items(),
                        key=lambda kv: kv[1], reverse=True)
    top_po  = _po_sorted[0]  if _po_sorted else None
    thin_po = _po_sorted[-1] if _po_sorted else None

    # POs below target
    below_pos = [p.get("po", "") for p in a.get("atr_rows", [])]
    n_po_eval = len(a.get("po_attainment", []))

    # SDG-CO top contributor and SDG-PO leaders
    _sdgco_sorted = sorted(a.get("sdgco_contrib", []),
                           key=lambda c: c.get("pct", 0), reverse=True)
    top_sdgco = _sdgco_sorted[0] if _sdgco_sorted else None
    _sdgpo_sorted = sorted(a.get("sdgpo_results", []),
                           key=lambda r: r.get("contribution", 0), reverse=True)
    lead_sdgs = [r.get("sdg", "") for r in _sdgpo_sorted[:3]
                 if (r.get("contribution") or 0) > 0]

    # ── Optional, data-specific sentences (blank when the data isn't present) ──
    co_spread_note = (
        f"Of the COs assessed, {high} sit at the upper levels (2-3) and {low} at the "
        "lower levels (0-1), so the headline is "
        + ("backed by broad depth across the class. "
           if high >= low else
           "carried by a smaller group of strong outcomes. ")
    ) if (high or low) else ""

    co_extremes_note = (
        f"In practice {top_co.get('co')} performed best at {(top_co.get('pct') or 0):.0f}%, "
        f"while {bottom_co.get('co')} was lowest at {(bottom_co.get('pct') or 0):.0f}%. "
    ) if top_co and bottom_co and top_co is not bottom_co else ""

    weak_co_note = (
        f"The outcomes below the {t3}% floor - {_names(weak_cos)} - are the clearest "
        "priorities for revised delivery or assessment redesign. "
    ) if weak_cos else ""

    strong_co_note = (
        f"{strong_co[0]} reaches across the most programme outcomes (average "
        f"{strong_co[1]}), so it carries a large share of the graduate attributes and "
        "a dip there would ripple into several POs. "
    ) if strong_co and strong_co[1] else ""

    po_extreme_note = (
        f"{top_po[0]} is the best-supported programme outcome (column average "
        f"{top_po[1]}), whereas {thin_po[0]} rests on the thinnest CO support "
        f"({thin_po[1]}) and is the most exposed if any feeding CO slips. "
    ) if top_po and thin_po and top_po[0] != thin_po[0] else ""

    below_po_note = (
        f"{len(below_pos)} of {n_po_eval} POs are below target ({_names(below_pos)}); "
        "each needs a documented corrective action in the final section. "
    ) if below_pos else (
        "Every programme outcome currently clears its target. " if n_po_eval else ""
    )

    atr_count_note = (
        f"This cycle {len(below_pos)} of {n_po_eval} POs ({_names(below_pos)}) fell "
        "short and appear below with their planned actions. "
    ) if below_pos else (
        "This cycle every programme outcome met its target, so no corrective actions "
        "are required - the section is kept as evidence that the check was carried "
        "out. " if n_po_eval else ""
    )

    top_sdgco_note = (
        f"{top_sdgco.get('co')} is the single largest contributor, accounting for "
        f"{(top_sdgco.get('pct') or 0):.0f}% of the alignment to "
        f"{a.get('target_sdg') or 'this goal'}. "
    ) if top_sdgco and top_sdgco.get("pct") else ""

    lead_sdg_note = (
        f"The goals this course speaks to most strongly are {_names(lead_sdgs)}, which "
        "is typical for an engineering subject where quality-education and industry-"
        "and-innovation themes dominate. "
    ) if lead_sdgs else ""

    return {
        "overview": (
            "This section sets the scope of the assessment - the number of course "
            f"outcomes ({a.get('n_cos', 0)}) and programme outcomes ({a.get('n_pos', 0)}) "
            "evaluated, together with the average level at which students achieved the "
            f"COs. A mean CO attainment of {mean}% means that, on balance, {mean_read}. "
            "Treat the average as a starting point rather than a verdict: the same "
            "figure can come from a tightly clustered cohort or from a mix of very "
            f"strong and very weak outcomes. {co_spread_note}"
            "The sections that follow break this headline down outcome by outcome and "
            "trace it through to programme outcomes and sustainability goals."
        ),
        "copo": (
            "Each value shows how strongly a course outcome supports a programme "
            "outcome: a 3 means the CO contributes heavily to that PO, while a 0 "
            "means it has little bearing on it. The mix of strong and blank cells "
            "reflects the natural fit between what each CO teaches and the broader "
            "competencies the programme expects - design- and skill-oriented COs "
            "tend to map strongly to more POs, which is why their rows look denser. A "
            "row that is mostly filled marks a CO that pulls real weight, while a "
            "column that is mostly blank points to a programme outcome few COs reach "
            f"and that the syllabus may under-serve. {strong_co_note}"
            "Read the table by rows to see which outcomes do the most work, and by "
            "columns to see where the curriculum is thin."
        ),
        "co_strength": (
            "The average here tells you how broadly each CO reaches across the "
            "programme outcomes. Outcomes marked 'Strong' touch many POs at a high "
            "level and so carry real weight in shaping graduate attributes, whereas "
            "'Low' ones are more narrowly focused - which is not a fault in itself, "
            "since some COs are deliberately specialised. What matters is the balance: "
            "a course that leans on one or two strong COs is fragile, because weakness "
            f"in those outcomes spreads quickly into the POs they feed. {strong_co_note}"
            "Use this to spot the load-bearing COs and protect them when the course is "
            "revised."
        ),
        "co_att": (
            "The attainment percentage reflects how the cohort actually performed on "
            "the work tied to each CO, and the level is set by comparing that figure "
            f"against the thresholds shown above (floor {t3}%). A higher level means "
            f"more students cleared the expected standard. {co_extremes_note}"
            f"{weak_co_note}"
            "When acting on a low CO, separate the two usual causes: the outcome may "
            "have been taught well but assessed too harshly, or the topic genuinely "
            "needs more time and practice. The remedy differs, so review the "
            "assessment design before concluding the teaching was at fault."
        ),
        "co_compare": (
            "Two independent readings of the same course sit side by side here. The "
            "marks-based column is direct evidence: it counts how many students actually "
            "cleared the target in each CO, separately for the internal tests (CIE) and "
            "the semester exam (SEE), then combines them as 20% CIE and 80% SEE on the "
            "0-3 attainment scale. The AI column is an indicative estimate made from the "
            "CO statements alone, before any marks are seen. Where the two agree, the "
            "estimate is corroborated by real performance; where they diverge, trust the "
            "marks-based figure for accreditation and read the gap as a sign that the CO "
            "was harder or easier in practice than its wording suggested. A CO that the "
            "AI rated highly but students attained poorly is the most important one to "
            "examine - the topic likely needs more time or a gentler assessment."
        ),
        "po_summary": (
            "PO attainment is carried over from CO performance: each PO draws its "
            "strength from the COs mapped to it and how well those COs were attained. "
            "POs supported by many well-attained COs therefore score higher, while "
            "those resting on a few weaker COs score lower - which is why these "
            f"values broadly track the CO results above. {po_extreme_note}"
            "A weak PO has two quite different causes: too few COs feed it, or the COs "
            "that do feed it underperformed. The first is a curriculum-design issue to "
            "fix when COs are next written; the second is a delivery issue handled "
            "through the action plan."
        ),
        "po_nba": (
            "This view weighs each PO's attainment against its target. 'Met' means "
            "the cohort cleared the expected bar for that competency; 'Below' flags a "
            f"graduate attribute that is not yet being fully achieved. {below_po_note}"
            "The summary at the top gives a one-glance read on overall programme "
            "health, but the POs marked 'Below' are what an accreditor examines first, "
            "so each should carry a concrete, dated action in the final section."
        ),
        "level_dist": (
            "This shows how the COs are spread across attainment levels. Here, "
            f"{dist_read} ({high} CO(s) at levels 2-3 against {low} at levels 0-1). A "
            "distribution weighted toward the higher levels signals a strong course, "
            "while a cluster at the lower levels points to outcomes that need "
            "attention. The shape matters as much as the average: a wide spread says "
            "the class is splitting into those who keep up and those who fall behind, "
            "which calls for differentiated support, whereas a tight cluster low down "
            "points to a problem with the outcome or its assessment rather than with "
            "individual students."
        ),
        "co_sdg": (
            "This shows how each course outcome contributes to the leading Sustainable "
            "Development Goal for the course. COs with a larger share are the ones "
            f"whose content and skills align most closely with that goal's themes. "
            f"{top_sdgco_note}"
            "A more even spread means the goal is supported across the whole course "
            "rather than by a single outcome, which is the healthier pattern - it "
            "shows sustainability is woven through the syllabus instead of parked in "
            "one topic. A highly concentrated share is a prompt to ask whether the "
            "theme could be reinforced in more outcomes."
        ),
        "po_sdg": (
            "This estimates how the programme outcomes collectively support each SDG. "
            "A higher contribution percentage means the attributes developed in this "
            "course align well with that goal, and the interpretation label puts that "
            f"percentage into plain language. {lead_sdg_note}"
            "The spread reflects which goals the course speaks to most naturally; a "
            "few strong goals and many near-zero ones is normal for a specialised "
            "engineering subject, so read the leaders as the course's real "
            "sustainability footprint rather than expecting broad coverage of every "
            "goal."
        ),
        "composite": (
            "The composite index rolls the individual SDG contributions into a single "
            f"figure for the course's overall sustainability alignment, indicating "
            f"{comp_band}. A higher value means the course, taken as a whole, supports "
            f"the development goals more strongly; here {n_sdg} SDG(s) feed into it. "
            "Because the index blends all the goals together, a course that is "
            "excellent on a few SDGs but silent on the rest will show a moderate "
            "figure - that is expected, and is best read as a measure of breadth and "
            "strength together rather than a pass mark. Use the band as a quick "
            "benchmark and the per-SDG rows above to see where the strength sits."
        ),
        "checklist": (
            "This checklist tracks whether the evidence required for OBE review "
            "is in place. 'Met' items are complete and documented; 'Needs Attention' "
            "items are gaps to close before an audit. It is a readiness self-check, "
            "not a formal score - the aim is to surface missing pieces early, while "
            "there is still time to gather the evidence. Work down the 'Needs "
            "Attention' items first: most are about producing or attaching a document "
            "rather than improving results, so they are usually the quickest gaps to "
            "close before submission."
        ),
        "atr": (
            "This section records the corrective actions planned for any programme "
            f"outcome that fell short of its target. {atr_count_note}"
            "Closing the loop in this way - naming the weak POs and stating what will "
            "be done about them - is a core expectation of outcome-based education and "
            "is how continuous improvement is demonstrated to accreditors. The "
            "strongest action statements are specific and verifiable: they say what "
            "will change, in which part of the course, and how the effect will be "
            "checked next cycle, rather than simply noting that performance will be "
            "monitored."
        ),
    }


# ── Analytics ─────────────────────────────────────────────────────────────────

def _norm_co(v):
    """Normalize a CO label so 'CO 1' / 'co1' / 'CO1' all compare equal."""
    return str(v or "").strip().upper().replace(" ", "")


def _build_marks_comparison(pomap, coatt_marks, co_results, po_attainment):
    """Build CO and PO comparison rows: marks-based (direct) vs AI-estimated.

    Returns (co_compare, po_compare, po_marks) or (None, None, None) when no
    marks data is present.
    """
    marks_co = (coatt_marks or {}).get("coResults") or []
    if not marks_co:
        return None, None, None

    # AI lookups keyed by normalized CO label.
    ai_co = {_norm_co(r.get("co")): r for r in (co_results or [])}

    co_compare = []
    for m in marks_co:
        key = _norm_co(m.get("co"))
        ai  = ai_co.get(key, {})
        ai_level = ai.get("level")
        co_compare.append({
            "co":         m.get("co"),
            "ia_pct":     m.get("ia_pct"),
            "see_pct":    m.get("see_pct"),
            "marks_level": m.get("level"),
            "marks_pct":  m.get("pct"),
            "ai_pct":     ai.get("pct"),
            "ai_level":   ai_level,
            "delta":      (round(m.get("level") - ai_level, 2)
                           if (m.get("level") is not None and ai_level is not None)
                           else None),
        })

    # Marks-based PO attainment: same weighted formula as the AI PO tool
    # (sum(CO_attainment x CO-PO weight) / sum(weight)), using marks CO levels.
    po_marks = []
    po_compare = None
    if pomap:
        level_map = {_norm_co(m.get("co")): float(m.get("level") or 0) for m in marks_co}
        po_keys = list(pomap[0].get("scores", {}).keys())
        for pk in po_keys:
            total_w = 0.0
            wsum    = 0.0
            for row in pomap:
                w = float(row.get("scores", {}).get(pk, 0) or 0)
                if w > 0:
                    wsum    += level_map.get(_norm_co(row.get("co")), 0.0) * w
                    total_w += w
            raw = round(wsum / total_w, 3) if total_w else 0.0
            po_marks.append({"po": pk, "level": raw, "pct": round(raw / 3 * 100, 1)})

        ai_po = {p.get("po"): p for p in (po_attainment or [])}
        po_compare = []
        for pm in po_marks:
            ai = ai_po.get(pm["po"], {})
            ai_pct = ai.get("pct")
            po_compare.append({
                "po":        pm["po"],
                "marks_pct": pm["pct"],
                "ai_pct":    ai_pct,
                "delta":     (round(pm["pct"] - ai_pct, 1)
                              if ai_pct is not None else None),
            })

    return co_compare, po_compare, po_marks


def compute_analytics(pomap_rows, coatt_data, poatt_data, sdgco_data, sdgpo_data,
                      coatt_marks=None):
    pomap  = pomap_rows  or []
    coatt  = coatt_data  or {}
    poatt  = poatt_data  or {}
    sdgco  = sdgco_data  or {}
    sdgpo  = sdgpo_data  or {}
    coattm = coatt_marks or {}

    co_results  = coatt.get("coResults", [])
    po_results  = coatt.get("poResults", [])
    thresholds  = coatt.get("thresholds", {"t1": 70, "t2": 60, "t3": 50})

    # CO names from pomap or coatt
    co_names = [r.get("co", "") for r in pomap] or [r.get("co", "") for r in co_results]
    n_cos = len(co_names) or len(co_results)

    # PO keys from pomap
    po_keys = list(pomap[0].get("scores", {}).keys()) if pomap else [f"PO{i}" for i in range(1, 13)]
    n_pos   = len(po_keys)

    # Average PO score per CO (strength)
    co_strength = {}
    for row in pomap:
        scores = row.get("scores", {})
        vals   = [v for v in scores.values() if v is not None]
        co_strength[row.get("co", "")] = round(sum(vals) / len(vals), 2) if vals else 0

    # Average score per PO (column average)
    po_avg = {}
    for pk in po_keys:
        vals = [row.get("scores", {}).get(pk, 0) for row in pomap]
        po_avg[pk] = round(sum(vals) / len(vals), 2) if vals else 0

    # Attainment level distribution
    level_dist = {3: 0, 2: 0, 1: 0, 0: 0}
    for r in co_results:
        lv = int(r.get("level", 0))
        level_dist[lv] = level_dist.get(lv, 0) + 1

    # Mean CO attainment %
    pcts = [r.get("pct", 0) for r in co_results]
    mean_att_pct = round(sum(pcts) / len(pcts), 2) if pcts else 0

    # SDG CO
    target_sdg     = sdgco.get("targetSdg", "")
    sdgco_contrib  = sdgco.get("contributions", [])
    sdgco_total    = sdgco.get("total", 0)

    # New all-SDG matrix format from /generate_sdg_co_all
    # ({matrix, cos, sdgs, co_totals, sdg_totals, top_sdgs}). Convert it into the
    # legacy {target_sdg, contributions, total} shape the renderers expect, using
    # the top-contributing SDG as the target.
    if not sdgco_contrib and sdgco.get("matrix"):
        matrix     = sdgco.get("matrix", {})
        top_sdgs   = sdgco.get("top_sdgs", [])
        sdg_totals = sdgco.get("sdg_totals", {})
        # CO statements aren't carried in the SDG payload; pull them from the
        # CO-PO mapping rows, which key on the same CO names ("CO1", ...).
        stmt_map = {r.get("co", ""): r.get("statement", "") for r in pomap}
        if top_sdgs:
            target_sdg  = top_sdgs[0]
            sdgco_total = sdg_totals.get(target_sdg, 0)
            for co_name in sdgco.get("cos", list(matrix.keys())):
                score = matrix.get(co_name, {}).get(target_sdg, 0)
                pct   = (score / sdgco_total * 100) if sdgco_total else 0
                sdgco_contrib.append({
                    "co": co_name,
                    "statement": stmt_map.get(co_name, ""),
                    "score": score,
                    "pct": pct,
                })

    # SDG PO
    sdgpo_results  = sdgpo.get("results", [])
    sdgpo_composite = sdgpo.get("composite", 0)
    sdgs_covered   = [r.get("sdg", "") for r in sdgpo_results]

    # PO Attainment (dedicated tool)
    po_attainment   = poatt.get("po_attainment", [])
    po_att_summary  = poatt.get("summary", {})
    atr_rows        = [p for p in po_attainment if not p.get("target_met")]

    # Marks-based attainment (Method-2 Tier-I) vs AI comparison.
    marks_summary = coattm.get("summary", {}) if coattm else {}
    co_compare, po_compare, po_marks = _build_marks_comparison(
        pomap, coattm, co_results, po_attainment)

    return dict(
        co_names=co_names,
        n_cos=n_cos,
        po_keys=po_keys,
        n_pos=n_pos,
        pomap_rows=pomap,
        co_strength=co_strength,
        po_avg=po_avg,
        co_results=co_results,
        po_results=po_results,
        thresholds=thresholds,
        level_dist=level_dist,
        mean_att_pct=mean_att_pct,
        target_sdg=target_sdg,
        sdgco_contrib=sdgco_contrib,
        sdgco_total=sdgco_total,
        sdgpo_results=sdgpo_results,
        sdgpo_composite=sdgpo_composite,
        sdgs_covered=sdgs_covered,
        po_attainment=po_attainment,
        po_att_summary=po_att_summary,
        atr_rows=atr_rows,
        marks_summary=marks_summary,
        co_compare=co_compare,
        po_compare=po_compare,
        po_marks=po_marks,
    )


# ── AI analysis ───────────────────────────────────────────────────────────────

def get_ai_analysis(client, all_data, analytics, code, title):
    co_lines   = "\n".join(f"  {r['co']}: {r['pct']:.1f}% (Level {r['level']})" for r in analytics["co_results"][:8])
    po_lines   = "\n".join(f"  {r['po']}: {r['attainment']}" for r in analytics["po_results"][:12])
    sdg_lines  = "\n".join(f"  {r['sdg']}: {r['contribution']:.1f}% ({r['interpretation']})" for r in analytics["sdgpo_results"][:5])
    target_sdg = analytics["target_sdg"]

    prompt = (
        f"You are an OBE accreditation expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({code})\n\n"
        f"CO Attainment (AI-estimated):\n{co_lines or '  Not available'}\n\n"
        f"PO Attainment:\n{po_lines or '  Not available'}\n\n"
        f"Target SDG (CO-level): {target_sdg or 'Not set'}\n\n"
        f"PO-SDG Contributions:\n{sdg_lines or '  Not available'}\n\n"
        "Provide a structured assessment in this exact JSON format:\n"
        "{\n"
        '  "attainment_quality": "2-sentence assessment of CO/PO attainment levels",\n'
        '  "mapping_strength": "1-sentence on CO-PO mapping robustness",\n'
        '  "sdg_integration": "1-sentence on SDG alignment quality",\n'
        '  "obe_compliance": "1-sentence OBE compliance observation",\n'
        '  "recommendations": ["rec 1", "rec 2", "rec 3"],\n'
        '  "readiness_score": <integer 0-100 overall assessment readiness>\n'
        "}"
    )

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=600,
            messages=[{"role": "user", "content": prompt}],
        )
        import json as _json
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        return _json.loads(raw.strip())
    except Exception:
        return {}


# ── Shared text builder ───────────────────────────────────────────────────────

def _build_txt(pomap_data, coatt_data, poatt_data, sdgco_data, sdgpo_data,
               analytics, ai, code, title, semester):
    a    = analytics
    now  = datetime.datetime.now().strftime("%d %b %Y")
    sem  = f"Semester {semester}  |  " if semester else ""
    expl = build_explanations(a, ai)
    lines = []

    def h1(t): lines.extend([t, "=" * len(t), ""])
    def h2(t): lines.extend([t, "-" * len(t), ""])
    def row(*cols): lines.append("  ".join(str(c) for c in cols))

    def note(key):
        text = expl.get(key)
        if not text:
            return
        wrapped = textwrap.wrap(text, width=76)
        lines.append("")
        lines.append("  Why this result:")
        for w in wrapped:
            lines.append(f"  {w}")
        lines.append("")

    h1(f"MODULE 4 COMPREHENSIVE ASSESSMENT REPORT")
    lines.append(f"Course: {title} ({code})  |  {sem}Generated: {now}")
    lines.append("")

    h2("1. COURSE OVERVIEW")
    lines.append(f"  Course Outcomes   : {a['n_cos']}")
    lines.append(f"  Programme Outcomes: {a['n_pos']}")
    lines.append(f"  Mean CO Attainment: {a['mean_att_pct']}%")
    lines.append(f"  SDG Focus (CO)    : {a['target_sdg'] or 'N/A'}")
    lines.append(f"  SDG Coverage (PO) : {len(a['sdgs_covered'])} SDGs")
    note("overview")

    h2("2. CO-PO MAPPING TABLE  (3=Strong  2=Moderate  1=Low  0=None)")
    if a["pomap_rows"]:
        pk = a["po_keys"]
        row("CO".ljust(8), *[k.ljust(4) for k in pk])
        lines.append("  " + "-" * (8 + 6 * len(pk)))
        for r in a["pomap_rows"]:
            sc = [str(r.get("scores", {}).get(k, 0)).ljust(4) for k in pk]
            row(r.get("co", "").ljust(8), *sc)
    else:
        lines.append("  No mapping data.")
    note("copo")

    h2("3. CO STRENGTH SUMMARY")
    for co, avg in a["co_strength"].items():
        bar = "#" * int(avg / 3 * 20)
        lines.append(f"  {co:<8} avg={avg:.2f}  {bar}")
    note("co_strength")

    h2("4. CO ATTAINMENT ANALYSIS")
    t = a["thresholds"]
    lines.append(f"  Thresholds: Level 3 >= {t['t1']}%  Level 2 >= {t['t2']}%  Level 1 >= {t['t3']}%")
    lines.append("")
    row("  CO".ljust(10), "Attainment%".ljust(14), "Level")
    lines.append("  " + "-" * 30)
    for r in a["co_results"]:
        row(f"  {r['co']}".ljust(10), f"{r['pct']:.2f}%".ljust(14), r["level"])
    lines.append(f"\n  Mean: {a['mean_att_pct']}%")
    note("co_att")

    if a.get("co_compare"):
        ms = a.get("marks_summary", {}) or {}
        h2("4b. CO ATTAINMENT: DIRECT (MARKS-BASED) vs AI-ESTIMATED")
        lines.append(f"  Method-2 Tier-I  |  Target {ms.get('target_pct', 60)}%  |  "
                     f"CIE {ms.get('cie_weight', 20)}% + SEE {ms.get('see_weight', 80)}%  |  "
                     f"Students: CIE {ms.get('n_students_cie', 0)} / SEE {ms.get('n_students_see', 0)}")
        lines.append("")
        row("  CO".ljust(8), "CIE%".ljust(8), "SEE%".ljust(8),
            "Marks Lv".ljust(10), "AI %".ljust(8), "AI Lv".ljust(7), "Delta")
        lines.append("  " + "-" * 56)
        for c in a["co_compare"]:
            ia = "-" if c["ia_pct"] is None else f"{c['ia_pct']:.0f}"
            se = "-" if c["see_pct"] is None else f"{c['see_pct']:.0f}"
            aip = "-" if c["ai_pct"] is None else f"{c['ai_pct']:.0f}"
            ail = "-" if c["ai_level"] is None else str(c["ai_level"])
            dl  = "-" if c["delta"] is None else f"{c['delta']:+.2f}"
            row(f"  {c['co']}".ljust(8), ia.ljust(8), se.ljust(8),
                f"{c['marks_level']:.2f}".ljust(10), aip.ljust(8), ail.ljust(7), dl)
        if a.get("po_compare"):
            lines.append("")
            lines.append("  PO Attainment %: Marks-Based vs AI")
            row("  PO".ljust(8), "Marks%".ljust(10), "AI%".ljust(10), "Delta")
            lines.append("  " + "-" * 38)
            for p in a["po_compare"]:
                aip = "-" if p["ai_pct"] is None else f"{p['ai_pct']:.1f}"
                dl  = "-" if p["delta"] is None else f"{p['delta']:+.1f}"
                row(f"  {p['po']}".ljust(8), f"{p['marks_pct']:.1f}".ljust(10), aip.ljust(10), dl)
        note("co_compare")

    h2("5. PO ATTAINMENT SUMMARY")
    if a["po_results"]:
        row("  " + "  ".join(f"{r['po']}".ljust(5) for r in a["po_results"]))
        row("  " + "  ".join(f"{r['attainment']:.2f}".ljust(5) for r in a["po_results"]))
    note("po_summary")

    h2("5b. PO ATTAINMENT ANALYSIS")
    if a["po_attainment"]:
        pa_sum = a["po_att_summary"]
        lines.append(f"  Mean PO Attainment: {pa_sum.get('mean_pct', 0):.1f}%  |  "
                     f"POs Meeting Target: {pa_sum.get('pos_meeting_target', 0)}/{len(a['po_attainment'])}")
        lines.append("")
        row("  PO".ljust(8), "Name".ljust(34), "Att%".ljust(8), "Level", "Target")
        lines.append("  " + "-" * 62)
        for p in a["po_attainment"]:
            row(f"  {p['po']}".ljust(8), p["name"][:34].ljust(34),
                f"{p['pct']:.1f}%".ljust(8), str(p["level"]), "Met" if p["target_met"] else "Below")
    else:
        lines.append("  Run PO Attainment tool to see this section.")
    note("po_nba")

    h2("6. ATTAINMENT LEVEL DISTRIBUTION")
    ld = a["level_dist"]
    for lv in [3, 2, 1, 0]:
        bar = "#" * ld.get(lv, 0)
        lines.append(f"  Level {lv}: {ld.get(lv, 0):>2} CO(s)  {bar}")
    note("level_dist")

    h2("7. CO TO SDG CONTRIBUTION ANALYSIS")
    if a["sdgco_contrib"]:
        lines.append(f"  Target SDG: {a['target_sdg']}")
        row("  CO".ljust(10), "Score".ljust(8), "Share%")
        lines.append("  " + "-" * 28)
        for c in a["sdgco_contrib"]:
            row(f"  {c['co']}".ljust(10), str(c["score"]).ljust(8), f"{c['pct']:.1f}%")
        lines.append(f"\n  Total score: {a['sdgco_total']}")
    else:
        lines.append("  No CO-SDG data.")
    note("co_sdg")

    h2("8. PO TO SDG CONTRIBUTION ANALYSIS")
    if a["sdgpo_results"]:
        for r in a["sdgpo_results"]:
            lines.append(f"  {r['sdg']:<48}  {r['contribution']:.2f}%  ({r['interpretation']})")
    else:
        lines.append("  No PO-SDG data.")
    note("po_sdg")

    h2("9. COMPOSITE SDG INDEX")
    lines.append(f"  Composite SDG Index: {a['sdgpo_composite']:.2f}%")
    interp = ("Excellent" if a["sdgpo_composite"] >= 85 else
              "Strong"    if a["sdgpo_composite"] >= 70 else
              "Moderate"  if a["sdgpo_composite"] >= 50 else "Weak")
    lines.append(f"  Interpretation: {interp}")
    note("composite")

    h2("10. OBE ASSESSMENT COMPLIANCE CHECKLIST")
    has_pomap   = bool(a["pomap_rows"])
    has_coatt   = bool(a["co_results"])
    has_sdgco   = bool(a["sdgco_contrib"])
    has_sdgpo   = bool(a["sdgpo_results"])
    good_att    = a["mean_att_pct"] >= 60
    multi_sdg   = len(a["sdgs_covered"]) >= 2
    has_atr     = (not a["atr_rows"]) or all(p.get("atr", "").strip() for p in a["atr_rows"])
    checks = [
        ("CO-PO Mapping generated",                    has_pomap),
        ("CO Attainment calculated",                   has_coatt),
        ("CO-level SDG contribution mapped",           has_sdgco),
        ("PO-level SDG contribution mapped",           has_sdgpo),
        ("Mean CO attainment >= 60%",                  good_att),
        ("Multiple SDGs covered",                      multi_sdg),
        ("Action Taken Report documented",             has_atr),
    ]
    for label, ok in checks:
        lines.append(f"  {'[OK]' if ok else '[  ]'}  {label}")
    note("checklist")

    h2("11. ACTION TAKEN REPORT")
    lines.append("  Corrective actions for Programme Outcomes below attainment target.")
    lines.append("")
    if not a["po_attainment"]:
        lines.append("  Run PO Attainment tool to generate this section.")
    elif not a["atr_rows"]:
        lines.append("  All Programme Outcomes met attainment target.")
        lines.append("  No corrective actions required.")
    else:
        lines.append(f"  {len(a['atr_rows'])} PO(s) below target:")
        lines.append("")
        for p in a["atr_rows"]:
            lines.append(f"  {p['po']} - {p['name'][:40]}  |  {p['pct']:.1f}%  (Level {p['level']})")
            lines.append(f"    Action Taken: {p.get('atr', '') or 'Not specified'}")
            lines.append("")
    note("atr")

    return "\n".join(lines)


# ── TXT output ────────────────────────────────────────────────────────────────

def build_txt(pomap_data, coatt_data, poatt_data, sdgco_data, sdgpo_data,
              analytics, ai, code, title, semester, output_path):
    txt = _build_txt(pomap_data, coatt_data, poatt_data, sdgco_data, sdgpo_data,
                     analytics, ai, code, title, semester)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(txt)


# ── DOCX output ───────────────────────────────────────────────────────────────

def build_docx(pomap_data, coatt_data, poatt_data, sdgco_data, sdgpo_data,
               analytics, ai, code, title, semester, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    a    = analytics
    now  = datetime.datetime.now().strftime("%d %b %Y")
    sem  = f"Semester {semester}  |  " if semester else ""
    expl = build_explanations(a, ai)

    doc = Document()

    # narrow margins
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    GREEN  = RGBColor(5, 150, 105)   # #059669
    DGREEN = RGBColor(1, 91, 63)
    BLACK  = RGBColor(17, 24, 39)
    GREY   = RGBColor(107, 114, 128)

    def _s(v): return str(v) if v is not None else ""

    def add_title(text):
        p = doc.add_heading(text, level=0)
        p.runs[0].font.color.rgb = DGREEN
        p.runs[0].font.size = Pt(16)
        p.paragraph_format.space_after = Pt(4)

    def add_h1(text):
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.color.rgb = GREEN
            run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)

    def add_para(text, size=10, bold=False, color=None, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = color or BLACK
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_kv(key, val):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(key + ": ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(_s(val))
        r2.font.size = Pt(10)

    def note(key):
        text = expl.get(key)
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)
        r1 = p.add_run("Why this result:  ")
        r1.font.bold  = True
        r1.font.italic = True
        r1.font.size  = Pt(9)
        r1.font.color.rgb = DGREEN
        r2 = p.add_run(text)
        r2.font.italic = True
        r2.font.size   = Pt(9)
        r2.font.color.rgb = GREY

    def _insert_chart(buf, width=5.2):
        if buf is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(_fname, width=Inches(width))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        except Exception:
            pass
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    # Title
    add_title("Module 4 Comprehensive Assessment Report")
    sub = doc.add_paragraph(f"{title} ({code})  |  {sem}Generated: {now}")
    sub.runs[0].font.size  = Pt(9)
    sub.runs[0].font.color.rgb = GREY
    sub.paragraph_format.space_after = Pt(8)

    # 1. Course Overview
    add_h1("1. Course Overview")
    add_kv("Course Code",          code)
    add_kv("Course Title",         title)
    add_kv("Course Outcomes",      a["n_cos"])
    add_kv("Programme Outcomes",   a["n_pos"])
    add_kv("Mean CO Attainment",   f"{a['mean_att_pct']}%")
    add_kv("SDG Focus (CO-level)", a["target_sdg"] or "N/A")
    add_kv("SDGs Covered (PO)",    len(a["sdgs_covered"]))
    note("overview")
    doc.add_paragraph()

    # 2. CO-PO Mapping Table
    add_h1("2. CO-PO-PSO Mapping Table")
    add_para("3 = Strong  |  2 = Moderate  |  1 = Low  |  0 = None",
             size=9, color=GREY, indent=True)
    if a["pomap_rows"]:
        pk   = a["po_keys"]
        tbl  = doc.add_table(rows=1 + len(a["pomap_rows"]), cols=1 + len(pk))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "CO"
        for i, k in enumerate(pk):
            hdr[i + 1].text = k
        for ridx, row in enumerate(a["pomap_rows"]):
            cells = tbl.rows[ridx + 1].cells
            cells[0].text = _s(row.get("co", ""))
            for i, k in enumerate(pk):
                cells[i + 1].text = _s(row.get("scores", {}).get(k, 0))
        # header row bold
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        doc.add_paragraph()
    if _rc:
        _insert_chart(_rc.copo_heatmap(a["pomap_rows"], a["po_keys"]))
    note("copo")

    # 3. CO Strength Summary
    add_h1("3. CO Strength Summary (Average PO Score)")
    if a["co_strength"]:
        tbl = doc.add_table(rows=1 + len(a["co_strength"]), cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["CO", "Avg PO Score", "Interpretation"]):
            hdr[i].text = h
        for i, (co, avg) in enumerate(a["co_strength"].items()):
            interp = "Strong" if avg >= 2 else "Moderate" if avg >= 1 else "Low"
            cells = tbl.rows[i + 1].cells
            cells[0].text = co
            cells[1].text = f"{avg:.2f}"
            cells[2].text = interp
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        doc.add_paragraph()
    note("co_strength")

    # 4. CO Attainment Analysis
    add_h1("4. CO Attainment Analysis")
    t = a["thresholds"]
    add_para(f"Thresholds: Level 3 >= {t['t1']}%  |  Level 2 >= {t['t2']}%  |  Level 1 >= {t['t3']}%",
             size=9, color=GREY, indent=True)
    if a["co_results"]:
        tbl = doc.add_table(rows=1 + len(a["co_results"]), cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["CO", "Attainment %", "Level"]):
            hdr[i].text = h
        for i, r in enumerate(a["co_results"]):
            cells = tbl.rows[i + 1].cells
            cells[0].text = _s(r.get("co", ""))
            cells[1].text = f"{r['pct']:.2f}%"
            cells[2].text = _s(r.get("level", 0))
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    add_kv("Mean CO Attainment", f"{a['mean_att_pct']}%")
    doc.add_paragraph()
    if _rc:
        _insert_chart(_rc.co_attainment_chart(a["co_results"], a["thresholds"]))
    note("co_att")

    # 4b. CO Attainment: Direct (Marks-Based) vs AI-Estimated
    if a.get("co_compare"):
        ms = a.get("marks_summary", {}) or {}
        add_h1("4b. CO Attainment: Direct (Marks-Based) vs AI-Estimated")
        add_para(f"Method-2 Tier-I  |  Target {ms.get('target_pct', 60)}%  |  "
                 f"CIE {ms.get('cie_weight', 20)}% + SEE {ms.get('see_weight', 80)}%  |  "
                 f"Students: CIE {ms.get('n_students_cie', 0)} / SEE {ms.get('n_students_see', 0)}",
                 size=9, color=GREY, indent=True)
        cc = a["co_compare"]
        tbl = doc.add_table(rows=1 + len(cc), cols=7)
        tbl.style = "Table Grid"
        for i, h in enumerate(["CO", "CIE %", "SEE %", "Marks Level",
                               "AI %", "AI Level", "Delta"]):
            tbl.rows[0].cells[i].text = h
        for i, c in enumerate(cc):
            cells = tbl.rows[i + 1].cells
            cells[0].text = _s(c["co"])
            cells[1].text = "-" if c["ia_pct"]  is None else f"{c['ia_pct']:.0f}%"
            cells[2].text = "-" if c["see_pct"] is None else f"{c['see_pct']:.0f}%"
            cells[3].text = f"{c['marks_level']:.2f}"
            cells[4].text = "-" if c["ai_pct"]   is None else f"{c['ai_pct']:.0f}%"
            cells[5].text = "-" if c["ai_level"] is None else _s(c["ai_level"])
            cells[6].text = "-" if c["delta"]    is None else f"{c['delta']:+.2f}"
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        doc.add_paragraph()
        if _rc:
            _insert_chart(_rc.co_attainment_compare_chart(cc))
        if a.get("po_compare"):
            add_para("PO Attainment %: Marks-Based vs AI", size=10, bold=True, indent=True)
            pc  = a["po_compare"]
            ptbl = doc.add_table(rows=1 + len(pc), cols=4)
            ptbl.style = "Table Grid"
            for i, h in enumerate(["PO", "Marks %", "AI %", "Delta"]):
                ptbl.rows[0].cells[i].text = h
            for i, p in enumerate(pc):
                cells = ptbl.rows[i + 1].cells
                cells[0].text = _s(p["po"])
                cells[1].text = f"{p['marks_pct']:.1f}%"
                cells[2].text = "-" if p["ai_pct"] is None else f"{p['ai_pct']:.1f}%"
                cells[3].text = "-" if p["delta"]  is None else f"{p['delta']:+.1f}"
            for cell in ptbl.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            doc.add_paragraph()
        note("co_compare")

    # 5. PO Attainment Summary
    add_h1("5. PO Attainment Summary")
    if a["po_results"]:
        tbl = doc.add_table(rows=2, cols=len(a["po_results"]))
        tbl.style = "Table Grid"
        for i, r in enumerate(a["po_results"]):
            tbl.rows[0].cells[i].text = _s(r.get("po", ""))
            tbl.rows[1].cells[i].text = f"{r['attainment']:.2f}"
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()
    note("po_summary")

    # 5b. PO Attainment Analysis
    add_h1("5b. PO Attainment Analysis")
    if a["po_attainment"]:
        pa_sum = a["po_att_summary"]
        add_kv("Mean PO Attainment",   f"{pa_sum.get('mean_pct', 0):.1f}%")
        add_kv("POs Meeting Target",   f"{pa_sum.get('pos_meeting_target', 0)} / {len(a['po_attainment'])}")
        tbl = doc.add_table(rows=1 + len(a["po_attainment"]), cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["PO", "Name", "Attainment %", "Level", "Target"]):
            hdr[i].text = h
        for ridx, p in enumerate(a["po_attainment"]):
            cells = tbl.rows[ridx + 1].cells
            cells[0].text = _s(p.get("po", ""))
            cells[1].text = _s(p.get("name", ""))
            cells[2].text = f"{p['pct']:.1f}%"
            cells[3].text = _s(p.get("level", 0))
            cells[4].text = "Met" if p.get("target_met") else "Below"
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    else:
        add_para("Run PO Attainment tool to see this section.", indent=True)
    doc.add_paragraph()
    if _rc:
        _insert_chart(_rc.po_attainment_chart(a["po_attainment"]))
    note("po_nba")

    # 6. Attainment Level Distribution
    add_h1("6. Attainment Level Distribution")
    ld  = a["level_dist"]
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    for i, lv in enumerate([3, 2, 1, 0]):
        tbl.rows[0].cells[i].text = f"Level {lv}"
        tbl.rows[1].cells[i].text = f"{ld.get(lv, 0)} CO(s)"
    for cell in tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    doc.add_paragraph()
    if _rc:
        _insert_chart(_rc.attainment_level_chart(a["level_dist"]))
    note("level_dist")

    # 7. CO to SDG Contribution
    add_h1("7. CO to SDG Contribution Analysis")
    if a["sdgco_contrib"]:
        add_kv("Target SDG", a["target_sdg"])
        tbl = doc.add_table(rows=1 + len(a["sdgco_contrib"]), cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["CO", "Statement", "Score", "Share %"]):
            hdr[i].text = h
        for i, c in enumerate(a["sdgco_contrib"]):
            cells = tbl.rows[i + 1].cells
            cells[0].text = _s(c.get("co", ""))
            cells[1].text = _s(c.get("statement", ""))
            cells[2].text = _s(c.get("score", 0))
            cells[3].text = f"{c.get('pct', 0):.1f}%"
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        add_kv("Total Score", a["sdgco_total"])
    else:
        add_para("No CO-SDG data available.", indent=True)
    doc.add_paragraph()
    note("co_sdg")

    # 8. PO to SDG Contribution
    add_h1("8. PO to SDG Contribution Analysis")
    if a["sdgpo_results"]:
        tbl = doc.add_table(rows=1 + len(a["sdgpo_results"]), cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["SDG", "Contribution %", "Interpretation"]):
            hdr[i].text = h
        for i, r in enumerate(a["sdgpo_results"]):
            cells = tbl.rows[i + 1].cells
            cells[0].text = _s(r.get("sdg", ""))
            cells[1].text = f"{r.get('contribution', 0):.2f}%"
            cells[2].text = _s(r.get("interpretation", ""))
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    else:
        add_para("No PO-SDG data available.", indent=True)
    doc.add_paragraph()
    note("po_sdg")

    # 9. Composite SDG Index
    add_h1("9. Composite SDG Index")
    interp = ("Excellent" if a["sdgpo_composite"] >= 85 else
              "Strong"    if a["sdgpo_composite"] >= 70 else
              "Moderate"  if a["sdgpo_composite"] >= 50 else "Weak")
    add_kv("Composite SDG Index", f"{a['sdgpo_composite']:.2f}%")
    add_kv("Interpretation",      interp)
    doc.add_paragraph()
    note("composite")

    # 10. OBE Compliance Checklist
    add_h1("10. OBE Assessment Compliance Checklist")
    has_atr = (not a["atr_rows"]) or all(p.get("atr", "").strip() for p in a["atr_rows"])
    checks = [
        ("CO-PO Mapping generated",                    bool(a["pomap_rows"])),
        ("CO Attainment calculated",                   bool(a["co_results"])),
        ("CO-level SDG contribution mapped",           bool(a["sdgco_contrib"])),
        ("PO-level SDG contribution mapped",           bool(a["sdgpo_results"])),
        ("Mean CO attainment >= 60%",                  a["mean_att_pct"] >= 60),
        ("Multiple SDGs covered",                      len(a["sdgs_covered"]) >= 2),
        ("Action Taken Report documented",             has_atr),
    ]
    tbl = doc.add_table(rows=1 + len(checks), cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Criterion"
    tbl.rows[0].cells[1].text = "Status"
    for i, (label, ok) in enumerate(checks):
        tbl.rows[i + 1].cells[0].text = label
        tbl.rows[i + 1].cells[1].text = "Met" if ok else "Needs Attention"
    for cell in tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    doc.add_paragraph()
    note("checklist")

    # 11. Action Taken Report
    add_h1("11. Action Taken Report")
    add_para("Corrective actions for Programme Outcomes below attainment target.",
             size=9, color=GREY, indent=True)
    if not a["po_attainment"]:
        add_para("Run PO Attainment tool to generate this section.", indent=True)
    elif not a["atr_rows"]:
        add_para("All Programme Outcomes met attainment target. No corrective actions required.",
                 indent=True, color=GREEN)
    else:
        add_para(f"{len(a['atr_rows'])} PO(s) below target — corrective actions recorded below.",
                 indent=True)
        tbl = doc.add_table(rows=1 + len(a["atr_rows"]), cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["PO", "Name", "Attainment %", "Level", "Action Taken"]):
            hdr[i].text = h
        for ridx, p in enumerate(a["atr_rows"]):
            cells = tbl.rows[ridx + 1].cells
            cells[0].text = _s(p.get("po", ""))
            cells[1].text = _s(p.get("name", ""))
            cells[2].text = f"{p['pct']:.1f}%"
            cells[3].text = _s(p.get("level", 0))
            cells[4].text = _s(p.get("atr", "")) or "Not specified"
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()
    note("atr")

    doc.save(output_path)


# ── PDF output ────────────────────────────────────────────────────────────────

def build_pdf(pomap_data, coatt_data, poatt_data, sdgco_data, sdgpo_data,
              analytics, ai, code, title, semester, output_path):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    a    = analytics
    now  = datetime.datetime.now().strftime("%d %b %Y")
    sem  = f"Semester {semester}  |  " if semester else ""
    expl = build_explanations(a, ai)

    def _s(v): return str(v) if v is not None else ""

    def _safe(s):
        # FPDF's core fonts are latin-1 only, so transliterate the common
        # Unicode punctuation (em/en dashes, smart quotes, ellipsis, bullet)
        # that appears in SDG names and AI-generated text to ASCII first -
        # otherwise latin-1 'replace' turns each one into a stray '?'.
        s = str(s)
        for uni, ascii_ in (
            ("—", "-"), ("–", "-"),          # em / en dash
            ("‘", "'"), ("’", "'"),          # single curly quotes
            ("“", '"'), ("”", '"'),          # double curly quotes
            ("…", "..."), ("•", "-"),        # ellipsis, bullet
            (" ", " "),                            # non-breaking space
        ):
            s = s.replace(uni, ascii_)
        return s.encode("latin-1", "replace").decode("latin-1")

    def _pdf_chart(buf, w=155):
        if buf is None or _rc is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            x = pdf.l_margin + (180 - w) / 2
            pdf.image(_fname, x=x, w=w)
            pdf.ln(3)
        except Exception:
            pdf.set_x(pdf.l_margin)
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def h1(text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(5, 150, 105)
        pdf.cell(0, 8, _safe(_s(text)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    def h2(text):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(1, 91, 63)
        pdf.cell(0, 7, _safe(_s(text)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    def kv(key, val):
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.cell(55, 6, _safe(_s(key)) + ":", border=0)
        pdf.set_font("Helvetica", "", 9.5)
        pdf.multi_cell(0, 6, _safe(_s(val)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def note(key):
        text = expl.get(key)
        if not text:
            return
        pdf.ln(1)
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "BI", 8.5)
        pdf.set_text_color(1, 91, 63)
        pdf.cell(0, 5, "Why this result:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "I", 8.5)
        pdf.set_text_color(107, 114, 128)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 4.6, _safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    # Cover
    h1("Module 4 Comprehensive Assessment Report")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 5, _safe(f"{title} ({code})  |  {sem}Generated: {now}"),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # 1. Overview
    h2("1. Course Overview")
    kv("Course Outcomes", a["n_cos"])
    kv("Programme Outcomes", a["n_pos"])
    kv("Mean CO Attainment", f"{a['mean_att_pct']}%")
    kv("SDG Focus (CO)", a["target_sdg"] or "N/A")
    kv("SDGs Covered (PO)", len(a["sdgs_covered"]))
    pdf.ln(3)
    note("overview")

    # 2. CO-PO Mapping
    h2("2. CO-PO Mapping Table")
    if a["pomap_rows"]:
        pk      = a["po_keys"]
        avail_w = pdf.w - pdf.l_margin - pdf.r_margin
        co_col  = 28
        cw      = max(8, int((avail_w - co_col) / max(len(pk), 1)))
        fs      = 7 if cw < 11 else 8
        pdf.set_font("Helvetica", "B", fs)
        pdf.set_fill_color(209, 250, 229)
        pdf.cell(co_col, 6, "CO", border=1, fill=True, align="C")
        for k in pk:
            pdf.cell(cw, 6, k, border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", fs)
        for row in a["pomap_rows"]:
            pdf.cell(co_col, 5, _s(row.get("co", "")), border=1)
            for k in pk:
                pdf.cell(cw, 5, _s(row.get("scores", {}).get(k, 0)), border=1, align="C")
            pdf.ln()
    pdf.ln(3)
    _pdf_chart(_rc.copo_heatmap(a["pomap_rows"], a["po_keys"]) if _rc else None)
    note("copo")

    # 4. CO Attainment
    h2("4. CO Attainment Analysis")
    t = a["thresholds"]
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 5, f"Level 3 >= {t['t1']}%  |  Level 2 >= {t['t2']}%  |  Level 1 >= {t['t3']}%",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if a["co_results"]:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("CO", 40), ("Attainment %", 55), ("Level", 40)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9)
        for r in a["co_results"]:
            pdf.cell(40, 6, _s(r.get("co", "")), border=1)
            pdf.cell(55, 6, f"{r['pct']:.2f}%", border=1, align="C")
            pdf.cell(40, 6, _s(r.get("level", 0)), border=1, align="C")
            pdf.ln()
    kv("Mean CO Attainment", f"{a['mean_att_pct']}%")
    pdf.ln(3)
    _pdf_chart(_rc.co_attainment_chart(a["co_results"], a["thresholds"]) if _rc else None)
    note("co_att")

    # 4b. CO Attainment: Direct (Marks-Based) vs AI-Estimated
    if a.get("co_compare"):
        ms = a.get("marks_summary", {}) or {}
        h2("4b. CO Attainment: Marks-Based vs AI-Estimated")
        pdf.set_font("Helvetica", "I", 8)
        pdf.multi_cell(0, 4.6,
            _safe(f"Method-2 Tier-I  |  Target {ms.get('target_pct', 60)}%  |  "
                  f"CIE {ms.get('cie_weight', 20)}% + SEE {ms.get('see_weight', 80)}%  |  "
                  f"Students: CIE {ms.get('n_students_cie', 0)} / SEE {ms.get('n_students_see', 0)}"),
            new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        cc = a["co_compare"]
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("CO", 25), ("CIE %", 25), ("SEE %", 25), ("Marks Lv", 30),
                       ("AI %", 25), ("AI Lv", 22), ("Delta", 28)]:
            pdf.cell(w, 7, lbl, border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", 8.5)
        for c in cc:
            ia  = "-" if c["ia_pct"]   is None else f"{c['ia_pct']:.0f}%"
            se  = "-" if c["see_pct"]  is None else f"{c['see_pct']:.0f}%"
            aip = "-" if c["ai_pct"]   is None else f"{c['ai_pct']:.0f}%"
            ail = "-" if c["ai_level"] is None else _s(c["ai_level"])
            dl  = "-" if c["delta"]    is None else f"{c['delta']:+.2f}"
            pdf.cell(25, 6, _s(c["co"]), border=1)
            pdf.cell(25, 6, ia,  border=1, align="C")
            pdf.cell(25, 6, se,  border=1, align="C")
            pdf.cell(30, 6, f"{c['marks_level']:.2f}", border=1, align="C")
            pdf.cell(25, 6, aip, border=1, align="C")
            pdf.cell(22, 6, ail, border=1, align="C")
            pdf.cell(28, 6, dl,  border=1, align="C")
            pdf.ln()
        pdf.ln(3)
        _pdf_chart(_rc.co_attainment_compare_chart(cc) if _rc else None)
        if a.get("po_compare"):
            pdf.set_font("Helvetica", "B", 9.5)
            pdf.cell(0, 6, "PO Attainment %: Marks-Based vs AI",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(209, 250, 229)
            for lbl, w in [("PO", 40), ("Marks %", 45), ("AI %", 45), ("Delta", 40)]:
                pdf.cell(w, 7, lbl, border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 9)
            for p in a["po_compare"]:
                aip = "-" if p["ai_pct"] is None else f"{p['ai_pct']:.1f}%"
                dl  = "-" if p["delta"]  is None else f"{p['delta']:+.1f}"
                pdf.cell(40, 6, _s(p["po"]), border=1)
                pdf.cell(45, 6, f"{p['marks_pct']:.1f}%", border=1, align="C")
                pdf.cell(45, 6, aip, border=1, align="C")
                pdf.cell(40, 6, dl,  border=1, align="C")
                pdf.ln()
            pdf.ln(3)
        note("co_compare")

    # 5. PO Attainment
    h2("5. PO Attainment Summary")
    if a["po_results"]:
        avail_w = pdf.w - pdf.l_margin - pdf.r_margin
        col_w = max(10, int(avail_w / max(len(a["po_results"]), 1)))
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(209, 250, 229)
        for r in a["po_results"]:
            pdf.cell(col_w, 6, _s(r["po"]), border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for r in a["po_results"]:
            pdf.cell(col_w, 5, f"{r['attainment']:.2f}", border=1, align="C")
        pdf.ln()
    pdf.ln(3)
    note("po_summary")

    # 5b. PO Attainment
    h2("5b. PO Attainment Analysis")
    if a["po_attainment"]:
        pa_sum = a["po_att_summary"]
        kv("Mean PO Attainment",  f"{pa_sum.get('mean_pct', 0):.1f}%")
        kv("POs Meeting Target",  f"{pa_sum.get('pos_meeting_target', 0)} / {len(a['po_attainment'])}")
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("PO", 15), ("Name", 70), ("Att%", 20), ("Level", 15), ("Target", 25)]:
            pdf.cell(w, 7, lbl, border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", 8.5)
        for p in a["po_attainment"]:
            pdf.cell(15, 5.5, _s(p.get("po", "")), border=1)
            pdf.cell(70, 5.5, _s(p.get("name", ""))[:38], border=1)
            pdf.cell(20, 5.5, f"{p['pct']:.1f}%", border=1, align="C")
            pdf.cell(15, 5.5, _s(p.get("level", 0)), border=1, align="C")
            pdf.cell(25, 5.5, "Met" if p.get("target_met") else "Below", border=1, align="C")
            pdf.ln()
    else:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, "Run PO Attainment tool to populate this section.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)
    _pdf_chart(_rc.po_attainment_chart(a["po_attainment"]) if _rc else None)
    note("po_nba")

    # 7. CO-SDG
    h2("7. CO to SDG Contribution")
    if a["sdgco_contrib"]:
        kv("Target SDG", a["target_sdg"])
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("CO", 25), ("Score", 25), ("Share %", 30)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9)
        for c in a["sdgco_contrib"]:
            pdf.cell(25, 6, _s(c.get("co", "")), border=1)
            pdf.cell(25, 6, _s(c.get("score", 0)), border=1, align="C")
            pdf.cell(30, 6, f"{c.get('pct', 0):.1f}%", border=1, align="C")
            pdf.ln()
        kv("Total Score", a["sdgco_total"])
    pdf.ln(3)
    note("co_sdg")

    # 8. PO-SDG
    h2("8. PO to SDG Contribution")
    if a["sdgpo_results"]:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("SDG", 100), ("Contribution %", 40), ("Interpretation", 40)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9)
        for r in a["sdgpo_results"]:
            pdf.cell(100, 6, _safe(_s(r.get("sdg", "")))[:55], border=1)
            pdf.cell(40,  6, f"{r.get('contribution', 0):.2f}%", border=1, align="C")
            pdf.cell(40,  6, _safe(_s(r.get("interpretation", ""))), border=1, align="C")
            pdf.ln()
    pdf.ln(3)
    note("po_sdg")

    # 9. Composite SDG
    h2("9. Composite SDG Index")
    interp = ("Excellent" if a["sdgpo_composite"] >= 85 else
              "Strong"    if a["sdgpo_composite"] >= 70 else
              "Moderate"  if a["sdgpo_composite"] >= 50 else "Weak")
    kv("Composite SDG Index", f"{a['sdgpo_composite']:.2f}%  ({interp})")
    pdf.ln(3)
    note("composite")

    # 10. Checklist
    h2("10. OBE Assessment Compliance Checklist")
    has_atr = (not a["atr_rows"]) or all(p.get("atr", "").strip() for p in a["atr_rows"])
    checks = [
        ("CO-PO Mapping generated",                    bool(a["pomap_rows"])),
        ("CO Attainment calculated",                   bool(a["co_results"])),
        ("CO-level SDG contribution mapped",           bool(a["sdgco_contrib"])),
        ("PO-level SDG contribution mapped",           bool(a["sdgpo_results"])),
        ("Mean CO attainment >= 60%",                  a["mean_att_pct"] >= 60),
        ("Multiple SDGs covered",                      len(a["sdgs_covered"]) >= 2),
        ("Action Taken Report documented",             has_atr),
    ]
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(209, 250, 229)
    pdf.cell(130, 7, "Criterion", border=1, fill=True)
    pdf.cell(50,  7, "Status",    border=1, fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    for label, ok in checks:
        pdf.cell(130, 6, label, border=1)
        pdf.cell(50,  6, "Met" if ok else "Needs Attention", border=1, align="C")
        pdf.ln()
    pdf.ln(3)
    note("checklist")

    # 11. Action Taken Report
    h2("11. Action Taken Report")
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 5, "Corrective actions for Programme Outcomes below attainment target.",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)
    if not a["po_attainment"]:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, "Run PO Attainment tool to generate this section.",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    elif not a["atr_rows"]:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(5, 150, 105)
        pdf.cell(0, 5, "All Programme Outcomes met attainment target. No corrective actions required.",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
    else:
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, f"{len(a['atr_rows'])} PO(s) below target:",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(209, 250, 229)
        for lbl, w in [("PO", 15), ("Name", 55), ("Att%", 18), ("Level", 15), ("Action Taken", 77)]:
            pdf.cell(w, 7, lbl, border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", 8.5)
        for p in a["atr_rows"]:
            atr_text = (_s(p.get("atr", "")) or "Not specified")[:42]
            pdf.cell(15, 5.5, _s(p.get("po", "")),           border=1)
            pdf.cell(55, 5.5, _s(p.get("name", ""))[:30],    border=1)
            pdf.cell(18, 5.5, f"{p['pct']:.1f}%",            border=1, align="C")
            pdf.cell(15, 5.5, _s(p.get("level", 0)),         border=1, align="C")
            pdf.cell(77, 5.5, atr_text,                       border=1)
            pdf.ln()
    pdf.ln(3)
    note("atr")

    pdf.output(output_path)
