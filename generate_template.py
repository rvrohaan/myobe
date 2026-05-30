"""
generate_template.py – Template question paper parser and generator.

Parses an existing PDF/DOCX question paper to extract header metadata
and question slot structure, then renders a new paper in the same
institutional table format with QB questions filling each slot.
"""

import re
from collections import OrderedDict


# ---------------------------------------------------------------------------
# Common regexes
# ---------------------------------------------------------------------------

_OR_RE   = re.compile(r'^\s*OR\s*$', re.IGNORECASE)
# Handle any separator between UNIT and the numeral: dash, en/em dash, replacement-char, space-only
_UNIT_RE = re.compile(r'^unit[\s\W]*[IVXivx\d]+', re.IGNORECASE)
# Detect column-header rows (CO … PO … Marks) to find where the question table starts
_COL_HDR_LINE_RE = re.compile(r'\bCO\b.{0,30}\bPO\b.{0,30}\bMarks\b', re.IGNORECASE)


# ---------------------------------------------------------------------------
# Header field extraction
# ---------------------------------------------------------------------------

def _parse_header_fields(lines):
    """Extract structured metadata from the free-text header lines."""
    h = {
        'college': '', 'sub_title': '', 'exam_name': '',
        'programme': '', 'semester': '', 'branch': '', 'duration': '',
        'course_code': '', 'max_marks': '', 'course': '',
        'instructions': [],
    }
    in_inst = False

    for raw in lines:
        s = raw.strip()
        if not s:
            continue
        sl = s.lower()

        # College / institution (first line containing key words)
        if not h['college'] and any(k in sl for k in (
                'college of engineering', 'college of technology',
                'university', 'institute of technology', 'b.m.s', 'bms')):
            h['college'] = s
            continue

        if not h['sub_title'] and any(k in sl for k in ('autonomous', 'affiliated to', 'deemed')):
            h['sub_title'] = s
            continue

        if not h['exam_name'] and any(k in sl for k in (
                'semester end', 'end examination', 'end main',
                'internal assessment', 'assessment test', 'examinations')):
            h['exam_name'] = s
            continue

        # Programme — stop before "Semester" if it appears on the same line
        m = re.search(r'Programme\s*[:\-]\s*(.+?)(?=\s+Semester\s*[:\-]|\s*$)', s, re.IGNORECASE)
        if m and not h['programme']:
            h['programme'] = m.group(1).strip()

        # Semester
        m = re.search(r'Semester\s*[:\-]\s*(\S+)', s, re.IGNORECASE)
        if m and not h['semester']:
            h['semester'] = m.group(1).strip()

        # Branch — stop before "Duration"
        m = re.search(r'Branch\s*[:\-]\s*(.+?)(?=\s+Duration\s*[:\-]|\s*$)', s, re.IGNORECASE)
        if m and not h['branch']:
            h['branch'] = m.group(1).strip()

        # Duration
        m = re.search(r'Duration\s*[:\-]\s*(.+?)(?=\s+Max|\s*$)', s, re.IGNORECASE)
        if m and not h['duration']:
            h['duration'] = m.group(1).strip()

        # Course Code — stop before "Max Marks"
        m = re.search(r'Course\s+Code\s*[:\-]\s*(.+?)(?=\s+Max[\s.]*Marks?|\s*$)', s, re.IGNORECASE)
        if m and not h['course_code']:
            h['course_code'] = m.group(1).strip()

        # Max Marks
        m = re.search(r'Max[\s.]*Marks?\s*[:\-]\s*(\d+)', s, re.IGNORECASE)
        if m and not h['max_marks']:
            h['max_marks'] = m.group(1).strip()

        # Course name (not "Course Code")
        m = re.match(r'Course\s*[:\-]\s*(.+)', s, re.IGNORECASE)
        if m and 'code' not in sl and not h['course']:
            h['course'] = m.group(1).strip()

        # Instructions
        if 'instruction' in sl:
            in_inst = True
        if in_inst and s and 'important note' not in sl:
            h['instructions'].append(s)

    return h


# ---------------------------------------------------------------------------
# Text-based question slot parser  (used for both PDF and DOCX fallback)
# ---------------------------------------------------------------------------

# Question line pattern: optional Q#, optional sub-label, text, CO#, PO#, optional BL digit, marks
_Q_LINE_RE = re.compile(
    r'^(?:(\d{1,2})\s+)?'
    r'(?:([a-e][\)\.:])\s+)?'
    r'(.+?)\s+'
    r'(C[O0]\d+)\s+'
    r'(PO\d+)\s+'
    r'(?:[1-6]\s+)?'          # optional Bloom Level digit (some papers have it)
    r'(\d{1,3})\s*$',
    re.IGNORECASE,
)

# Column-header lines to skip
_COL_HDR_RE = re.compile(r'^.{0,15}C[O0]\s+PO\s+(BL\s+)?Marks\s*$', re.IGNORECASE)


def _parse_text_slots(text):
    """Parse question slots from plain extracted text (line by line)."""
    slots = []
    current_unit = ''
    p_qnum = p_sub = p_co = p_po = p_marks = None
    p_parts = []

    def flush():
        nonlocal p_qnum, p_sub, p_co, p_po, p_marks, p_parts
        if p_marks is not None and p_parts:
            slots.append({
                'is_or':         False,
                'unit':          current_unit or 'General',
                'qnum':          p_qnum or '',
                'sub':           p_sub  or '',
                'marks':         p_marks,
                'co':            p_co   or '',
                'po':            p_po   or '',
                'original_text': ' '.join(p_parts).strip()[:400],
            })
        p_qnum = p_sub = p_co = p_po = p_marks = None
        p_parts = []

    for raw in text.split('\n'):
        s = raw.strip()
        if not s:
            continue

        # Skip very short lines that are watermark/scan artifacts (but keep "OR")
        if len(s) <= 2 and not _OR_RE.match(s):
            continue

        # Strip leading watermark chars: 1-4 lowercase letter groups that appear
        # before actual content (digit, uppercase letter, or sub-label like 'a)')
        # Handles patterns like 'n 1 a) ...'  and 'a p 1 a) ...' and 'n ia b) ...'
        stripped = re.sub(r'^([a-z]\s*){1,4}(?=\d|[A-Z]|[a-e][\)\.:])', '', s)
        if stripped and stripped != s:
            s = stripped

        # OR divider
        if _OR_RE.match(s):
            flush()
            slots.append({'is_or': True})
            continue

        # Unit header — handles "UNIT - I", "UNIT – II", "UNIT I TO V", "UNIT � I" etc.
        um = re.match(
            r'^(UNIT[\s\W]*[IVXivx\d]+(?:[\s\W]+(?:TO|AND)[\s\W]+[IVXivx\d]+)?)',
            s, re.IGNORECASE)
        if um:
            flush()
            current_unit = um.group(1).strip()
            continue

        # Skip column-header rows
        if _COL_HDR_RE.match(s):
            continue

        # Question row
        m = _Q_LINE_RE.match(s)
        if m:
            flush()
            p_qnum  = m.group(1) or ''
            p_sub   = m.group(2) or ''
            p_marks = int(m.group(6))
            p_co    = m.group(4)
            p_po    = m.group(5)
            txt     = m.group(3).strip()
            p_parts = [txt] if txt else []
            continue

        # Continuation of question text
        if p_marks is not None:
            if not re.match(r'^(C[O0]|PO|Marks|UNIT|OR)\b', s, re.IGNORECASE):
                p_parts.append(s)

    flush()

    # Post-process: if questions appear before any named unit AND the next
    # named unit is UNIT-II (or higher), infer the first group is UNIT-I.
    named_units = [
        s.get('unit', '') for s in slots
        if not s.get('is_or') and s.get('unit') and s.get('unit') != 'General'
    ]
    first_is_generic = (
        slots
        and not slots[0].get('is_or')
        and slots[0].get('unit') in ('', 'General')
    )
    if first_is_generic and named_units:
        first_named = named_units[0]
        m2 = re.match(r'UNIT[\s\W]*([IVXivx2-9\d]+)', first_named, re.IGNORECASE)
        if m2:
            roman = m2.group(1).upper()
            if roman not in ('I', '1'):
                for s in slots:
                    if not s.get('is_or') and s.get('unit') in ('', 'General'):
                        s['unit'] = 'UNIT - I'

    return slots


# ---------------------------------------------------------------------------
# Q-block / Experiment paragraph-format parser (assignment/lab manuals)
# ---------------------------------------------------------------------------

_Q_BLOCK_LINE_RE   = re.compile(r'^\s*Q\s*[\.\-]?\s*(\d+)\s*\.?\s*$', re.IGNORECASE)
_EXP_BLOCK_LINE_RE = re.compile(r'^Experiment\s*[\W]+\s*(\d+)', re.IGNORECASE)
_ASSIGN_UNIT_RE    = re.compile(r'^Assignment\s*[\W\s]+(?:[IVXivx]+|\d+)', re.IGNORECASE)
_CO_LABEL_RE       = re.compile(r'^CO\s*[:\-]\s*(C[O0]\d+)', re.IGNORECASE)
_MARKS_LABEL_RE    = re.compile(r'^Marks\s*[:\-]\s*(\d+)', re.IGNORECASE)
_CO_BL_MARKS_ROW_RE = re.compile(r'^(C[O0]\d+)\t(L?\d+)\t(\d+)$', re.IGNORECASE)


def _parse_qblock_text(text, _pats: dict | None = None):
    """Parse assignment / lab-manual paragraph-format question blocks.

    _pats: optional compiled-pattern dict from _compile_qblock_pats().
    Without _pats, uses the default built-in patterns (all block types).

    Recognises:
      • Q-blocks:         a lone 'Q1.' / 'Q 1' paragraph, followed by multi-line
                          question text, then 'CO: CO1', 'BL: L4', 'Marks: 10' lines.
      • Experiment-blocks:'Experiment − 1' header, title on next line,
                          followed by content paragraphs, then a serialised
                          CO|BL|Marks table row ('CO1\\tL3\\t10').
      • Unit markers:     'Assignment − I' lines set the current unit.
    """
    p = _pats or {}
    # Block-start patterns: each dict has 'regex', 'type', 'num_group'
    if 'block_starts' in p:
        block_patterns = p['block_starts']
    else:
        block_patterns = [
            {'regex': _Q_BLOCK_LINE_RE,   'type': 'q',    'num_group': 1},
            {'regex': _EXP_BLOCK_LINE_RE, 'type': 'exp',  'num_group': 1},
            {'regex': _ASSIGN_UNIT_RE,    'type': 'unit',  'num_group': 0},
        ]
    co_re     = p.get('co_re',       _CO_LABEL_RE)
    marks_re  = p.get('marks_re',    _MARKS_LABEL_RE)
    bl_re     = p.get('bl_re',       re.compile(r'^BL\s*[:\-]', re.IGNORECASE))
    co_tbl_re = p.get('co_table_re', _CO_BL_MARKS_ROW_RE)

    lines = text.split('\n')
    slots = []
    current_unit = ''

    # Locate every block start
    blocks = []
    for i, line in enumerate(lines):
        s = line.strip()
        for bp in block_patterns:
            m = bp['regex'].match(s)
            if m:
                ng   = bp.get('num_group', 1)
                bval = m.group(ng) if ng > 0 else s
                blocks.append((i, bp['type'], bval))
                break

    if not blocks:
        return []

    for bi, (start, btype, bval) in enumerate(blocks):
        end = blocks[bi + 1][0] if bi + 1 < len(blocks) else len(lines)
        block_lines = [l.strip() for l in lines[start:end] if l.strip()]

        if btype == 'unit':
            current_unit = bval
            continue

        co_val    = ''
        marks_val = None
        text_parts = []

        for line in block_lines[1:]:
            mc  = co_re.match(line)
            mm  = marks_re.match(line)
            mct = co_tbl_re.match(line)
            if mct:
                if not co_val:
                    co_val = mct.group(1)
                if marks_val is None:
                    marks_val = int(mct.group(3))
            elif mc:
                co_val = mc.group(1)
            elif mm:
                marks_val = int(mm.group(1))
            elif bl_re.match(line):
                pass  # skip BL metadata lines
            elif btype == 'q':
                text_parts.append(line)

        if btype == 'exp':
            title  = block_lines[1] if len(block_lines) > 1 else f'Experiment {bval}'
            q_text = f'Experiment {bval}: {title}'
        else:
            q_text = ' '.join(text_parts).strip()

        if q_text and marks_val:
            slots.append({
                'is_or':         False,
                'unit':          current_unit or 'General',
                'qnum':          str(bval),
                'sub':           '',
                'marks':         marks_val,
                'co':            co_val,
                'po':            '',
                'original_text': q_text[:400],
            })

    return slots


def _build_docx_ordered_text(doc):
    """Return DOCX body as a single string with paragraphs and table rows
    interleaved in their original document order."""
    paras  = list(doc.paragraphs)
    tables = list(doc.tables)
    pi = ti = 0
    lines = []
    for child in doc.element.body:
        tag = child.tag
        if tag.endswith('}p'):
            if pi < len(paras):
                txt = paras[pi].text.strip()
                if txt:
                    lines.append(txt)
                pi += 1
        elif tag.endswith('}tbl'):
            if ti < len(tables):
                tbl = tables[ti]
                for row in tbl.rows:
                    seen: set = set()
                    cells = []
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in seen:
                            seen.add(t)
                            cells.append(t)
                    if cells:
                        lines.append('\t'.join(cells))
                ti += 1
    return '\n'.join(lines)


# ---------------------------------------------------------------------------
# Profile-pattern compilation helpers
# ---------------------------------------------------------------------------

_DEFAULT_CO_RE   = re.compile(r'^C[O0]\d+',         re.IGNORECASE)
_DEFAULT_PO_RE   = re.compile(r'^PO\d+',            re.IGNORECASE)
_DEFAULT_BL_RE   = re.compile(r'^[1-6]$')
_DEFAULT_QNUM_RE = re.compile(r'^\d{1,2}\.?$')
_DEFAULT_SUB_RE  = re.compile(r'^[a-e][\)\.:]{0,1}$', re.IGNORECASE)


def _compile_table_pats(profile: dict) -> dict:
    """Compile a dict of regex objects from a table-type format profile."""
    pats = profile.get('patterns', {})
    opts = profile.get('options', {})
    return {
        'unit_re':   re.compile(pats.get('unit_header',  r'^unit[\s\W]*[IVXivx\d]+'), re.IGNORECASE),
        'colhdr_re': re.compile(pats.get('col_header',   r'^.{0,15}C[O0]\s+PO\s+(BL\s+)?Marks\s*$'), re.IGNORECASE),
        'or_re':     re.compile(pats.get('or_separator', r'^\s*OR\s*$'), re.IGNORECASE),
        'co_re':     re.compile(pats.get('co_cell',      r'^C[O0]\d+'), re.IGNORECASE),
        'po_re':     re.compile(pats.get('po_cell',      r'^PO\d+'), re.IGNORECASE),
        'bl_re':     re.compile(pats.get('bl_cell',      r'^[1-6]$')),
        'qnum_re':   re.compile(pats.get('qnum_cell',    r'^\d{1,2}\.?$')),
        'sub_re':    re.compile(pats.get('sub_cell',     r'^[a-e][\)\.:]{0,1}$'), re.IGNORECASE),
        'marks_min': opts.get('marks_min', 1),
        'marks_max': opts.get('marks_max', 100),
        'min_text':  opts.get('min_text_length', 5),
    }


def _compile_qblock_pats(profile: dict) -> dict:
    """Compile a dict of regex objects from a qblock-type format profile."""
    pats = profile.get('patterns', {})
    raw_blocks = pats.get('block_starts', [])
    compiled_blocks = []
    for b in raw_blocks:
        compiled_blocks.append({
            'regex':     re.compile(b['regex'], re.IGNORECASE),
            'type':      b['type'],
            'num_group': b.get('num_group', 1),
        })
    return {
        'block_starts':  compiled_blocks,
        'co_re':         re.compile(pats.get('co_line',      r'^CO\s*[:\-]\s*(C[O0]\d+)'), re.IGNORECASE),
        'marks_re':      re.compile(pats.get('marks_line',   r'^Marks\s*[:\-]\s*(\d+)'), re.IGNORECASE),
        'bl_re':         re.compile(pats.get('bl_line',      r'^BL\s*[:\-]'), re.IGNORECASE),
        'co_table_re':   re.compile(pats.get('co_table_row', r'^(C[O0]\d+)\t(L?\d+)\t(\d+)$'), re.IGNORECASE),
    }


def _extract_pdf_text(path: str) -> str:
    """Extract all text from a PDF using pdfplumber."""
    import pdfplumber
    text = ''
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or '') + '\n'
    return text


# ---------------------------------------------------------------------------
# DOCX table-based slot extraction (reliable for DOCX files)
# ---------------------------------------------------------------------------

def _parse_table_rows(tables, _pats: dict | None = None):
    """Extract question slots from table data (list of list of list of str).

    _pats: optional compiled-pattern dict from _compile_table_pats().
    Each table is evaluated independently: if every content slot produced
    by a table has an empty CO value the whole table is discarded as a
    rubric / info table rather than a question table.
    """
    p = _pats or {}
    or_re     = p.get('or_re',     _OR_RE)
    unit_re   = p.get('unit_re',   _UNIT_RE)
    colhdr_re = p.get('colhdr_re', _COL_HDR_RE)
    co_re     = p.get('co_re',     _DEFAULT_CO_RE)
    po_re     = p.get('po_re',     _DEFAULT_PO_RE)
    bl_re     = p.get('bl_re',     _DEFAULT_BL_RE)
    qnum_re   = p.get('qnum_re',   _DEFAULT_QNUM_RE)
    sub_re    = p.get('sub_re',    _DEFAULT_SUB_RE)
    marks_min = p.get('marks_min', 1)
    marks_max = p.get('marks_max', 100)
    min_text  = p.get('min_text',  5)

    all_slots    = []
    current_unit = ''

    for table in tables:
        tbl_slots = []

        for row in table:
            if not row:
                continue
            # De-duplicate adjacent identical merged-cell values
            cells = []
            for c in [str(x or '').strip() for x in row]:
                if not cells or c != cells[-1]:
                    cells.append(c)
            while cells and not cells[-1]:
                cells.pop()
            if not cells:
                continue

            joined = ' '.join(cells)

            # OR divider
            if or_re.match(joined.strip()):
                tbl_slots.append({'is_or': True})
                continue

            # Unit header row (1-3 cells, first matches UNIT pattern)
            if unit_re.match(cells[0]) and len(cells) <= 4:
                current_unit = cells[0]
                continue
            if len(set(cells)) <= 2 and unit_re.match(cells[0]):
                current_unit = cells[0]
                continue

            # Skip column-header rows
            jl = joined.lower()
            if re.search(r'\bco\b.*\bpo\b', jl) and len(joined) < 30:
                continue
            if colhdr_re.match(joined.strip()):
                continue

            # Must end with a numeric marks value within configured range
            marks_val = marks_idx = None
            for i in range(len(cells) - 1, max(len(cells) - 4, -1), -1):
                m = re.match(r'^(\d{1,3})$', cells[i])
                if m and marks_min <= int(m.group(1)) <= marks_max:
                    marks_val = int(m.group(1))
                    marks_idx = i
                    break
            if marks_val is None:
                continue

            remaining = cells[:marks_idx]
            after     = cells[marks_idx + 1:]

            # CO and PO: try traditional position (before marks) first,
            # then alternate position (after marks) for Q# | text | Marks | CO | PO format.
            po_val = ''
            if remaining and po_re.match(remaining[-1]):
                po_val = remaining.pop()

            co_val = ''
            if remaining and co_re.match(remaining[-1]):
                co_val = remaining.pop()

            if not co_val and after:
                m_co = next((c for c in after if co_re.match(c)), None)
                m_po = next((c for c in after if po_re.match(c)), None)
                co_val = m_co or ''
                po_val = m_po or po_val

            # BL (Bloom level digit)
            if remaining and bl_re.match(remaining[-1]):
                remaining.pop()

            qnum = ''
            if remaining and qnum_re.match(remaining[0]):
                qnum = remaining.pop(0).rstrip('.')

            sub = ''
            if remaining and sub_re.match(remaining[0]):
                sub = remaining.pop(0)

            text = ' '.join(remaining).strip()
            text = re.sub(r'\*+', '', text).strip()

            if text and len(text) > min_text:
                tbl_slots.append({
                    'is_or':         False,
                    'unit':          current_unit or 'General',
                    'qnum':          qnum,
                    'sub':           sub,
                    'marks':         marks_val,
                    'co':            co_val,
                    'po':            po_val,
                    'original_text': text[:400],
                })

        # Discard tables that have content slots but none have a CO value —
        # these are rubric / info tables (e.g. Criteria | Marks), not question tables.
        content_slots = [s for s in tbl_slots if not s.get('is_or')]
        if content_slots and not any(s.get('co') for s in content_slots):
            continue

        all_slots.extend(tbl_slots)

    return all_slots


# ---------------------------------------------------------------------------
# PDF / DOCX reader
# ---------------------------------------------------------------------------

def parse_template_paper(path: str, ext: str) -> dict:
    """Parse a template question paper using all loaded format profiles.

    Tries every registered profile in priority order and returns the result
    with the highest confidence score.  If all profile parsers produce a
    low-confidence result (< 0.3), and ANTHROPIC_API_KEY is set, falls back
    to an LLM extraction pass using Claude Haiku.

    Returns
    -------
    dict with keys: header, slots, profile_id, confidence
    """
    import parse_profiles as pp
    from docx import Document

    if ext not in ('.pdf', '.docx', '.doc'):
        raise ValueError(f"Unsupported format: {ext}")

    # ── Extract raw content ──────────────────────────────────────────────────
    if ext == '.pdf':
        raw_text   = _extract_pdf_text(path)
        raw_tables = None          # PDF tables come via line-by-line text
    else:
        doc        = Document(path)
        raw_text   = _build_docx_ordered_text(doc)
        raw_tables = [[[ cell.text for cell in row.cells]
                       for row in tbl.rows]
                      for tbl in doc.tables]

    # ── Parse header (shared across profiles) ───────────────────────────────
    if ext == '.pdf':
        lines = raw_text.split('\n')
        tbl_start = len(lines)
        for i, line in enumerate(lines):
            s = line.strip()
            if _UNIT_RE.match(s) or _COL_HDR_LINE_RE.search(s):
                tbl_start = i
                break
        header_src = '\n'.join(lines[:tbl_start])
        pdf_table_text = '\n'.join(lines[tbl_start:])
    else:
        para_lines  = [p.text for p in doc.paragraphs if p.text.strip()]
        extra_hdr: list[str] = []
        for tbl in doc.tables[:2]:
            for row in tbl.rows[:8]:
                seen: set[str] = set()
                for cell in row.cells:
                    for txt in cell.text.split('\n'):
                        txt = txt.strip()
                        if txt and txt not in seen:
                            seen.add(txt)
                            extra_hdr.append(txt)
        header_src = '\n'.join(para_lines + extra_hdr)
        pdf_table_text = ''

    header = _parse_header_fields(header_src.split('\n'))

    # ── Try each profile ─────────────────────────────────────────────────────
    best_slots:    list  = []
    best_score:    float = -1.0
    best_pid:      str   = 'none'

    for profile in pp.get_profiles():
        ptype = profile.get('parser_type', 'table')
        try:
            if ptype == 'table':
                if raw_tables is not None:
                    # DOCX: use table-row parser with profile patterns.
                    # No qblock fallback here — qblock profiles handle those documents.
                    tpats = _compile_table_pats(profile)
                    slots = _parse_table_rows(raw_tables, _pats=tpats)
                else:
                    # PDF: use line-by-line text parser (respects profile unit pattern)
                    unit_re_p = re.compile(
                        profile.get('patterns', {}).get('unit_header', r'^unit[\s\W]*[IVXivx\d]+'),
                        re.IGNORECASE,
                    )
                    lines_p = raw_text.split('\n')
                    ts = len(lines_p)
                    for i, line in enumerate(lines_p):
                        s = line.strip()
                        if unit_re_p.match(s) or _COL_HDR_LINE_RE.search(s):
                            ts = i
                            break
                    slots = _parse_text_slots('\n'.join(lines_p[ts:]))

            elif ptype == 'qblock':
                qpats = _compile_qblock_pats(profile)
                slots = _parse_qblock_text(raw_text, _pats=qpats)

            else:
                continue

            sc = pp.score_slots(slots)
            if sc > best_score:
                best_score = sc
                best_slots = slots
                best_pid   = profile['id']

        except Exception:
            continue

    # ── LLM fallback (opt-in, requires ANTHROPIC_API_KEY) ───────────────────
    if best_score < 0.3:
        llm_slots = pp.parse_with_llm(raw_text)
        if llm_slots:
            llm_sc = pp.score_slots(llm_slots)
            if llm_sc > best_score:
                best_slots = llm_slots
                best_score = llm_sc
                best_pid   = 'llm-fallback'

    return {
        'header':     header,
        'slots':      best_slots,
        'profile_id': best_pid,
        'confidence': round(max(best_score, 0.0), 3),
    }


# ---------------------------------------------------------------------------
# CO coverage helpers
# ---------------------------------------------------------------------------

def _co_coverage_from_slots(slots):
    """Return OrderedDict {co: {'count': int, 'marks': int}} from slot list."""
    tally = {}
    for slot in slots:
        if slot.get('is_or'):
            continue
        co = (slot.get('replacement_co') or slot.get('co', '')).strip()
        if not co:
            continue
        m = re.match(r'C[O0](\d+)', co, re.IGNORECASE)
        if not m:
            continue
        key = f"CO{m.group(1)}"
        marks = int(slot.get('marks') or 0)
        if key not in tally:
            tally[key] = {'count': 0, 'marks': 0}
        tally[key]['count'] += 1
        tally[key]['marks'] += marks
    return dict(sorted(tally.items(), key=lambda kv: int(re.search(r'\d+', kv[0]).group())))


# ---------------------------------------------------------------------------
# DOCX generator — BMS institutional table format
# ---------------------------------------------------------------------------

def build_docx(header, slots, output_path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(0.75)

    def _p(text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=1):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    def _shade(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(_qn('w:val'),   'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'),  hex_color)
        tcPr.append(shd)

    # ── Header ──
    if header.get('college'):
        _p(header['college'], bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    if header.get('sub_title'):
        _p(header['sub_title'], bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    if header.get('exam_name'):
        _p(header['exam_name'], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    pairs = [
        (f"Programme: {header.get('programme','')}", f"Semester: {header.get('semester','')}"),
        (f"Branch: {header.get('branch','')}", f"Duration: {header.get('duration','')}"),
        (f"Course Code: {header.get('course_code','')}", f"Max Marks: {header.get('max_marks','')}"),
    ]
    if header.get('course'):
        pairs.append((f"Course:  {header.get('course','')}", ''))

    for left, right in pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(left)
        r1.bold = True
        r1.font.size = Pt(10)
        if right:
            p.add_run('\t\t')
            r2 = p.add_run(right)
            r2.bold = True
            r2.font.size = Pt(10)

    if header.get('instructions'):
        doc.add_paragraph()
        for line in header['instructions']:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(9)
            r.italic = 'important' in line.lower()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)

    doc.add_paragraph()

    # ── Question table ──
    # Cols: Q# | Sub | Question | CO | PO | Marks
    COL_W = [Inches(0.35), Inches(0.35), Inches(4.4), Inches(0.7), Inches(0.7), Inches(0.65)]
    tbl   = doc.add_table(rows=0, cols=6)
    tbl.style = 'Table Grid'

    def _set_widths(row):
        for i, cell in enumerate(row.cells):
            if i < len(COL_W):
                cell.width = COL_W[i]

    def _cell_fmt(cell, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        for para in cell.paragraphs:
            para.alignment = align
            for run in para.runs:
                run.font.size = Pt(size)
                run.bold = bold

    def _add_header_row():
        row = tbl.add_row()
        _set_widths(row)
        for cell, txt in zip(row.cells, ['', '', '', 'CO', 'PO', 'Marks']):
            cell.text = txt
            _cell_fmt(cell, bold=True)
            _shade(cell, 'D9E1F2')

    def _add_unit_row(label):
        row = tbl.add_row()
        merged = row.cells[0]
        for i in range(1, 6):
            merged = merged.merge(row.cells[i])
        merged.text = label
        _cell_fmt(merged, size=10, bold=True)
        _shade(merged, 'BDD7EE')

    def _add_or_row():
        row = tbl.add_row()
        merged = row.cells[0]
        for i in range(1, 6):
            merged = merged.merge(row.cells[i])
        merged.text = 'OR'
        _cell_fmt(merged, size=10, bold=True)
        _shade(merged, 'F2F2F2')

    def _add_question_row(slot):
        row = tbl.add_row()
        _set_widths(row)
        q_text = slot.get('replacement_text') or slot.get('original_text', '')
        co_val = slot.get('replacement_co')   or slot.get('co', '')
        vals   = [slot.get('qnum',''), slot.get('sub',''), q_text,
                  co_val, slot.get('po',''), str(slot.get('marks',''))]
        for i, (cell, val) in enumerate(zip(row.cells, vals)):
            cell.text = val
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
            _cell_fmt(cell, align=align)

    _add_header_row()
    current_unit = None
    for slot in slots:
        if slot.get('is_or'):
            _add_or_row()
            continue
        unit = slot.get('unit', '')
        if unit != current_unit:
            _add_unit_row(unit)
            current_unit = unit
        _add_question_row(slot)

    # ── CO Coverage table ──
    co_tally = _co_coverage_from_slots(slots)
    if co_tally:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("CO Coverage")
        r.bold = True
        r.font.size = Pt(11)

        cov_tbl = doc.add_table(rows=1, cols=3)
        cov_tbl.style = 'Table Grid'
        hdr_cells = cov_tbl.rows[0].cells
        for i, h in enumerate(["CO", "Questions", "Total Marks"]):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(hdr_cells[i], "1F497D")

        tot_count = tot_marks = 0
        for co, v in co_tally.items():
            row_cells = cov_tbl.add_row().cells
            row_cells[0].text = co
            row_cells[1].text = str(v['count'])
            row_cells[2].text = str(v['marks'])
            tot_count += v['count']
            tot_marks += v['marks']

        tot_cells = cov_tbl.add_row().cells
        _shade(tot_cells[0], "DCE6F1")
        _shade(tot_cells[1], "DCE6F1")
        _shade(tot_cells[2], "DCE6F1")
        tot_cells[0].text = "Total"
        tot_cells[1].text = str(tot_count)
        tot_cells[2].text = str(tot_marks)
        for cell in tot_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF generator — fixed-height rows to avoid blank-page overflow
# ---------------------------------------------------------------------------

_UNICODE_MAP = str.maketrans({
    '—': '--', '–': '-', '‘': "'", '’': "'",
    '“': '"', '”': '"', '…': '...', '•': '-',
    ' ': ' ', '→': '->', '≤': '<=', '≥': '>=',
    '×': 'x', '÷': '/',
})

def _safe(text):
    text = str(text).translate(_UNICODE_MAP)
    return text.encode('latin-1', errors='replace').decode('latin-1')


def build_pdf(header, slots, output_path):
    from fpdf import FPDF, XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)   # we manage page breaks manually
    pdf.add_page()
    pdf.set_margins(20, 15, 15)
    pw   = pdf.w - pdf.l_margin - pdf.r_margin
    bm   = pdf.b_margin          # bottom margin
    ph   = pdf.h                 # page height

    # ── Header ──
    if header.get('college'):
        pdf.set_font("Helvetica", "B", 14)
        pdf.multi_cell(0, 8, _safe(header['college']), align='C',
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if header.get('sub_title'):
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(0, 5, _safe(header['sub_title']), align='C',
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if header.get('exam_name'):
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 7, _safe(header['exam_name']), align='C',
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    meta = [
        (f"Programme: {header.get('programme','')}", f"Semester: {header.get('semester','')}"),
        (f"Branch: {header.get('branch','')}", f"Duration: {header.get('duration','')}"),
        (f"Course Code: {header.get('course_code','')}", f"Max Marks: {header.get('max_marks','')}"),
    ]
    if header.get('course'):
        meta.append((f"Course: {header.get('course','')}", ''))

    pdf.set_font("Helvetica", "B", 10)
    for left, right in meta:
        if right:
            pdf.cell(pw / 2, 6, _safe(left))
            pdf.cell(pw / 2, 6, _safe(right), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.cell(0, 6, _safe(left), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    if header.get('instructions'):
        pdf.set_font("Helvetica", "", 8)
        for line in header['instructions']:
            pdf.multi_cell(0, 4, _safe(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

    # ── Question table ──
    # Col widths: Q# | Sub | Question | CO | PO | Marks
    cw = [8, 8, pw - 8 - 8 - 18 - 18 - 18, 18, 18, 18]
    LH = 5   # line height in pts for question text

    def _need_page(row_h):
        """Return True if a new page is required before drawing a row of height row_h."""
        return pdf.get_y() + row_h > ph - bm - 5

    def _col_header():
        pdf.set_fill_color(217, 225, 242)
        pdf.set_font("Helvetica", "B", 8)
        for w, h in zip(cw, ['', '', '', 'CO', 'PO', 'Marks']):
            pdf.cell(w, 7, _safe(h), border=1, align='C', fill=True)
        pdf.ln()

    def _unit_row(label):
        if _need_page(8):
            pdf.add_page()
            _col_header()
        pdf.set_fill_color(189, 215, 238)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(sum(cw), 8, _safe(label), border=1, align='C', fill=True)
        pdf.ln()

    def _or_row():
        if _need_page(6):
            pdf.add_page()
            _col_header()
        pdf.set_fill_color(242, 242, 242)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(sum(cw), 6, 'OR', border=1, align='C', fill=True)
        pdf.ln()

    def _question_row(slot):
        q_text = _safe(slot.get('replacement_text') or slot.get('original_text', ''))
        co_val = _safe(slot.get('replacement_co') or slot.get('co', ''))

        # Split question text into lines that fit inside the cell
        chars_per_line = max(1, int(cw[2] / 2.1))
        lines_q = []
        remaining = q_text
        while remaining:
            lines_q.append(remaining[:chars_per_line])
            remaining = remaining[chars_per_line:]
        if not lines_q:
            lines_q = ['']
        rh = max(7, len(lines_q) * LH + 2)

        # Manual page break
        if _need_page(rh):
            pdf.add_page()
            _col_header()

        y0 = pdf.get_y()
        x0 = pdf.l_margin

        # Q# and Sub (left of question text)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_xy(x0, y0)
        pdf.cell(cw[0], rh, _safe(slot.get('qnum', '')), border=1, align='C')
        pdf.cell(cw[1], rh, _safe(slot.get('sub', '')),  border=1, align='C')

        # Question text box
        x_q = x0 + cw[0] + cw[1]
        pdf.set_draw_color(0, 0, 0)
        pdf.rect(x_q, y0, cw[2], rh)          # draw border
        pdf.set_font("Helvetica", "", 8)
        for i, line in enumerate(lines_q):
            pdf.set_xy(x_q + 1, y0 + 1 + i * LH)
            pdf.cell(cw[2] - 2, LH, line, border=0, align='L')

        # CO | PO | Marks
        side_vals = [co_val, _safe(slot.get('po', '')), _safe(str(slot.get('marks', '')))]
        pdf.set_font("Helvetica", "", 8)
        x_cur = x_q + cw[2]
        for w, val in zip(cw[3:], side_vals):
            pdf.set_xy(x_cur, y0)
            pdf.cell(w, rh, val, border=1, align='C')
            x_cur += w

        pdf.set_xy(x0, y0 + rh)

    # Draw initial column header row
    _col_header()

    current_unit = None
    for slot in slots:
        if slot.get('is_or'):
            _or_row()
            continue
        unit = slot.get('unit', '')
        if unit != current_unit:
            _unit_row(unit)
            current_unit = unit
        _question_row(slot)

    # ── CO Coverage table ──
    co_tally = _co_coverage_from_slots(slots)
    if co_tally:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(220, 230, 241)
        pdf.multi_cell(0, 8, "CO Coverage", fill=True,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        col_w = [pw * 0.25, pw * 0.35, pw * 0.40]
        headers = ["CO", "Questions", "Total Marks"]

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(31, 73, 125)
        pdf.set_text_color(255, 255, 255)
        for w, h in zip(col_w, headers):
            pdf.cell(w, 8, h, border=1, fill=True, align='C')
        pdf.ln()
        pdf.set_text_color(0, 0, 0)

        tot_count = tot_marks = 0
        for idx, (co, v) in enumerate(co_tally.items()):
            fill = idx % 2 == 0
            pdf.set_fill_color(240, 245, 255) if fill else pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Helvetica", "", 9)
            for w, val in zip(col_w, [co, str(v['count']), str(v['marks'])]):
                pdf.cell(w, 7, val, border=1, fill=fill, align='C')
            pdf.ln()
            tot_count += v['count']
            tot_marks += v['marks']

        pdf.set_fill_color(220, 230, 241)
        pdf.set_font("Helvetica", "B", 9)
        for w, val in zip(col_w, ["Total", str(tot_count), str(tot_marks)]):
            pdf.cell(w, 8, val, border=1, fill=True, align='C')
        pdf.ln()

    pdf.output(output_path)


# ---------------------------------------------------------------------------
# TXT generator
# ---------------------------------------------------------------------------

def build_txt(header, slots, output_path):
    lines = []
    for key in ('college', 'sub_title', 'exam_name'):
        if header.get(key):
            lines.append(header[key])
    lines.append('')
    for key, label in [
        ('programme', 'Programme'), ('semester', 'Semester'),
        ('branch', 'Branch'),       ('duration', 'Duration'),
        ('course_code', 'Course Code'), ('max_marks', 'Max Marks'),
        ('course', 'Course'),
    ]:
        if header.get(key):
            lines.append(f"{label}: {header[key]}")
    lines.append('')
    for inst in (header.get('instructions') or []):
        lines.append(inst)
    lines.append('')
    lines.append('=' * 80)

    current_unit = None
    for slot in slots:
        if slot.get('is_or'):
            lines.append('  OR')
            continue
        unit = slot.get('unit', '')
        if unit != current_unit:
            lines.append(f'\n  {unit}')
            current_unit = unit
        q = slot.get('qnum', '')
        s = slot.get('sub', '')
        t = slot.get('replacement_text') or slot.get('original_text', '')
        co = slot.get('replacement_co') or slot.get('co', '')
        po = slot.get('po', '')
        m  = slot.get('marks', '')
        prefix = f"  {q:2s} {s:3s}".rstrip()
        lines.append(f"{prefix}  {t}  {co}  {po}  {m} marks")

    co_tally = _co_coverage_from_slots(slots)
    if co_tally:
        lines.append('\n' + '=' * 80)
        lines.append('CO Coverage')
        lines.append(f"  {'CO':<8} {'Questions':>10} {'Total Marks':>12}")
        lines.append(f"  {'-'*8} {'-'*10} {'-'*12}")
        tot_count = tot_marks = 0
        for co, v in co_tally.items():
            lines.append(f"  {co:<8} {v['count']:>10} {v['marks']:>12}")
            tot_count += v['count']
            tot_marks += v['marks']
        lines.append(f"  {'-'*8} {'-'*10} {'-'*12}")
        lines.append(f"  {'Total':<8} {tot_count:>10} {tot_marks:>12}")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
