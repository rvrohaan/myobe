"""Module 3 Comprehensive Report Generator.

Builds a DOCX/PDF/TXT report covering all Module-3 deliverables:
  1.  Course Delivery Overview
  2.  Course Outcomes & Bloom's Distribution
  3.  CO-PO-SDG Alignment Summary
  4.  Lesson Plan Coverage Analysis
  5.  Session Delivery Plan
  6.  Teaching Diary Analysis
  7.  CO-Session Coverage Matrix
  8.  Teaching Method Distribution
  9.  SDG & AI Integration
 10.  NBA/NAAC Compliance Checklist
 11.  Course Delivery Readiness Dashboard
"""

import re
import json
import datetime

try:
    import report_charts as _rc
except Exception:
    _rc = None


# ── Rule-based analytics ──────────────────────────────────────────────────────

def compute_analytics(lp_data: dict, td_data: dict) -> dict:
    """Derive all quantitative metrics from lesson-plan and teaching-diary dicts."""
    lp = lp_data or {}
    td = td_data or {}

    cos            = lp.get("cos", []) or td.get("cos", [])
    session_plan   = lp.get("session_plan", [])
    co_po_sdg      = lp.get("co_po_sdg_mapping", [])
    units          = lp.get("units", [])
    sdg_framework  = lp.get("sdg_framework", [])
    industry_ai    = lp.get("industry_ai", [])

    theory_sessions = td.get("theory_sessions", [])
    tutorials       = td.get("tutorials", [])
    lab_sessions    = td.get("lab_sessions", [])
    co_attainment   = td.get("co_attainment", [])
    integration     = td.get("integration_summary", {})
    nba_alignment   = td.get("nba_alignment", [])

    n_cos = len(cos)

    # Bloom's distribution from LP COs
    bloom_dist: dict[str, int] = {}
    for co in cos:
        bl = (co.get("bloom") or "Unknown").capitalize()
        bloom_dist[bl] = bloom_dist.get(bl, 0) + 1

    # CO coverage in session_plan
    cos_in_sessions: set[str] = set()
    for s in session_plan:
        tag = (s.get("co") or "").strip()
        if tag:
            cos_in_sessions.add(tag.upper())
    lp_co_coverage_pct = round(len(cos_in_sessions) / n_cos * 100, 1) if n_cos else 0

    # Session-level CO tally {CO1: N, CO2: M, ...}
    co_session_tally: dict[str, int] = {}
    for s in session_plan:
        tag = (s.get("co") or "").strip().upper()
        if tag:
            co_session_tally[tag] = co_session_tally.get(tag, 0) + 1

    # Teaching method distribution
    method_dist: dict[str, int] = {}
    for s in session_plan:
        m = (s.get("method") or "Lecture").strip()
        method_dist[m] = method_dist.get(m, 0) + 1

    # SDG distribution across sessions
    sdg_dist: dict[str, int] = {}
    for s in session_plan:
        sdg = (s.get("sdg") or "").strip()
        if sdg:
            sdg_dist[sdg] = sdg_dist.get(sdg, 0) + 1

    # PO coverage: list of POs that have any score > 0 across all COs
    active_pos: set[str] = set()
    for entry in co_po_sdg:
        for po, score in (entry.get("po_scores") or {}).items():
            if score and score > 0:
                active_pos.add(po)

    # PO strength table: for each PO, sum of scores across all COs
    po_strength: dict[str, int] = {}
    for entry in co_po_sdg:
        for po, score in (entry.get("po_scores") or {}).items():
            po_strength[po] = po_strength.get(po, 0) + int(score or 0)

    # TD session counts
    n_theory   = len(theory_sessions)
    n_tutorial = len(tutorials)
    n_lab      = len(lab_sessions)
    n_td_total = n_theory + n_tutorial + n_lab

    # CO attainment targets
    att_targets = {a.get("co", "?"): a.get("target", 60) for a in co_attainment}

    # SDG unique goals in session plan + sdg_framework
    all_sdgs: set[str] = set(sdg_dist.keys())
    for sf in sdg_framework:
        g = (sf.get("sdg_goal") or "").strip()
        if g:
            all_sdgs.add(g)

    # Industry/AI integration count from TD
    ai_learning = td.get("ai_learning", [])
    industry_esg = td.get("industry_esg", [])

    return {
        "cos":                cos,
        "n_cos":              n_cos,
        "bloom_dist":         bloom_dist,
        "cos_in_sessions":    sorted(cos_in_sessions),
        "lp_co_coverage_pct": lp_co_coverage_pct,
        "co_session_tally":   co_session_tally,
        "n_sessions":         len(session_plan),
        "method_dist":        method_dist,
        "sdg_dist":           sdg_dist,
        "all_sdgs":           sorted(all_sdgs),
        "active_pos":         sorted(active_pos),
        "po_strength":        po_strength,
        "n_theory":           n_theory,
        "n_tutorial":         n_tutorial,
        "n_lab":              n_lab,
        "n_td_total":         n_td_total,
        "att_targets":        att_targets,
        "n_units":            len(units),
        "n_mini_projects":    len(lp.get("mini_projects", [])),
        "n_ai_learning":      len(ai_learning),
        "n_industry_esg":     len(industry_esg),
        "nba_alignment":      nba_alignment,
        "integration":        integration,
        "final_summary":      td.get("final_summary", {}),
    }


# ── AI qualitative analysis ───────────────────────────────────────────────────

def get_ai_analysis(client, lp_data: dict, td_data: dict, analytics: dict,
                    code: str, title: str) -> dict:
    """Single Claude Haiku call returning JSON for qualitative report sections."""
    cos = analytics["cos"]
    co_lines = "\n".join(
        f"  CO{c['num']}: {c['statement']}  [Bloom: {c.get('bloom','?')}]"
        for c in cos
    )
    method_str  = ", ".join(f"{k}({v})" for k, v in analytics["method_dist"].items())
    sdg_str     = ", ".join(analytics["all_sdgs"]) or "None"
    bloom_str   = ", ".join(f"{k}:{v}" for k, v in analytics["bloom_dist"].items())

    prompt = f"""You are an NBA/NAAC accreditation expert for Indian engineering colleges.

Course: {title} ({code})
Course Outcomes ({analytics['n_cos']} total):
{co_lines}

Bloom's Distribution: {bloom_str}
Lesson Plan: {analytics['n_sessions']} sessions across {analytics['n_units']} units
CO Coverage in Sessions: {analytics['lp_co_coverage_pct']}% ({len(analytics['cos_in_sessions'])} of {analytics['n_cos']} COs)
Teaching Methods: {method_str}
SDG Goals Addressed: {sdg_str}
TD Sessions: {analytics['n_theory']} theory, {analytics['n_tutorial']} tutorial, {analytics['n_lab']} lab

Provide a structured analysis in EXACT JSON (no markdown, no extra text):
{{
  "co_quality": {{
    "assessment": "...(2-3 sentences on CO statement quality and Bloom alignment)",
    "rating": "Excellent|Good|Satisfactory|Needs Improvement"
  }},
  "delivery_plan_quality": {{
    "assessment": "...(2-3 sentences on lesson plan completeness and session coverage)",
    "rating": "Excellent|Good|Satisfactory|Needs Improvement",
    "strengths": ["...", "..."],
    "gaps": ["...", "..."]
  }},
  "sdg_integration": {{
    "assessment": "...(2 sentences on SDG integration quality)",
    "rating": "Excellent|Good|Satisfactory|Needs Improvement"
  }},
  "obe_compliance": {{
    "co_po_alignment": "Excellent|Good|Satisfactory|Needs Improvement",
    "delivery_coverage": "Excellent|Good|Satisfactory|Needs Improvement",
    "overall_obe_score": 82,
    "compliance_level": "Fully Compliant|Substantially Compliant|Partially Compliant",
    "strengths": ["...", "..."],
    "recommendations": ["...", "..."]
  }},
  "accreditation_readiness": {{
    "overall_score": 80,
    "readiness_level": "Ready|Mostly Ready|Needs Work",
    "key_metrics": [
      {{"metric": "CO Formulation", "score": 85, "status": "Strong"}},
      {{"metric": "Session Coverage", "score": 80, "status": "Good"}},
      {{"metric": "Bloom Alignment", "score": 75, "status": "Good"}},
      {{"metric": "SDG Integration", "score": 70, "status": "Adequate"}},
      {{"metric": "Teaching Methods", "score": 78, "status": "Good"}}
    ],
    "action_items": ["...", "..."],
    "strengths": ["...", "..."]
  }}
}}"""

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = re.sub(r"```(?:json)?\s*|\s*```", "", resp.content[0].text).strip()
        start = raw.find('{')
        end   = raw.rfind('}')
        if start >= 0 and end > start:
            raw = raw[start:end + 1]
        return json.loads(raw)
    except Exception:
        return {}


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(lp_data: dict, td_data: dict, analytics: dict, ai: dict,
               code: str, title: str, semester, output_path: str):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    INDIGO = RGBColor(0x31, 0x2E, 0x81)
    TEAL   = RGBColor(0x06, 0x4E, 0x5B)
    GREEN  = RGBColor(0x06, 0x65, 0x28)
    AMBER  = RGBColor(0x78, 0x35, 0x00)
    RED_C  = RGBColor(0x7F, 0x1D, 0x1D)
    GRAY   = RGBColor(0x37, 0x41, 0x51)

    STATUS_COLOR = {
        "Excellent": GREEN, "Strong": GREEN, "Good": GREEN, "Ready": GREEN,
        "Satisfactory": AMBER, "Adequate": AMBER, "Mostly Ready": AMBER,
        "Needs Improvement": RED_C, "Needs Work": RED_C, "Poor": RED_C,
        "Fully Compliant": GREEN, "Substantially Compliant": AMBER,
        "Partially Compliant": RED_C,
    }

    def _add_heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color and p.runs:
            p.runs[0].font.color.rgb = color
        return p

    def _add_para(text, bold=False, color=None, indent=0, size=10):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent * 0.25)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(2)
        return p

    def _add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
        for ri, row_data in enumerate(rows):
            r = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = r.cells[ci]
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        if col_widths:
            for ri2, row2 in enumerate(t.rows):
                for ci2, cell2 in enumerate(row2.cells):
                    if ci2 < len(col_widths):
                        cell2.width = Inches(col_widths[ci2])
        return t

    def _spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    def _insert_chart(buf, width=5.2):
        if buf is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(_fname, width=Inches(width))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        except Exception:
            pass
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    # ── Cover ──
    sem_label = f"  |  Semester {semester}" if semester else ""
    p = doc.add_heading(f"{code}  -  {title}", 0)
    if p.runs:
        p.runs[0].font.color.rgb = INDIGO
    doc.add_heading("Module 3 Comprehensive Report", 1).runs[0].font.color.rgb = INDIGO
    _add_para(f"Course Delivery Management Report{sem_label}", color=GRAY)
    _add_para(f"Generated: {datetime.date.today().strftime('%d %B %Y')}", color=GRAY, size=9)
    _spacer()

    lp = lp_data or {}
    td = td_data or {}

    # ── Section 1: Course Delivery Overview ──
    _add_heading("1. Course Delivery Overview", 2, color=INDIGO)
    meta_lp = lp.get("meta", {}) or {}
    meta_td = td.get("meta", {}) or {}
    meta    = {**meta_td, **meta_lp}

    overview_rows = [
        ("Academic Year",  meta.get("academic_year", "-")),
        ("Semester",       meta.get("semester", "-")),
        ("Regulation",     meta.get("regulation", "-")),
        ("Program",        meta.get("program", "-")),
        ("Course Type",    meta.get("course_type", "-")),
        ("Credits / LTP",  f"{meta.get('credits', '-')} / {meta.get('ltp', '-')}"),
        ("Faculty",        meta.get("faculty_name", "-")),
        ("Department",     meta.get("department", "-")),
    ]
    _add_table(["Field", "Value"], overview_rows, col_widths=[2.0, 4.0])
    _spacer()

    overview_text = lp.get("course_overview") or ""
    if overview_text:
        _add_para(overview_text, size=10)
        _spacer()

    # ── Section 2: Course Outcomes & Bloom's Distribution ──
    _add_heading("2. Course Outcomes & Bloom's Taxonomy Distribution", 2, color=INDIGO)
    cos = analytics["cos"]
    if cos:
        _add_table(
            ["CO", "Statement", "Bloom's Level", "Knowledge Level"],
            [(f"CO{c['num']}", c["statement"], c.get("bloom", "-"), c.get("knowledge", "-"))
             for c in cos],
            col_widths=[0.6, 3.6, 1.0, 1.0],
        )
        _spacer()
        bloom_rows = sorted(analytics["bloom_dist"].items(),
                            key=lambda x: x[1], reverse=True)
        _add_para("Bloom's Level Distribution:", bold=True, size=10)
        _add_table(
            ["Bloom's Level", "Count", "% of COs"],
            [(level, count, f"{round(count / analytics['n_cos'] * 100)}%")
             for level, count in bloom_rows],
            col_widths=[2.0, 1.0, 1.0],
        )
        if _rc:
            _insert_chart(_rc.bloom_distribution_chart(analytics["bloom_dist"]))
    _spacer()

    # ── Section 3: CO-PO-SDG Alignment ──
    _add_heading("3. CO-PO-SDG Alignment Summary", 2, color=INDIGO)
    co_po_sdg = lp.get("co_po_sdg_mapping", [])
    if co_po_sdg:
        pos = sorted({po for entry in co_po_sdg for po in (entry.get("po_scores") or {})})
        headers = ["CO"] + pos + ["PSO1", "SDG"]
        table_rows = []
        for entry in co_po_sdg:
            scores = entry.get("po_scores") or {}
            row = [entry.get("co", "-")]
            row += [str(scores.get(po, 0)) for po in pos]
            row += [str(entry.get("pso1", 0)), entry.get("sdg", "-")]
            table_rows.append(row)
        total_cols = len(headers)
        col_w = [0.55] + [0.4] * (total_cols - 3) + [0.45, 0.65]
        _add_table(headers, table_rows, col_widths=col_w)
        _spacer()
        _add_para(
            f"Active POs (score > 0): {', '.join(analytics['active_pos']) or 'None'}",
            color=TEAL, size=9,
        )
    _spacer()

    # ── Section 4: Lesson Plan Coverage ──
    _add_heading("4. Lesson Plan Coverage Analysis", 2, color=INDIGO)
    units = lp.get("units", [])
    if units:
        _add_table(
            ["Unit", "Title", "Mapped CO", "POs", "SDG", "AI Tool"],
            [(f"Unit {u['num']}", u.get("title", "-"), u.get("co", "-"),
              u.get("pos", "-"), u.get("sdg", "-"), u.get("ai_tool", "-"))
             for u in units],
            col_widths=[0.5, 1.8, 0.7, 1.0, 0.7, 0.9],
        )
        _spacer()

    n_sess = analytics["n_sessions"]
    cov_pct = analytics["lp_co_coverage_pct"]
    _add_para(
        f"Sessions planned: {n_sess}  |  CO coverage in sessions: {cov_pct}%  "
        f"({len(analytics['cos_in_sessions'])} of {analytics['n_cos']} COs addressed)",
        bold=True, color=TEAL if cov_pct >= 80 else AMBER, size=10,
    )
    _spacer()

    # ── Section 5: Session Delivery Plan ──
    _add_heading("5. Session Delivery Plan", 2, color=INDIGO)
    session_plan = lp.get("session_plan", [])
    if session_plan:
        _add_table(
            ["#", "Topic", "CO", "Method", "SDG", "Assessment"],
            [(str(s.get("session", i + 1)), s.get("topic", "-"),
              s.get("co", "-"), s.get("method", "-"),
              s.get("sdg", "-"), s.get("assessment", "-"))
             for i, s in enumerate(session_plan)],
            col_widths=[0.3, 2.5, 0.55, 1.0, 0.65, 1.0],
        )
    _spacer()

    # ── Section 6: Teaching Diary Analysis ──
    _add_heading("6. Teaching Diary Analysis", 2, color=INDIGO)
    td_counts = [
        ("Theory Sessions",   analytics["n_theory"]),
        ("Tutorial Sessions", analytics["n_tutorial"]),
        ("Lab Sessions",      analytics["n_lab"]),
        ("Total TD Sessions", analytics["n_td_total"]),
    ]
    _add_table(["Session Type", "Count"], td_counts, col_widths=[2.5, 1.0])
    _spacer()

    if td.get("theory_sessions"):
        _add_para("Theory Sessions (sample):", bold=True, size=10)
        sample = td["theory_sessions"][:8]
        _add_table(
            ["#", "Unit", "Topic", "CO", "Method", "Bloom"],
            [(str(s.get("no", i + 1)), s.get("unit", "-"), s.get("topic", "-"),
              s.get("co", "-"), s.get("method", "-"), s.get("bloom", "-"))
             for i, s in enumerate(sample)],
            col_widths=[0.3, 0.8, 2.2, 0.55, 0.9, 0.85],
        )
        _spacer()

    final = analytics["final_summary"]
    if final:
        _add_para("Delivery Summary:", bold=True, size=10)
        fs_rows = [
            ("Syllabus Coverage",   final.get("syllabus_coverage", "-")),
            ("SDG Activities",      final.get("sdg_activities", "-")),
            ("AI Activities",       final.get("ai_activities", "-")),
            ("Industry Activities", final.get("industry_activities", "-")),
            ("Innovations",         final.get("innovations", "-")),
        ]
        _add_table(["Metric", "Value"], fs_rows, col_widths=[2.5, 1.5])
        _spacer()

    # ── Section 7: CO-Session Coverage Matrix ──
    _add_heading("7. CO-Session Coverage Matrix", 2, color=INDIGO)
    co_tally = analytics["co_session_tally"]
    if cos and co_tally:
        cov_rows = []
        for c in cos:
            key  = f"CO{c['num']}"
            n    = co_tally.get(key, 0)
            pct  = round(n / analytics["n_sessions"] * 100) if analytics["n_sessions"] else 0
            status = "Covered" if n > 0 else "Not Covered"
            cov_rows.append((key, c["statement"][:55], str(n), f"{pct}%", status))
        _add_table(
            ["CO", "Statement", "Sessions", "% of Plan", "Status"],
            cov_rows,
            col_widths=[0.55, 2.8, 0.7, 0.7, 0.9],
        )
        if _rc:
            _insert_chart(_rc.co_session_chart(cos, analytics["co_session_tally"]))
    _spacer()

    # ── Section 8: Teaching Method Distribution ──
    _add_heading("8. Teaching Method Distribution", 2, color=INDIGO)
    if analytics["method_dist"]:
        total_sess = sum(analytics["method_dist"].values())
        method_rows = sorted(analytics["method_dist"].items(), key=lambda x: x[1], reverse=True)
        _add_table(
            ["Teaching Method", "Sessions", "% of Plan"],
            [(m, str(n), f"{round(n / total_sess * 100)}%") for m, n in method_rows],
            col_widths=[2.5, 1.0, 1.0],
        )
        if _rc:
            _insert_chart(_rc.teaching_method_chart(analytics["method_dist"]))
    _spacer()

    # ── Section 9: SDG & AI Integration ──
    _add_heading("9. SDG & AI Integration", 2, color=INDIGO)
    if analytics["sdg_dist"]:
        sdg_rows = sorted(analytics["sdg_dist"].items(), key=lambda x: x[1], reverse=True)
        _add_table(
            ["SDG Goal", "Sessions Covering"],
            sdg_rows,
            col_widths=[1.5, 1.5],
        )
        _spacer()

    industry_ai_list = lp.get("industry_ai", [])
    if industry_ai_list:
        _add_para("Industry & AI Integration (from Lesson Plan):", bold=True, size=10)
        _add_table(
            ["Topic", "Industry Application", "AI Integration"],
            [(ia.get("topic", "-"), ia.get("industry", "-"), ia.get("ai_integration", "-"))
             for ia in industry_ai_list[:6]],
            col_widths=[1.5, 2.0, 2.0],
        )
        _spacer()

    # ── Section 10: NBA/NAAC Compliance ──
    _add_heading("10. NBA/NAAC Compliance Checklist", 2, color=INDIGO)

    # AI analysis sections
    obe = ai.get("obe_compliance", {})
    compliance_level = obe.get("compliance_level", "")
    cl_color = STATUS_COLOR.get(compliance_level, GRAY)
    if compliance_level:
        _add_para(f"OBE Compliance Level: {compliance_level}", bold=True, color=cl_color, size=10)
        _spacer()

    checklist_items = [
        ("CO-PO Alignment",        obe.get("co_po_alignment", "-")),
        ("Delivery Coverage",      obe.get("delivery_coverage", "-")),
        ("Overall OBE Score",      f"{obe.get('overall_obe_score', '-')}/100"),
    ]
    # Add NBA frameworks from TD
    for nb in analytics["nba_alignment"]:
        checklist_items.append((nb.get("framework", "-"), nb.get("evidence", "-")))
    _add_table(["Item", "Status / Evidence"], checklist_items, col_widths=[2.0, 4.0])
    _spacer()

    strengths_obe = obe.get("strengths", [])
    if strengths_obe:
        _add_para("Strengths:", bold=True, size=10)
        for s in strengths_obe:
            _add_para(f"  + {s}", color=GREEN, indent=1, size=9)
    recs_obe = obe.get("recommendations", [])
    if recs_obe:
        _add_para("Recommendations:", bold=True, size=10)
        for r in recs_obe:
            _add_para(f"  - {r}", color=AMBER, indent=1, size=9)
    _spacer()

    # ── Section 11: Readiness Dashboard ──
    _add_heading("11. Course Delivery Readiness Dashboard", 2, color=INDIGO)
    rd = ai.get("accreditation_readiness", {})
    overall = rd.get("overall_score", "-")
    rl      = rd.get("readiness_level", "-")
    rl_color = STATUS_COLOR.get(rl, GRAY)

    _add_para(
        f"Overall Readiness Score: {overall}/100  |  Level: {rl}",
        bold=True, color=rl_color, size=11,
    )
    _spacer()

    metrics = rd.get("key_metrics", [])
    if metrics:
        _add_table(
            ["Metric", "Score", "Status"],
            [(m["metric"], f"{m['score']}/100", m["status"]) for m in metrics],
            col_widths=[2.5, 0.8, 1.2],
        )
        _spacer()

    ai_items = rd.get("action_items", [])
    if ai_items:
        _add_para("Action Items:", bold=True, size=10)
        for item in ai_items:
            _add_para(f"  -> {item}", color=AMBER, indent=1, size=9)

    rd_strengths = rd.get("strengths", [])
    if rd_strengths:
        _add_para("Strengths:", bold=True, size=10)
        for s in rd_strengths:
            _add_para(f"  + {s}", color=GREEN, indent=1, size=9)

    doc.save(output_path)


# ── PDF builder ───────────────────────────────────────────────────────────────

def build_pdf(lp_data: dict, td_data: dict, analytics: dict, ai: dict,
              code: str, title: str, semester, output_path: str):
    """Build a PDF version of the Module 3 report using fpdf2."""
    from fpdf import FPDF, XPos, YPos

    def _safe(text: str) -> str:
        if not text:
            return ""
        return (str(text)
                .encode("latin-1", errors="replace")
                .decode("latin-1"))

    class PDF(FPDF):
        def header(self):
            pass

        def _section_heading(self, text, level=1):
            self.ln(4)
            sz = 13 if level == 1 else 11 if level == 2 else 10
            self.set_font("Helvetica", "B", sz)
            self.set_text_color(49, 46, 129)
            self.multi_cell(0, 6, _safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            self.set_text_color(0, 0, 0)
            self.ln(1)

        def _para(self, text, bold=False, color=None, size=10):
            self.set_font("Helvetica", "B" if bold else "", size)
            if color:
                self.set_text_color(*color)
            self.multi_cell(0, 5, _safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            self.set_text_color(0, 0, 0)
            self.ln(1)

        def _table(self, headers, rows, col_ratios=None):
            if not headers:
                return
            page_w = self.w - 2 * self.l_margin
            if col_ratios:
                total = sum(col_ratios)
                col_ws = [page_w * r / total for r in col_ratios]
            else:
                col_ws = [page_w / len(headers)] * len(headers)

            def _row(cells, bold=False, fill=False):
                line_h = 5
                self.set_font("Helvetica", "B" if bold else "", 8.5)
                if fill:
                    self.set_fill_color(220, 220, 235)
                else:
                    self.set_fill_color(255, 255, 255)
                x0 = self.l_margin
                y0 = self.get_y()
                max_lines = 1
                for ci, cell_text in enumerate(cells):
                    safe = _safe(str(cell_text))
                    cw = col_ws[ci] if ci < len(col_ws) else col_ws[-1]
                    if cw > 0:
                        sw = self.get_string_width(safe)
                        max_lines = max(max_lines, max(1, int(sw / max(cw - 4, 1)) + 1))
                max_h = max_lines * line_h + 2
                if y0 + max_h > self.h - 15:
                    self.add_page()
                    y0 = self.get_y()
                for ci, cell_text in enumerate(cells):
                    safe = _safe(str(cell_text))
                    cw = col_ws[ci] if ci < len(col_ws) else col_ws[-1]
                    self.set_xy(x0 + sum(col_ws[:ci]), y0)
                    self.multi_cell(cw, line_h, safe, border=1, fill=fill,
                                    new_x=XPos.RIGHT, new_y=YPos.TOP)
                self.set_xy(x0, y0 + max_h)

            _row(headers, bold=True, fill=True)
            for row in rows:
                _row(row)
            self.ln(2)

    def _pdf_chart(buf, w=155):
        if buf is None or _rc is None:
            return
        import tempfile as _tf, os as _os2
        _fname = None
        try:
            with _tf.NamedTemporaryFile(delete=False, suffix='.png') as _f:
                _f.write(buf.read())
                _fname = _f.name
            x = pdf.l_margin + (180 - w) / 2
            pdf.image(_fname, x=x, w=w)
            pdf.ln(3)
        except Exception:
            pass
        finally:
            if _fname:
                try:
                    _os2.unlink(_fname)
                except Exception:
                    pass

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    sem_label = f"  |  Semester {semester}" if semester else ""
    lp = lp_data or {}
    td = td_data or {}

    pdf._section_heading(f"{code}  -  {title}", 1)
    pdf._section_heading("Module 3 Comprehensive Report", 2)
    pdf._para(f"Course Delivery Management Report{sem_label}")
    pdf._para(f"Generated: {datetime.date.today().strftime('%d %B %Y')}", size=9)
    pdf.ln(3)

    # 1. Overview
    pdf._section_heading("1. Course Delivery Overview", 2)
    meta = {**(td.get("meta") or {}), **(lp.get("meta") or {})}
    ov = [
        ("Academic Year",  meta.get("academic_year", "-")),
        ("Semester",       meta.get("semester", "-")),
        ("Regulation",     meta.get("regulation", "-")),
        ("Program",        meta.get("program", "-")),
        ("Course Type",    meta.get("course_type", "-")),
        ("Credits / LTP",  f"{meta.get('credits', '-')} / {meta.get('ltp', '-')}"),
    ]
    pdf._table(["Field", "Value"], ov, col_ratios=[1, 2])

    overview_text = lp.get("course_overview", "")
    if overview_text:
        pdf._para(overview_text)

    # 2. COs & Bloom's
    pdf._section_heading("2. Course Outcomes & Bloom's Distribution", 2)
    cos = analytics["cos"]
    if cos:
        pdf._table(
            ["CO", "Statement", "Bloom", "Knowledge"],
            [(f"CO{c['num']}", c["statement"][:60], c.get("bloom","-"), c.get("knowledge","-"))
             for c in cos],
            col_ratios=[0.5, 3.5, 1, 1],
        )
        bloom_rows = sorted(analytics["bloom_dist"].items(), key=lambda x: x[1], reverse=True)
        pdf._table(
            ["Bloom Level", "Count", "% of COs"],
            [(lv, str(n), f"{round(n/analytics['n_cos']*100)}%") for lv, n in bloom_rows],
            col_ratios=[2, 1, 1],
        )
        _pdf_chart(_rc.bloom_distribution_chart(analytics["bloom_dist"]) if _rc else None)

    # 3. CO-PO-SDG
    pdf._section_heading("3. CO-PO-SDG Alignment", 2)
    co_po_sdg = lp.get("co_po_sdg_mapping", [])
    if co_po_sdg:
        pos = sorted({po for e in co_po_sdg for po in (e.get("po_scores") or {})})
        hdrs = ["CO"] + pos + ["PSO1", "SDG"]
        t_rows = []
        for e in co_po_sdg:
            sc = e.get("po_scores") or {}
            t_rows.append([e.get("co","-")] + [str(sc.get(p,0)) for p in pos]
                          + [str(e.get("pso1",0)), e.get("sdg","-")])
        ratios = [0.5] + [0.35] * len(pos) + [0.35, 0.6]
        pdf._table(hdrs, t_rows, col_ratios=ratios)

    # 4. Lesson Plan Coverage
    pdf._section_heading("4. Lesson Plan Coverage Analysis", 2)
    units = lp.get("units", [])
    if units:
        pdf._table(
            ["Unit", "Title", "CO", "POs", "SDG"],
            [(f"U{u['num']}", u.get("title","-")[:30], u.get("co","-"),
              u.get("pos","-"), u.get("sdg","-"))
             for u in units],
            col_ratios=[0.4, 2.5, 0.6, 1.2, 0.6],
        )
    pdf._para(
        f"Sessions: {analytics['n_sessions']}  |  CO coverage: "
        f"{analytics['lp_co_coverage_pct']}% ({len(analytics['cos_in_sessions'])} of {analytics['n_cos']} COs)",
        bold=True,
    )

    # 5. Session Plan
    pdf._section_heading("5. Session Delivery Plan", 2)
    sp = lp.get("session_plan", [])
    if sp:
        pdf._table(
            ["#", "Topic", "CO", "Method", "SDG"],
            [(str(s.get("session", i+1)), s.get("topic","-")[:40],
              s.get("co","-"), s.get("method","-")[:60], s.get("sdg","-"))
             for i, s in enumerate(sp)],
            col_ratios=[0.3, 3, 0.5, 1.5, 0.6],
        )

    # 6. TD Analysis
    pdf._section_heading("6. Teaching Diary Analysis", 2)
    pdf._table(
        ["Session Type", "Count"],
        [("Theory", str(analytics["n_theory"])),
         ("Tutorial", str(analytics["n_tutorial"])),
         ("Lab", str(analytics["n_lab"])),
         ("Total", str(analytics["n_td_total"]))],
        col_ratios=[2, 1],
    )
    final = analytics["final_summary"]
    if final:
        pdf._table(
            ["Metric", "Value"],
            [(k.replace("_", " ").title(), str(v)) for k, v in final.items()],
            col_ratios=[2, 1],
        )

    # 7. CO-Session Coverage Matrix
    pdf._section_heading("7. CO-Session Coverage Matrix", 2)
    co_tally = analytics["co_session_tally"]
    if cos and co_tally:
        def _co_pdf_row(c):
            key = f"CO{c['num']}"
            n   = co_tally.get(key, 0)
            pct = round(n / analytics["n_sessions"] * 100) if analytics["n_sessions"] else 0
            return (key, c["statement"][:45], str(n), f"{pct}%",
                    "Covered" if n > 0 else "Not Covered")
        pdf._table(
            ["CO", "Statement", "Sessions", "% Plan", "Status"],
            [_co_pdf_row(c) for c in cos],
            col_ratios=[0.5, 3, 0.6, 0.6, 0.9],
        )
        _pdf_chart(_rc.co_session_chart(cos, analytics["co_session_tally"]) if _rc else None)

    # 8. Method Distribution
    pdf._section_heading("8. Teaching Method Distribution", 2)
    if analytics["method_dist"]:
        total = sum(analytics["method_dist"].values())
        pdf._table(
            ["Method", "Sessions", "%"],
            [(m[:80], str(n), f"{round(n/total*100)}%")
             for m, n in sorted(analytics["method_dist"].items(), key=lambda x: x[1], reverse=True)],
            col_ratios=[2.5, 1, 1],
        )
        _pdf_chart(_rc.teaching_method_chart(analytics["method_dist"]) if _rc else None)

    # 9. SDG & AI
    pdf._section_heading("9. SDG & AI Integration", 2)
    if analytics["sdg_dist"]:
        pdf._table(
            ["SDG Goal", "Sessions"],
            [(k, str(v)) for k, v in sorted(analytics["sdg_dist"].items(), key=lambda x: x[1], reverse=True)],
            col_ratios=[2, 1],
        )

    # 10. NBA/NAAC
    pdf._section_heading("10. NBA/NAAC Compliance Checklist", 2)
    obe = ai.get("obe_compliance", {})
    cl = obe.get("compliance_level", "-")
    pdf._para(f"Compliance Level: {cl}  |  OBE Score: {obe.get('overall_obe_score','-')}/100", bold=True)
    if analytics["nba_alignment"]:
        pdf._table(
            ["Framework", "Evidence"],
            [(nb.get("framework","-"), nb.get("evidence","-"))
             for nb in analytics["nba_alignment"]],
            col_ratios=[1, 3],
        )
    for s in obe.get("strengths", []):
        pdf._para(f"  + {s}", size=9)
    for r in obe.get("recommendations", []):
        pdf._para(f"  -> {r}", size=9)

    # 11. Readiness Dashboard
    pdf._section_heading("11. Course Delivery Readiness Dashboard", 2)
    rd = ai.get("accreditation_readiness", {})
    pdf._para(
        f"Overall Score: {rd.get('overall_score','-')}/100  |  Level: {rd.get('readiness_level','-')}",
        bold=True,
    )
    if rd.get("key_metrics"):
        pdf._table(
            ["Metric", "Score", "Status"],
            [(m["metric"], f"{m['score']}/100", m["status"])
             for m in rd["key_metrics"]],
            col_ratios=[2.5, 0.8, 1.2],
        )
    for item in rd.get("action_items", []):
        pdf._para(f"  -> {item}", size=9)
    for s in rd.get("strengths", []):
        pdf._para(f"  + {s}", size=9)

    pdf.output(output_path)


# ── TXT builder ───────────────────────────────────────────────────────────────

def build_txt(lp_data: dict, td_data: dict, analytics: dict, ai: dict,
              code: str, title: str, semester, output_path: str):
    lp   = lp_data or {}
    td   = td_data or {}
    meta = {**(td.get("meta") or {}), **(lp.get("meta") or {})}
    sem_label = f"  |  Semester {semester}" if semester else ""
    cos = analytics["cos"]

    lines = [
        f"{code}  -  {title}",
        "Module 3 Comprehensive Report  -  Course Delivery Management",
        f"Generated: {datetime.date.today().strftime('%d %B %Y')}{sem_label}",
        "=" * 70,
        "",
        "1. COURSE DELIVERY OVERVIEW",
        "-" * 40,
        f"  Academic Year : {meta.get('academic_year', '-')}",
        f"  Semester      : {meta.get('semester', '-')}",
        f"  Regulation    : {meta.get('regulation', '-')}",
        f"  Program       : {meta.get('program', '-')}",
        f"  Course Type   : {meta.get('course_type', '-')}",
        f"  Credits / LTP : {meta.get('credits', '-')} / {meta.get('ltp', '-')}",
        "",
        "2. COURSE OUTCOMES & BLOOM'S DISTRIBUTION",
        "-" * 40,
    ]
    for c in cos:
        lines.append(f"  CO{c['num']}: {c['statement']}  [{c.get('bloom','-')} / {c.get('knowledge','-')}]")
    lines += [
        "",
        "  Bloom's Distribution:",
    ]
    for lv, n in sorted(analytics["bloom_dist"].items(), key=lambda x: x[1], reverse=True):
        pct = round(n / analytics["n_cos"] * 100) if analytics["n_cos"] else 0
        lines.append(f"    {lv:<14} {n}  ({pct}%)")

    lines += [
        "",
        "3. CO-PO-SDG ALIGNMENT",
        "-" * 40,
        f"  Active POs: {', '.join(analytics['active_pos']) or 'None'}",
        "",
        "4. LESSON PLAN COVERAGE",
        "-" * 40,
        f"  Sessions planned       : {analytics['n_sessions']}",
        f"  CO coverage in sessions: {analytics['lp_co_coverage_pct']}% ({len(analytics['cos_in_sessions'])} of {analytics['n_cos']} COs)",
        "",
        "7. CO-SESSION COVERAGE MATRIX",
        "-" * 40,
    ]
    co_tally = analytics["co_session_tally"]
    for c in cos:
        key = f"CO{c['num']}"
        n   = co_tally.get(key, 0)
        pct = round(n / analytics["n_sessions"] * 100) if analytics["n_sessions"] else 0
        status = "Covered" if n > 0 else "NOT COVERED"
        lines.append(f"  {key}: {n} session(s) ({pct}%)  [{status}]")

    lines += [
        "",
        "8. TEACHING METHOD DISTRIBUTION",
        "-" * 40,
    ]
    total_sess = sum(analytics["method_dist"].values()) or 1
    for m, n in sorted(analytics["method_dist"].items(), key=lambda x: x[1], reverse=True):
        lines.append(f"  {m:<24} {n} ({round(n/total_sess*100)}%)")

    lines += [
        "",
        "6. TEACHING DIARY SUMMARY",
        "-" * 40,
        f"  Theory sessions  : {analytics['n_theory']}",
        f"  Tutorial sessions: {analytics['n_tutorial']}",
        f"  Lab sessions     : {analytics['n_lab']}",
        f"  Total            : {analytics['n_td_total']}",
    ]
    final = analytics["final_summary"]
    if final:
        for k, v in final.items():
            lines.append(f"  {k.replace('_',' ').title():<24}: {v}")

    # AI sections
    obe = ai.get("obe_compliance", {})
    lines += [
        "",
        "10. NBA/NAAC COMPLIANCE",
        "-" * 40,
        f"  Compliance Level : {obe.get('compliance_level', '-')}",
        f"  OBE Score        : {obe.get('overall_obe_score', '-')}/100",
        f"  CO-PO Alignment  : {obe.get('co_po_alignment', '-')}",
        f"  Delivery Coverage: {obe.get('delivery_coverage', '-')}",
    ]
    for s in obe.get("strengths", []):
        lines.append(f"  + {s}")
    for r in obe.get("recommendations", []):
        lines.append(f"  -> {r}")

    rd = ai.get("accreditation_readiness", {})
    lines += [
        "",
        "11. COURSE DELIVERY READINESS DASHBOARD",
        "-" * 40,
        f"  Overall Score  : {rd.get('overall_score', '-')}/100",
        f"  Readiness Level: {rd.get('readiness_level', '-')}",
        "",
        "  Key Metrics:",
    ]
    for m in rd.get("key_metrics", []):
        lines.append(f"    {m['metric']:<24} {m['score']}/100  [{m['status']}]")
    if rd.get("action_items"):
        lines.append("  Action Items:")
        for item in rd["action_items"]:
            lines.append(f"    -> {item}")
    if rd.get("strengths"):
        lines.append("  Strengths:")
        for s in rd["strengths"]:
            lines.append(f"    + {s}")

    lines.append("")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
