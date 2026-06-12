import os
import re
import json
import sys
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

_SYSTEM = """\
You are an expert in Outcome-Based Education (OBE), NBA accreditation, NAAC, NEP 2020,
and SDG integration for engineering courses. You specialize in creating comprehensive
Teaching Diaries that comply with NBA/NAAC accreditation standards.

Generate a complete Teaching Diary as a single JSON object. The JSON must be valid,
complete, and follow the exact structure specified. Keep ALL string field values
SHORT and CONCISE (3-6 words max per field). No markdown fences, no explanation —
output ONLY the JSON object.
"""


def _build_prompt(course_code: str, course_title: str, course_text: str,
                  num_cos: int, meta: dict, existing_cos: list | None = None,
                  is_lab: bool = False) -> str:
    ltp = meta.get("ltp", "3-0-2")
    try:
        parts = ltp.split("-")
        n_theory = int(parts[0]) * 15
        n_tutorial = int(parts[1]) * 15 if len(parts) > 1 and parts[1] != "0" else 0
        n_lab = (int(parts[2]) // 2) * 15 if len(parts) > 2 and parts[2] != "0" else 0
    except Exception:
        n_theory, n_tutorial, n_lab = 45, 15, 15

    n_theory = min(n_theory, 20)
    n_tutorial = min(n_tutorial, 10)
    n_lab = min(n_lab, 10) if is_lab else 0

    if existing_cos:
        num_cos = len(existing_cos)
        co_block = "\n".join(f"CO{c['num']}: {c['statement']}" for c in existing_cos)
        co_section = f"""PRE-DEFINED COURSE OUTCOMES (use EXACTLY as given):
{co_block}"""
        co_req = f"- Use the {num_cos} pre-defined COs above EXACTLY"
    else:
        co_section = ""
        co_req = f"- Generate exactly {num_cos} Course Outcomes (CO1–CO{num_cos})"
    cos_addressed = ", ".join(f"CO{i}" for i in range(1, num_cos + 1))

    return f"""Generate a complete Teaching Diary for the following engineering course.

COURSE METADATA:
- Course Code: {course_code}
- Course Title: {course_title}
- Academic Year: {meta.get('academic_year', '2026-2027')}
- Semester: {meta.get('semester', 'III Semester')}
- Regulation: {meta.get('regulation', 'R26')}
- Program: {meta.get('program', 'B.Tech – CSE')}
- Course Type: {meta.get('course_type', 'Theory + Lab')}
- Credits: {meta.get('credits', '4')}
- L-T-P: {ltp}

SYLLABUS CONTENT:
{course_text[:4000]}

{co_section}

REQUIREMENTS:
{co_req}
- Generate exactly {n_theory} representative theory session entries (spread evenly across all units)
- Generate exactly {n_tutorial} tutorial entries (0 if no tutorial hours)
- Generate exactly {n_lab} lab session entries (0 if no lab hours)
- CO-PO scores: 3=Major, 2=Moderate, 1=Minor, 0=None
- SDG goals: SDG4 (Quality Education), SDG9 (Industry Innovation), SDG13 (Climate Action)
- Bloom levels: Remember, Understand, Apply, Analyze, Evaluate, Create
- Keep ALL string values SHORT (3-6 words)

OUTPUT — return ONLY this JSON (no markdown, no explanation):
{{
  "cos": [
    {{"num": 1, "statement": "<CO statement>", "bloom": "<Bloom level>", "knowledge": "<type>"}}
  ],
  "integration_summary": {{
    "cos_addressed": "{cos_addressed}",
    "pos_addressed": "PO1, PO2, PO3, PO4, PO5",
    "psos": "PSO1, PSO2",
    "cep": "<complex engineering problem type>",
    "sdg_goals": "SDG4, SDG9",
    "ai_ict_tools": "<tool names>",
    "industry_relevance": "<industry context>",
    "sustainability_relevance": "<sustainability aspect>",
    "innovation_focus": "<innovation type>"
  }},
  "co_mapping": [
    {{"co": "CO1", "po": "PO1, PO2", "ga": "<GA name>", "sdg": "SDG4", "sustainability": "<focus>"}}
  ],
  "theory_sessions": [
    {{
      "no": 1, "unit": "Unit 1", "topic": "<topic>",
      "co": "CO1", "po": "PO1", "ga": "<GA>", "sdg": "SDG4",
      "bloom": "Remember", "method": "Lecture", "ai_tool": "ChatGPT",
      "industry": "<application>", "sustainability": "<focus>", "assessment": "Quiz"
    }}
  ],
  "tutorials": [
    {{
      "no": 1, "topic": "<topic>", "co": "CO1", "po": "PO2",
      "cep": "<CEP type>", "sdg": "SDG4", "activity_type": "Problem Solving",
      "ai_tool": "ChatGPT", "assessment": "Viva", "evidence": "Written work"
    }}
  ],
  "lab_sessions": [
    {{
      "no": 1, "title": "<experiment title>", "co": "CO2", "po": "PO3",
      "cep": "<CEP>", "sdg": "SDG4", "software": "Python",
      "ai_integration": "<AI use>", "industry": "<application>",
      "sustainability": "<focus>", "rubric": "<criteria>", "evidence": "Code repo"
    }}
  ],
  "ai_learning": [
    {{"session": "<session/week>", "ai_tool": "<tool>", "purpose": "<purpose>", "outcome": "<outcome>"}}
  ],
  "industry_esg": [
    {{"topic": "<topic>", "industry": "<application>", "esg": "<ESG relevance>"}}
  ],
  "co_attainment": [
    {{"co": "CO1", "target": 60}}
  ],
  "nba_alignment": [
    {{"framework": "NBA", "evidence": "<evidence>"}},
    {{"framework": "NAAC", "evidence": "<evidence>"}},
    {{"framework": "NEP 2020", "evidence": "<evidence>"}},
    {{"framework": "SDGs", "evidence": "<evidence>"}},
    {{"framework": "THE Impact", "evidence": "<evidence>"}}
  ],
  "final_summary": {{
    "syllabus_coverage": "100%",
    "sdg_activities": "5",
    "ai_activities": "3",
    "industry_activities": "2",
    "innovations": "2"
  }}
}}"""


def _extract_json(raw: str) -> dict | None:
    cleaned = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    start = cleaned.find('{')
    end = cleaned.rfind('}')
    if start != -1 and end > start:
        try:
            return json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError:
            pass
    return None


def generate_teaching_diary(client, course_code: str, course_title: str,
                             course_text: str, meta: dict,
                             num_cos: int = 5,
                             existing_cos: list | None = None,
                             is_lab: bool = False) -> dict:
    if existing_cos:
        num_cos = len(existing_cos)
    prompt = _build_prompt(course_code, course_title, course_text, num_cos,
                           meta, existing_cos=existing_cos, is_lab=is_lab)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        system=_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text
    data = _extract_json(raw)
    if not data:
        raise ValueError(
            f"AI returned invalid JSON. Raw (first 300 chars): {raw[:300]}"
        )
    return data


def generate_teaching_diary_stream(client, course_code: str, course_title: str,
                                    course_text: str, meta: dict,
                                    num_cos: int = 5,
                                    existing_cos: list | None = None,
                                    is_lab: bool = False):
    """Yield SSE-style progress dicts. Last item has key 'done' and 'data'.

    Streams the model response so the SSE connection keeps receiving bytes.
    A single blocking call sends nothing for tens of seconds while the large
    JSON is generated, which a hosting proxy treats as an idle connection and
    drops — surfacing client-side as a generic "Connection error". The
    periodic 'heartbeat' items keep the connection alive (the route emits them
    as SSE comments that the browser ignores).
    """
    if existing_cos:
        num_cos = len(existing_cos)
        yield {"progress": f"Using {len(existing_cos)} uploaded COs — generating teaching diary…"}
    else:
        yield {"progress": "Generating COs and teaching diary sections…"}

    prompt = _build_prompt(course_code, course_title, course_text, num_cos,
                           meta, existing_cos=existing_cos, is_lab=is_lab)
    raw = ""
    try:
        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        ) as stream:
            for chunk in stream.text_stream:
                raw += chunk
                yield {"heartbeat": True}
    except Exception as e:
        yield {"error": str(e)}
        return

    data = _extract_json(raw)
    if not data:
        yield {"error": f"AI returned invalid JSON. Raw (first 300 chars): {raw[:300]}"}
        return
    yield {"done": True, "data": data}


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11 if level == 1 else 10)
        run.bold = True
    return p


def _cell_bg(cell, rgb):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r, g, b = rgb
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, size_pt=8, center=False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = str(text)
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.bold = bold


def _make_table(doc, headers, rows_data, header_bg=(180, 198, 231), col_widths=None):
    from docx.shared import Inches, Pt
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    tbl.style = 'Table Grid'

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _cell_bg(hdr[i], header_bg)
        _set_cell_text(hdr[i], h, bold=True, size_pt=8)

    for ri, row in enumerate(rows_data, 1):
        cells = tbl.rows[ri].cells
        for ci, val in enumerate(row):
            _set_cell_text(cells[ci], val, size_pt=8)

    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)
    return tbl


def _make_2col_table(doc, rows, header_bg=(180, 198, 231)):
    from docx.shared import Inches
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for ri, (label, val) in enumerate(rows):
        cells = tbl.rows[ri].cells
        _cell_bg(cells[0], header_bg)
        _set_cell_text(cells[0], label, bold=True, size_pt=9)
        _set_cell_text(cells[1], val, size_pt=9)
        cells[0].width = Inches(2.5)
        cells[1].width = Inches(6.5)
    return tbl


def _set_landscape(doc):
    from docx.shared import Inches
    for section in doc.sections:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)


def build_docx(data: dict, meta: dict, course_code: str, course_title: str,
               output_path: str):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_landscape(doc)

    # ── Title ──────────────────────────────────────────────────────────────────
    def center_bold(text, size):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        return p

    center_bold("TEACHING DIARY (THEORY)", 14)
    center_bold(f"{course_code} — {course_title}", 12)
    dept = meta.get("department", "Department of Computer Science and Engineering")
    center_bold(dept, 11)
    doc.add_paragraph()

    # ── Course Header Table ────────────────────────────────────────────────────
    _heading(doc, "COURSE DETAILS", level=1)
    header_rows = [
        ("Item",                     "Details"),
        ("Department",               meta.get("department", "")),
        ("Program",                  meta.get("program", "")),
        ("Semester",                 meta.get("semester", "")),
        ("Section",                  meta.get("section", "")),
        ("Academic Year",            meta.get("academic_year", "")),
        ("Course Code & Name",       f"{course_code} — {course_title}"),
        ("Credits",                  str(meta.get("credits", "4"))),
        ("L-T-P",                    meta.get("ltp", "3-0-2")),
        ("Regulation",               meta.get("regulation", "R26")),
        ("Faculty Name",             meta.get("faculty_name", "")),
        ("Total Contact Hours",      str(meta.get("total_hours", "45"))),
        ("Course Coordinator",       meta.get("coordinator", "")),
        ("HoD",                      meta.get("hod", "")),
    ]
    _make_2col_table(doc, header_rows[1:])
    doc.add_paragraph()

    # ── 1. Course Integration Summary ─────────────────────────────────────────
    _heading(doc, "1. COURSE INTEGRATION SUMMARY")
    si = data.get("integration_summary", {})
    summary_rows = [
        ("Course Outcomes (COs) Addressed",      si.get("cos_addressed", "")),
        ("Program Outcomes (POs) Addressed",     si.get("pos_addressed", "")),
        ("Program Specific Outcomes (PSOs)",     si.get("psos", "")),
        ("Complex Engineering Problems (CEP)",   si.get("cep", "")),
        ("SDG Goals Integrated",                 si.get("sdg_goals", "")),
        ("AI/ICT Tools Used",                    si.get("ai_ict_tools", "")),
        ("Industry Relevance",                   si.get("industry_relevance", "")),
        ("Sustainability Relevance",             si.get("sustainability_relevance", "")),
        ("Innovation Focus",                     si.get("innovation_focus", "")),
    ]
    _make_2col_table(doc, summary_rows)
    doc.add_paragraph()

    # ── 2. CO–PO–GA–SDG Mapping ───────────────────────────────────────────────
    _heading(doc, "2. CO–PO–SDG–GA MAPPING SUMMARY")
    co_map = data.get("co_mapping", [])
    _make_table(doc,
        ["CO", "PO", "GA", "SDG", "Sustainability Focus"],
        [(r.get("co",""), r.get("po",""), r.get("ga",""),
          r.get("sdg",""), r.get("sustainability","")) for r in co_map],
        col_widths=[0.6, 1.5, 1.5, 0.7, 2.5]
    )
    doc.add_paragraph()

    # ── 3. Theory Teaching Delivery Plan ──────────────────────────────────────
    _heading(doc, "3. THEORY TEACHING DELIVERY PLAN")
    sessions = data.get("theory_sessions", [])
    th_headers = [
        "Lesson No", "Date", "Time", "Unit", "Topic Covered", "CO", "PO",
        "GA Addressed", "SDG Goal", "Bloom's Level", "Teaching Method",
        "AI Tool Used", "Industry Relevance", "Sustainability Context",
        "Assessment Method", "Evidence Generated", "Student Engagement",
        "Faculty Reflection", "Initials"
    ]
    th_rows = []
    for s in sessions:
        th_rows.append((
            str(s.get("no", "")), "", "", s.get("unit", ""),
            s.get("topic", ""), s.get("co", ""), s.get("po", ""),
            s.get("ga", ""), s.get("sdg", ""), s.get("bloom", ""),
            s.get("method", ""), s.get("ai_tool", ""), s.get("industry", ""),
            s.get("sustainability", ""), s.get("assessment", ""), "", "", "", ""
        ))
    th_widths = [0.45, 0.5, 0.45, 0.5, 1.1, 0.35, 0.35,
                 0.6, 0.45, 0.55, 0.65, 0.55, 0.65, 0.65,
                 0.6, 0.6, 0.55, 0.65, 0.4]
    _make_table(doc, th_headers, th_rows, col_widths=th_widths)
    doc.add_paragraph()

    # ── 4. Tutorial / Activity Diary ──────────────────────────────────────────
    tutorials = data.get("tutorials", [])
    if tutorials:
        _heading(doc, "4. TUTORIAL / ACTIVITY DIARY")
        tut_headers = [
            "Tutorial No", "Topic", "CO", "PO", "CEP",
            "SDG", "Activity Type", "AI Tool Used", "Assessment", "Evidence"
        ]
        tut_rows = []
        for t in tutorials:
            tut_rows.append((
                str(t.get("no", "")), t.get("topic", ""), t.get("co", ""),
                t.get("po", ""), t.get("cep", ""), t.get("sdg", ""),
                t.get("activity_type", ""), t.get("ai_tool", ""),
                t.get("assessment", ""), t.get("evidence", "")
            ))
        _make_table(doc, tut_headers, tut_rows,
                    col_widths=[0.6, 1.5, 0.5, 0.5, 0.8, 0.5, 0.9, 0.8, 0.8, 0.8])
        doc.add_paragraph()

    # ── 5. Lab Teaching Diary ─────────────────────────────────────────────────
    lab_sessions = data.get("lab_sessions", [])
    if lab_sessions:
        _heading(doc, "5. LAB TEACHING DIARY")
        lab_headers = [
            "Lab No", "Experiment Title", "CO", "PO", "CEP", "SDG",
            "Software/Tool", "AI Integration", "Industry Application",
            "Sustainability Context", "Assessment Rubric", "Evidence Generated"
        ]
        lab_rows = []
        for l in lab_sessions:
            lab_rows.append((
                str(l.get("no", "")), l.get("title", ""), l.get("co", ""),
                l.get("po", ""), l.get("cep", ""), l.get("sdg", ""),
                l.get("software", ""), l.get("ai_integration", ""),
                l.get("industry", ""), l.get("sustainability", ""),
                l.get("rubric", ""), l.get("evidence", "")
            ))
        _make_table(doc, lab_headers, lab_rows,
                    col_widths=[0.4, 1.2, 0.4, 0.4, 0.7, 0.4,
                                0.8, 0.85, 0.85, 0.85, 0.85, 0.85])
        doc.add_paragraph()

    # ── 6. Student Learning Analytics ─────────────────────────────────────────
    _heading(doc, "6. STUDENT LEARNING ANALYTICS")
    analytics_rows = [
        ("Attendance Trend",       "", ""),
        ("Quiz Performance",       "", ""),
        ("Coding Performance",     "", ""),
        ("AI Tool Utilization",    "", ""),
        ("Innovation Participation","", ""),
        ("SDG Awareness",          "", ""),
        ("Team Collaboration",     "", ""),
    ]
    _make_table(doc,
        ["Indicator", "Observation", "Action Taken"],
        analytics_rows,
        col_widths=[2.0, 3.5, 3.5]
    )
    doc.add_paragraph()

    # ── 7. AI & Digital Learning Tracker ──────────────────────────────────────
    _heading(doc, "7. AI & DIGITAL LEARNING TRACKER")
    ai_rows = [(r.get("session",""), r.get("ai_tool",""),
                r.get("purpose",""), r.get("outcome",""))
               for r in data.get("ai_learning", [])]
    _make_table(doc,
        ["Session", "AI Tool Used", "Purpose", "Student Outcome"],
        ai_rows,
        col_widths=[1.5, 2.0, 2.5, 2.5]
    )
    doc.add_paragraph()

    # ── 8. Industry 4.0 / ESG Integration ─────────────────────────────────────
    _heading(doc, "8. INDUSTRY 4.0 / ESG INTEGRATION")
    esg_rows = [(r.get("topic",""), r.get("industry",""), r.get("esg",""))
                for r in data.get("industry_esg", [])]
    _make_table(doc,
        ["Topic", "Industry Application", "ESG/Sustainability Relevance"],
        esg_rows,
        col_widths=[2.5, 3.5, 3.5]
    )
    doc.add_paragraph()

    # ── 9. Continuous Quality Improvement (CQI) ───────────────────────────────
    _heading(doc, "9. CONTINUOUS QUALITY IMPROVEMENT (CQI)")
    _make_table(doc,
        ["Observation", "Gap Identified", "Corrective Action", "Expected Improvement"],
        [("", "", "", ""), ("", "", "", ""), ("", "", "", "")],
        col_widths=[2.25, 2.25, 2.25, 2.25]
    )
    doc.add_paragraph()

    # ── 10. Course Outcome Attainment Analysis ────────────────────────────────
    _heading(doc, "10. COURSE OUTCOME ATTAINMENT ANALYSIS")
    co_att = data.get("co_attainment", [])
    att_rows = [(r.get("co",""), str(r.get("target",60))+"%", "", "", "")
                for r in co_att]
    if not att_rows:
        cos = data.get("cos", [])
        att_rows = [(f"CO{c['num']}", "60%", "", "", "") for c in cos]
    _make_table(doc,
        ["CO", "Target (%)", "Attainment (%)", "Gap", "Action Plan"],
        att_rows,
        col_widths=[0.6, 1.2, 1.2, 1.2, 4.8]
    )
    doc.add_paragraph()

    # ── 11. Student Reflections ───────────────────────────────────────────────
    _heading(doc, "11. STUDENT REFLECTIONS")
    _make_table(doc,
        ["Student Reflection Aspect", "Observation"],
        [
            ("Real-world learning gained", ""),
            ("Sustainability understanding", ""),
            ("AI tool effectiveness", ""),
            ("Innovation exposure", ""),
            ("Industry relevance", ""),
        ],
        col_widths=[3.5, 6.0]
    )
    doc.add_paragraph()

    # ── 12. Faculty Reflective Journal ────────────────────────────────────────
    _heading(doc, "12. FACULTY REFLECTIVE JOURNAL")
    _make_table(doc,
        ["Reflection Area", "Observation"],
        [
            ("Student engagement", ""),
            ("Teaching effectiveness", ""),
            ("AI integration effectiveness", ""),
            ("Sustainability integration", ""),
            ("Improvement needed", ""),
        ],
        col_widths=[3.5, 6.0]
    )
    doc.add_paragraph()

    # ── 13. Digital Evidence Repository ──────────────────────────────────────
    _heading(doc, "13. DIGITAL EVIDENCE REPOSITORY")
    _make_table(doc,
        ["Evidence Type", "Repository Link / Storage"],
        [
            ("Assignments", ""),
            ("Quiz Reports", ""),
            ("Attendance", ""),
            ("AI Activity Reports", ""),
            ("GitHub Repositories", ""),
            ("SDG Projects", ""),
            ("Videos / Photos", ""),
            ("Rubrics", ""),
        ],
        col_widths=[3.0, 6.5]
    )
    doc.add_paragraph()

    # ── 14. NBA / NAAC / NEP 2020 / THE Impact Alignment ─────────────────────
    _heading(doc, "14. NBA / NAAC / NEP 2020 / THE IMPACT ALIGNMENT")
    nba = data.get("nba_alignment", [
        {"framework": "NBA",         "evidence": "CO-PO attainment, CEP"},
        {"framework": "NAAC",        "evidence": "ICT-enabled teaching, innovation"},
        {"framework": "NEP 2020",    "evidence": "Skill-based learning"},
        {"framework": "SDGs",        "evidence": "Sustainability integration"},
        {"framework": "THE Impact",  "evidence": "Societal impact evidence"},
    ])
    _make_table(doc,
        ["Framework", "Alignment Evidence"],
        [(r.get("framework",""), r.get("evidence","")) for r in nba],
        col_widths=[2.0, 7.5]
    )
    doc.add_paragraph()

    # ── 15. Final Semester Summary ────────────────────────────────────────────
    _heading(doc, "15. FINAL SEMESTER SUMMARY")
    fs = data.get("final_summary", {})
    _make_table(doc,
        ["Parameter", "Status"],
        [
            ("Syllabus Coverage (%)",          fs.get("syllabus_coverage", "")),
            ("CO Attainment Level",            ""),
            ("Student Participation",          ""),
            ("SDG Activities Conducted",       fs.get("sdg_activities", "")),
            ("AI-Based Activities Conducted",  fs.get("ai_activities", "")),
            ("Industry Activities Conducted",  fs.get("industry_activities", "")),
            ("Innovations Implemented",        fs.get("innovations", "")),
        ],
        col_widths=[4.0, 5.5]
    )
    doc.add_paragraph()

    # ── 16. Signatures ────────────────────────────────────────────────────────
    _heading(doc, "16. SIGNATURES")
    _make_table(doc,
        ["Role", "Signature", "Date"],
        [
            ("Faculty", "", ""),
            ("Course Coordinator", "", ""),
            ("Program Coordinator", "", ""),
            ("HoD", "", ""),
            ("IQAC Coordinator", "", ""),
        ],
        col_widths=[3.0, 4.0, 2.5]
    )

    doc.save(output_path)


# ── PDF builder ────────────────────────────────────────────────────────────────

def build_pdf(data: dict, meta: dict, course_code: str, course_title: str,
              output_path: str):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle('T', parent=styles['Title'],   fontSize=13, alignment=TA_CENTER, spaceAfter=4)
    sub_s   = ParagraphStyle('S', parent=styles['Normal'],  fontSize=10, alignment=TA_CENTER, spaceAfter=4,
                              fontName='Helvetica-Bold')
    head_s  = ParagraphStyle('H', parent=styles['Heading2'], fontSize=10, spaceAfter=4, spaceBefore=8,
                              fontName='Helvetica-Bold')
    body_s  = ParagraphStyle('B', parent=styles['Normal'],  fontSize=8,  spaceAfter=2)
    cell_s  = ParagraphStyle('C', parent=styles['Normal'],  fontSize=7,  leading=9)

    HDR  = colors.Color(180/255, 198/255, 231/255)
    ALT  = colors.Color(242/255, 242/255, 242/255)

    def _ts(nrows):
        return TableStyle([
            ('BACKGROUND',  (0,0), (-1,0), HDR),
            ('FONTNAME',    (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',    (0,0), (-1,-1), 7),
            ('GRID',        (0,0), (-1,-1), 0.5, colors.grey),
            ('VALIGN',      (0,0), (-1,-1), 'TOP'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, ALT]),
            ('LEFTPADDING',  (0,0), (-1,-1), 3),
            ('RIGHTPADDING', (0,0), (-1,-1), 3),
            ('TOPPADDING',   (0,0), (-1,-1), 2),
            ('BOTTOMPADDING',(0,0), (-1,-1), 2),
        ])

    def _w(text):
        return Paragraph(str(text), cell_s)

    PAGE = landscape(A4)
    doc = SimpleDocTemplate(output_path, pagesize=PAGE,
                             leftMargin=0.75*inch, rightMargin=0.75*inch,
                             topMargin=0.75*inch, bottomMargin=0.75*inch)
    W = PAGE[0] - 1.5 * inch

    story = []
    dept = meta.get("department", "Department of Computer Science and Engineering")
    story += [
        Paragraph("TEACHING DIARY (THEORY)", title_s),
        Paragraph(f"{course_code} — {course_title}", sub_s),
        Paragraph(dept, sub_s),
        Spacer(1, 6),
    ]

    # Course Details
    story.append(Paragraph("COURSE DETAILS", head_s))
    hdr_data = [
        [_w("Department"),           _w(meta.get("department",""))],
        [_w("Program"),              _w(meta.get("program",""))],
        [_w("Semester"),             _w(meta.get("semester",""))],
        [_w("Section"),              _w(meta.get("section",""))],
        [_w("Academic Year"),        _w(meta.get("academic_year",""))],
        [_w("Course Code & Name"),   _w(f"{course_code} — {course_title}")],
        [_w("Credits"),              _w(str(meta.get("credits","4")))],
        [_w("L-T-P"),                _w(meta.get("ltp","3-0-2"))],
        [_w("Regulation"),           _w(meta.get("regulation","R26"))],
        [_w("Faculty Name"),         _w(meta.get("faculty_name",""))],
        [_w("Total Contact Hours"),  _w(str(meta.get("total_hours","45")))],
        [_w("Course Coordinator"),   _w(meta.get("coordinator",""))],
        [_w("HoD"),                  _w(meta.get("hod",""))],
    ]
    t = Table(hdr_data, colWidths=[2.0*inch, W - 2.0*inch])
    t.setStyle(TableStyle([
        ('FONTNAME',    (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,0), (-1,-1), 8),
        ('GRID',        (0,0), (-1,-1), 0.5, colors.grey),
        ('VALIGN',      (0,0), (-1,-1), 'TOP'),
        ('BACKGROUND',  (0,0), (0,-1), HDR),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('RIGHTPADDING',(0,0), (-1,-1), 3),
        ('TOPPADDING',  (0,0), (-1,-1), 2),
        ('BOTTOMPADDING',(0,0),(-1,-1), 2),
    ]))
    story += [t, Spacer(1,6)]

    # 1. Integration Summary
    story.append(Paragraph("1. COURSE INTEGRATION SUMMARY", head_s))
    si = data.get("integration_summary", {})
    sum_data = [
        [_w("COs Addressed"),             _w(si.get("cos_addressed",""))],
        [_w("POs Addressed"),             _w(si.get("pos_addressed",""))],
        [_w("PSOs"),                       _w(si.get("psos",""))],
        [_w("Complex Engineering Problems"),_w(si.get("cep",""))],
        [_w("SDG Goals"),                  _w(si.get("sdg_goals",""))],
        [_w("AI/ICT Tools"),               _w(si.get("ai_ict_tools",""))],
        [_w("Industry Relevance"),         _w(si.get("industry_relevance",""))],
        [_w("Sustainability Relevance"),   _w(si.get("sustainability_relevance",""))],
        [_w("Innovation Focus"),           _w(si.get("innovation_focus",""))],
    ]
    t = Table(sum_data, colWidths=[2.5*inch, W-2.5*inch])
    t.setStyle(TableStyle([
        ('FONTNAME',    (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,0), (-1,-1), 8),
        ('GRID',        (0,0), (-1,-1), 0.5, colors.grey),
        ('BACKGROUND',  (0,0), (0,-1), HDR),
        ('VALIGN',      (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('TOPPADDING',  (0,0), (-1,-1), 2),
        ('BOTTOMPADDING',(0,0),(-1,-1), 2),
    ]))
    story += [t, Spacer(1,6)]

    # 2. CO Mapping
    story.append(Paragraph("2. CO–PO–SDG–GA MAPPING SUMMARY", head_s))
    co_map = data.get("co_mapping", [])
    map_rows = [[_w("CO"), _w("PO"), _w("GA"), _w("SDG"), _w("Sustainability Focus")]]
    for r in co_map:
        map_rows.append([_w(r.get("co","")), _w(r.get("po","")), _w(r.get("ga","")),
                          _w(r.get("sdg","")), _w(r.get("sustainability",""))])
    cw5 = W/5
    t = Table(map_rows, colWidths=[cw5]*5)
    t.setStyle(_ts(len(map_rows)))
    story += [t, Spacer(1,6)]

    # 3. Theory Delivery Plan
    story.append(Paragraph("3. THEORY TEACHING DELIVERY PLAN", head_s))
    sessions = data.get("theory_sessions", [])
    th_hdr = ["#", "Unit", "Topic", "CO", "PO", "GA", "SDG",
              "Bloom's", "Method", "AI Tool", "Industry", "Sustainability", "Assessment"]
    th_rows = [th_hdr]
    for s in sessions:
        th_rows.append([
            _w(str(s.get("no",""))), _w(s.get("unit","")), _w(s.get("topic","")),
            _w(s.get("co","")), _w(s.get("po","")), _w(s.get("ga","")),
            _w(s.get("sdg","")), _w(s.get("bloom","")), _w(s.get("method","")),
            _w(s.get("ai_tool","")), _w(s.get("industry","")),
            _w(s.get("sustainability","")), _w(s.get("assessment",""))
        ])
    cw13 = [0.3, 0.55, 1.2, 0.4, 0.5, 0.65, 0.45, 0.55, 0.7, 0.6, 0.65, 0.65, 0.65]
    cw13 = [c * inch for c in cw13]
    t = Table(th_rows, colWidths=cw13)
    t.setStyle(_ts(len(th_rows)))
    story += [t, Spacer(1,6)]

    # 4. Tutorial Diary
    tutorials = data.get("tutorials", [])
    if tutorials:
        story.append(Paragraph("4. TUTORIAL / ACTIVITY DIARY", head_s))
        tut_hdr = ["#", "Topic", "CO", "PO", "CEP", "SDG", "Activity Type", "AI Tool", "Assessment", "Evidence"]
        tut_rows = [tut_hdr]
        for t2 in tutorials:
            tut_rows.append([
                _w(str(t2.get("no",""))), _w(t2.get("topic","")), _w(t2.get("co","")),
                _w(t2.get("po","")), _w(t2.get("cep","")), _w(t2.get("sdg","")),
                _w(t2.get("activity_type","")), _w(t2.get("ai_tool","")),
                _w(t2.get("assessment","")), _w(t2.get("evidence",""))
            ])
        cw_tut = [0.3, 1.2, 0.45, 0.45, 0.75, 0.45, 0.9, 0.75, 0.75, 0.75]
        tbl = Table(tut_rows, colWidths=[c*inch for c in cw_tut])
        tbl.setStyle(_ts(len(tut_rows)))
        story += [tbl, Spacer(1,6)]

    # 5. Lab Diary
    lab_sessions = data.get("lab_sessions", [])
    if lab_sessions:
        story.append(Paragraph("5. LAB TEACHING DIARY", head_s))
        lab_hdr = ["#", "Experiment", "CO", "PO", "CEP", "SDG",
                   "Software", "AI Integration", "Industry", "Sustainability", "Rubric", "Evidence"]
        lab_rows = [lab_hdr]
        for l in lab_sessions:
            lab_rows.append([
                _w(str(l.get("no",""))), _w(l.get("title","")), _w(l.get("co","")),
                _w(l.get("po","")), _w(l.get("cep","")), _w(l.get("sdg","")),
                _w(l.get("software","")), _w(l.get("ai_integration","")),
                _w(l.get("industry","")), _w(l.get("sustainability","")),
                _w(l.get("rubric","")), _w(l.get("evidence",""))
            ])
        cw_lab = [0.3, 1.2, 0.4, 0.4, 0.65, 0.4, 0.7, 0.8, 0.8, 0.8, 0.7, 0.75]
        tbl = Table(lab_rows, colWidths=[c*inch for c in cw_lab])
        tbl.setStyle(_ts(len(lab_rows)))
        story += [tbl, Spacer(1,6)]

    # 6–16: remaining sections in compact form
    story.append(Paragraph("6. STUDENT LEARNING ANALYTICS", head_s))
    an_data = [
        [_w("Indicator"), _w("Observation"), _w("Action Taken")],
        *[[_w(ind), _w(""), _w("")] for ind in [
            "Attendance Trend","Quiz Performance","Coding Performance",
            "AI Tool Utilization","Innovation Participation","SDG Awareness","Team Collaboration"
        ]]
    ]
    t = Table(an_data, colWidths=[2.0*inch, 3.5*inch, W-5.5*inch])
    t.setStyle(_ts(len(an_data)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("7. AI & DIGITAL LEARNING TRACKER", head_s))
    ai_rows = [[_w("Session"), _w("AI Tool"), _w("Purpose"), _w("Student Outcome")]]
    for r in data.get("ai_learning", []):
        ai_rows.append([_w(r.get("session","")), _w(r.get("ai_tool","")),
                        _w(r.get("purpose","")), _w(r.get("outcome",""))])
    cw4 = W/4
    t = Table(ai_rows, colWidths=[cw4]*4)
    t.setStyle(_ts(len(ai_rows)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("8. INDUSTRY 4.0 / ESG INTEGRATION", head_s))
    esg_rows = [[_w("Topic"), _w("Industry Application"), _w("ESG/Sustainability Relevance")]]
    for r in data.get("industry_esg", []):
        esg_rows.append([_w(r.get("topic","")), _w(r.get("industry","")), _w(r.get("esg",""))])
    cw3 = W/3
    t = Table(esg_rows, colWidths=[cw3]*3)
    t.setStyle(_ts(len(esg_rows)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("9. CONTINUOUS QUALITY IMPROVEMENT (CQI)", head_s))
    cqi = [[_w("Observation"), _w("Gap Identified"), _w("Corrective Action"), _w("Expected Improvement")]]
    for _ in range(3):
        cqi.append([_w("")]*4)
    t = Table(cqi, colWidths=[W/4]*4)
    t.setStyle(_ts(len(cqi)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("10. COURSE OUTCOME ATTAINMENT ANALYSIS", head_s))
    co_att = data.get("co_attainment", [])
    att_rows = [[_w("CO"), _w("Target (%)"), _w("Attainment (%)"), _w("Gap"), _w("Action Plan")]]
    for r in co_att:
        att_rows.append([_w(r.get("co","")), _w(str(r.get("target",60))+"%"), _w(""), _w(""), _w("")])
    if len(att_rows) == 1:
        for c in data.get("cos", []):
            att_rows.append([_w(f"CO{c['num']}"), _w("60%"), _w(""), _w(""), _w("")])
    t = Table(att_rows, colWidths=[0.6*inch, 0.8*inch, 0.8*inch, 0.6*inch, W-2.8*inch])
    t.setStyle(_ts(len(att_rows)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("11–12. STUDENT REFLECTIONS & FACULTY JOURNAL", head_s))
    sr = [[_w("Reflection Aspect"), _w("Student Observation"), _w("Faculty Observation")]]
    for asp in ["Real-world learning","Sustainability understanding","AI tool effectiveness",
                "Innovation exposure","Industry relevance"]:
        sr.append([_w(asp), _w(""), _w("")])
    cw_sr = [2.0*inch, (W-2.0*inch)/2, (W-2.0*inch)/2]
    t = Table(sr, colWidths=cw_sr)
    t.setStyle(_ts(len(sr)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("14. NBA / NAAC / NEP 2020 / THE IMPACT ALIGNMENT", head_s))
    nba_rows = [[_w("Framework"), _w("Alignment Evidence")]]
    for r in data.get("nba_alignment", [
        {"framework":"NBA","evidence":"CO-PO attainment, CEP"},
        {"framework":"NAAC","evidence":"ICT-enabled teaching"},
        {"framework":"NEP 2020","evidence":"Skill-based learning"},
        {"framework":"SDGs","evidence":"Sustainability integration"},
        {"framework":"THE Impact","evidence":"Societal impact evidence"},
    ]):
        nba_rows.append([_w(r.get("framework","")), _w(r.get("evidence",""))])
    t = Table(nba_rows, colWidths=[1.5*inch, W-1.5*inch])
    t.setStyle(_ts(len(nba_rows)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("15. FINAL SEMESTER SUMMARY", head_s))
    fs = data.get("final_summary", {})
    fs_data = [
        [_w("Parameter"), _w("Status")],
        [_w("Syllabus Coverage (%)"),          _w(fs.get("syllabus_coverage",""))],
        [_w("CO Attainment Level"),            _w("")],
        [_w("SDG Activities Conducted"),       _w(fs.get("sdg_activities",""))],
        [_w("AI-Based Activities Conducted"),  _w(fs.get("ai_activities",""))],
        [_w("Industry Activities Conducted"),  _w(fs.get("industry_activities",""))],
        [_w("Innovations Implemented"),        _w(fs.get("innovations",""))],
    ]
    t = Table(fs_data, colWidths=[3.0*inch, W-3.0*inch])
    t.setStyle(_ts(len(fs_data)))
    story += [t, Spacer(1,6)]

    story.append(Paragraph("16. SIGNATURES", head_s))
    sig_data = [
        [_w("Role"), _w("Signature"), _w("Date")],
        *[[_w(r), _w(""), _w("")] for r in [
            "Faculty","Course Coordinator","Program Coordinator","HoD","IQAC Coordinator"
        ]]
    ]
    t = Table(sig_data, colWidths=[2.5*inch, 4.0*inch, 2.0*inch])
    t.setStyle(_ts(len(sig_data)))
    story.append(t)

    doc.build(story)


# ── TXT builder ────────────────────────────────────────────────────────────────

def build_txt(data: dict, meta: dict, course_code: str, course_title: str,
              output_path: str):
    lines = []
    sep = "=" * 80

    def sec(title):
        lines.extend(["", sep, title, sep])

    def tbl(headers, rows):
        widths = [max(len(str(r[i])) for r in [headers]+list(rows))+2
                  for i in range(len(headers))]
        div = "+" + "+".join("-"*w for w in widths) + "+"
        def fmt(row):
            return "|" + "|".join((" "+str(v)).ljust(w) for v, w in zip(row, widths)) + "|"
        lines.append(div)
        lines.append(fmt(headers))
        lines.append(div)
        for row in rows:
            lines.append(fmt(row))
        lines.append(div)

    lines.extend([
        "TEACHING DIARY (THEORY)",
        f"Course: {course_code} — {course_title}",
        f"Faculty: {meta.get('faculty_name','')}",
        f"Academic Year: {meta.get('academic_year','')}",
        f"Semester: {meta.get('semester','')}",
    ])

    sec("1. COURSE INTEGRATION SUMMARY")
    si = data.get("integration_summary", {})
    for k, v in [
        ("COs Addressed",             si.get("cos_addressed","")),
        ("POs Addressed",             si.get("pos_addressed","")),
        ("SDG Goals",                  si.get("sdg_goals","")),
        ("AI/ICT Tools",               si.get("ai_ict_tools","")),
        ("Industry Relevance",         si.get("industry_relevance","")),
    ]:
        lines.append(f"  {k}: {v}")

    sec("2. CO–PO–SDG–GA MAPPING")
    co_map = data.get("co_mapping", [])
    tbl(["CO","PO","GA","SDG","Sustainability"],
        [(r.get("co",""),r.get("po",""),r.get("ga",""),r.get("sdg",""),r.get("sustainability",""))
         for r in co_map])

    sec("3. THEORY TEACHING DELIVERY PLAN")
    sessions = data.get("theory_sessions", [])
    tbl(["#","Unit","Topic","CO","PO","SDG","Bloom's","Method","Assessment"],
        [(str(s.get("no","")),s.get("unit",""),s.get("topic","")[:30],
          s.get("co",""),s.get("po",""),s.get("sdg",""),
          s.get("bloom",""),s.get("method","")[:15],s.get("assessment",""))
         for s in sessions])

    tutorials = data.get("tutorials", [])
    if tutorials:
        sec("4. TUTORIAL / ACTIVITY DIARY")
        tbl(["#","Topic","CO","PO","Activity Type","Assessment"],
            [(str(t.get("no","")),t.get("topic","")[:30],t.get("co",""),
              t.get("po",""),t.get("activity_type",""),t.get("assessment",""))
             for t in tutorials])

    lab_sessions = data.get("lab_sessions", [])
    if lab_sessions:
        sec("5. LAB TEACHING DIARY")
        tbl(["#","Experiment Title","CO","PO","Software","Assessment"],
            [(str(l.get("no","")),l.get("title","")[:30],l.get("co",""),
              l.get("po",""),l.get("software",""),l.get("rubric",""))
             for l in lab_sessions])

    sec("10. COURSE OUTCOME ATTAINMENT TARGETS")
    co_att = data.get("co_attainment", [])
    tbl(["CO","Target (%)","Attainment (%)"],
        [(r.get("co",""),str(r.get("target",60))+"%","") for r in co_att])

    sec("14. NBA / NAAC ALIGNMENT")
    for r in data.get("nba_alignment", []):
        lines.append(f"  {r.get('framework','')} : {r.get('evidence','')}")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
