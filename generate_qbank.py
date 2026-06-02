"""
generate_qbank.py  –  Question Bank Generator for Engineering Courses

Usage:
  python generate_qbank.py <syllabus_file>
  python generate_qbank.py <syllabus_file> <co_file.txt>

Generates a comprehensive question bank aligned to the course outcomes.
"""

import sys
import os
import re
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

try:
    from generate_cos import (
        extract, split_into_courses, generate_cos_for_course,
        parse_co_selection, filter_cos,
    )
except ImportError:
    print("Error: generate_cos.py must be in the same directory.")
    sys.exit(1)


# ---------------------------------------------------------------------------
# CO parsing
# ---------------------------------------------------------------------------

def parse_cos_from_text(text):
    """Return list of CO dicts parsed from a TXT or DOCX CO file.

    Each dict: {num, statement, unit, kdim, bloom}
    Handles pipe-delimited markdown table rows produced by generate_cos.py.
    """
    co_stmts = {}
    co_meta  = {}

    for line in text.split('\n'):
        s = line.strip()

        # CO statement: "CO1: The students will be able to …"
        m = re.match(r'CO(\d+):\s*(.*)', s)
        if m:
            co_stmts[int(m.group(1))] = m.group(2).strip()
            continue

        # Pipe-delimited table row: | CO1 | Unit 1 | Conceptual | L3 - Apply |
        if s.startswith('|') and s.endswith('|') and not re.match(r'\|[-| ]+\|', s):
            cells = [c.strip() for c in s.strip('|').split('|')]
            if len(cells) == 4:
                m2 = re.match(r'CO(\d+)', cells[0])
                if m2 and cells[1]:
                    n = int(m2.group(1))
                    co_meta[n] = {'unit': cells[1], 'kdim': cells[2], 'bloom': cells[3]}

    cos = []
    for n in sorted(co_stmts):
        meta = co_meta.get(n, {})
        cos.append({
            'num':       n,
            'statement': co_stmts[n],
            'unit':      meta.get('unit', ''),
            'kdim':      meta.get('kdim', ''),
            'bloom':     meta.get('bloom', ''),
        })
    return cos


def load_cos_from_pdf(path, course_code=None):
    """Load COs from a PDF CO file using pdfplumber's table extractor.

    pdfplumber.extract_tables() reliably returns bordered table cells as
    structured rows, unlike extract_text() which concatenates them.
    CO statements are taken from the text layer; metadata from the table.
    """
    import pdfplumber

    co_stmts = {}
    co_meta  = {}

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # --- Table rows → metadata ---
            for table in (page.extract_tables() or []):
                for row in table:
                    if not row or len(row) < 4:
                        continue
                    cells = [str(c).strip() if c else '' for c in row]
                    m = re.match(r'CO(\d+)', cells[0])
                    if m and cells[1] and cells[1].lower() != 'unit':
                        n = int(m.group(1))
                        co_meta[n] = {
                            'unit':  cells[1],
                            'kdim':  cells[2],
                            'bloom': cells[3],
                        }

            # --- Text layer → CO statements ---
            text = page.extract_text() or ''
            for line in text.split('\n'):
                s = line.strip()
                m = re.match(r'CO(\d+):\s*(\S.*)', s)
                if m:
                    n = int(m.group(1))
                    if n not in co_stmts:
                        co_stmts[n] = m.group(2).strip()

    cos = []
    for n in sorted(co_stmts):
        meta = co_meta.get(n, {})
        cos.append({
            'num':       n,
            'statement': co_stmts[n],
            'unit':      meta.get('unit', ''),
            'kdim':      meta.get('kdim', ''),
            'bloom':     meta.get('bloom', ''),
        })
    return cos


def _extract_co_section(co_text, course_code):
    """If co_text has multiple ## sections, return just the one for course_code."""
    pattern = re.compile(
        rf'## {re.escape(course_code)}.*?\n(.*?)(?=\n## |\Z)', re.DOTALL
    )
    m = pattern.search(co_text)
    return m.group(1) if m else co_text


# ---------------------------------------------------------------------------
# Unit grouping
# ---------------------------------------------------------------------------

def _unit_sort_key(label):
    m = re.search(r'\d+', label)
    return int(m.group()) if m else 999


def extract_unit_titles(client, course_text):
    """Ask Claude to extract unit numbers and titles from syllabus text.

    Returns {unit_label: title}, e.g. {'Unit 1': 'Network Security Fundamentals'}.
    Falls back to an empty dict on any error.
    """
    prompt = (
        "From the syllabus text below, extract every unit/module heading. "
        "Output ONLY lines in this exact format, one per line:\n"
        "Unit N: Title\n\n"
        "Rules:\n"
        "- N is the unit number (integer)\n"
        "- Title is the unit heading, exactly as written in the syllabus\n"
        "- Do not include any explanation, numbering, or extra text\n"
        "- If no unit headings are found, output nothing\n\n"
        f"Syllabus:\n{course_text[:4000]}"
    )
    try:
        resp = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        titles = {}
        for line in resp.content[0].text.strip().split('\n'):
            m = re.match(r'Unit\s+(\d+)\s*[:\-]\s*(.+)', line.strip(), re.IGNORECASE)
            if m:
                titles[f"Unit {m.group(1)}"] = m.group(2).strip()[:80]
        return titles
    except Exception as e:
        print(f"  (unit title extraction failed: {e})")
        return {}


def build_unit_plan(client, cos, course_text):
    """Return ordered dict {unit_label: {'title': str, 'cos': [...]}}

    Only includes units that have at least one CO mapped to them.
    COs with no unit info are grouped as 'Unmapped' at the end.
    """
    def _normalise_unit_label(raw: str) -> str:
        """Convert '1', 'unit 1', 'Unit1' etc. → 'Unit 1'."""
        raw = raw.strip()
        m = re.match(r'^(?:unit\s*)?(\d+)$', raw, re.IGNORECASE)
        if m:
            return f"Unit {m.group(1)}"
        # already like 'Unit 1 – something' — keep as-is but normalise prefix
        m2 = re.match(r'^unit\s+(\d+)', raw, re.IGNORECASE)
        if m2:
            return f"Unit {m2.group(1)}"
        return raw

    # CO groups keyed by normalised "Unit N" label
    co_groups: dict = {}
    unmapped: list  = []
    for co in cos:
        raw_unit = co['unit'] if co['unit'] and co['unit'] not in ('N/A', '') else ''
        if raw_unit:
            label = _normalise_unit_label(raw_unit)
            co_groups.setdefault(label, []).append(co)
        else:
            unmapped.append(co)

    # Unit titles via Claude
    print("  Extracting unit titles from syllabus...")
    unit_titles = extract_unit_titles(client, course_text)
    if unit_titles:
        for lbl, ttl in sorted(unit_titles.items(), key=lambda x: _unit_sort_key(x[0])):
            print(f"    {lbl}: {ttl}")
    else:
        print("    (no unit titles found — units will be numbered only)")

    # Build plan: all syllabus units, sorted by unit number.
    # Units without COs use the content-only prompt (empty cos list).
    all_labels = set(co_groups.keys()) | set(unit_titles.keys())
    plan = {}
    for label in sorted(all_labels, key=_unit_sort_key):
        title = unit_titles.get(label, '')
        plan[label] = {'title': title, 'cos': co_groups.get(label, [])}

    # Distribute unmapped COs across existing units instead of creating an 'Unmapped' bucket.
    # This ensures every CO gets questions generated for it.
    if unmapped:
        if plan:
            unit_labels = list(plan.keys())
            for i, co in enumerate(unmapped):
                target = unit_labels[i % len(unit_labels)]
                plan[target]['cos'].append(co)
                print(f"    Assigned CO{co['num']} (unmapped) → {target}")
        else:
            plan['Unit 1'] = {'title': 'Course Content', 'cos': unmapped}

    return plan


# ---------------------------------------------------------------------------
# Question generation  (per unit, COs shown as tags on each question)
# ---------------------------------------------------------------------------

# CEP and SDG criteria are embedded here as internal design constraints only.
# They never appear in the generated questions or output documents.
_CEP_SDG_RULES = """\
Mandatory internal design criteria — apply rigorously but NEVER mention, label, \
reference, or allude to any of the following in the output:

1. COMPLEX ENGINEERING PROBLEMS
   * Frame 5-mark and 10-mark questions around authentic scenarios that involve \
multiple interacting variables: cost vs performance, safety vs throughput, \
resource constraints, scalability, regulatory compliance, or conflicting \
stakeholder requirements.
   * Present problems with no single correct answer; the student must select \
an approach and justify it with engineering reasoning.
   * For 10-mark questions provide a brief real-world context paragraph, then \
ask the student to analyse, design, optimise, or evaluate.

2. SUSTAINABLE DEVELOPMENT
   * Weave sustainability dimensions naturally into higher-mark question contexts \
where the subject matter allows: energy efficiency, waste reduction, climate \
adaptation, equitable access, long-term economic viability, or social impact.
   * Do NOT use the words "sustainable", "SDG", "green", "UN goals", or any \
framework name. Embed the dimension as part of the engineering scenario.

3. QUALITY RULES
   * Each question is self-contained and unambiguous.
   * No sub-questions, hints, model answers, or annotations beyond what is \
explicitly required by the format below.
   * Do NOT reference Bloom's taxonomy, NBA, CEP, SDG, OBE, or any \
accreditation framework anywhere in the questions.\
"""

_MARK_LEGEND = """\
Question length expectations:
  2 marks  — precise recall, one-sentence definition, or brief enumeration.
  5 marks  — comparative analysis, worked example, or explanation requiring \
4-6 sentences or clearly shown steps.
  10 marks — design problem, case-based scenario, or evaluation task requiring \
multi-step reasoning and justification of choices.

ABSOLUTE RULE: Output questions ONLY. Do NOT include answers, solutions, model \
answers, answer keys, hints, marking schemes, sample responses, or any \
explanatory text after a question. Every line after a Q<n>. stem must be \
part of the question itself, never an answer.\
"""

# --- Prompt: unit WITH mapped COs (questions carry [CON] tags) ---
_SYSTEM_QBANK_WITH_CO = """\
You are a senior examiner for accredited engineering programs, specialising in \
Outcome-Based Education (OBE) assessment design.

IMPORTANT: The Course Outcomes listed below are authoritative and correct. \
Do NOT question, validate, cross-check, or comment on them. \
Do NOT ask for clarification. Generate questions immediately.

Generate a question bank for the given unit. The unit covers {num_cos} Course Outcome(s):

{co_list}

Produce exactly {n2} questions worth 2 marks, {n5} worth 5 marks, \
and {n10} worth 10 marks ({total} total).
Distribute questions across all COs — every CO must appear at least once.

At the END of each question line append the CO tag: [CO1], [CO2], etc.
Use only the CO numbers listed above.

""" + _MARK_LEGEND + "\n\n" + _CEP_SDG_RULES + """

Output format — follow exactly, with NO preamble, explanation, or commentary:

[2 Marks]
Q1. <question text>  [CO1]
Q2. <question text>  [CO2]

[5 Marks]
Q<n>. <question text>  [CO1]

[10 Marks]
Q<n>. <question stem — keep to one line>  [CO2]
<Optional: multi-line scenario/context here. No CO tag on these lines.>

CRITICAL: The [CON] tag MUST appear on the same line as Q<n>., immediately \
after the question stem, before any scenario text.
Start your response directly with [2 Marks] — no introduction.
Number questions sequentially (Q1, Q2, Q3 ...).\
"""

# --- Prompt: unit WITHOUT mapped COs (questions have no CO tags) ---
_SYSTEM_QBANK_NO_CO = """\
You are a senior examiner for accredited engineering programs.

Generate a question bank for the given unit based on the syllabus content provided.
Produce exactly {n2} questions worth 2 marks, {n5} worth 5 marks, \
and {n10} worth 10 marks ({total} total).

""" + _MARK_LEGEND + "\n\n" + _CEP_SDG_RULES + """

Output format — follow exactly, with NO preamble, explanation, or commentary:

[2 Marks]
Q1. <question text>
Q2. <question text>

[5 Marks]
Q<n>. <question text>

[10 Marks]
Q<n>. <question text>

Start your response directly with [2 Marks] — no introduction.
Number questions sequentially (Q1, Q2, Q3 ...).\
"""

# --- Assignment prompt templates ---
_SYSTEM_ASSIGNMENT_WITH_CO = """\
You are a senior examiner designing assignment tasks for accredited engineering programs.

IMPORTANT: The Course Outcomes listed below are authoritative and correct. \
Do NOT question, validate, or comment on them. Generate assignments immediately.

Generate assignment tasks for the given unit. The unit covers {num_cos} Course Outcome(s):

{co_list}

Produce exactly {n_assign} assignment task(s). Each must be a substantial, open-ended \
task requiring 2–4 hours of student effort.
Distribute assignments across all COs — every CO must appear at least once.
Append the CO tag [CO1], [CO2] etc. immediately after the A<n>. title line.

""" + _CEP_SDG_RULES + """

Output format — follow exactly, no preamble or commentary:

[Assignment]
A1. <one-line problem title or stem>  [CO1]
Context: <2–3 sentences of real-world scenario>
Tasks:
  (i)  <specific sub-task or deliverable>
  (ii) <specific sub-task or deliverable>
  (iii) <specific sub-task or deliverable>
Evaluation Criteria: <measurable performance standard>

A2. <one-line problem title or stem>  [CO2]
...

Start directly with [Assignment]. No introduction or explanation.\
"""

_SYSTEM_ASSIGNMENT_NO_CO = """\
You are a senior examiner designing assignment tasks for accredited engineering programs.

Generate exactly {n_assign} assignment task(s) based on the unit syllabus content. \
Each must be a substantial, open-ended task requiring 2–4 hours of student effort.

""" + _CEP_SDG_RULES + """

Output format — follow exactly, no preamble or commentary:

[Assignment]
A1. <one-line problem title or stem>
Context: <2–3 sentences of real-world scenario>
Tasks:
  (i)  <specific sub-task or deliverable>
  (ii) <specific sub-task or deliverable>
  (iii) <specific sub-task or deliverable>
Evaluation Criteria: <measurable performance standard>

Start directly with [Assignment]. No introduction or explanation.\
"""

# --- Quiz prompt templates ---
_SYSTEM_QUIZ_WITH_CO = """\
You are a senior examiner designing quiz questions for accredited engineering programs.

IMPORTANT: The Course Outcomes listed below are authoritative. Generate quiz questions immediately.

Generate quiz questions for the given unit. The unit covers {num_cos} Course Outcome(s):

{co_list}

Produce exactly {n_quiz} quiz questions — a balanced mix of:
  - Multiple Choice (MCQ): question stem + exactly 4 options labelled (a)–(d)
  - True or False: prefix the stem with "True or False:"
  - Fill in the Blank: use ___ to mark the missing term, prefix with "Fill in the blank:"

Distribute questions across all COs — every CO must appear at least once.
Append [CO1], [CO2] etc. at the end of each Q<n>. line (not on option lines).

ABSOLUTE RULE: Output questions and options ONLY. No answers, answer keys, hints, \
or explanations whatsoever.

Output format — follow exactly, no preamble:

[Quiz]
Q1. <MCQ stem>  [CO1]
   (a) <option>
   (b) <option>
   (c) <option>
   (d) <option>

Q2. True or False: <statement>  [CO2]

Q3. Fill in the blank: <sentence with ___ for the missing term>  [CO1]

Start directly with [Quiz]. Number questions sequentially (Q1, Q2, Q3 ...).\
"""

_SYSTEM_QUIZ_NO_CO = """\
You are a senior examiner designing quiz questions for accredited engineering programs.

Generate exactly {n_quiz} quiz questions based on the unit syllabus content — a balanced mix of:
  - Multiple Choice (MCQ): question stem + exactly 4 options labelled (a)–(d)
  - True or False: prefix the stem with "True or False:"
  - Fill in the Blank: use ___ to mark the missing term, prefix with "Fill in the blank:"

ABSOLUTE RULE: Output questions and options ONLY. No answers, answer keys, hints, \
or explanations whatsoever.

Output format — follow exactly, no preamble:

[Quiz]
Q1. <MCQ stem>
   (a) <option>
   (b) <option>
   (c) <option>
   (d) <option>

Q2. True or False: <statement>

Q3. Fill in the blank: <sentence with ___ for the missing term>

Start directly with [Quiz]. Number questions sequentially (Q1, Q2, Q3 ...).\
"""

# --- Lab experiment prompt templates ---
_SYSTEM_LAB_WITH_CO = """\
You are a senior laboratory instructor designing experiment write-ups for an accredited \
engineering laboratory course.

IMPORTANT: The Course Outcomes listed below are authoritative and correct. \
Do NOT question, validate, or comment on them. Generate experiments immediately.

Generate {n_exp} laboratory experiment(s) for the given unit. \
The unit covers {num_cos} Course Outcome(s):

{co_list}

Every CO must appear at least once. Append the CO tag [CO1], [CO2] etc. on the same line \
as "Exp <n>. <Title>", immediately after the title.

Each experiment must contain ALL of these sections in this exact order:
  Aim        — one clear sentence stating what the student will do or demonstrate.
  Equipment/Tools — comma-separated list of required hardware, software, or materials.
  Theory     — 2-3 sentences of relevant background theory.
  Procedure  — numbered steps (at least 5) describing exactly what the student must do.
  Expected Observations/Output — what the student should record, observe, or measure.
  Result     — what the experiment proves or demonstrates.
  Viva Questions — exactly 5 short questions (V1. … V5.) testing conceptual understanding.

ABSOLUTE RULES:
  - Do NOT provide answers to the viva questions.
  - Do NOT include model answers, hints, or sample responses anywhere.
  - Each experiment must be completely self-contained and unambiguous.
  - Do NOT reference accreditation frameworks, OBE, Bloom's taxonomy, CEP, or SDG.

Output format — follow exactly, no preamble or commentary:

[Experiments]
Exp 1. <Experiment Title>  [CO1]
Aim: <one-sentence objective>
Equipment/Tools: <comma-separated list>
Theory:
<2-3 sentences of background>
Procedure:
  1. <step>
  2. <step>
  3. <step>
  4. <step>
  5. <step>
Expected Observations/Output:
<description of what to record or observe>
Result:
<what this experiment proves or demonstrates>
Viva Questions:
  V1. <question>
  V2. <question>
  V3. <question>
  V4. <question>
  V5. <question>

Exp 2. <Experiment Title>  [CO2]
...

Start directly with [Experiments]. Number experiments Exp 1, Exp 2, ... sequentially.\
"""

_SYSTEM_LAB_NO_CO = """\
You are a senior laboratory instructor designing experiment write-ups for an accredited \
engineering laboratory course.

Generate {n_exp} laboratory experiment(s) for the given unit based on the syllabus content.

Each experiment must contain ALL of these sections in this exact order:
  Aim        — one clear sentence stating what the student will do or demonstrate.
  Equipment/Tools — comma-separated list of required hardware, software, or materials.
  Theory     — 2-3 sentences of relevant background theory.
  Procedure  — numbered steps (at least 5) describing exactly what the student must do.
  Expected Observations/Output — what the student should record, observe, or measure.
  Result     — what the experiment proves or demonstrates.
  Viva Questions — exactly 5 short questions (V1. … V5.) testing conceptual understanding.

ABSOLUTE RULES:
  - Do NOT provide answers to the viva questions.
  - Do NOT include model answers, hints, or sample responses anywhere.
  - Each experiment must be completely self-contained and unambiguous.

Output format — follow exactly, no preamble or commentary:

[Experiments]
Exp 1. <Experiment Title>
Aim: <one-sentence objective>
Equipment/Tools: <comma-separated list>
Theory:
<2-3 sentences of background>
Procedure:
  1. <step>
  2. <step>
  3. <step>
  4. <step>
  5. <step>
Expected Observations/Output:
<description of what to record or observe>
Result:
<what this experiment proves or demonstrates>
Viva Questions:
  V1. <question>
  V2. <question>
  V3. <question>
  V4. <question>
  V5. <question>

Start directly with [Experiments]. Number experiments Exp 1, Exp 2, ... sequentially.\
"""


def generate_questions_for_unit(client, unit_label, unit_title, cos,
                                 course_code, course_title, course_syllabus,
                                 n2=4, n5=4, n10=4):
    """Stream question generation for one unit.

    Uses CO-aware prompt when COs are mapped; falls back to content-only
    prompt when no COs are available for the unit.
    """
    total    = n2 + n5 + n10
    full_unit = f"{unit_label}: {unit_title}" if unit_title else unit_label

    if cos:
        co_list = "\n".join(
            f"  CO{co['num']}: {co['statement']}"
            + (f"  [{co['bloom']}]" if co['bloom'] else "")
            for co in cos
        )
        system = _SYSTEM_QBANK_WITH_CO.format(
            num_cos=len(cos), co_list=co_list,
            n2=n2, n5=n5, n10=n10, total=total,
        )
    else:
        system = _SYSTEM_QBANK_NO_CO.format(n2=n2, n5=n5, n10=n10, total=total)

    unit_excerpt = _extract_unit_text(unit_label, course_syllabus)

    user_msg = (
        f"Course Code: {course_code}\n"
        f"Course Title: {course_title}\n"
        f"Unit: {full_unit}\n\n"
        f"Syllabus content for this unit:\n{unit_excerpt}"
    )

    max_tokens = 600 + total * 180
    result = ""
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_msg}]
    ) as stream:
        for chunk in stream.text_stream:
            print(chunk, end="", flush=True)
            result += chunk

    print("\n")
    return result


def _extract_unit_text(unit_label, course_syllabus):
    """Return the syllabus excerpt for a specific unit (up to 3000 chars)."""
    unit_num_m = re.search(r'\d+', unit_label)
    if unit_num_m:
        unum = unit_num_m.group()
        for pat in [
            rf'(?i)(unit\s*{unum}[\s\S]{{0,3000}}?)(?=unit\s*{int(unum)+1}|\Z)',
            rf'(?i)(module\s*{unum}[\s\S]{{0,3000}}?)(?=module\s*{int(unum)+1}|\Z)',
        ]:
            m = re.search(pat, course_syllabus)
            if m:
                return m.group(1)[:3000]
    return course_syllabus[:3000]


def generate_assignments_for_unit(client, unit_label, unit_title, cos,
                                   course_code, course_title, course_syllabus,
                                   n_assign=2):
    """Stream assignment task generation for one unit."""
    full_unit = f"{unit_label}: {unit_title}" if unit_title else unit_label
    if cos:
        co_list = "\n".join(
            f"  CO{co['num']}: {co['statement']}"
            + (f"  [{co['bloom']}]" if co['bloom'] else "")
            for co in cos
        )
        system = _SYSTEM_ASSIGNMENT_WITH_CO.format(
            num_cos=len(cos), co_list=co_list, n_assign=n_assign,
        )
    else:
        system = _SYSTEM_ASSIGNMENT_NO_CO.format(n_assign=n_assign)

    unit_excerpt = _extract_unit_text(unit_label, course_syllabus)
    user_msg = (
        f"Course Code: {course_code}\n"
        f"Course Title: {course_title}\n"
        f"Unit: {full_unit}\n\n"
        f"Syllabus content for this unit:\n{unit_excerpt}"
    )
    max_tokens = 300 + n_assign * 280
    result = ""
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_msg}]
    ) as stream:
        for chunk in stream.text_stream:
            print(chunk, end="", flush=True)
            result += chunk
    print("\n")
    return result


def generate_quiz_for_unit(client, unit_label, unit_title, cos,
                            course_code, course_title, course_syllabus,
                            n_quiz=5):
    """Stream quiz question generation for one unit."""
    full_unit = f"{unit_label}: {unit_title}" if unit_title else unit_label
    if cos:
        co_list = "\n".join(
            f"  CO{co['num']}: {co['statement']}"
            + (f"  [{co['bloom']}]" if co['bloom'] else "")
            for co in cos
        )
        system = _SYSTEM_QUIZ_WITH_CO.format(
            num_cos=len(cos), co_list=co_list, n_quiz=n_quiz,
        )
    else:
        system = _SYSTEM_QUIZ_NO_CO.format(n_quiz=n_quiz)

    unit_excerpt = _extract_unit_text(unit_label, course_syllabus)
    user_msg = (
        f"Course Code: {course_code}\n"
        f"Course Title: {course_title}\n"
        f"Unit: {full_unit}\n\n"
        f"Syllabus content for this unit:\n{unit_excerpt}"
    )
    max_tokens = 200 + n_quiz * 150
    result = ""
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_msg}]
    ) as stream:
        for chunk in stream.text_stream:
            print(chunk, end="", flush=True)
            result += chunk
    print("\n")
    return result


def generate_lab_experiments_for_unit(client, unit_label, unit_title, cos,
                                       course_code, course_title, course_syllabus,
                                       n_exp=3):
    """Stream lab experiment generation for one unit."""
    full_unit = f"{unit_label}: {unit_title}" if unit_title else unit_label
    if cos:
        co_list = "\n".join(
            f"  CO{co['num']}: {co['statement']}"
            + (f"  [{co['bloom']}]" if co['bloom'] else "")
            for co in cos
        )
        system = _SYSTEM_LAB_WITH_CO.format(
            num_cos=len(cos), co_list=co_list, n_exp=n_exp,
        )
    else:
        system = _SYSTEM_LAB_NO_CO.format(n_exp=n_exp)

    unit_excerpt = _extract_unit_text(unit_label, course_syllabus)
    user_msg = (
        f"Course Code: {course_code}\n"
        f"Course Title: {course_title}\n"
        f"Unit: {full_unit}\n\n"
        f"Syllabus content for this unit:\n{unit_excerpt}"
    )
    max_tokens = 400 + n_exp * 650
    result = ""
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_msg}]
    ) as stream:
        for chunk in stream.text_stream:
            print(chunk, end="", flush=True)
            result += chunk
    print("\n")
    return result


# ---------------------------------------------------------------------------
# Post-processing: normalise AI-generated question blocks
# ---------------------------------------------------------------------------

_BOLD_RE    = re.compile(r'\*\*(.+?)\*\*')
_ITALIC_RE  = re.compile(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)')
_MARK_HDR   = re.compile(r'^\s*#{0,3}\s*\[(\d+)\s*[Mm]arks?\]')
_ASSIGN_HDR = re.compile(r'^\s*#{0,3}\s*\[Assignment\b', re.IGNORECASE)
_QUIZ_HDR   = re.compile(r'^\s*#{0,3}\s*\[Quiz\b', re.IGNORECASE)
_LAB_HDR    = re.compile(r'^\s*#{0,3}\s*\[Experiments?\b', re.IGNORECASE)


def _normalize_qblock(text, unit=""):
    """Sanitise raw AI output for consistent rendering.

    - Discards any preamble the model added before the first [N Marks] line
    - Canonicalises mark headers: '## [2 marks]' / ' [2 Marks]' → '[2 Marks]  |  Unit X'
    - Converts '---' dividers to blank lines
    - Strips markdown bold (**text**) and italic (*text*) markers
    """
    lines = text.split('\n')

    # Find the first real [N Marks] header and discard everything before it.
    # If none found the model refused or produced unusable output — return empty.
    start = None
    for i, line in enumerate(lines):
        if _MARK_HDR.match(line):
            start = i
            break
    if start is None:
        return ""

    unit_suffix = f"  |  {unit}" if unit else ""
    _CO_TAG_ONLY = re.compile(r'^\[CO\d+\]\s*$')
    _POSTAMBLE_RE = re.compile(
        r'^(end of|note[:\s]|these questions|the above questions|'
        r'this completes|all questions|question bank complete|total questions)',
        re.IGNORECASE,
    )
    _ANSWER_HDR_RE = re.compile(
        r'^(answer[:\s]|solution[:\s]|model answer|answer key|'
        r'marking scheme|sample (answer|response)|explanation[:\s])',
        re.IGNORECASE,
    )

    out = []
    skip_answer = False
    for line in lines[start:]:
        s = line.strip()

        # Resume normal processing at the next question or mark header
        if skip_answer:
            if _MARK_HDR.match(s) or re.match(r'Q\d+\.', s):
                skip_answer = False
            else:
                continue

        if not s or s in ('---', '***', '___'):
            out.append('')
            continue

        m = _MARK_HDR.match(s)
        if m:
            out.append(f"[{m.group(1)} Marks]{unit_suffix}")
            continue

        # If model put the [CON] tag on its own line, attach it to the last Q line
        # (even if continuation lines appear between the Q line and the tag)
        if _CO_TAG_ONLY.match(s):
            for j in range(len(out) - 1, -1, -1):
                if re.match(r'Q\d+\.', out[j]) and not re.search(r'\[CO\d+\]', out[j]):
                    out[j] = out[j].rstrip() + '  ' + s
                    break
            continue

        # Drop answer/solution blocks the model may generate despite instructions
        if _ANSWER_HDR_RE.match(s):
            skip_answer = True
            continue

        # Strip markdown bold and italic
        s = _BOLD_RE.sub(r'\1', s)
        s = _ITALIC_RE.sub(r'\1', s)
        out.append(s)

    # Strip AI postamble: trailing lines that match common sign-off patterns or are empty
    last_meaningful = len(out) - 1
    while last_meaningful >= 0:
        s = out[last_meaningful].strip()
        if s and not _POSTAMBLE_RE.match(s):
            break
        last_meaningful -= 1
    out = out[:last_meaningful + 1]

    return '\n'.join(out)


def _normalize_assignment_block(text, unit=""):
    """Sanitise raw assignment output: find [Assignment] header, strip preamble/postamble."""
    lines = text.split('\n')
    start = None
    for i, line in enumerate(lines):
        if _ASSIGN_HDR.match(line):
            start = i
            break
    if start is None:
        return ""

    unit_suffix = f"  |  {unit}" if unit else ""
    _CO_TAG_ONLY = re.compile(r'^\[CO\d+\]\s*$')
    _POSTAMBLE_RE = re.compile(
        r'^(end of|note[:\s]|these assignments|this completes|all assignments)',
        re.IGNORECASE,
    )

    out = []
    for line in lines[start:]:
        s = line.strip()
        if not s:
            out.append('')
            continue
        if _ASSIGN_HDR.match(s):
            out.append(f"[Assignment]{unit_suffix}")
            continue
        # CO tag on its own line → attach to last A-line
        if _CO_TAG_ONLY.match(s):
            for j in range(len(out) - 1, -1, -1):
                if re.match(r'A\d+\.', out[j]) and not re.search(r'\[CO\d+\]', out[j]):
                    out[j] = out[j].rstrip() + '  ' + s
                    break
            continue
        s = _BOLD_RE.sub(r'\1', s)
        s = _ITALIC_RE.sub(r'\1', s)
        out.append(s)

    last_meaningful = len(out) - 1
    while last_meaningful >= 0:
        s = out[last_meaningful].strip()
        if s and not _POSTAMBLE_RE.match(s):
            break
        last_meaningful -= 1
    return '\n'.join(out[:last_meaningful + 1])


def _normalize_quiz_block(text, unit=""):
    """Sanitise raw quiz output: find [Quiz] header, strip preamble/postamble."""
    lines = text.split('\n')
    start = None
    for i, line in enumerate(lines):
        if _QUIZ_HDR.match(line):
            start = i
            break
    if start is None:
        return ""

    unit_suffix = f"  |  {unit}" if unit else ""
    _CO_TAG_ONLY = re.compile(r'^\[CO\d+\]\s*$')
    _POSTAMBLE_RE = re.compile(
        r'^(end of|note[:\s]|these questions|this completes|all questions|answer[:\s]|'
        r'solution[:\s]|correct answer)',
        re.IGNORECASE,
    )

    out = []
    skip_answer = False
    for line in lines[start:]:
        s = line.strip()
        if skip_answer:
            if _QUIZ_HDR.match(s) or re.match(r'Q\d+\.', s):
                skip_answer = False
            else:
                continue
        if not s:
            out.append('')
            continue
        if _QUIZ_HDR.match(s):
            out.append(f"[Quiz]{unit_suffix}")
            continue
        if _POSTAMBLE_RE.match(s):
            skip_answer = True
            continue
        if _CO_TAG_ONLY.match(s):
            for j in range(len(out) - 1, -1, -1):
                if re.match(r'Q\d+\.', out[j]) and not re.search(r'\[CO\d+\]', out[j]):
                    out[j] = out[j].rstrip() + '  ' + s
                    break
            continue
        s = _BOLD_RE.sub(r'\1', s)
        s = _ITALIC_RE.sub(r'\1', s)
        out.append(s)

    last_meaningful = len(out) - 1
    while last_meaningful >= 0:
        s = out[last_meaningful].strip()
        if s and not _POSTAMBLE_RE.match(s):
            break
        last_meaningful -= 1
    return '\n'.join(out[:last_meaningful + 1])


def _normalize_lab_block(text, unit=""):
    """Sanitise raw lab experiment output: find [Experiments] header, strip preamble/postamble."""
    lines = text.split('\n')
    start = None
    for i, line in enumerate(lines):
        if _LAB_HDR.match(line):
            start = i
            break
    if start is None:
        return ""

    unit_suffix = f"  |  {unit}" if unit else ""
    _CO_TAG_ONLY = re.compile(r'^\[CO\d+\]\s*$')
    _POSTAMBLE_RE = re.compile(
        r'^(end of|note[:\s]|these experiments|this completes|all experiments)',
        re.IGNORECASE,
    )
    _ANSWER_RE = re.compile(
        r'^(answer[:\s]|solution[:\s]|sample answer|viva answer)',
        re.IGNORECASE,
    )

    out = []
    skip_answer = False
    for line in lines[start:]:
        s = line.strip()
        if skip_answer:
            if _LAB_HDR.match(s) or re.match(r'Exp\s+\d+\.', s, re.IGNORECASE):
                skip_answer = False
            else:
                continue
        if not s:
            out.append('')
            continue
        if _LAB_HDR.match(s):
            out.append(f"[Experiments]{unit_suffix}")
            continue
        if _CO_TAG_ONLY.match(s):
            for j in range(len(out) - 1, -1, -1):
                if re.match(r'Exp\s+\d+\.', out[j], re.IGNORECASE) and not re.search(r'\[CO\d+\]', out[j]):
                    out[j] = out[j].rstrip() + '  ' + s
                    break
            continue
        if _ANSWER_RE.match(s):
            skip_answer = True
            continue
        s = _BOLD_RE.sub(r'\1', s)
        s = _ITALIC_RE.sub(r'\1', s)
        out.append(s)

    last_meaningful = len(out) - 1
    while last_meaningful >= 0:
        s = out[last_meaningful].strip()
        if s and not _POSTAMBLE_RE.match(s):
            break
        last_meaningful -= 1
    return '\n'.join(out[:last_meaningful + 1])


# ---------------------------------------------------------------------------
# Lab course detection
# ---------------------------------------------------------------------------

_LAB_TITLE_RE = re.compile(
    r'\b(lab(oratory)?|practical[s]?|workshop[s]?)\b', re.IGNORECASE
)
_LAB_CODE_RE = re.compile(
    r'(LAB|PRAC|PRACT)(\d|$)|[A-Z]\d*L$', re.IGNORECASE
)


def is_lab_course(code, title, has_lab=False):
    """Return True if the course is a lab/practical course.

    Detects via three signals:
      - has_lab: P > 0 in the syllabus L-T-P line (most reliable for integrated lab courses)
      - title keywords: lab, laboratory, practical(s), workshop(s)
      - code patterns: ends in L, or contains LAB/PRAC/PRACT
    """
    return has_lab or bool(_LAB_TITLE_RE.search(title) or _LAB_CODE_RE.search(code))


# ---------------------------------------------------------------------------
# Statistics helper
# ---------------------------------------------------------------------------

def question_summary(qtext):
    """Return (total, {2: n, 5: n, 10: n}) counts from a questions block."""
    counts = {2: 0, 5: 0, 10: 0}
    current_marks = None
    for line in qtext.split('\n'):
        s = line.strip()
        m = re.match(r'\[(\d+) Marks?\]', s)
        if m:
            current_marks = int(m.group(1))
        elif re.match(r'Q\d+\.', s) and current_marks in counts:
            counts[current_marks] += 1
    return sum(counts.values()), counts


def count_co_questions(blocks):
    """Count questions per CO broken down by tier (2-mark, 5-mark, 10-mark, assign, quiz).

    Returns {co_num_str: {2: int, 5: int, 10: int, 'assign': int, 'quiz': int}}.
    Only counts lines that carry a [CON] tag.
    """
    tally: dict = {}
    current_marks = None   # int for exam tiers, 'assign' or 'quiz' for others

    def _blank():
        return {2: 0, 5: 0, 10: 0, 'assign': 0, 'quiz': 0}

    for block in blocks:
        for line in block.split('\n'):
            s = line.strip()
            m = re.match(r'\[(\d+) Marks?\]', s)
            if m:
                current_marks = int(m.group(1))
                continue
            if _ASSIGN_HDR.match(s):
                current_marks = 'assign'
                continue
            if _QUIZ_HDR.match(s):
                current_marks = 'quiz'
                continue

            is_exam_q   = re.match(r'Q\d+\.', s) and current_marks in (2, 5, 10)
            is_assign_q = re.match(r'A\d+\.', s) and current_marks == 'assign'
            is_quiz_q   = re.match(r'Q\d+\.', s) and current_marks == 'quiz'

            if is_exam_q or is_assign_q or is_quiz_q:
                co_m = re.search(r'\[CO(\d+)\]', s)
                if co_m:
                    cn = co_m.group(1)
                    if cn not in tally:
                        tally[cn] = _blank()
                    tally[cn][current_marks] += 1

    return dict(sorted(tally.items(), key=lambda kv: int(kv[0])))


def count_co_lab(blocks):
    """Count experiments per CO from lab blocks.

    Returns {co_num_str: {'exp': int}}.
    Only counts Exp N. lines that carry a [CON] tag.
    """
    tally: dict = {}
    in_experiments = False
    for block in blocks:
        for line in block.split('\n'):
            s = line.strip()
            if _LAB_HDR.match(s):
                in_experiments = True
                continue
            if in_experiments and re.match(r'Exp\s+\d+\.', s, re.IGNORECASE):
                co_m = re.search(r'\[CO(\d+)\]', s)
                if co_m:
                    cn = co_m.group(1)
                    if cn not in tally:
                        tally[cn] = {'exp': 0}
                    tally[cn]['exp'] += 1
    return dict(sorted(tally.items(), key=lambda kv: int(kv[0])))


# ---------------------------------------------------------------------------
# Save helpers
# ---------------------------------------------------------------------------

def _co_summary_txt(tally):
    """Return a plain-text CO summary table string."""
    if not tally:
        return ""
    rows = []
    tot2 = tot5 = tot10 = tota = totq = 0
    for cn, c in tally.items():
        rows.append((f"CO{cn}", c[2], c[5], c[10], c['assign'], c['quiz'],
                     c[2]+c[5]+c[10]+c['assign']+c['quiz']))
        tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
        tota += c['assign']; totq += c['quiz']
    rows.append(("Total", tot2, tot5, tot10, tota, totq, tot2+tot5+tot10+tota+totq))

    hdr = (f"  {'CO':<8} {'2-mark':>8} {'5-mark':>8} {'10-mark':>9}"
           f" {'Assign':>8} {'Quiz':>8} {'Total':>8}")
    sep = ("  " + "-"*8 + (" " + "-"*8)*3 + " " + "-"*9 + " " + "-"*8 + " " + "-"*8)
    lines = ["\n## CO Coverage Summary\n", hdr, sep]
    for i, (co, m2, m5, m10, ma, mq, tot) in enumerate(rows):
        if i == len(rows) - 1:
            lines.append(sep)
        lines.append(f"  {co:<8} {m2:>8} {m5:>8} {m10:>9} {ma:>8} {mq:>8} {tot:>8}")
    return '\n'.join(lines) + '\n'


def _co_lab_summary_txt(tally):
    """Return a plain-text CO summary for lab question banks."""
    if not tally:
        return ""
    rows = []
    tot_exp = 0
    for cn, c in tally.items():
        rows.append((f"CO{cn}", c['exp']))
        tot_exp += c['exp']
    rows.append(("Total", tot_exp))

    hdr = f"  {'CO':<8} {'Experiments':>13}"
    sep = "  " + "-"*8 + " " + "-"*13
    lines = ["\n## CO Coverage Summary\n", hdr, sep]
    for i, (co, exp) in enumerate(rows):
        if i == len(rows) - 1:
            lines.append(sep)
        lines.append(f"  {co:<8} {exp:>13}")
    return '\n'.join(lines) + '\n'


def save_txt(blocks, output_path, co_tally=None, lab_tally=None):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(blocks))
        if lab_tally:
            f.write(_co_lab_summary_txt(lab_tally))
        elif co_tally:
            f.write(_co_summary_txt(co_tally))


def save_docx(blocks, output_path, co_tally=None, lab_tally=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx.oxml.ns as qns
    from docx.oxml import OxmlElement

    doc = Document()
    # Tighten default margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    def _shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qns.qn('w:val'),   'clear')
        shd.set(qns.qn('w:color'), 'auto')
        shd.set(qns.qn('w:fill'),  hex_color)
        tcPr.append(shd)

    for block in blocks:
        for raw in block.split('\n'):
            line = _BOLD_RE.sub(r'\1', _ITALIC_RE.sub(r'\1', raw.strip()))
            if not line or line in ('---', '***', '___'):
                doc.add_paragraph()
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)

            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)

            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)

            elif re.match(r'\[\d+ Marks?\]', line):
                p    = doc.add_paragraph()
                run  = p.add_run(line)
                run.bold            = True
                run.font.size       = Pt(10)
                run.font.color.rgb  = RGBColor(0x1F, 0x49, 0x7D)

            elif re.match(r'Q\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                q_label, _, rest = line.partition('.')
                run_label = p.add_run(q_label + '.  ')
                run_label.bold = True
                # Separate trailing [CON] tag
                co_m = re.search(r'\s*(\[CO\d+\])\s*$', rest)
                if co_m:
                    p.add_run(rest[:co_m.start()].strip() + '  ')
                    co_run = p.add_run(co_m.group(1))
                    co_run.bold = True
                    co_run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
                else:
                    p.add_run(rest.strip())

            elif re.match(r'\[Assignment\b', line, re.IGNORECASE):
                p   = doc.add_paragraph()
                run = p.add_run(line)
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = RGBColor(0xBF, 0x47, 0x00)  # dark amber

            elif re.match(r'\[Quiz\b', line, re.IGNORECASE):
                p   = doc.add_paragraph()
                run = p.add_run(line)
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = RGBColor(0x37, 0x5C, 0x23)  # dark green

            elif re.match(r'A\d+\.', line):
                p = doc.add_paragraph()
                a_label, _, rest = line.partition('.')
                run_label = p.add_run(a_label + '.  ')
                run_label.bold = True
                run_label.font.color.rgb = RGBColor(0xBF, 0x47, 0x00)
                co_m = re.search(r'\s*(\[CO\d+\])\s*$', rest)
                if co_m:
                    p.add_run(rest[:co_m.start()].strip() + '  ')
                    co_run = p.add_run(co_m.group(1))
                    co_run.bold = True
                    co_run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
                else:
                    p.add_run(rest.strip())

            elif re.match(r'\[Experiments?\b', line, re.IGNORECASE):
                p   = doc.add_paragraph()
                run = p.add_run(line)
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = RGBColor(0x4B, 0x00, 0x82)  # indigo

            elif re.match(r'Exp\s+\d+\.', line, re.IGNORECASE):
                p = doc.add_paragraph()
                m_exp = re.match(r'(Exp\s+\d+)(\.)\s*(.*)', line, re.IGNORECASE)
                if m_exp:
                    run_label = p.add_run(m_exp.group(1) + '.  ')
                    run_label.bold = True
                    run_label.font.color.rgb = RGBColor(0x4B, 0x00, 0x82)
                    rest = m_exp.group(3).strip()
                    co_m = re.search(r'\s*(\[CO\d+\])\s*$', rest)
                    if co_m:
                        p.add_run(rest[:co_m.start()].strip() + '  ')
                        co_run = p.add_run(co_m.group(1))
                        co_run.bold = True
                        co_run.font.color.rgb = RGBColor(0x00, 0x70, 0x70)
                    else:
                        p.add_run(rest)
                else:
                    p.add_run(line)

            elif re.match(r'^(Aim|Equipment/Tools|Theory|Procedure|Expected Observations/Output|Result|Viva Questions)\s*:', line):
                p = doc.add_paragraph()
                label, _, rest = line.partition(':')
                run_lbl = p.add_run(label + ': ')
                run_lbl.bold = True
                if rest.strip():
                    p.add_run(rest.strip())

            elif re.match(r'^\s+V\d+\.', line):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.add_run(line.strip())

            elif re.match(r'^\([a-d]\)', line):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.add_run(line)

            else:
                doc.add_paragraph(line)

    if lab_tally:
        doc.add_heading("CO Coverage Summary", level=1)
        headers = ["CO", "Experiments"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(hdr_cells[i], "4B0082")
        tot_exp = 0
        for cn, c in lab_tally.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f"CO{cn}"
            row_cells[1].text = str(c['exp'])
            tot_exp += c['exp']
        tot_cells = table.add_row().cells
        for cell in tot_cells:
            _shade_cell(cell, "E8D5F5")
        tot_cells[0].text = "Total"
        tot_cells[1].text = str(tot_exp)
        for cell in tot_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    elif co_tally:
        doc.add_heading("CO Coverage Summary", level=1)
        headers = ["CO", "2-mark", "5-mark", "10-mark", "Assign", "Quiz", "Total"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(hdr_cells[i], "1F497D")

        tot2 = tot5 = tot10 = tota = totq = 0
        for cn, c in co_tally.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f"CO{cn}"
            row_cells[1].text = str(c[2])
            row_cells[2].text = str(c[5])
            row_cells[3].text = str(c[10])
            row_cells[4].text = str(c['assign'])
            row_cells[5].text = str(c['quiz'])
            row_cells[6].text = str(c[2]+c[5]+c[10]+c['assign']+c['quiz'])
            tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
            tota += c['assign']; totq += c['quiz']

        # Total row
        tot_cells = table.add_row().cells
        for cell in tot_cells:
            _shade_cell(cell, "DCE6F1")
        tot_cells[0].text = "Total"
        tot_cells[1].text = str(tot2)
        tot_cells[2].text = str(tot5)
        tot_cells[3].text = str(tot10)
        tot_cells[4].text = str(tota)
        tot_cells[5].text = str(totq)
        tot_cells[6].text = str(tot2+tot5+tot10+tota+totq)
        for cell in tot_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    doc.save(output_path)


_UNICODE_SUBS = str.maketrans({
    '—': '--',   # em dash
    '–': '-',    # en dash
    '‘': "'",    # left single quote
    '’': "'",    # right single quote
    '“': '"',    # left double quote
    '”': '"',    # right double quote
    '…': '...',  # ellipsis
    '•': '-',    # bullet
    ' ': ' ',    # non-breaking space
    '→': '->',   # right arrow
    '←': '<-',   # left arrow
    '≤': '<=',   # less-or-equal
    '≥': '>=',   # greater-or-equal
    '×': 'x',    # multiplication sign
    '÷': '/',    # division sign
    'α': 'alpha',
    'β': 'beta',
    'γ': 'gamma',
    'δ': 'delta',
    'π': 'pi',
    'Σ': 'Sigma',
    'Ω': 'Omega',
})


def _pdf_safe(text):
    """Map common Unicode chars to Latin-1 equivalents; replace anything remaining."""
    text = text.translate(_UNICODE_SUBS)
    return text.encode('latin-1', errors='replace').decode('latin-1')


def save_pdf(blocks, output_path, co_tally=None, lab_tally=None):
    from fpdf import FPDF, XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    def _h1(text):
        pdf.set_font("Helvetica", "B", 15)
        pdf.set_fill_color(31, 73, 125)
        pdf.set_text_color(255, 255, 255)
        pdf.multi_cell(0, 10, text, fill=True,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def _h2(text):
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(220, 230, 241)
        pdf.multi_cell(0, 8, text, fill=True,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _h3(text):
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_fill_color(240, 245, 240)
        pdf.multi_cell(0, 7, text, fill=True,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _mark_hdr(text):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(200, 220, 200)
        pdf.cell(0, 7, "  " + text, fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _assign_hdr(text):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(255, 210, 160)  # amber
        pdf.cell(0, 7, "  " + text, fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _quiz_hdr(text):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(180, 225, 180)  # light green
        pdf.cell(0, 7, "  " + text, fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _labelled_q(text, label_re, label_r, label_g, label_b):
        """Render a labelled question line (Q or A) with optional CO tag."""
        m = re.match(label_re + r'\.\s*(.*)', text, re.DOTALL)
        label = (m.group(1) + ".  ") if m else ""
        body  = m.group(2).strip() if m else text

        co_tag = ""
        co_m = re.search(r'\s*(\[CO\d+\])\s*$', body)
        if co_m:
            co_tag = co_m.group(1)
            body   = body[:co_m.start()].strip()

        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(label_r, label_g, label_b)
        lw = pdf.get_string_width(label)
        pdf.cell(lw, 6, label)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)
        full_line = (body + "  " + co_tag) if co_tag else body
        pdf.multi_cell(0, 6, full_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _question(text):
        _labelled_q(text, r'(Q\d+)', 0, 0, 0)

    def _assignment_q(text):
        _labelled_q(text, r'(A\d+)', 191, 71, 0)  # amber label

    def _lab_exp(text):
        _labelled_q(text, r'(Exp\s+\d+)', 75, 0, 130)  # indigo label

    def _lab_hdr_pdf(text):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(216, 191, 216)  # light purple
        pdf.cell(0, 7, "  " + text, fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def _lab_section_label(text):
        label, _, rest = text.partition(':')
        pdf.set_font("Helvetica", "B", 9)
        lw = pdf.get_string_width(label + ': ')
        pdf.cell(lw, 5, label + ': ')
        pdf.set_font("Helvetica", "", 9)
        if rest.strip():
            pdf.multi_cell(0, 5, rest.strip(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.ln(5)

    for block in blocks:
        for raw in block.split('\n'):
            line = _pdf_safe(raw.strip())
            if not line or line in ('---', '***', '___'):
                pdf.ln(2)
                continue
            if line.startswith('# '):
                _h1(line[2:])
            elif line.startswith('## '):
                _h2(line[3:])
            elif line.startswith('### '):
                _h3(line[4:])
            elif re.match(r'\[\d+ Marks?\]', line):
                _mark_hdr(line)
            elif re.match(r'\[Assignment\b', line, re.IGNORECASE):
                _assign_hdr(line)
            elif re.match(r'\[Quiz\b', line, re.IGNORECASE):
                _quiz_hdr(line)
            elif re.match(r'\[Experiments?\b', line, re.IGNORECASE):
                _lab_hdr_pdf(line)
            elif re.match(r'Exp\s+\d+\.', line, re.IGNORECASE):
                _lab_exp(line)
            elif re.match(r'^(Aim|Equipment/Tools|Theory|Procedure|Expected Observations/Output|Result|Viva Questions)\s*:', line):
                _lab_section_label(line)
            elif re.match(r'^V\d+\.', line):
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(8, 5, "")  # indent
                pdf.multi_cell(0, 5, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            elif re.match(r'A\d+\.', line):
                _assignment_q(line)
            elif re.match(r'Q\d+\.', line):
                _question(line)
            elif re.match(r'^\([a-d]\)', line):
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(8, 5, "")  # indent
                pdf.multi_cell(0, 5, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(0, 5, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if lab_tally:
        pdf.add_page()
        _h2("CO Coverage Summary")
        pdf.ln(2)
        col_widths = [30, 40]
        headers    = ["CO", "Experiments"]
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(75, 0, 130)
        pdf.set_text_color(255, 255, 255)
        for w, h in zip(col_widths, headers):
            pdf.cell(w, 8, h, border=1, fill=True)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        tot_exp = 0
        for idx, (cn, c) in enumerate(lab_tally.items()):
            fill = idx % 2 == 0
            pdf.set_fill_color(240, 230, 255) if fill else pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Helvetica", "", 9)
            for w, v in zip(col_widths, [f"CO{cn}", str(c['exp'])]):
                pdf.cell(w, 7, v, border=1, fill=fill)
            pdf.ln()
            tot_exp += c['exp']
        pdf.set_fill_color(216, 191, 216)
        pdf.set_font("Helvetica", "B", 9)
        for w, v in zip(col_widths, ["Total", str(tot_exp)]):
            pdf.cell(w, 8, v, border=1, fill=True)
        pdf.ln()

    elif co_tally:
        pdf.add_page()
        _h2("CO Coverage Summary")
        pdf.ln(2)

        col_widths = [22, 24, 24, 28, 24, 24, 24]
        headers    = ["CO", "2-mark", "5-mark", "10-mark", "Assign", "Quiz", "Total"]

        # Header row
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(31, 73, 125)
        pdf.set_text_color(255, 255, 255)
        for w, h in zip(col_widths, headers):
            pdf.cell(w, 8, h, border=1, fill=True)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)

        tot2 = tot5 = tot10 = tota = totq = 0
        for idx, (cn, c) in enumerate(co_tally.items()):
            fill = idx % 2 == 0
            pdf.set_fill_color(240, 245, 255) if fill else pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Helvetica", "", 9)
            vals = [f"CO{cn}", str(c[2]), str(c[5]), str(c[10]),
                    str(c['assign']), str(c['quiz']),
                    str(c[2]+c[5]+c[10]+c['assign']+c['quiz'])]
            for w, v in zip(col_widths, vals):
                pdf.cell(w, 7, v, border=1, fill=fill)
            pdf.ln()
            tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
            tota += c['assign']; totq += c['quiz']

        # Total row
        pdf.set_fill_color(220, 230, 241)
        pdf.set_font("Helvetica", "B", 9)
        for w, v in zip(col_widths, ["Total", str(tot2), str(tot5), str(tot10),
                                      str(tota), str(totq),
                                      str(tot2+tot5+tot10+tota+totq)]):
            pdf.cell(w, 8, v, border=1, fill=True)
        pdf.ln()

    pdf.output(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python generate_qbank.py <syllabus_file>")
        print("  python generate_qbank.py <syllabus_file> <co_file.txt>")
        sys.exit(1)

    syllabus_path = sys.argv[1]
    co_path       = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(syllabus_path):
        print(f"Error: File not found: {syllabus_path}")
        sys.exit(1)
    if co_path and not os.path.exists(co_path):
        print(f"Error: CO file not found: {co_path}")
        sys.exit(1)

    # ---------- Extract and split syllabus ----------
    print(f"\nExtracting text from: {syllabus_path}")
    syllabus_text = extract(syllabus_path)
    if not syllabus_text.strip():
        print("Error: No text could be extracted.")
        sys.exit(1)

    courses = split_into_courses(syllabus_text)
    if not courses:
        co_file_m = re.search(
            r'[_\-]([A-Z0-9]{4,20})_COs\.(pdf|docx|txt)$', syllabus_path, re.IGNORECASE
        )
        if co_file_m:
            print(f"Error: '{os.path.basename(syllabus_path)}' is a CO file, not a syllabus.")
            print(f"Usage: python generate_qbank.py <syllabus_file> [co_file]")
        else:
            print("Error: No course codes detected in the syllabus.")
        sys.exit(1)

    # ---------- Auto-detect course code from CO file ----------
    auto_code   = None
    auto_source = ""

    if co_path:
        # Strategy 1: filename pattern  …_23CS8PENES_COs.pdf
        m = re.search(r'[_\-]([A-Z0-9]{4,20})_COs\.', os.path.basename(co_path), re.IGNORECASE)
        if m:
            candidate = m.group(1).upper()
            if candidate in courses:
                auto_code   = candidate
                auto_source = "filename"

        # Strategy 2: scan CO file content
        if not auto_code:
            try:
                co_ext = os.path.splitext(co_path)[1].lower()
                if co_ext == '.pdf':
                    import pdfplumber
                    raw_peek = ""
                    with pdfplumber.open(co_path) as _pdf:
                        for _pg in _pdf.pages[:3]:
                            _t = _pg.extract_text()
                            if _t:
                                raw_peek += _t + "\n"
                else:
                    with open(co_path, encoding='utf-8', errors='replace') as _f:
                        raw_peek = _f.read(8000)

                # 2a: explicit label patterns
                for pat in [
                    r'Course\s+Code\s*[:\-]\s*([A-Z0-9]{4,20})',
                    r'##\s+([A-Z0-9]{4,20})\s*[:\-]',
                    r'Course\s*:\s*([A-Z0-9]{4,20})\b',
                ]:
                    cm = re.search(pat, raw_peek, re.IGNORECASE)
                    if cm:
                        candidate = cm.group(1).upper()
                        if candidate in courses:
                            auto_code   = candidate
                            auto_source = "file content (label)"
                            break

                # 2b: any known course code appearing in the text
                if not auto_code:
                    for code in courses:
                        if re.search(r'\b' + re.escape(code) + r'\b', raw_peek):
                            auto_code   = code
                            auto_source = "file content (code match)"
                            break
            except Exception:
                pass

    # ---------- List and select course ----------
    print("\n" + "=" * 60)
    print("  Available Courses")
    print("=" * 60)
    for code, info in courses.items():
        sem = f"  [Sem {info['semester']}]" if info['semester'] else ""
        print(f"  {code:<20} {info['title'][:45]}{sem}")
    print("=" * 60)

    if auto_code:
        matched = auto_code
        print(f"\nAuto-detected course [{auto_source}]: {matched}")
    else:
        raw_code = input("\nEnter Course Code: ").strip().upper()
        if raw_code in courses:
            matched = raw_code
        else:
            hits = [c for c in courses if raw_code in c]
            if len(hits) == 1:
                matched = hits[0]
            elif len(hits) > 1:
                print(f"Multiple matches: {', '.join(hits)}. Please be more specific.")
                sys.exit(1)
            else:
                print(f"Course code '{raw_code}' not found.")
                sys.exit(1)

    info      = courses[matched]
    sem_label = f"  [Semester {info['semester']}]" if info['semester'] else ""
    print(f"\nSelected: {matched} — {info['title']}{sem_label}")

    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"], timeout=120.0)

    # ---------- Load or generate COs ----------
    if co_path:
        print(f"\nLoading CO file: {co_path}")
        co_ext = os.path.splitext(co_path)[1].lower()
        if co_ext == '.pdf':
            cos = load_cos_from_pdf(co_path, matched)
        else:
            raw_co  = extract(co_path)
            co_text = _extract_co_section(raw_co, matched)
            cos     = parse_cos_from_text(co_text)
        print(f"Loaded from: {os.path.basename(co_path)}")
    else:
        NUM_COS = 15
        print(f"\nNo CO file provided — generating {NUM_COS} COs for {matched}...")
        print("=" * 60)
        co_text = generate_cos_for_course(
            client, matched, info['title'], info['text'], num_cos=NUM_COS
        )
        print("=" * 60)
        print("Enter CO numbers to keep (e.g.  1,3,5  |  1-5  |  1,3-7,10  |  all) [default: all]:")
        raw_sel  = input("Selection: ").strip()
        selected = parse_co_selection(raw_sel, NUM_COS)
        if not selected:
            print("No valid COs selected. Exiting.")
            sys.exit(1)
        co_text = filter_cos(co_text, selected)
        cos = parse_cos_from_text(co_text)

    if not cos:
        print("Error: No COs could be parsed. Check the CO file format.")
        sys.exit(1)

    print(f"\nFound {len(cos)} Course Outcome(s):")
    print(f"  {'CO':<6} {'Unit':<12} {'Knowledge Dim.':<18} {'Bloom':<22} Statement")
    print(f"  {'-'*6} {'-'*12} {'-'*18} {'-'*22} {'-'*40}")
    for co in cos:
        u = co['unit']  or '(none)'
        k = co['kdim']  or '(none)'
        b = co['bloom'] or '(none)'
        print(f"  CO{co['num']:<4} {u:<12} {k:<18} {b:<22} {co['statement'][:45]}...")

    # Build unit plan (all syllabus units, with or without CO mapping)
    unit_plan = build_unit_plan(client, cos, info['text'])
    num_units = len(unit_plan)
    print(f"\nOrganised into {num_units} unit(s):")
    for ulabel, udata in unit_plan.items():
        co_nums = ', '.join(f"CO{c['num']}" for c in udata['cos'])
        co_tag  = f"  [{co_nums}]" if udata['cos'] else "  [no COs mapped]"
        title_tag = f" — {udata['title']}" if udata['title'] else ""
        print(f"  {ulabel}{title_tag}{co_tag}")

    # ---------- Lab course detection ----------
    lab_mode = is_lab_course(matched, info['title'])
    if lab_mode:
        print(f"\n  Lab/Practical course detected — switching to experiment generation mode.")

    def ask_int(prompt, default):
        raw = input(f"  {prompt} [default {default}]: ").strip()
        try:
            v = int(raw)
            return max(0, v)
        except ValueError:
            return default

    # ---------- Output format ----------
    print("\nOutput format:")
    print("  1. txt")
    print("  2. docx  (Word)")
    print("  3. pdf")
    fmt_input = input("Enter 1, 2, or 3 [default: 1]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_input, "txt")
    if fmt_input and fmt_input not in ("1", "2", "3"):
        print("Invalid choice — using txt.")

    # ==========================================================================
    # LAB COURSE PATH
    # ==========================================================================
    if lab_mode:
        print("\n" + "=" * 60)
        print("Lab Experiment Distribution per Unit  (press Enter for defaults)")
        print("=" * 60)
        n_exp = ask_int("Experiments per unit", 3)
        if n_exp < 1:
            n_exp = 3

        print(f"\n  → {n_exp} experiment(s) × {num_units} unit(s) = {n_exp * num_units} experiments total")
        print(f"\nGenerating lab question bank (organised by unit)...")
        print("=" * 60)

        all_blocks = []
        all_blocks.append(
            f"# Lab Question Bank\n"
            f"Course: {matched} — {info['title']}{sem_label}\n"
            f"Experiments per unit: {n_exp}  |  {n_exp * num_units} total\n"
        )

        for i, (unit_label, udata) in enumerate(unit_plan.items(), 1):
            unit_title = udata['title']
            unit_cos   = udata['cos']

            full_unit  = f"{unit_label}: {unit_title}" if unit_title else unit_label
            co_summary = "  |  ".join(f"CO{c['num']}" for c in unit_cos) or "no COs mapped"

            header = f"\n## {full_unit}\n"
            if unit_cos:
                header += f"### Course Outcomes: {co_summary}\n"
                for co in unit_cos:
                    header += f"### CO{co['num']}: {co['statement']}\n"
            else:
                header += f"### Experiments based on syllabus content\n"
            all_blocks.append(header)

            print(f"\n[{i}/{num_units}] {full_unit}  ({co_summary})")
            print("-" * 60)

            etext = generate_lab_experiments_for_unit(
                client, unit_label, unit_title, unit_cos,
                matched, info['title'], info['text'], n_exp=n_exp,
            )
            etext = _normalize_lab_block(etext, unit="")
            if not etext.strip():
                print(f"  WARNING: no experiments generated for {full_unit} — model may have refused. Skipping.")
            all_blocks.append(etext + "\n")

        # ---------- Save (lab) ----------
        base_path   = os.path.splitext(syllabus_path)[0]
        output_path = f"{base_path}_{matched}_LabQuestionBank.{fmt}"

        lab_tally = count_co_lab(all_blocks)

        if fmt == "txt":
            save_txt(all_blocks, output_path, lab_tally=lab_tally)
        elif fmt == "docx":
            save_docx(all_blocks, output_path, lab_tally=lab_tally)
        elif fmt == "pdf":
            save_pdf(all_blocks, output_path, lab_tally=lab_tally)

        total_exp = sum(
            1 for b in all_blocks
            for line in b.split('\n')
            if re.match(r'Exp\s+\d+\.', line.strip(), re.IGNORECASE)
        )
        print(f"\nSaved {total_exp} experiment(s) to: {output_path}")

        if lab_tally:
            print("\nCO Coverage Summary:")
            print(f"  {'CO':<8} {'Experiments':>13}")
            print(f"  {'-'*8} {'-'*13}")
            tot_exp = 0
            for cn, c in lab_tally.items():
                print(f"  {'CO'+cn:<8} {c['exp']:>13}")
                tot_exp += c['exp']
            print(f"  {'-'*8} {'-'*13}")
            print(f"  {'Total':<8} {tot_exp:>13}")
        return

    # ==========================================================================
    # REGULAR COURSE PATH
    # ==========================================================================
    print("\n" + "=" * 60)
    print("Question Distribution per Unit  (press Enter for defaults)")
    print("=" * 60)

    n2       = ask_int("2-mark  questions per unit", 4)
    n5       = ask_int("5-mark  questions per unit", 4)
    n10      = ask_int("10-mark questions per unit", 4)
    n_assign = ask_int("Assignment tasks  per unit (0 to skip)", 2)
    n_quiz   = ask_int("Quiz questions    per unit (0 to skip)", 5)
    if n2 < 1: n2 = 1
    if n5 < 1: n5 = 1
    if n10 < 1: n10 = 1
    total_per_unit = n2 + n5 + n10

    print(f"\n  → {total_per_unit} exam questions × {num_units} unit(s) = "
          f"{total_per_unit * num_units} exam questions total"
          + (f"  |  {n_assign} assignment(s) per unit" if n_assign else "")
          + (f"  |  {n_quiz} quiz question(s) per unit" if n_quiz else ""))

    # ---------- Generate ----------
    print(f"\nGenerating question bank (organised by unit)...")
    print("=" * 60)

    all_blocks = []
    extras = []
    if n_assign:
        extras.append(f"{n_assign} assignment task(s)")
    if n_quiz:
        extras.append(f"{n_quiz} quiz question(s)")
    extra_line = ("  |  " + " + ".join(extras) + " per unit\n") if extras else "\n"
    all_blocks.append(
        f"# Question Bank\n"
        f"Course: {matched} — {info['title']}{sem_label}\n"
        f"Exam questions per unit: {n2} x 2-mark + {n5} x 5-mark + {n10} x 10-mark"
        f" = {total_per_unit} per unit  |  {total_per_unit * num_units} total\n"
        + extra_line
    )

    for i, (unit_label, udata) in enumerate(unit_plan.items(), 1):
        unit_title = udata['title']
        unit_cos   = udata['cos']

        full_unit  = f"{unit_label}: {unit_title}" if unit_title else unit_label
        co_summary = "  |  ".join(f"CO{c['num']}" for c in unit_cos) or "no COs mapped"

        header = f"\n## {full_unit}\n"
        if unit_cos:
            header += f"### Course Outcomes: {co_summary}\n"
            for co in unit_cos:
                header += f"### CO{co['num']}: {co['statement']}\n"
        else:
            header += f"### Questions based on syllabus content\n"
        all_blocks.append(header)

        print(f"\n[{i}/{num_units}] {full_unit}  ({co_summary})")
        print("-" * 60)

        qtext = generate_questions_for_unit(
            client, unit_label, unit_title, unit_cos,
            matched, info['title'], info['text'],
            n2=n2, n5=n5, n10=n10,
        )
        qtext = _normalize_qblock(qtext, unit="")
        if not qtext.strip():
            print(f"  WARNING: no questions generated for {full_unit} — model may have refused. Skipping.")
        all_blocks.append(qtext + "\n")

        if n_assign:
            print(f"\n  Generating {n_assign} assignment task(s)...")
            print("-" * 60)
            atext = generate_assignments_for_unit(
                client, unit_label, unit_title, unit_cos,
                matched, info['title'], info['text'], n_assign=n_assign,
            )
            atext = _normalize_assignment_block(atext)
            if not atext.strip():
                print(f"  WARNING: no assignments generated for {full_unit}. Skipping.")
            else:
                all_blocks.append(atext + "\n")

        if n_quiz:
            print(f"\n  Generating {n_quiz} quiz question(s)...")
            print("-" * 60)
            qztext = generate_quiz_for_unit(
                client, unit_label, unit_title, unit_cos,
                matched, info['title'], info['text'], n_quiz=n_quiz,
            )
            qztext = _normalize_quiz_block(qztext)
            if not qztext.strip():
                print(f"  WARNING: no quiz questions generated for {full_unit}. Skipping.")
            else:
                all_blocks.append(qztext + "\n")

    # ---------- Save ----------
    base_path   = os.path.splitext(syllabus_path)[0]
    output_path = f"{base_path}_{matched}_QuestionBank.{fmt}"

    co_tally = count_co_questions(all_blocks)

    if fmt == "txt":
        save_txt(all_blocks, output_path, co_tally=co_tally)
    elif fmt == "docx":
        save_docx(all_blocks, output_path, co_tally=co_tally)
    elif fmt == "pdf":
        save_pdf(all_blocks, output_path, co_tally=co_tally)

    total_q = sum(question_summary(b)[0] for b in all_blocks)
    print(f"\nSaved {total_q} questions to: {output_path}")

    if co_tally:
        print("\nCO Coverage Summary:")
        print(f"  {'CO':<8} {'2-mark':>8} {'5-mark':>8} {'10-mark':>9} {'Assign':>8} {'Quiz':>8} {'Total':>8}")
        print(f"  {'-'*8} {'-'*8} {'-'*8} {'-'*9} {'-'*8} {'-'*8} {'-'*8}")
        tot2 = tot5 = tot10 = tota = totq = 0
        for cn, c in co_tally.items():
            row_tot = c[2]+c[5]+c[10]+c['assign']+c['quiz']
            print(f"  {'CO'+cn:<8} {c[2]:>8} {c[5]:>8} {c[10]:>9} {c['assign']:>8} {c['quiz']:>8} {row_tot:>8}")
            tot2 += c[2]; tot5 += c[5]; tot10 += c[10]
            tota += c['assign']; totq += c['quiz']
        print(f"  {'-'*8} {'-'*8} {'-'*8} {'-'*9} {'-'*8} {'-'*8} {'-'*8}")
        print(f"  {'Total':<8} {tot2:>8} {tot5:>8} {tot10:>9} {tota:>8} {totq:>8} {tot2+tot5+tot10+tota+totq:>8}")


if __name__ == "__main__":
    main()
