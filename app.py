import os
import re
import json
import uuid
import time
import secrets
import hashlib
import tempfile
import threading
import resend
from functools import wraps

from flask import Flask, render_template, request, jsonify, Response, session, redirect, url_for, flash, make_response, after_this_request, stream_with_context
from werkzeug.security import check_password_hash, generate_password_hash
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

from generate_cos import (
    extract, split_into_courses,
    parse_co_selection, filter_cos,
    bloom_level_summary, build_taxonomy_grid,
    save_txt as cos_save_txt,
    save_docx as cos_save_docx,
    save_pdf as cos_save_pdf,
    generate_cos_stream,
    generate_cos_for_course,
)

from generate_qbank import (
    parse_cos_from_text, load_cos_from_pdf, build_unit_plan,
    generate_questions_for_unit, generate_assignments_for_unit, generate_quiz_for_unit,
    generate_lab_experiments_for_unit,
    _normalize_qblock, _normalize_assignment_block, _normalize_quiz_block,
    _normalize_lab_block,
    is_lab_course,
    count_co_questions, count_co_lab,
    save_txt as qb_save_txt,
    save_docx as qb_save_docx,
    save_pdf as qb_save_pdf,
)

from generate_qpaper import (
    parse_qbank, _flat_list,
    is_lab_qbank, parse_lab_qbank, _flat_lab_list,
    build_assignment_paper, build_quiz_paper, build_exam_paper,
    build_lab_practical_paper, build_lab_viva_paper,
    save_txt as qp_save_txt,
    save_docx as qp_save_docx,
    save_pdf as qp_save_pdf,
    _load_qbank_text,
    _count_co_coverage,
    _count_co_lab_coverage,
    _pdf_safe,
)

from generate_po_mapping import (
    parse_cos_for_mapping, parse_course_header, parse_pos_from_text,
    generate_mapping_stream, parse_mapping_response,
    POS as _ENG_POS,           # kept for backward compat
    POS_BY_TYPE, COLLEGE_TYPE_META, get_pos_for_type,
    save_txt as po_save_txt,
    save_docx as po_save_docx,
    save_pdf  as po_save_pdf,
)
# Dynamic helper — always reads from current user's session
def _get_standard_pos():
    return get_pos_for_type(session.get("college_type", "engineering"))

STANDARD_POS = _ENG_POS   # legacy alias used in a few places; prefer _get_standard_pos()

from generate_template import (
    parse_template_paper,
    build_docx as tpl_build_docx,
    build_pdf  as tpl_build_pdf,
    build_txt  as tpl_build_txt,
)

from generate_lesson_plan import (
    generate_lesson_plan_stream,
    build_docx as lp_build_docx,
    build_pdf  as lp_build_pdf,
    build_txt  as lp_build_txt,
)

from generate_teaching_diary import (
    generate_teaching_diary_stream,
    build_docx as td_build_docx,
    build_pdf  as td_build_pdf,
    build_txt  as td_build_txt,
)

import generate_m1_report as _m1r

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", os.urandom(24))

# In-process store keyed by session ID  -  fine for single-user local tool
_store: dict = {}

_DATA_DIR   = os.environ.get("DATA_DIR", os.path.dirname(os.path.abspath(__file__)))
os.makedirs(_DATA_DIR, exist_ok=True)
USERS_FILE  = os.path.join(_DATA_DIR, "users.json")
TOKENS_FILE = os.path.join(_DATA_DIR, "reset_tokens.json")


def load_users() -> dict:
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, encoding="utf-8") as f:
            raw = json.load(f)
        needs_save = False
        for k, v in list(raw.items()):
            if isinstance(v, str):
                raw[k] = {"password": v, "tokens": 5, "college_type": "engineering"}
                needs_save = True
            elif isinstance(v, dict):
                if "tokens" not in v:
                    raw[k]["tokens"] = 5
                    needs_save = True
                if "college_type" not in v:
                    raw[k]["college_type"] = "engineering"
                    needs_save = True
        if needs_save:
            _save_users(raw)
        return raw
    return {}


# ── College context helpers ───────────────────────────────────────────────────

_COLLEGE_CONTEXTS = {
    "engineering": {
        "type": "Engineering", "discipline": "engineering",
        "accreditation": "NBA/NAAC (National Board of Accreditation)",
        "professional": "engineer", "program": "B.Tech / B.E.",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "NBA", "framework": "OBE/NBA",
    },
    "medical": {
        "type": "Medical", "discipline": "medicine",
        "accreditation": "NMC (National Medical Commission)",
        "professional": "physician/doctor", "program": "MBBS / MD",
        "outcomes_term": "Graduate Attributes / Programme Outcomes",
        "body": "NMC", "framework": "CBME/NMC",
    },
    "dental": {
        "type": "Dental", "discipline": "dentistry",
        "accreditation": "DCI (Dental Council of India)",
        "professional": "dental surgeon / dentist", "program": "BDS / MDS",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "DCI", "framework": "OBE/DCI",
    },
    "law": {
        "type": "Law", "discipline": "law",
        "accreditation": "BCI (Bar Council of India)",
        "professional": "advocate / lawyer", "program": "LLB / LLM",
        "outcomes_term": "Programme Outcomes / Graduate Attributes",
        "body": "BCI", "framework": "OBE/BCI",
    },
    "pharmacy": {
        "type": "Pharmacy", "discipline": "pharmacy",
        "accreditation": "PCI (Pharmacy Council of India)",
        "professional": "pharmacist", "program": "B.Pharm / M.Pharm",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "PCI", "framework": "OBE/PCI",
    },
    "management": {
        "type": "Management", "discipline": "management / business",
        "accreditation": "AICTE/NAAC",
        "professional": "manager / business professional", "program": "BBA / MBA",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "AICTE", "framework": "OBE/AICTE",
    },
    "architecture": {
        "type": "Architecture & Planning", "discipline": "architecture",
        "accreditation": "COA (Council of Architecture)",
        "professional": "architect / planner", "program": "B.Arch / M.Arch",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "COA", "framework": "OBE/COA",
    },
    "agriculture": {
        "type": "Agriculture & Allied Sciences", "discipline": "agriculture",
        "accreditation": "ICAR (Indian Council of Agricultural Research)",
        "professional": "agronomist / agricultural scientist", "program": "B.Sc Ag / M.Sc Ag",
        "outcomes_term": "Programme Outcomes (POs)",
        "body": "ICAR", "framework": "OBE/ICAR",
    },
}


def _get_college_context() -> dict:
    """Return college context dict for the current user's college_type."""
    ct = session.get("college_type", "engineering")
    return _COLLEGE_CONTEXTS.get(ct, _COLLEGE_CONTEXTS["engineering"])


def _ctx_prompt() -> str:
    """One-liner system-prompt prefix for AI calls: role + accreditation body."""
    c = _get_college_context()
    return (
        f"You are an OBE and accreditation expert for Indian {c['type']} colleges "
        f"({c['accreditation']}). You specialise in {c['framework']} compliance, "
        f"Bloom's taxonomy-aligned {c['outcomes_term']}, and quality education."
    )


def login_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login", next=request.path))
        return fn(*args, **kwargs)
    return wrapper


TOPUP_PACKAGES = [
    {"id": "starter", "label": "Starter", "tokens": 10,  "price": 100},
    {"id": "popular", "label": "Popular", "tokens": 60,  "price": 500},
    {"id": "pro",     "label": "Pro",     "tokens": 150, "price": 1000},
]


def tokens_required(cost=1):
    """Deducts `cost` tokens before running the view; requires login."""
    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            if not session.get("logged_in"):
                return redirect(url_for("login", next=request.path))
            username = session["username"]
            users = load_users()
            balance = users.get(username, {}).get("tokens", 0)
            if balance < cost:
                err = f"Not enough tokens. You need {cost} token(s) but have {balance}. Top up in Settings."
                if "text/event-stream" in request.headers.get("Accept", ""):
                    def _err_stream():
                        yield f"data: {json.dumps({'error': err})}\n\n"
                    return Response(_err_stream(), mimetype="text/event-stream",
                                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})
                return jsonify({"error": err}), 403
            users[username]["tokens"] = balance - cost
            _save_users(users)
            session["tokens"] = users[username]["tokens"]
            new_balance = users[username]["tokens"]

            @after_this_request
            def _inject_token_header(response):
                if response.content_type != "text/event-stream":
                    response.headers["X-Tokens-Remaining"] = str(new_balance)
                return response

            return fn(*args, **kwargs)
        return wrapper
    return decorator


def _refund_tokens(cost):
    """Refund tokens to the current user when SSE generation fails mid-stream."""
    username = session.get("username")
    if not username or cost <= 0:
        return
    users = load_users()
    if username in users:
        users[username]["tokens"] = users[username].get("tokens", 0) + cost
        _save_users(users)


def _charge_report(module, sig_material, cost=2):
    """Charge `cost` tokens for generating a comprehensive module report.

    A comprehensive report is generated once for the editable preview and again
    each time the user downloads it (DOCX/PDF). We charge only the first time a
    given report is produced: re-rendering the same report (identical data
    signature) in the same session — whether re-previewing or downloading in
    any format — is free. Changing the underlying data (which requires paid
    regeneration of the deliverables) produces a new signature and is charged.

    Returns ``(charged, error_response)`` — ``error_response`` is non-None only
    when the user has too few tokens (the view should return it directly).
    """
    sig = hashlib.md5(sig_material.encode("utf-8", "ignore")).hexdigest()
    key = f"{module}_report_sig"
    if session.get(key) == sig:
        return False, None
    username = session.get("username")
    users = load_users()
    balance = users.get(username, {}).get("tokens", 0)
    if balance < cost:
        err = f"Not enough tokens. You need {cost} token(s) but have {balance}. Top up in Settings."
        return False, (jsonify({"error": err}), 403)
    users[username]["tokens"] = balance - cost
    _save_users(users)
    session["tokens"] = users[username]["tokens"]
    session[key] = sig
    return True, None


def _sid():
    if "sid" not in session:
        session["sid"] = str(uuid.uuid4())
    return session["sid"]


def _get_store():
    return _store.setdefault(_sid(), {})


# â"€â"€ Auth routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

def _save_users(users: dict) -> None:
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, indent=2)


# â"€â"€ Password-reset token helpers â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

def _load_tokens() -> dict:
    if os.path.exists(TOKENS_FILE):
        with open(TOKENS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _save_tokens(tokens: dict) -> None:
    with open(TOKENS_FILE, "w", encoding="utf-8") as f:
        json.dump(tokens, f)


def _purge_expired(tokens: dict) -> dict:
    now = time.time()
    return {k: v for k, v in tokens.items() if v["expires"] > now}


def _send_reset_email(to_email: str, reset_url: str) -> bool:
    api_key = os.environ.get("RESEND_API_KEY", "").strip()
    frm     = os.environ.get("SMTP_FROM", "MyOBE <help@myobe.in>").strip()
    if not api_key:
        app.logger.error("RESEND_API_KEY not set")
        return False

    plain = (
        "You requested a password reset for your MyOBE account.\n\n"
        f"Reset link (expires in 1 hour):\n{reset_url}\n\n"
        "If you didn't request this, you can safely ignore this email."
    )
    html = (
        "<p>You requested a password reset for your <strong>MyOBE</strong> account.</p>"
        f'<p><a href="{reset_url}" style="display:inline-block;padding:10px 24px;'
        'background:#4f46e5;color:#fff;text-decoration:none;border-radius:8px;font-weight:700;">'
        "Reset my password</a></p>"
        '<p style="color:#64748b;font-size:.85em;">Link expires in 1 hour. '
        "If you didn't request this, ignore this email.</p>"
    )

    try:
        resend.api_key = api_key
        resend.Emails.send({
            "from":    frm,
            "to":      [to_email],
            "subject": "Reset your MyOBE password",
            "html":    html,
            "text":    plain,
        })
        return True
    except Exception as e:
        app.logger.error("Resend error: %s", e)
        return False


@app.route("/login", methods=["GET", "POST"])
def login():
    if session.get("logged_in"):
        return redirect(url_for("dashboard"))

    error = None
    # Tab and pre-fill state passed back from /register on validation failure
    tab          = request.args.get("tab", "signin")
    reg_error    = request.args.get("reg_error")
    reg_username = request.args.get("reg_username", "")
    reg_email    = request.args.get("reg_email", "")
    reset_ok     = request.args.get("reset") == "1"

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        users    = load_users()

        if not users:
            error = "No users yet  -  use the Register tab to create an account."
        elif username not in users:
            error = "Invalid username or password."
        elif not check_password_hash(users[username]["password"], password):
            error = "Invalid username or password."
        else:
            session.clear()
            session["logged_in"]    = True
            session["username"]     = username
            session["tokens"]       = users[username].get("tokens", 0)
            session["college_type"] = users[username].get("college_type", "engineering")
            return redirect(request.args.get("next") or url_for("dashboard"))

    resp = make_response(render_template(
        "login.html",
        error=error,
        tab=tab,
        reg_error=reg_error,
        reg_username=reg_username,
        reg_email=reg_email,
        reset_ok=reset_ok,
    ))
    resp.headers["Cache-Control"] = "no-store"
    return resp


@app.route("/register", methods=["POST"])
def register():
    username = request.form.get("username", "").strip()
    email    = request.form.get("email",    "").strip().lower()
    password = request.form.get("password", "")
    confirm  = request.form.get("confirm",  "")

    def fail(msg):
        return redirect(url_for("login", tab="register", reg_error=msg,
                                reg_username=username, reg_email=email))

    if not username:
        return fail("Username cannot be empty.")
    if len(username) < 3:
        return fail("Username must be at least 3 characters.")
    if not all(c.isalnum() or c in "-_" for c in username):
        return fail("Username may only contain letters, numbers, - and _.")
    if not email or "@" not in email or "." not in email.split("@")[-1]:
        return fail("Please enter a valid email address.")
    if len(password) < 6:
        return fail("Password must be at least 6 characters.")
    if password != confirm:
        return fail("Passwords do not match.")

    users = load_users()
    if username in users:
        return fail(f"Username '{username}' is already taken.")
    if any(d.get("email", "").lower() == email for d in users.values() if isinstance(d, dict)):
        return fail("An account with that email address already exists.")

    college_type = request.form.get("college_type", "engineering").strip()
    if college_type not in _COLLEGE_CONTEXTS:
        college_type = "engineering"

    users[username] = {
        "password":     generate_password_hash(password),
        "tokens":       5,
        "email":        email,
        "college_type": college_type,
    }
    _save_users(users)

    session.clear()
    session["logged_in"]    = True
    session["username"]     = username
    session["tokens"]       = 5
    session["college_type"] = college_type
    return redirect(url_for("dashboard"))


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# -- Module 1: CO Generation routes ------------------------------------------

@app.route("/")
def home():
    return render_template("home.html", logged_in=session.get("logged_in", False))


@app.route("/app")
@login_required
def dashboard():
    ct = session.get("college_type", "engineering")
    ct_meta = COLLEGE_TYPE_META.get(ct, COLLEGE_TYPE_META["engineering"])
    return render_template("index.html", username=session.get("username"),
                           college_type=ct,
                           college_type_label=ct_meta["label"],
                           college_type_accreditation=ct_meta["accreditation"])


@app.route("/upload", methods=["POST"])
@login_required
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".doc", ".txt"):
        return jsonify({"error": f"Unsupported file type '{ext}'"}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    text = None
    try:
        f.save(tmp.name)
        tmp.close()
        text = extract(tmp.name)
    except BaseException:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        return jsonify({"error": "Failed to read the file. It may be corrupted or use an unsupported format."}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not text or not text.strip():
        return jsonify({"error": "No text could be extracted from the file"}), 400

    courses = split_into_courses(text)

    store = _get_store()
    store["raw_syllabus_text"] = text  # keep for AI fallback

    if not courses:
        return jsonify({"error": "No course codes detected in the file.", "ai_fallback": True}), 400

    store["courses"] = courses

    return jsonify({
        "courses": [
            {
                "code":     code,
                "title":    info["title"],
                "semester": info["semester"],
                "has_lab":  info.get("has_lab", False),
            }
            for code, info in courses.items()
        ]
    })


@app.route("/upload_ai_extract", methods=["POST"])
@login_required
def upload_ai_extract():
    store = _get_store()
    text = store.get("raw_syllabus_text", "")
    if not text:
        return jsonify({"error": "No syllabus text found. Please re-upload your file."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "AI service is not configured."}), 500

    truncated = text[:15000]
    prompt = (
        "You are a university syllabus parser. Extract every course/subject listed in this syllabus document.\n\n"
        "For each course output exactly one line:\n"
        "CODE | Course Title | Semester Number\n\n"
        "Rules:\n"
        "- CODE: the alphanumeric course/subject code (e.g. CS301, 23CS32, MAT-101)\n"
        "- Course Title: the subject name (max 80 chars)\n"
        "- Semester Number: integer 1-10, or leave blank if not stated\n"
        "- Output ONLY data lines — no headers, no explanations, no extra text\n\n"
        "Syllabus text:\n" + truncated
    )

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key, timeout=30.0)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
    except Exception as e:
        return jsonify({"error": f"AI extraction failed: {str(e)}"}), 500

    courses = {}
    for line in raw.splitlines():
        line = line.strip()
        if not line or "|" not in line:
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 2 or not parts[0] or not parts[1]:
            continue
        code  = parts[0][:20]
        title = parts[1][:100]
        semester = None
        if len(parts) > 2 and parts[2]:
            try:
                semester = int(parts[2])
            except ValueError:
                pass
        courses[code] = {"title": title, "semester": semester, "text": text, "has_lab": False}

    if not courses:
        return jsonify({"error": "No courses could be detected in this file. Please check the file and try again."}), 400

    store["courses"] = courses
    return jsonify({
        "courses": [
            {"code": code, "title": info["title"], "semester": info["semester"], "has_lab": False}
            for code, info in courses.items()
        ]
    })


def _parse_cos_from_raw(raw):
    """Parse CO statements from raw LLM output, joining wrapped continuation lines."""
    cos, seen = [], set()
    lines = [l.strip() for l in raw.replace('\r\n', '\n').split('\n')]
    i = 0
    while i < len(lines):
        line = re.sub(r'^\|?\s*', '', lines[i]).strip()
        line = re.sub(r'\s*\|.*$', '', line).strip()
        m = re.match(r'^CO(\d+)\s*[:\.\-–]?\s*(.+)', line, re.IGNORECASE)
        if m:
            num  = int(m.group(1))
            stmt = m.group(2).strip()
            j = i + 1
            while j < len(lines):
                nxt = re.sub(r'^\|?\s*', '', lines[j]).strip()
                nxt = re.sub(r'\s*\|.*$', '', nxt).strip()
                if (not nxt
                        or re.match(r'^CO\d+\s*[:\.\-]', nxt, re.IGNORECASE)
                        or re.match(r'^[-|=\s]+$', nxt)
                        or re.match(r'^(Unit|Bloom|BT\s*Level|Knowledge|PO|SDG|Mapping)', nxt, re.IGNORECASE)):
                    break
                stmt += ' ' + nxt
                j += 1
            stmt = re.sub(r'\s+', ' ', stmt).strip()
            if num not in seen and len(stmt) > 15:
                seen.add(num)
                cos.append({'num': num, 'statement': stmt})
            i = j
        else:
            i += 1
    return cos


@app.route("/generate")
@tokens_required(1)
def generate():
    store = _get_store()
    courses = store.get("courses", {})

    code = request.args.get("code", "").strip().upper()
    num_cos = max(3, min(15, int(request.args.get("num_cos", 5))))

    if code not in courses:
        return jsonify({"error": "Course not found"}), 400

    info = courses[code]
    store["current_code"] = code
    store["current_result"] = ""

    def event_stream():
        import anthropic
        client = anthropic.Anthropic(
            api_key=os.environ["ANTHROPIC_API_KEY"],
            timeout=90.0,   # 90 s per read  -  aborts if Anthropic stops sending chunks
            max_retries=3,
        )
        _success = False
        try:
            chunks = []
            try:
                for chunk in generate_cos_stream(
                    client, code, info["title"], info["text"], num_cos
                ):
                    chunks.append(chunk)
                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            except anthropic.APITimeoutError:
                yield f"data: {json.dumps({'error': 'Request timed out - check your internet connection and try again.'})}\n\n"
                return
            except anthropic.APIConnectionError:
                yield f"data: {json.dumps({'error': 'Could not reach the AI service - check your internet connection and try again.'})}\n\n"
                return
            except Exception as e:
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
                return

            result = "".join(chunks)
            store["current_result"] = result
            # Cache per-CO taxonomy so /export never needs to re-parse text
            _co_tax = {}
            for _line in result.split('\n'):
                _s = _line.strip()
                if not (_s.startswith('|') and _s.endswith('|')):
                    continue
                if re.match(r'\|[-| ]+\|', _s):
                    continue
                _cells = [c.strip() for c in _s.strip('|').split('|')]
                if len(_cells) >= 4:
                    _cm = re.match(r'^CO(\d+)$', _cells[0])
                    if _cm:
                        _bm = re.match(r'(L\d+)', _cells[-1])
                        if _bm:
                            _co_tax[int(_cm.group(1))] = {
                                'kdim': _cells[-2].strip(),
                                'bloom': _bm.group(1)
                            }
            store["co_taxonomy"] = _co_tax
            summary = bloom_level_summary(result)
            cos_parsed = _parse_cos_from_raw(result)
            _success = True
            yield f"data: {json.dumps({'done': True, 'summary': summary, 'cos': cos_parsed})}\n\n"
        finally:
            if not _success:
                _refund_tokens(1)

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/export", methods=["POST"])
@login_required
def export():
    store = _get_store()
    data = request.get_json()

    result_text = store.get("current_result", "")
    code = store.get("current_code", "COURSE")
    courses = store.get("courses", {})
    info = courses.get(code, {"title": "", "semester": None})

    if not result_text:
        return jsonify({"error": "No generated result. Please generate COs first."}), 400

    selected_raw = data.get("selected", "all")
    fmt          = data.get("fmt", "txt")
    co_rows_raw  = data.get("coRows", [])   # [{num, kdim, bloom}, ...] from frontend
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    total = len(re.findall(r"^CO\d+:", result_text, re.MULTILINE))
    selected = parse_co_selection(selected_raw, total)
    if not selected:
        return jsonify({"error": "No valid COs selected"}), 400

    filtered = filter_cos(result_text, selected)
    summary  = bloom_level_summary(filtered)

    # Build taxonomy grid — primary: cached server-side taxonomy (most reliable)
    from generate_cos import _KDIMS, _LEVELS, _KDIM_NORM
    grid = {kd: {lv: [] for lv in _LEVELS} for kd in _KDIMS}
    co_taxonomy = store.get("co_taxonomy", {})
    if co_taxonomy:
        for i, orig in enumerate(sorted(selected), 1):
            entry = co_taxonomy.get(orig)
            if not entry:
                continue
            kdim = _KDIM_NORM.get(entry['kdim'].lower(), entry['kdim'])
            bloom_lv = entry['bloom']
            if kdim in grid and bloom_lv in grid[kdim]:
                grid[kdim][bloom_lv].append(f'CO{i}')
    # Fallback 1: frontend coRows
    if not any(grid[kd][lv] for kd in grid for lv in grid[kd]):
        if co_rows_raw:
            new_num_map = {orig: i for i, orig in enumerate(sorted(selected), 1)}
            for row in co_rows_raw:
                orig     = row.get('num', 0)
                kdim_raw = row.get('kdim', '').strip()
                bloom    = row.get('bloom', '').strip()
                m        = re.match(r'(L\d+)', bloom)
                if not m:
                    continue
                kdim = _KDIM_NORM.get(kdim_raw.lower(), kdim_raw)
                if kdim in grid:
                    new_n = new_num_map.get(orig, orig)
                    grid[kdim][m.group(1)].append(f'CO{new_n}')
    # Fallback 2: re-parse from filtered/raw text
    if not any(grid[kd][lv] for kd in grid for lv in grid[kd]):
        grid = build_taxonomy_grid(filtered)
    if not any(grid[kd][lv] for kd in grid for lv in grid[kd]):
        grid = build_taxonomy_grid(result_text)

    sem_label = f"  [Semester {info['semester']}]" if info.get("semester") else ""
    all_output = [
        f"# Course Outcomes\nGenerated from syllabus\nCOs exported: {len(selected)}\n",
        f"## {code}: {info.get('title', '')}{sem_label}\n",
        (summary + "\n\n" + filtered if summary else filtered) + "\n",
    ]

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime = {
        "txt": "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            cos_save_txt(all_output, tmp.name, taxonomy_grid=grid)
        elif fmt == "docx":
            cos_save_docx(all_output, tmp.name, taxonomy_grid=grid)
        elif fmt == "pdf":
            cos_save_pdf(all_output, tmp.name, taxonomy_grid=grid)

        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{code}_COs{suffix}"'},
    )


# -- Module 1: Question Bank routes ------------------------------------------

@app.route("/upload_co_qbank", methods=["POST"])
@login_required
def upload_co_qbank():
    """Parse an uploaded CO file (txt/docx/pdf) and return its COs for review."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        if ext == ".pdf":
            cos     = load_cos_from_pdf(tmp.name)
            raw_txt = "\n".join(
                f"CO{co['num']}: {co['statement']}"
                + (f"\n| CO{co['num']} | {co['unit']} | {co['kdim']} | {co['bloom']} |"
                   if co.get("unit") else "")
                for co in cos
            )
        else:
            raw_txt = extract(tmp.name)
            cos     = parse_cos_from_text(raw_txt)
    except Exception as e:
        return jsonify({"error": f"Failed to parse CO file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not cos:
        return jsonify({"error": "No COs found in the uploaded file. Make sure it was generated by this tool."}), 400

    store = _get_store()
    store["qbank_co_raw"] = raw_txt   # used by /generate_qbank when selected_cos provided

    return jsonify({
        "cos": [
            {"num": co["num"], "statement": co["statement"],
             "unit": co.get("unit", ""), "kdim": co.get("kdim", ""), "bloom": co.get("bloom", "")}
            for co in cos
        ]
    })


@app.route("/export_qbank_cos", methods=["POST"])
@login_required
def export_qbank_cos():
    """Download the COs shown at the Select COs step (auto-generated or uploaded)."""
    store = _get_store()
    data  = request.get_json()
    fmt   = data.get("fmt", "txt")
    code  = (data.get("code") or store.get("current_code") or "COURSE").strip().upper()

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    # Prefer auto-generated result; fall back to uploaded raw text
    co_text = store.get("current_result") or store.get("qbank_co_raw", "")
    if not co_text.strip():
        return jsonify({"error": "No COs available. Please generate or upload COs first."}), 400

    # Filter to only the COs the user selected (comma-separated nums, e.g. "1,3,5")
    selected_cos_raw = data.get("selected_cos", "")
    if selected_cos_raw:
        try:
            nums = [int(x) for x in selected_cos_raw.split(",") if x.strip().isdigit()]
        except ValueError:
            nums = []
        if nums:
            co_text = filter_cos(co_text, nums)

    courses  = store.get("courses", {})
    info     = courses.get(code, {"title": "", "semester": None})
    sem_label = f"  [Semester {info['semester']}]" if info.get("semester") else ""
    summary  = bloom_level_summary(co_text)
    grid     = build_taxonomy_grid(co_text)

    all_output = [
        f"# Course Outcomes\nCourse: {code}  -  {info.get('title', '')}{sem_label}\n",
        (summary + "\n\n" + co_text if summary else co_text) + "\n",
    ]

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            cos_save_txt(all_output, tmp.name, taxonomy_grid=grid)
        elif fmt == "docx":
            cos_save_docx(all_output, tmp.name, taxonomy_grid=grid)
        elif fmt == "pdf":
            cos_save_pdf(all_output, tmp.name, taxonomy_grid=grid)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{code}_COs{suffix}"'},
    )


@app.route("/generate_qbank")
@tokens_required(3)
def generate_qbank_route():
    store = _get_store()
    courses = store.get("courses", {})

    code             = request.args.get("code", "").strip().upper()
    selected_cos_raw = request.args.get("selected_cos", "").strip()   # "1,2,3" or ""
    try:
        n2       = max(1, min(20, int(request.args.get("n2",       4))))
        n5       = max(1, min(20, int(request.args.get("n5",       4))))
        n10      = max(1, min(20, int(request.args.get("n10",      4))))
        n_assign = max(0, min(10, int(request.args.get("n_assign", 2))))
        n_quiz   = max(0, min(20, int(request.args.get("n_quiz",   5))))
        n_exp    = max(1, min(10, int(request.args.get("n_exp",    3))))
    except (ValueError, TypeError):
        return jsonify({"error": "Invalid parameters"}), 400

    if code not in courses:
        return jsonify({"error": "Course not found"}), 400

    info    = courses[code]
    lab_mode = is_lab_course(code, info["title"], has_lab=info.get("has_lab", False))

    store["qbank_code"]       = code
    store["qbank_all_blocks"] = []
    store["qbank_co_tally"]   = {}
    store["qbank_is_lab"]     = lab_mode
    store["qbank_lab_tally"]  = {}

    # Resolve pre-selected COs (from uploaded file or from /generate CO stream)
    preloaded_cos = None
    co_source_key = store.get("qbank_co_raw") or store.get("current_result", "")
    if selected_cos_raw and co_source_key:
        selected_nums = set()
        for part in selected_cos_raw.split(","):
            try:
                selected_nums.add(int(part.strip()))
            except ValueError:
                pass
        if selected_nums:
            all_cos = parse_cos_from_text(co_source_key)
            preloaded_cos = [co for co in all_cos if co["num"] in selected_nums]

    def event_stream():
        import anthropic
        client = anthropic.Anthropic(
            api_key=os.environ["ANTHROPIC_API_KEY"],
            timeout=120.0,  # 120 s per read  -  QB generates more text per call
            max_retries=3,
        )
        all_blocks = []

        def _ev(payload):
            return f"data: {json.dumps(payload)}\n\n"

        def _timeout_msg():
            return _ev({"type": "error",
                        "message": "Request timed out - check your internet connection and try again."})

        def _conn_msg():
            return _ev({"type": "error",
                        "message": "Could not reach the AI service - check your internet connection and try again."})

        try:
            sem_label = f"  [Semester {info['semester']}]" if info.get("semester") else ""

            # Notify frontend whether this is a lab course
            yield _ev({"type": "lab_mode", "is_lab": lab_mode})

            if preloaded_cos:
                cos = preloaded_cos
                yield _ev({"type": "status",
                           "message": f"Using {len(cos)} pre-selected Course Outcome(s)",
                           "stage": "cos_done"})
            else:
                yield _ev({"type": "status", "message": "Generating Course Outcomes...", "stage": "cos"})
                co_text = generate_cos_for_course(client, code, info["title"], info["text"], num_cos=15)
                cos     = parse_cos_from_text(co_text)
                if not cos:
                    _refund_tokens(3)
                    yield _ev({"type": "error", "message": "No COs could be generated"})
                    return
                yield _ev({"type": "status", "message": f"Generated {len(cos)} Course Outcomes", "stage": "cos_done"})

            yield _ev({"type": "status", "message": "Extracting unit structure from syllabus...", "stage": "units"})
            unit_plan = build_unit_plan(client, cos, info["text"])
            num_units = len(unit_plan)
            yield _ev({"type": "status", "message": f"Found {num_units} unit(s)", "stage": "units_done", "num_units": num_units})

            # â"€â"€ Lab course path â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
            if lab_mode:
                all_blocks.append(
                    f"# Lab Question Bank\n"
                    f"Course: {code}  -  {info['title']}{sem_label}\n"
                    f"Experiments per unit: {n_exp}  |  {n_exp * num_units} total\n"
                )
                for idx, (unit_label, udata) in enumerate(unit_plan.items(), 1):
                    unit_title = udata["title"]
                    unit_cos   = udata["cos"]
                    full_unit  = f"{unit_label}: {unit_title}" if unit_title else unit_label
                    co_summary = "  |  ".join(f"CO{c['num']}" for c in unit_cos) or "no COs mapped"

                    header = f"\n## {full_unit}\n"
                    if unit_cos:
                        header += f"### Course Outcomes: {co_summary}\n"
                        for co in unit_cos:
                            header += f"### CO{co['num']}: {co['statement']}\n"
                    else:
                        header += "### Experiments based on syllabus content\n"
                    all_blocks.append(header)

                    yield _ev({"type": "status",
                               "message": f"[{idx}/{num_units}] {full_unit} - generating experiments...",
                               "stage": "unit", "unit_idx": idx, "unit_total": num_units})
                    etext = generate_lab_experiments_for_unit(
                        client, unit_label, unit_title, unit_cos,
                        code, info["title"], info["text"], n_exp=n_exp,
                    )
                    etext = _normalize_lab_block(etext, unit="")
                    if etext.strip():
                        all_blocks.append(etext + "\n")
                    yield _ev({"type": "unit_done", "unit": unit_label, "unit_idx": idx})

                store["qbank_all_blocks"] = all_blocks
                lab_tally = count_co_lab(all_blocks)
                store["qbank_lab_tally"] = lab_tally
                yield _ev({"type": "done", "num_units": num_units, "is_lab": True})
                return

            # â"€â"€ Regular course path â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
            total_per_unit = n2 + n5 + n10
            extras = []
            if n_assign: extras.append(f"{n_assign} assignment task(s)")
            if n_quiz:   extras.append(f"{n_quiz} quiz question(s)")
            extra_line = ("  |  " + " + ".join(extras) + " per unit\n") if extras else "\n"

            all_blocks.append(
                f"# Question Bank\n"
                f"Course: {code}  -  {info['title']}{sem_label}\n"
                f"Exam questions per unit: {n2} x 2-mark + {n5} x 5-mark + {n10} x 10-mark"
                f" = {total_per_unit} per unit  |  {total_per_unit * num_units} total\n"
                + extra_line
            )

            for idx, (unit_label, udata) in enumerate(unit_plan.items(), 1):
                unit_title = udata["title"]
                unit_cos   = udata["cos"]
                full_unit  = f"{unit_label}: {unit_title}" if unit_title else unit_label
                co_summary = "  |  ".join(f"CO{c['num']}" for c in unit_cos) or "no COs mapped"

                header = f"\n## {full_unit}\n"
                if unit_cos:
                    header += f"### Course Outcomes: {co_summary}\n"
                    for co in unit_cos:
                        header += f"### CO{co['num']}: {co['statement']}\n"
                else:
                    header += "### Questions based on syllabus content\n"
                all_blocks.append(header)

                yield _ev({"type": "status",
                           "message": f"[{idx}/{num_units}] {full_unit} - generating exam questions...",
                           "stage": "unit", "unit_idx": idx, "unit_total": num_units})
                qtext = generate_questions_for_unit(
                    client, unit_label, unit_title, unit_cos,
                    code, info["title"], info["text"], n2=n2, n5=n5, n10=n10,
                )
                qtext = _normalize_qblock(qtext, unit="")
                if qtext.strip():
                    all_blocks.append(qtext + "\n")

                if n_assign:
                    yield _ev({"type": "status",
                               "message": f"[{idx}/{num_units}] {unit_label} - generating assignments...",
                               "stage": "unit_assign", "unit_idx": idx})
                    atext = generate_assignments_for_unit(
                        client, unit_label, unit_title, unit_cos,
                        code, info["title"], info["text"], n_assign=n_assign,
                    )
                    atext = _normalize_assignment_block(atext)
                    if atext.strip():
                        all_blocks.append(atext + "\n")

                if n_quiz:
                    yield _ev({"type": "status",
                               "message": f"[{idx}/{num_units}] {unit_label} - generating quiz questions...",
                               "stage": "unit_quiz", "unit_idx": idx})
                    qztext = generate_quiz_for_unit(
                        client, unit_label, unit_title, unit_cos,
                        code, info["title"], info["text"], n_quiz=n_quiz,
                    )
                    qztext = _normalize_quiz_block(qztext)
                    if qztext.strip():
                        all_blocks.append(qztext + "\n")

                yield _ev({"type": "unit_done", "unit": unit_label, "unit_idx": idx})

        except anthropic.APITimeoutError:
            _refund_tokens(3)
            yield _timeout_msg()
            return
        except anthropic.APIConnectionError:
            _refund_tokens(3)
            yield _conn_msg()
            return
        except Exception as e:
            _refund_tokens(3)
            yield _ev({"type": "error", "message": str(e)})
            return

        store["qbank_all_blocks"] = all_blocks
        co_tally = count_co_questions(all_blocks)
        store["qbank_co_tally"] = co_tally
        yield _ev({"type": "done", "num_units": num_units})

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/export_qbank", methods=["POST"])
@login_required
def export_qbank():
    store      = _get_store()
    all_blocks = store.get("qbank_all_blocks", [])
    code       = store.get("qbank_code", "COURSE")
    is_lab     = store.get("qbank_is_lab", False)
    co_tally   = store.get("qbank_co_tally", {})
    lab_tally  = store.get("qbank_lab_tally", {})

    if not all_blocks:
        return jsonify({"error": "No question bank generated. Please generate one first."}), 400

    data = request.get_json()
    fmt  = data.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    filename_base = f"{code}_LabQuestionBank" if is_lab else f"{code}_QuestionBank"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if is_lab:
            if fmt == "txt":
                qb_save_txt(all_blocks, tmp.name, lab_tally=lab_tally)
            elif fmt == "docx":
                qb_save_docx(all_blocks, tmp.name, lab_tally=lab_tally)
            elif fmt == "pdf":
                qb_save_pdf(all_blocks, tmp.name, lab_tally=lab_tally)
        else:
            if fmt == "txt":
                qb_save_txt(all_blocks, tmp.name, co_tally=co_tally)
            elif fmt == "docx":
                qb_save_docx(all_blocks, tmp.name, co_tally=co_tally)
            elif fmt == "pdf":
                qb_save_pdf(all_blocks, tmp.name, co_tally=co_tally)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename_base}{suffix}"'},
    )


# -- Module 2: Question Paper routes -----------------------------------------

@app.route("/parse_generated_qbank", methods=["POST"])
@login_required
def parse_generated_qbank():
    """Turn store["qbank_all_blocks"] into qb_data for the paper builder."""
    store      = _get_store()
    all_blocks = store.get("qbank_all_blocks", [])
    code       = store.get("qbank_code", "COURSE")
    courses    = store.get("courses", {})
    info       = courses.get(code, {"title": "", "semester": None})

    if not all_blocks:
        return jsonify({"error": "No question bank in session. Please generate one first."}), 400

    raw_text = "\n".join(all_blocks)
    try:
        qb_is_lab = is_lab_qbank(raw_text)
        data = parse_lab_qbank(raw_text) if qb_is_lab else parse_qbank(raw_text)
    except Exception as e:
        return jsonify({"error": f"Failed to parse question bank: {e}"}), 500

    if not data["units"]:
        return jsonify({"error": "Question bank structure could not be parsed."}), 400

    store["qb_data"]     = data
    store["qb_is_lab"]   = qb_is_lab
    store["qb_filename"] = f"{code}_LabQuestionBank" if qb_is_lab else f"{code}_QuestionBank"

    if qb_is_lab:
        exp_count = sum(len(u["experiments"]) for u in data["units"].values())
        totals = {"experiments": exp_count}
    else:
        totals = {
            "2":      sum(len(u["exam"][2])  for u in data["units"].values()),
            "5":      sum(len(u["exam"][5])  for u in data["units"].values()),
            "10":     sum(len(u["exam"][10]) for u in data["units"].values()),
            "assign": sum(len(u["assign"])   for u in data["units"].values()),
            "quiz":   sum(len(u["quiz"])     for u in data["units"].values()),
        }

    return jsonify({
        "course_code":  data["course_code"],
        "course_title": data["course_title"],
        "semester":     data["semester"],
        "num_units":    len(data["units"]),
        "is_lab":       qb_is_lab,
        "units": [{"label": lbl, "title": ud["title"]} for lbl, ud in data["units"].items()],
        "totals": totals,
    })


@app.route("/get_question_list")
@login_required
def get_question_list():
    store = _get_store()
    data  = store.get("qb_data")
    if not data or not data.get("units"):
        return jsonify({"2": [], "5": [], "10": [], "assign": [], "quiz": []})

    if store.get("qb_is_lab"):
        flat = _flat_lab_list(data["units"])
        experiments = [
            {
                "idx":     i,
                "unit":    unit,
                "co":      exp.get("co") or "",
                "text":    exp.get("title") or "",
                "preview": (exp.get("title") or "")[:90],
                "more":    len(exp.get("title") or "") > 90,
                "has_viva": bool(exp.get("viva")),
            }
            for i, (unit, exp) in enumerate(flat)
        ]
        return jsonify({"is_lab": True, "experiments": experiments})

    result = {}
    for mark in [2, 5, 10]:
        flat = _flat_list(data["units"], "exam", mark)
        result[str(mark)] = [
            {
                "idx":     i,
                "unit":    unit,
                "co":      q.get("co") or "",
                "text":    q.get("text") or "",
                "preview": (q.get("text") or "")[:90],
                "more":    len(q.get("text") or "") > 90,
            }
            for i, (unit, q) in enumerate(flat)
        ]
    for section in ("assign", "quiz"):
        flat = _flat_list(data["units"], section)
        text_key = "title" if section == "assign" else "text"
        result[section] = [
            {
                "idx":     i,
                "unit":    unit,
                "co":      q.get("co") or "",
                "text":    q.get(text_key) or q.get("text") or "",
                "preview": (q.get(text_key) or q.get("text") or "")[:90],
                "more":    len(q.get(text_key) or q.get("text") or "") > 90,
            }
            for i, (unit, q) in enumerate(flat)
        ]
    return jsonify(result)


@app.route("/upload_qb", methods=["POST"])
@login_required
def upload_qb():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".pdf"):
        return jsonify({"error": f"Unsupported file type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        raw_text = _load_qbank_text(tmp.name)
        if not raw_text.strip():
            return jsonify({"error": "No text could be extracted from the file"}), 400
        qb_is_lab = is_lab_qbank(raw_text)
        data = parse_lab_qbank(raw_text) if qb_is_lab else parse_qbank(raw_text)
        if not data["units"]:
            return jsonify({"error": "No question bank structure detected. Please upload a valid question bank file."}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to parse question bank: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    store = _get_store()
    store["qb_data"]     = data
    store["qb_is_lab"]   = qb_is_lab
    store["qb_filename"] = os.path.splitext(f.filename)[0]

    if qb_is_lab:
        exp_count = sum(len(u["experiments"]) for u in data["units"].values())
        totals = {"experiments": exp_count}
    else:
        totals = {
            "2":      sum(len(u["exam"][2])  for u in data["units"].values()),
            "5":      sum(len(u["exam"][5])  for u in data["units"].values()),
            "10":     sum(len(u["exam"][10]) for u in data["units"].values()),
            "assign": sum(len(u["assign"])   for u in data["units"].values()),
            "quiz":   sum(len(u["quiz"])     for u in data["units"].values()),
        }

    return jsonify({
        "course_code":  data["course_code"],
        "course_title": data["course_title"],
        "semester":     data["semester"],
        "num_units":    len(data["units"]),
        "is_lab":       qb_is_lab,
        "units": [{"label": lbl, "title": ud["title"]} for lbl, ud in data["units"].items()],
        "totals": totals,
    })


def _build_lab_paper_blocks(body, data):
    """Build a lab paper (practical or viva) from request body. Returns (blocks, co_tally, label)."""
    ptype   = body.get("ptype")
    show_co = bool(body.get("show_co", False))

    if ptype not in ("practical", "viva"):
        raise ValueError("Invalid lab paper type  -  expected 'practical' or 'viva'")

    cfg = {
        "title":     body.get("title", "Practical Examination" if ptype == "practical" else "Viva Voce Examination"),
        "date":      body.get("date", ""),
        "duration":  body.get("duration", ""),
        "max_marks": body.get("max_marks", ""),
        "show_co":   show_co,
    }

    flat    = _flat_lab_list(data["units"])
    chosen  = body.get("selected_indices")
    items   = [flat[i] for i in chosen if i < len(flat)] if chosen else flat
    if not items:
        raise ValueError("No experiments selected")

    if ptype == "practical":
        cfg["instructions"] = [
            "Perform the given experiment and record your observations.",
            "Write the aim, procedure, observations, and result in your lab record.",
            "A brief viva will be conducted at the end of the session.",
        ]
        blocks = build_lab_practical_paper(items, data, cfg)
        label  = "Practical"
    else:
        cfg["instructions"] = ["Answer all questions clearly and concisely."]
        blocks = build_lab_viva_paper(items, data, cfg)
        label  = "Viva"

    return blocks, _count_co_lab_coverage(items), label


def _build_qpaper_blocks(body, data):
    """Build paper blocks from request body and qbank data. Returns (blocks, co_tally, label)."""
    ptype   = body.get("ptype")
    show_co = bool(body.get("show_co", False))

    if ptype not in ("assign", "quiz", "internal", "final"):
        raise ValueError("Invalid paper type")

    cfg = {
        "title":     body.get("title", "Examination Paper"),
        "date":      body.get("date", ""),
        "duration":  body.get("duration", ""),
        "max_marks": body.get("max_marks", ""),
        "show_co":   show_co,
    }

    if ptype == "assign":
        avail = _flat_list(data["units"], "assign")
        if not avail:
            raise ValueError("No assignment tasks found in the question bank")
        chosen = body.get("selected_indices")
        items  = [avail[i] for i in chosen if i < len(avail)] if chosen else avail
        if not items:
            raise ValueError("Selected indices produced no questions")
        cfg["instructions"] = [
            "Complete all tasks independently.",
            "Cite all references. Plagiarism will result in a zero grade.",
        ]
        blocks   = build_assignment_paper(items, data, cfg)
        co_tally = _count_co_coverage(sel_assign=items)

    elif ptype == "quiz":
        avail = _flat_list(data["units"], "quiz")
        if not avail:
            raise ValueError("No quiz questions found in the question bank")
        chosen = body.get("selected_indices")
        items  = [avail[i] for i in chosen if i < len(avail)] if chosen else avail
        if not items:
            raise ValueError("Selected indices produced no questions")
        marks_each = max(1, int(body.get("marks_each", 1)))
        cfg["marks_each"] = marks_each
        cfg["max_marks"]  = str(len(items) * marks_each)
        cfg["instructions"] = [
            "For MCQ questions, choose the single best answer.",
            "For Fill in the Blank, write the exact missing term.",
            "No negative marking.",
        ]
        blocks   = build_quiz_paper(items, data, cfg)
        co_tally = _count_co_coverage(sel_quiz=items)

    else:
        cfg["parts"] = {}
        instructions = []
        parts_config = body.get("parts", {})
        sel_2 = sel_5 = sel_10 = []

        for marks in (2, 5, 10):
            pc         = parts_config.get(str(marks), {})
            n          = int(pc.get("n", 0))
            answer_any = int(pc.get("answer_any", n))
            part_label = pc.get("label", {2: "PART A", 5: "PART B", 10: "PART C"}[marks])
            if n <= 0:
                continue
            avail = _flat_list(data["units"], "exam", marks)
            chosen = pc.get("selected_indices")
            if chosen:
                sel = [avail[i] for i in chosen if i < len(avail)]
            else:
                sel = avail[:n]
            if not sel:
                continue
            answer_any = min(max(1, answer_any), len(sel))
            cfg["parts"][marks] = {
                "label":      part_label,
                "answer_any": answer_any if answer_any < len(sel) else None,
            }
            instr = ("Answer all questions" if answer_any == len(sel)
                     else f"Answer any {answer_any} of {len(sel)}")
            instructions.append(f"{part_label}: {instr} ({marks} marks each)")
            if marks == 2:   sel_2  = sel
            elif marks == 5: sel_5  = sel
            else:            sel_10 = sel

        if not sel_2 and not sel_5 and not sel_10:
            raise ValueError("No questions configured for the paper")

        cfg["instructions"] = instructions
        blocks   = build_exam_paper(sel_2, sel_5, sel_10, data, cfg)
        co_tally = _count_co_coverage(sel_2=sel_2, sel_5=sel_5, sel_10=sel_10)

    label = {"assign": "Assignment", "quiz": "Quiz", "internal": "Internal", "final": "Final"}[ptype]
    return blocks, co_tally, label


def _edited_text_to_docx(text, output_path):
    """Convert plain edited paper text to a basic .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)
    for line in text.split('\n'):
        stripped = line.rstrip()
        if not stripped:
            p = doc.add_paragraph()
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(6)
        else:
            indent = len(stripped) - len(stripped.lstrip())
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip())
            run.font.size = Pt(10)
            if indent:
                p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
    doc.save(output_path)


def _edited_text_to_pdf(text, output_path):
    """Convert plain edited paper text to a basic .pdf file."""
    from fpdf import FPDF, XPos, YPos
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(15, 15, 15)
    for line in text.split('\n'):
        stripped = line.rstrip()
        if not stripped:
            pdf.ln(3)
        elif stripped.startswith('## '):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, _pdf_safe(stripped[3:].strip()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
        else:
            indent = len(stripped) - len(stripped.lstrip())
            pdf.set_font("Helvetica", "", 10)
            if indent:
                pdf.set_x(pdf.l_margin + min(indent * 1.5, 20))
            pdf.multi_cell(0, 5, _pdf_safe(stripped.strip()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.output(output_path)


@app.route("/export_qpaper", methods=["POST"])
@tokens_required(2)
def export_qpaper():
    store    = _get_store()
    data     = store.get("qb_data")
    filename = store.get("qb_filename", "course")

    if not data or not data.get("units"):
        return jsonify({"error": "No question bank loaded. Please upload one first."}), 400

    body = request.get_json()
    fmt  = body.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    try:
        builder = _build_lab_paper_blocks if store.get("qb_is_lab") else _build_qpaper_blocks
        blocks, co_tally, label = builder(body, data)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            qp_save_txt(blocks, tmp.name, co_tally=co_tally)
        elif fmt == "docx":
            qp_save_docx(blocks, tmp.name, co_tally=co_tally)
        elif fmt == "pdf":
            qp_save_pdf(blocks, tmp.name, co_tally=co_tally)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}_{label}_Paper{suffix}"'},
    )


@app.route("/preview_qpaper", methods=["POST"])
@login_required
def preview_qpaper():
    """Generate paper and return text content for editable preview."""
    store    = _get_store()
    data     = store.get("qb_data")
    filename = store.get("qb_filename", "course")

    if not data or not data.get("units"):
        return jsonify({"error": "No question bank loaded. Please upload one first."}), 400

    body = request.get_json()
    try:
        builder = _build_lab_paper_blocks if store.get("qb_is_lab") else _build_qpaper_blocks
        blocks, co_tally, label = builder(body, data)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode='w', encoding='utf-8')
    tmp.close()
    try:
        qp_save_txt(blocks, tmp.name, co_tally=co_tally)
        with open(tmp.name, encoding='utf-8') as fh:
            text = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return jsonify({"text": text, "label": label, "filename": filename, "is_lab": store.get("qb_is_lab", False)})


@app.route("/download_paper_from_text", methods=["POST"])
@login_required
def download_paper_from_text():
    """Download edited paper text as txt, docx, or pdf."""
    body     = request.get_json()
    text     = body.get("text", "")
    fmt      = body.get("fmt", "txt")
    label    = body.get("label", "Paper")
    filename = body.get("filename", "paper")

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            with open(tmp.name, 'w', encoding='utf-8') as fh:
                fh.write(text)
        elif fmt == "docx":
            _edited_text_to_docx(text, tmp.name)
        elif fmt == "pdf":
            _edited_text_to_pdf(text, tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}_{label}_Paper{suffix}"'},
    )


def _html_to_docx(html, output_path):
    """Convert mammoth HTML (paragraphs + tables) to a properly structured DOCX."""
    from html.parser import HTMLParser
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self._stack   = []
            self._bold    = False
            self._italic  = False
            self._para    = None
            self._in_tbl  = False
            self._rows    = []
            self._cur_row = []
            self._cell    = []

        def handle_starttag(self, tag, attrs):
            self._stack.append(tag)
            if tag in ('strong', 'b'):   self._bold   = True
            elif tag in ('em', 'i'):     self._italic = True
            elif tag == 'table':         self._in_tbl = True;  self._rows = []
            elif tag == 'tr':            self._cur_row = []
            elif tag in ('td', 'th'):    self._cell   = []
            elif tag == 'br' and self._para and not self._in_tbl:
                self._para.add_run('\n')
            elif not self._in_tbl:
                if tag in ('h1', 'h2', 'h3'):
                    self._para = doc.add_heading('', level=int(tag[1]))
                elif tag == 'p':
                    self._para = doc.add_paragraph()
                    self._para.paragraph_format.space_after = Pt(2)

        def handle_endtag(self, tag):
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()
            if tag in ('strong', 'b'):   self._bold   = False
            elif tag in ('em', 'i'):     self._italic = False
            elif tag in ('td', 'th'):
                self._cur_row.append(''.join(self._cell).strip())
                self._cell = []
            elif tag == 'tr':
                if self._cur_row:
                    self._rows.append(list(self._cur_row))
                self._cur_row = []
            elif tag == 'table':
                self._in_tbl = False
                if not self._rows:
                    return
                ncols = max(len(r) for r in self._rows)
                tbl = doc.add_table(rows=0, cols=ncols)
                tbl.style = 'Table Grid'
                # Smart column widths: widest-content column gets 55 %, rest split equally
                col_max = [0] * ncols
                for row in self._rows:
                    for ci, cell in enumerate(row):
                        if ci < ncols:
                            col_max[ci] = max(col_max[ci], len(cell))
                wide = col_max.index(max(col_max)) if any(col_max) else 0
                page_w = Inches(6.5)
                wide_w = int(page_w * 0.55)
                other_w = int(page_w * 0.45 / max(ncols - 1, 1))
                for ri, row_data in enumerate(self._rows):
                    row = tbl.add_row()
                    for ci in range(ncols):
                        cell_text = row_data[ci] if ci < len(row_data) else ''
                        cell = row.cells[ci]
                        cell.width = wide_w if ci == wide else other_w
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.size = Pt(9.5)
                        if ri == 0:
                            run.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                self._rows = []

        def handle_data(self, data):
            if self._in_tbl:
                self._cell.append(data)
            elif self._para is not None:
                run = self._para.add_run(data)
                run.bold   = self._bold
                run.italic = self._italic
                run.font.size = Pt(11)

    _P().feed(html)
    doc.save(output_path)


def _html_to_pdf(html, output_path):
    """Convert mammoth HTML (paragraphs + tables) to a PDF file."""
    from html.parser import HTMLParser
    from fpdf import FPDF, XPos, YPos

    # â"€â"€ parse into a flat list of blocks â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    blocks = []

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self._stack   = []
            self._bold    = False
            self._in_tbl  = False
            self._rows    = []
            self._cur_row = []
            self._cell    = []
            self._para    = []

        def handle_starttag(self, tag, attrs):
            self._stack.append(tag)
            if tag in ('strong', 'b'):  self._bold = True
            elif tag == 'table':        self._in_tbl = True; self._rows = []
            elif tag == 'tr':           self._cur_row = []
            elif tag in ('td', 'th'):   self._cell = []
            elif not self._in_tbl and tag in ('p', 'h1', 'h2', 'h3'):
                self._para = []

        def handle_endtag(self, tag):
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()
            if tag in ('strong', 'b'):  self._bold = False
            elif tag in ('td', 'th'):
                self._cur_row.append(''.join(self._cell).strip())
                self._cell = []
            elif tag == 'tr':
                if self._cur_row:
                    self._rows.append(list(self._cur_row))
                self._cur_row = []
            elif tag == 'table':
                self._in_tbl = False
                if self._rows:
                    blocks.append(('table', self._rows[:]))
                self._rows = []
            elif not self._in_tbl and tag in ('p', 'h1', 'h2', 'h3'):
                text = ''.join(self._para).strip()
                if text:
                    blocks.append(('text', text, tag in ('h1', 'h2', 'h3')))
                self._para = []

        def handle_data(self, data):
            if self._in_tbl:
                self._cell.append(data)
            else:
                self._para.append(data)

    _P().feed(html)

    # â"€â"€ render with fpdf2 â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)   # manual page breaks  -  avoids mid-row breaks
    pdf.add_page()
    pdf.set_margins(15, 15, 15)
    page_w  = pdf.w - 2 * pdf.l_margin   # ~180 mm on A4
    page_h  = pdf.h - 2 * 15             # usable height

    for block in blocks:
        if block[0] == 'text':
            _, text, is_heading = block
            # Manual page break for paragraphs
            if pdf.get_y() + 10 > pdf.h - 15:
                pdf.add_page()
            pdf.set_font('Helvetica', 'B' if is_heading else '', 12 if is_heading else 10.5)
            pdf.multi_cell(0, 6, _pdf_safe(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)

        elif block[0] == 'table':
            rows = block[1]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            if ncols == 0:
                continue

            # Smart column widths: longest-content column gets 52 %
            col_max = [0] * ncols
            for row in rows:
                for ci, cell in enumerate(row):
                    if ci < ncols:
                        col_max[ci] = max(col_max[ci], len(cell))
            wide    = col_max.index(max(col_max)) if any(col_max) else 0
            wide_w  = page_w * 0.52
            other_w = page_w * 0.48 / max(ncols - 1, 1)
            col_widths = [wide_w if i == wide else other_w for i in range(ncols)]

            LH      = 5   # line height per text line in multi-line cells
            x_start = pdf.l_margin

            for ri, row in enumerate(rows):
                is_hdr = ri == 0

                # Build per-cell line lists and determine row height
                cell_data = []
                for ci in range(ncols):
                    raw  = row[ci] if ci < len(row) else ''
                    safe = _pdf_safe(raw)
                    cw   = col_widths[ci]
                    cpl  = max(1, int(cw / 2.1))  # chars per line
                    lines = [safe[j:j+cpl] for j in range(0, len(safe), cpl)] if safe else ['']
                    cell_data.append((cw, lines))

                max_lines = max(len(lines) for _, lines in cell_data)
                rh = max(7, max_lines * LH + 2)

                # Manual page break before this row if needed
                if pdf.get_y() + rh > pdf.h - 15:
                    pdf.add_page()

                y_row = pdf.get_y()
                pdf.set_font('Helvetica', 'B' if is_hdr else '', 9)
                if is_hdr:
                    pdf.set_fill_color(220, 220, 220)
                else:
                    pdf.set_fill_color(255, 255, 255)

                for ci, (cw, lines) in enumerate(cell_data):
                    x_cell = x_start + sum(col_widths[:ci])
                    if len(lines) == 1:
                        pdf.set_xy(x_cell, y_row)
                        pdf.cell(cw, rh, lines[0], border=1, fill=is_hdr)
                    else:
                        # Multi-line: rect for border/fill, then draw each line
                        pdf.rect(x_cell, y_row, cw, rh, style='FD' if is_hdr else 'D')
                        for li, line in enumerate(lines):
                            pdf.set_xy(x_cell + 1, y_row + 1 + li * LH)
                            pdf.cell(cw - 2, LH, line, border=0, align='L')

                # Advance to the next row explicitly
                pdf.set_xy(x_start, y_row + rh)

            pdf.ln(3)

    pdf.output(output_path)


@app.route("/ai_rephrase", methods=["POST"])
@tokens_required(1)
def ai_rephrase():
    """Return a rephrased version of the given question text."""
    import anthropic as _ant

    body          = request.get_json() or {}
    question_text = (body.get("question_text") or "").strip()
    co            = (body.get("co")    or "").strip()
    marks         = str(body.get("marks") or "").strip()

    if not question_text:
        return jsonify({"error": "No question text provided"})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"})

    co_ctx    = f" (mapped to {co})"    if co    else ""
    marks_ctx = f" ({marks} marks)"     if marks else ""

    prompt = (
        f"Rephrase the following engineering exam question to be clearer, "
        f"more precise, and academically well-worded{co_ctx}{marks_ctx}.\n\n"
        f'Original question:\n"{question_text}"\n\n'
        f"Return ONLY the rephrased question  -  no explanation, no quotes, no preamble."
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 300,
            messages   = [{"role": "user", "content": prompt}],
        )
        rephrased = response.content[0].text.strip().strip('"').strip("'")
        return jsonify({"rephrased": rephrased})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route("/ai_apply_suggestion", methods=["POST"])
@login_required
def ai_apply_suggestion():
    """Rewrite a question to address one specific AI suggestion."""
    import anthropic as _ant

    body               = request.get_json() or {}
    question_text      = (body.get("question_text")      or "").strip()
    co                 = (body.get("co")                 or "").strip()
    marks              = str(body.get("marks")           or "").strip()
    suggestion_message = (body.get("suggestion_message") or "").strip()

    if not question_text or not suggestion_message:
        return jsonify({"error": "Missing question or suggestion"})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"})

    co_ctx    = f" (mapped to {co})"  if co    else ""
    marks_ctx = f" ({marks} marks)"   if marks else ""

    prompt = (
        f"You are improving an engineering exam question based on specific feedback.\n\n"
        f"Original question{co_ctx}{marks_ctx}:\n\"{question_text}\"\n\n"
        f"Feedback to address: {suggestion_message}\n\n"
        f"Rewrite the question to address this feedback while preserving the original intent "
        f"and scope. Return ONLY the improved question  -  no explanation, no quotes, no preamble."
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 400,
            messages   = [{"role": "user", "content": prompt}],
        )
        improved = response.content[0].text.strip().strip('"').strip("'")
        return jsonify({"improved": improved})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route("/ai_review_edit", methods=["POST"])
@tokens_required(1)
def ai_review_edit():
    """Call Claude Haiku to review the question under the cursor."""
    import anthropic as _ant
    import json as _json

    body          = request.get_json() or {}
    question_text = (body.get("question_text") or "").strip()
    co            = (body.get("co")    or "").strip()
    marks         = str(body.get("marks") or "").strip()

    if not question_text or len(question_text) < 10:
        return jsonify({"suggestions": []})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured", "suggestions": []})

    ctx_parts = []
    if co:    ctx_parts.append(f"mapped to {co}")
    if marks: ctx_parts.append(f"{marks} marks")
    ctx_str = f" ({', '.join(ctx_parts)})" if ctx_parts else ""

    marks_label = f"{marks} marks" if marks else "the marks assigned"
    co_label    = co if co else "its Course Outcome"

    prompt = f"""You are reviewing an engineering college exam question.

Question{ctx_str}:
"{question_text}"

Return ONLY valid JSON (no markdown, no explanation) with this exact structure:
{{
  "suggestions": [
    {{"type": "clarity",      "message": "...", "level": "ok|tip|warning"}},
    {{"type": "co_alignment", "message": "...", "level": "ok|tip|warning"}},
    {{"type": "blooms",       "message": "...", "level": "ok|tip|warning"}}
  ]
}}

Rules:
- clarity: Is the question clear and unambiguous? Suggest a sharper phrasing if needed.
- co_alignment: Does the question align with {co_label}? If no CO info, note that briefly.
- blooms: Identify the Bloom's taxonomy level (Remember/Understand/Apply/Analyze/Evaluate/Create). Is it appropriate for {marks_label}?
- level: "ok" = no issues, "tip" = minor improvement, "warning" = significant problem.
- Each message: 1-2 sentences. Be concise and constructive."""

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 400,
            messages   = [{"role": "user", "content": prompt}],
        )
        text = response.content[0].text.strip()
        # Strip markdown code fences if the model wraps in ```json
        text = re.sub(r"```(?:json)?\s*|\s*```", "", text).strip()
        data = _json.loads(text)
        return jsonify(data)
    except _json.JSONDecodeError:
        return jsonify({"error": "Model returned invalid JSON", "suggestions": []})
    except _ant.APITimeoutError:
        return jsonify({"error": "timeout", "suggestions": []}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "connection", "suggestions": []}), 503
    except Exception as e:
        return jsonify({"error": str(e), "suggestions": []})


@app.route("/m1_report_apply_suggestion", methods=["POST"])
@login_required
def m1_report_apply_suggestion():
    """Rewrite a selected report section to directly incorporate AI advice."""
    import anthropic as _ant

    body     = request.get_json() or {}
    original = (body.get("original") or "").strip()
    advice   = (body.get("advice")   or "").strip()

    if not original or not advice:
        return jsonify({"error": "Missing original text or advice"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    prompt = (
        "You are editing a section of a Module 1 comprehensive academic report for an Indian "
        "engineering college course.\n\n"
        f"Original section:\n\"\"\"\n{original[:1200]}\n\"\"\"\n\n"
        f"Improvement advice to incorporate:\n{advice[:800]}\n\n"
        "Rewrite the section so it directly reflects the improvements described in the advice. "
        "Keep the same writing style, tone, and approximate length as the original. "
        "Do NOT add headers, bullet points, or labels — return only the rewritten prose, "
        "ready to replace the original in the report."
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=30.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 600,
            messages   = [{"role": "user", "content": prompt}],
        )
        rewritten = response.content[0].text.strip().strip('"').strip("'")
        return jsonify({"rewritten": rewritten})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route("/m1_report_ai_suggest", methods=["POST"])
@tokens_required(1)
def m1_report_ai_suggest():
    """Answer a free-form question about the Module 1 report using selected report text as context."""
    import anthropic as _ant

    body     = request.get_json() or {}
    question = (body.get("question") or "").strip()
    context  = (body.get("context")  or "").strip()

    if not question:
        return jsonify({"error": "No question provided"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    ctx_part = f"\n\nSelected section from the report:\n\"{context[:800]}\"" if context else ""

    prompt = (
        "You are an expert OBE/NBA academic quality advisor reviewing a Module 1 comprehensive "
        "report for an Indian engineering college course. The report covers Course Outcomes, "
        "Question Bank quality, Bloom's taxonomy compliance, CO coverage, scenario analysis, "
        "QQI scores, marks distribution, and accreditation readiness."
        f"{ctx_part}\n\n"
        f"Question: {question}\n\n"
        "Give a concise, practical, and actionable answer. If the question is about improving "
        "a metric, explain specific techniques the faculty can apply to the question bank or "
        "COs. Keep the response under 200 words and use plain text (no markdown)."
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=25.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 450,
            messages   = [{"role": "user", "content": prompt}],
        )
        answer = response.content[0].text.strip()
        return jsonify({"answer": answer})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


# -- Module 1 "Apply improvement tip" (regenerate COs / QB) -------------------

def _m1_apply_tip_to_cos(client, store, code, title, tip):
    """Revise the stored COs to incorporate an improvement tip.

    May rewrite statements and — when the tip calls for it (e.g. raising higher-order
    thinking or broadening knowledge dimensions) — change a CO's Bloom level and/or
    knowledge dimension. The CO statement lines, the markdown taxonomy table, and the
    cached co_taxonomy are all updated so the rebuilt report stays consistent.
    """
    import json as _json
    from generate_qbank import parse_cos_from_text
    from generate_cos import _LNAMES, _KDIM_NORM

    co_text = store.get("current_result", "")
    cos     = parse_cos_from_text(co_text)
    if not cos:
        raise ValueError("No Course Outcomes found to improve")

    def _bloom_idx(s):
        m = re.search(r"L\s*([1-6])", str(s), re.I)
        if m:
            return int(m.group(1))
        sl = str(s).lower()
        for i, nm in enumerate(_LNAMES, 1):
            if nm.lower() in sl:
                return i
        return None

    co_block = "\n".join(
        f'CO{c["num"]}: {c["statement"]}  (level: {c.get("bloom") or "?"}, '
        f'dimension: {c.get("kdim") or "?"})'
        for c in cos
    )
    levels_help = " | ".join(f"L{i} - {n}" for i, n in enumerate(_LNAMES, 1))
    prompt = (
        "You are an NBA/NAAC accreditation expert for an Indian engineering college.\n"
        f"Course: {title} ({code})\n\n"
        f"Current Course Outcomes (with their Bloom level and knowledge dimension):\n"
        f"{co_block}\n\n"
        f"Improvement to apply: {tip}\n\n"
        "Revise the Course Outcomes to incorporate this improvement. Keep the SAME number "
        "of COs and the SAME CO numbers. Only change a CO's Bloom level and/or knowledge "
        "dimension if the improvement calls for it; otherwise keep them unchanged. Whenever "
        "you raise a CO's level, use a stronger action verb in its statement.\n"
        f"Valid Bloom levels: {levels_help}\n"
        "Valid knowledge dimensions: Factual | Conceptual | Procedural | Meta-Cognitive\n\n"
        'Return ONLY a JSON object mapping CO number to an object, e.g. '
        '{"1": {"statement": "...", "bloom": "L4 - Analyze", "kdim": "Conceptual"}}. '
        "No markdown, no commentary."
    )
    resp = client.messages.create(
        model="claude-haiku-4-5-20251001", max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"```[a-zA-Z]*", "", raw).replace("```", "").strip()
    start, end = raw.find("{"), raw.rfind("}")
    data = _json.loads(raw[start:end + 1]) if (start >= 0 and end > start) else {}
    if not data:
        raise ValueError("AI did not return improved Course Outcomes")

    improved = {}
    for n, obj in data.items():
        ns = re.sub(r"\D", "", str(n))
        if not ns or not isinstance(obj, dict):
            continue
        rec = {}
        if obj.get("statement"):
            rec["statement"] = str(obj["statement"]).strip()
        bi = _bloom_idx(obj.get("bloom", ""))
        if bi:
            rec["bloom_code"] = f"L{bi}"
            rec["bloom_full"] = f"L{bi} - {_LNAMES[bi - 1]}"
        kd = _KDIM_NORM.get(str(obj.get("kdim", "")).strip().lower())
        if kd:
            rec["kdim"] = kd
        if rec:
            improved[ns] = rec
    if not improved:
        raise ValueError("AI did not return usable Course Outcomes")

    # Rewrite statement lines + taxonomy table cells in current_result
    out = []
    for line in co_text.split("\n"):
        s = line.strip()
        m = re.match(r"(CO(\d+):\s*)(.*)", s)
        if m and m.group(2) in improved and improved[m.group(2)].get("statement"):
            out.append(f'{m.group(1)}{improved[m.group(2)]["statement"]}')
            continue
        if s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if len(cells) == 4:
                mm = re.match(r"CO(\d+)$", cells[0])
                if mm and mm.group(1) in improved:
                    rec = improved[mm.group(1)]
                    if rec.get("kdim"):
                        cells[2] = rec["kdim"]
                    if rec.get("bloom_full"):
                        cells[3] = rec["bloom_full"]
                    out.append("| " + " | ".join(cells) + " |")
                    continue
        out.append(line)
    store["current_result"] = "\n".join(out)

    # Update cached co_taxonomy so the rebuilt report's RBT grid reflects the changes
    tax = dict(store.get("co_taxonomy", {}))
    for ns, rec in improved.items():
        ni  = int(ns)
        ent = dict(tax.get(ni, {}))
        if rec.get("bloom_code"):
            ent["bloom"] = rec["bloom_code"]
        if rec.get("kdim"):
            ent["kdim"] = rec["kdim"]
        if ent.get("bloom") and ent.get("kdim"):
            tax[ni] = ent
    store["co_taxonomy"] = tax


def _m1_apply_tip_to_qbank(client, store, code, title, tip, metric):
    """Append targeted questions addressing an improvement tip to the Question Bank."""
    from generate_qbank import parse_cos_from_text, count_co_questions

    blocks  = list(store.get("qbank_all_blocks", []))
    co_text = store.get("qbank_co_raw") or store.get("current_result", "")
    cos     = parse_cos_from_text(co_text)
    co_list = "\n".join(f'CO{c["num"]}: {c["statement"]}' for c in cos) or "CO1, CO2, CO3"

    prompt = (
        "You are an NBA/NAAC question-bank expert for an Indian engineering college.\n"
        f"Course: {title} ({code})\n\n"
        f"Course Outcomes:\n{co_list}\n\n"
        f"Improvement to apply: {tip}\n\n"
        "Generate 6 ADDITIONAL examination questions that directly address this "
        "improvement, distributed across the COs above. Use EXACTLY this plain-text "
        "format and output nothing else:\n"
        "[2 Marks]\n"
        "Q1. <question text> [CO1]\n"
        "Q2. <question text> [CO2]\n"
        "[5 Marks]\n"
        "Q1. <question text> [CO3]\n"
        "Q2. <question text> [CO4]\n"
        "[10 Marks]\n"
        "Q1. <question text> [CO1]\n"
        "Q2. <question text> [CO5]\n"
        "Every question MUST end with a [COn] tag matching one of the COs above. "
        "No markdown fences, no commentary."
    )
    resp = client.messages.create(
        model="claude-haiku-4-5-20251001", max_tokens=1300,
        messages=[{"role": "user", "content": prompt}],
    )
    add_text = resp.content[0].text.strip()
    add_text = re.sub(r"```[a-zA-Z]*", "", add_text).replace("```", "").strip()
    if not re.search(r"\[CO\d+\]", add_text):
        raise ValueError("AI did not return questions in the expected format")

    header = f"\n## Improvement Additions{(' - ' + metric) if metric else ''}\n"
    blocks.append(header + add_text + "\n")
    store["qbank_all_blocks"] = blocks
    store["qbank_co_tally"]   = count_co_questions(blocks)


def _m1_weak_for_panel(store):
    """Current weak actions filtered to the frozen baseline, so the improvement
    panel only ever shrinks as fixes are applied (regeneration side-effects on
    other metrics are not surfaced as new to-do items)."""
    allw     = store.get("m1_weak_actions_all", [])
    baseline = store.get("m1_weak_baseline")
    if baseline is None:
        return allw
    bset = set(baseline)
    return [w for w in allw if w.get("metric") in bset]


@app.route("/m1_report_meta", methods=["GET"])
@login_required
def m1_report_meta():
    """Return the (baseline-filtered) weak-metric actions for the preview panel.
    Used when refreshing after an Apply/Undo — does NOT change the baseline."""
    store = _get_store()
    return jsonify({
        "weak_actions": _m1_weak_for_panel(store),
        "can_undo":     bool(store.get("m1_undo")),
    })


@app.route("/m1_improve_start", methods=["POST"])
@login_required
def m1_improve_start():
    """Freeze the current weak metrics as the improvement baseline (called when the
    user opens the 'Suggest ways to improve' panel) and return them."""
    store = _get_store()
    allw  = store.get("m1_weak_actions_all", [])
    store["m1_weak_baseline"] = [w.get("metric") for w in allw]
    return jsonify({
        "weak_actions": allw,
        "can_undo":     bool(store.get("m1_undo")),
    })


@app.route("/m1_apply_tip", methods=["POST"])
@tokens_required(2)
def m1_apply_tip():
    """Apply an improvement tip by regenerating the COs or the Question Bank, then
    let the client rebuild the whole report. Keeps a one-level undo snapshot."""
    import copy
    import anthropic as _ant

    store  = _get_store()
    body   = request.get_json() or {}
    target = (body.get("target") or "").strip()
    tip    = (body.get("tip")    or "").strip()
    metric = (body.get("metric") or "").strip()

    if target not in ("cos", "qbank"):
        return jsonify({"error": "Invalid target"}), 400
    if not tip:
        return jsonify({"error": "Missing improvement tip"}), 400
    if target == "qbank" and store.get("qbank_is_lab"):
        return jsonify({"error": "Applying tips to lab question banks is not supported yet."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    code  = (store.get("qbank_code") or store.get("current_code", "COURSE")).strip().upper()
    info  = store.get("courses", {}).get(code, {})
    title = info.get("title", "Untitled")

    if target == "cos" and not store.get("current_result", "").strip():
        return jsonify({"error": "No COs in session. Generate Module 1 first."}), 400
    if target == "qbank" and not store.get("qbank_all_blocks"):
        return jsonify({"error": "No Question Bank in session. Generate Module 1 first."}), 400

    # One-level undo snapshot of everything the report is built from
    store["m1_undo"] = {
        "current_result":   store.get("current_result", ""),
        "co_taxonomy":      copy.deepcopy(store.get("co_taxonomy", {})),
        "qbank_all_blocks": list(store.get("qbank_all_blocks", [])),
        "qbank_co_tally":   copy.deepcopy(store.get("qbank_co_tally", {})),
        "metric":           metric,
    }

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=60.0)
        if target == "cos":
            _m1_apply_tip_to_cos(client, store, code, title, tip)
        else:
            _m1_apply_tip_to_qbank(client, store, code, title, tip, metric)
    except _ant.APITimeoutError:
        store.pop("m1_undo", None); _refund_tokens(2)
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        store.pop("m1_undo", None); _refund_tokens(2)
        return jsonify({"error": "Could not reach the AI service"}), 503
    except Exception as e:
        store.pop("m1_undo", None); _refund_tokens(2)
        return jsonify({"error": str(e)}), 500

    return jsonify({"ok": True, "target": target})


@app.route("/m1_undo_tip", methods=["POST"])
@login_required
def m1_undo_tip():
    """Restore the COs/QB to the snapshot taken before the last applied tip."""
    store = _get_store()
    snap  = store.get("m1_undo")
    if not snap:
        return jsonify({"error": "Nothing to undo"}), 400
    store["current_result"]   = snap["current_result"]
    store["co_taxonomy"]      = snap["co_taxonomy"]
    store["qbank_all_blocks"] = snap["qbank_all_blocks"]
    store["qbank_co_tally"]   = snap["qbank_co_tally"]
    store.pop("m1_undo", None)
    return jsonify({"ok": True})


# -- Lesson Plan AI suggestion routes -----------------------------------------

@app.route("/lp_ai_review", methods=["POST"])
@tokens_required(1)
def lp_ai_review():
    """Analyse a selected CO or session item and return structured suggestions."""
    import anthropic as _ant
    import json as _json

    body      = request.get_json() or {}
    item_type = (body.get("item_type") or "co").strip()   # 'co' | 'session'
    text      = (body.get("text")      or "").strip()
    bloom     = (body.get("bloom")     or "").strip()
    co_ref    = (body.get("co_ref")    or "").strip()
    sdg       = (body.get("sdg")       or "").strip()
    method    = (body.get("method")    or "").strip()

    if len(text) < 5:
        return jsonify({"suggestions": []})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured", "suggestions": []})

    if item_type == "co":
        prompt = f"""You are an OBE/NBA expert reviewing a Course Outcome statement.

CO statement: "{text}"
Stated Bloom's Level: {bloom or "not specified"}

Return ONLY valid JSON (no markdown):
{{
  "suggestions": [
    {{"type": "clarity",         "message": "...", "level": "ok|tip|warning"}},
    {{"type": "blooms_alignment","message": "...", "level": "ok|tip|warning"}},
    {{"type": "obe_compliance",  "message": "...", "level": "ok|tip|warning"}}
  ]
}}

Rules:
- clarity: Is it specific, free of vague verbs (know, understand, learn)? 1 - 2 sentences.
- blooms_alignment: Does the action verb match {bloom or 'the stated Bloom level'}? 1 - 2 sentences.
- obe_compliance: Does it express ONE measurable competency in NBA OBE format? 1 - 2 sentences.
- level: "ok"=no issue, "tip"=minor improvement, "warning"=significant problem."""
    else:
        prompt = f"""You are an OBE curriculum expert reviewing a lesson plan session entry.

Session topic: "{text}"
Mapped CO: {co_ref or "not specified"}
SDG: {sdg or "not specified"}
Teaching method: {method or "not specified"}

Return ONLY valid JSON (no markdown):
{{
  "suggestions": [
    {{"type": "co_alignment",  "message": "...", "level": "ok|tip|warning"}},
    {{"type": "sdg_relevance", "message": "...", "level": "ok|tip|warning"}},
    {{"type": "method_fit",    "message": "...", "level": "ok|tip|warning"}}
  ]
}}

Rules:
- co_alignment: Does the topic logically map to {co_ref or 'the mapped CO'}? 1 - 2 sentences.
- sdg_relevance: Is the {sdg or 'SDG'} connection meaningful for this topic? 1 - 2 sentences.
- method_fit: Is the teaching method appropriate for this topic and CO level? 1 - 2 sentences.
- level: "ok"=no issue, "tip"=minor improvement, "warning"=significant problem."""

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 400,
            messages   = [{"role": "user", "content": prompt}],
        )
        raw  = re.sub(r"```(?:json)?\s*|\s*```", "", response.content[0].text).strip()
        data = _json.loads(raw)
        return jsonify(data)
    except _json.JSONDecodeError:
        return jsonify({"error": "Model returned invalid JSON", "suggestions": []})
    except _ant.APITimeoutError:
        return jsonify({"error": "timeout", "suggestions": []}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "connection", "suggestions": []}), 503
    except Exception as e:
        return jsonify({"error": str(e), "suggestions": []})


@app.route("/lp_ai_rephrase", methods=["POST"])
@tokens_required(1)
def lp_ai_rephrase():
    """Return a rephrased version of a CO statement or session topic."""
    import anthropic as _ant

    body      = request.get_json() or {}
    item_type = (body.get("item_type") or "co").strip()
    text      = (body.get("text")      or "").strip()
    bloom     = (body.get("bloom")     or "").strip()
    co_ref    = (body.get("co_ref")    or "").strip()

    if not text:
        return jsonify({"error": "No text provided"})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"})

    if item_type == "co":
        prompt = (
            f"Rephrase the following Course Outcome statement to be clearer, more specific, "
            f"and fully NBA/OBE-compliant using a strong Bloom's action verb at the "
            f"{bloom or 'appropriate'} level.\n\n"
            f'Original CO: "{text}"\n\n'
            f"Return ONLY the rephrased CO statement  -  no explanation, no label, no quotes."
        )
    else:
        co_ctx = f" (mapped to {co_ref})" if co_ref else ""
        prompt = (
            f"Rephrase the following lesson plan session topic{co_ctx} to be more specific "
            f"and clearly aligned to the course outcome.\n\n"
            f'Original topic: "{text}"\n\n'
            f"Return ONLY the rephrased topic  -  no explanation, no label, no quotes."
        )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 200,
            messages   = [{"role": "user", "content": prompt}],
        )
        rephrased = response.content[0].text.strip().strip('"').strip("'")
        return jsonify({"rephrased": rephrased})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route("/lp_ai_apply", methods=["POST"])
@login_required
def lp_ai_apply():
    """Rewrite a CO or session topic to address one specific suggestion."""
    import anthropic as _ant

    body               = request.get_json() or {}
    item_type          = (body.get("item_type")          or "co").strip()
    text               = (body.get("text")               or "").strip()
    bloom              = (body.get("bloom")               or "").strip()
    suggestion_message = (body.get("suggestion_message") or "").strip()

    if not text or not suggestion_message:
        return jsonify({"error": "Missing text or suggestion"})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"})

    if item_type == "co":
        prompt = (
            f"You are improving a Course Outcome statement based on specific feedback.\n\n"
            f'Original CO (Bloom level: {bloom or "unspecified"}): "{text}"\n\n'
            f"Feedback to address: {suggestion_message}\n\n"
            f"Rewrite the CO to address this feedback while preserving its original intent and scope. "
            f"Return ONLY the improved CO statement  -  no explanation, no label, no quotes."
        )
    else:
        prompt = (
            f"You are improving a lesson plan session topic based on specific feedback.\n\n"
            f'Original topic: "{text}"\n\n'
            f"Feedback to address: {suggestion_message}\n\n"
            f"Rewrite the topic to address this feedback. "
            f"Return ONLY the improved topic  -  no explanation, no label, no quotes."
        )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=20.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 300,
            messages   = [{"role": "user", "content": prompt}],
        )
        improved = response.content[0].text.strip().strip('"').strip("'")
        return jsonify({"improved": improved})
    except _ant.APITimeoutError:
        return jsonify({"error": "Request timed out"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route("/lp_patch_item", methods=["POST"])
@login_required
def lp_patch_item():
    """Persist an edited CO or session item back into the stored lesson plan."""
    store     = _get_store()
    lp_data   = store.get("lp_data")
    if not lp_data:
        return jsonify({"error": "No lesson plan in session"}), 400

    body      = request.get_json() or {}
    item_type = body.get("item_type", "co")
    num       = body.get("num")          # CO num (int) or session num (int)
    new_text  = (body.get("new_text") or "").strip()

    if not new_text:
        return jsonify({"error": "new_text is required"}), 400

    if item_type == "co":
        for co in lp_data.get("cos", []):
            if co.get("num") == num:
                co["statement"] = new_text
                break
    elif item_type == "session":
        for s in lp_data.get("session_plan", []):
            if s.get("session") == num:
                s["topic"] = new_text
                break
    else:
        return jsonify({"error": "Unknown item_type"}), 400

    store["lp_data"] = lp_data
    return jsonify({"ok": True})


@app.route("/download_paper_from_html", methods=["POST"])
@login_required
def download_paper_from_html():
    """Download the mammoth-rendered HTML (with tables) as docx or pdf."""
    body     = request.get_json()
    html     = body.get("html", "")
    fmt      = body.get("fmt", "docx")
    label    = body.get("label", "Paper")
    filename = body.get("filename", "paper")

    if fmt not in ("docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = ".docx" if fmt == "docx" else ".pdf"
    mime   = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              if fmt == "docx" else "application/pdf")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "docx":
            _html_to_docx(html, tmp.name)
        else:
            _html_to_pdf(html, tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}_{label}_Paper{suffix}"'},
    )


@app.route("/upload_template_paper", methods=["POST"])
@login_required
def upload_template_paper():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".doc"):
        return jsonify({"error": "Unsupported format. Use PDF or DOCX."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        result = parse_template_paper(tmp.name, ext)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not result.get("slots"):
        return jsonify({"error": "No question table detected in the uploaded paper."}), 400

    store = _get_store()
    store["template_data"] = result
    return jsonify(result)


@app.route("/export_template_paper", methods=["POST"])
@tokens_required(1)
def export_template_paper():
    body   = request.get_json()
    fmt    = body.get("fmt", "docx")
    header = body.get("header", {})
    slots  = body.get("slots", [])

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400
    if not slots:
        return jsonify({"error": "No question slots provided"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "docx":
            tpl_build_docx(header, slots, tmp.name)
        elif fmt == "pdf":
            tpl_build_pdf(header, slots, tmp.name)
        else:
            tpl_build_txt(header, slots, tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    course = header.get("course_code", "paper").replace("/", "-")
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{course}_Paper{suffix}"'},
    )


# â"€â"€ Format Profile management routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/api/profiles", methods=["GET"])
@login_required
def api_list_profiles():
    import parse_profiles as pp
    return jsonify(pp.list_profiles())


@app.route("/api/profiles", methods=["POST"])
@login_required
def api_add_profile():
    import parse_profiles as pp
    profile = request.get_json()
    if not profile:
        return jsonify({"error": "No JSON body"}), 400
    try:
        pp.add_custom_profile(profile)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    return jsonify({"ok": True, "id": profile["id"]})


@app.route("/api/profiles/<profile_id>", methods=["DELETE"])
@login_required
def api_delete_profile(profile_id):
    import parse_profiles as pp
    deleted = pp.delete_custom_profile(profile_id)
    if not deleted:
        return jsonify({"error": f"Profile '{profile_id}' not found or is built-in"}), 404
    return jsonify({"ok": True})


# â"€â"€ CO - PO Mapping routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/upload_co_pomap", methods=["POST"])
@login_required
def upload_co_pomap():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        if ext == ".pdf":
            from generate_qbank import load_cos_from_pdf
            from generate_cos import extract_pdf
            cos     = load_cos_from_pdf(tmp.name)
            co_text = extract_pdf(tmp.name)
        else:
            co_text = extract(tmp.name)
            cos     = parse_cos_for_mapping(co_text)
    except Exception as e:
        return jsonify({"error": f"Failed to parse CO file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not cos:
        return jsonify({"error": "No COs found in the uploaded file."}), 400

    course_code, course_title = parse_course_header(co_text)

    store = _get_store()
    store["pomap_cos"]          = cos
    store["pomap_course_code"]  = course_code
    store["pomap_course_title"] = course_title

    return jsonify({
        "course_code":  course_code,
        "course_title": course_title,
        "cos": [{"num": co["num"], "statement": co["statement"]} for co in cos],
    })


@app.route("/upload_po_pomap", methods=["POST"])
@login_required
def upload_po_pomap():
    """Parse an uploaded Programme Outcome file into a custom PO list.

    Used by Module 4's "Upload Syllabus + PO File" option so the user's own
    POs drive the whole pipeline instead of the standard NBA POs. The parsed
    POs are returned to the client, which then configures them via
    /configure_po_mapping (same path as manually-entered custom POs)."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        if ext == ".pdf":
            from generate_cos import extract_pdf
            po_text = extract_pdf(tmp.name)
        else:
            po_text = extract(tmp.name)
    except Exception as e:
        return jsonify({"error": f"Failed to parse PO file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    pos = parse_pos_from_text(po_text)
    if not pos:
        return jsonify({"error": "No Programme Outcomes (PO1, PO2, ...) found in the file."}), 400

    return jsonify({"pos": pos})


@app.route("/configure_po_mapping", methods=["POST"])
@login_required
def configure_po_mapping():
    """Store the user's selected PO list (standard + custom) in the session."""
    data              = request.get_json() or {}
    selected_standard = data.get("selected_standard", [])   # ["PO1", "PO2", ...]
    custom_pos        = data.get("custom_pos", [])           # [{key, name, desc}, ...]

    std_map = {p[0]: p for p in STANDARD_POS}
    pos = [std_map[k] for k in selected_standard if k in std_map]
    for cp in custom_pos:
        key  = (cp.get("key") or "").strip().upper()
        name = (cp.get("name") or "").strip()
        desc = (cp.get("desc") or "").strip()
        if key and name:
            pos.append((key, name, desc))

    if not pos:
        return jsonify({"error": "Please select at least one PO."}), 400

    store = _get_store()
    store["pomap_pos"] = pos
    return jsonify({"ok": True, "count": len(pos)})


@app.route("/generate_po_mapping")
@tokens_required(1)
def generate_po_mapping_route():
    store = _get_store()
    cos          = store.get("pomap_cos", [])
    course_code  = store.get("pomap_course_code", "COURSE")
    course_title = store.get("pomap_course_title", "Untitled")
    pos          = store.get("pomap_pos")   # None = use all standard POs
    po_keys      = [p[0] for p in pos] if pos else None

    if not cos:
        return jsonify({"error": "No COs loaded. Please upload a CO file first."}), 400

    def event_stream():
        import anthropic
        client = anthropic.Anthropic(
            api_key=os.environ["ANTHROPIC_API_KEY"],
            timeout=90.0,
            max_retries=3,
        )
        chunks = []
        try:
            for chunk in generate_mapping_stream(
                client, cos, course_code, course_title, pos=pos
            ):
                chunks.append(chunk)
        except anthropic.APITimeoutError:
            yield f"data: {json.dumps({'error': 'Request timed out - try again.'})}\n\n"
            return
        except anthropic.APIConnectionError:
            yield f"data: {json.dumps({'error': 'Could not reach AI service - check your connection.'})}\n\n"
            return
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            return

        raw  = "".join(chunks)
        rows = parse_mapping_response(raw, cos, po_keys=po_keys)
        store["pomap_rows"] = rows
        yield f"data: {json.dumps({'done': True, 'rows': rows})}\n\n"

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/export_po_mapping", methods=["POST"])
@login_required
def export_po_mapping():
    store = _get_store()
    data         = request.get_json() or {}
    rows         = data.get("rows") or store.get("pomap_rows") or []
    course_code  = store.get("pomap_course_code", "COURSE")
    course_title = store.get("pomap_course_title", "Untitled")

    if not rows:
        return jsonify({"error": "No mapping generated. Please generate one first."}), 400

    fmt  = data.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    configured_pos = store.get("pomap_pos") or []
    has_pso = any(k.upper().startswith("PSO") for k, *_ in configured_pos)
    map_title_base = "CO-PO-PSO Mapping Table" if has_pso else "CO-PO Mapping Table"
    map_title_docx = map_title_base.replace("CO-PO", "CO - PO")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            po_save_txt(rows, course_code, course_title, tmp.name, title=map_title_base)
        elif fmt == "docx":
            po_save_docx(rows, course_code, course_title, tmp.name, title=map_title_docx)
        elif fmt == "pdf":
            po_save_pdf(rows, course_code, course_title, tmp.name, title=map_title_base)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{course_code}_CO_PO_Mapping{suffix}"'},
    )


# â"€â"€ Lesson Plan routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/upload_lp", methods=["POST"])
@login_required
def upload_lp():
    """Upload a syllabus for lesson plan generation  -  reuses existing course parsing."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".doc", ".txt"):
        return jsonify({"error": f"Unsupported file type '{ext}'"}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        text = extract(tmp.name)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not text.strip():
        return jsonify({"error": "No text could be extracted from the file"}), 400

    courses = split_into_courses(text)
    if not courses:
        return jsonify({"error": "No course codes detected in the file"}), 400

    store = _get_store()
    store["lp_courses"] = courses

    return jsonify({
        "courses": [
            {"code": code, "title": info["title"], "semester": info.get("semester")}
            for code, info in courses.items()
        ]
    })


@app.route("/upload_co_lp", methods=["POST"])
@login_required
def upload_co_lp():
    """Parse an uploaded CO file for lesson plan generation."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        if ext == ".pdf":
            from generate_qbank import load_cos_from_pdf
            cos = load_cos_from_pdf(tmp.name)
        else:
            raw_txt = extract(tmp.name)
            cos = parse_cos_for_mapping(raw_txt)
    except Exception as e:
        return jsonify({"error": f"Failed to parse CO file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not cos:
        return jsonify({"error": "No COs found in the uploaded file."}), 400

    store = _get_store()
    store["lp_uploaded_cos"] = cos
    return jsonify({"cos": [{"num": co["num"], "statement": co["statement"]} for co in cos]})


@app.route("/generate_lp")
@tokens_required(3)
def generate_lp():
    """SSE stream: generate all lesson plan sections for the selected course."""
    store = _get_store()
    courses = store.get("lp_courses", {})

    code = request.args.get("code", "").strip().upper()
    num_cos = max(3, min(10, int(request.args.get("num_cos", 5))))

    if code not in courses:
        return jsonify({"error": "Course not found"}), 400

    info = courses[code]

    meta = {
        "academic_year": request.args.get("academic_year", "2026 - 2027"),
        "semester":      request.args.get("semester", f"Semester {info.get('semester','III')}"),
        "regulation":    request.args.get("regulation", "R26"),
        "program":       request.args.get("program", "B.Tech  -  CSE"),
        "course_type":   request.args.get("course_type", "Theory + Lab"),
        "credits":       request.args.get("credits", "4"),
        "ltp":           request.args.get("ltp", "3-0-2"),
        "faculty_name":  request.args.get("faculty_name", "Dr. Faculty"),
        "department":    request.args.get("department", "Department of Computer Science and Engineering"),
    }

    store["lp_code"]  = code
    store["lp_title"] = info["title"]
    store["lp_meta"]  = meta
    store["lp_data"]  = None

    # Use pre-uploaded COs if available and requested
    use_uploaded = request.args.get("use_uploaded_cos", "false").lower() == "true"
    uploaded_cos = store.get("lp_uploaded_cos") if use_uploaded else None

    def event_stream():
        import anthropic
        client = anthropic.Anthropic(
            api_key=os.environ["ANTHROPIC_API_KEY"],
            timeout=120.0,
            max_retries=3,
        )
        try:
            for item in generate_lesson_plan_stream(
                client, code, info["title"], info["text"], meta,
                num_cos=num_cos, existing_cos=uploaded_cos
            ):
                if "error" in item:
                    _refund_tokens(3)
                    yield f"data: {json.dumps({'error': item['error']})}\n\n"
                    return
                if "heartbeat" in item:
                    # SSE comment: keeps the connection alive during generation,
                    # ignored by the browser's EventSource.
                    yield ": ping\n\n"
                    continue
                if "progress" in item:
                    yield f"data: {json.dumps({'progress': item['progress']})}\n\n"
                if "done" in item:
                    store["lp_data"] = item["data"]
                    yield f"data: {json.dumps({'done': True, 'data': item['data']})}\n\n"
        except Exception as e:
            _refund_tokens(3)
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/export_lp", methods=["POST"])
@login_required
def export_lp():
    """Export the generated lesson plan to DOCX, PDF, or TXT."""
    store = _get_store()
    data  = store.get("lp_data")
    code  = store.get("lp_code", "COURSE")
    title = store.get("lp_title", "Untitled")
    meta  = store.get("lp_meta", {})

    if not data:
        return jsonify({"error": "No lesson plan generated. Please generate one first."}), 400

    req  = request.get_json() or {}
    fmt  = req.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            lp_build_txt(data, meta, code, title, tmp.name)
        elif fmt == "docx":
            lp_build_docx(data, meta, code, title, tmp.name)
        elif fmt == "pdf":
            lp_build_pdf(data, meta, code, title, tmp.name)

        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    filename = f"{code}_LessonPlan{suffix}"
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# â"€â"€ Teaching Diary routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/upload_td", methods=["POST"])
@login_required
def upload_td():
    """Upload a syllabus for teaching diary generation."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".doc", ".txt"):
        return jsonify({"error": f"Unsupported file type '{ext}'"}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        text = extract(tmp.name)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not text.strip():
        return jsonify({"error": "No text could be extracted from the file"}), 400

    courses = split_into_courses(text)
    if not courses:
        return jsonify({"error": "No course codes detected in the file"}), 400

    store = _get_store()
    store["td_courses"] = courses

    return jsonify({
        "courses": [
            {"code": code, "title": info["title"], "semester": info.get("semester")}
            for code, info in courses.items()
        ]
    })


@app.route("/upload_co_td", methods=["POST"])
@login_required
def upload_co_td():
    """Parse an uploaded CO file for teaching diary generation."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        if ext == ".pdf":
            from generate_qbank import load_cos_from_pdf
            cos = load_cos_from_pdf(tmp.name)
        else:
            raw_txt = extract(tmp.name)
            cos = parse_cos_for_mapping(raw_txt)
    except Exception as e:
        return jsonify({"error": f"Failed to parse CO file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not cos:
        return jsonify({"error": "No COs found in the uploaded file."}), 400

    store = _get_store()
    store["td_uploaded_cos"] = cos
    return jsonify({"cos": [{"num": co["num"], "statement": co["statement"]} for co in cos]})


@app.route("/generate_td")
@tokens_required(3)
def generate_td():
    """SSE stream: generate all teaching diary sections for the selected course."""
    store = _get_store()
    courses = store.get("td_courses", {})

    code = request.args.get("code", "").strip().upper()
    num_cos = max(3, min(10, int(request.args.get("num_cos", 5))))

    if code not in courses:
        return jsonify({"error": "Course not found"}), 400

    info = courses[code]

    ltp = request.args.get("ltp", "3-0-2")
    course_type = request.args.get("course_type", "Theory + Lab")
    is_lab = "lab" in course_type.lower() or any(
        x in ltp.split("-")[2] if len(ltp.split("-")) > 2 else "0"
        for x in ["1","2","3","4"]
    )

    meta = {
        "academic_year":  request.args.get("academic_year", "2026 - 2027"),
        "semester":       request.args.get("semester", f"Semester {info.get('semester','III')}"),
        "regulation":     request.args.get("regulation", "R26"),
        "program":        request.args.get("program", "B.Tech  -  CSE"),
        "course_type":    course_type,
        "credits":        request.args.get("credits", "4"),
        "ltp":            ltp,
        "faculty_name":   request.args.get("faculty_name", "Dr. Faculty"),
        "department":     request.args.get("department", "Department of Computer Science and Engineering"),
        "section":        request.args.get("section", "A"),
        "total_hours":    request.args.get("total_hours", "45"),
        "coordinator":    request.args.get("coordinator", ""),
        "hod":            request.args.get("hod", ""),
    }

    store["td_code"]  = code
    store["td_title"] = info["title"]
    store["td_meta"]  = meta
    store["td_data"]  = None

    use_uploaded = request.args.get("use_uploaded_cos", "false").lower() == "true"
    uploaded_cos = store.get("td_uploaded_cos") if use_uploaded else None

    def event_stream():
        import anthropic
        client = anthropic.Anthropic(
            api_key=os.environ["ANTHROPIC_API_KEY"],
            timeout=150.0,
            max_retries=3,
        )
        try:
            for item in generate_teaching_diary_stream(
                client, code, info["title"], info["text"], meta,
                num_cos=num_cos, existing_cos=uploaded_cos, is_lab=is_lab
            ):
                if "error" in item:
                    _refund_tokens(3)
                    yield f"data: {json.dumps({'error': item['error']})}\n\n"
                    return
                if "heartbeat" in item:
                    # SSE comment: keeps the connection alive during generation,
                    # ignored by the browser's EventSource.
                    yield ": ping\n\n"
                    continue
                if "progress" in item:
                    yield f"data: {json.dumps({'progress': item['progress']})}\n\n"
                if "done" in item:
                    store["td_data"] = item["data"]
                    yield f"data: {json.dumps({'done': True, 'data': item['data']})}\n\n"
        except Exception as e:
            _refund_tokens(3)
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/export_td", methods=["POST"])
@login_required
def export_td():
    """Export the generated teaching diary to DOCX, PDF, or TXT."""
    store = _get_store()
    data  = store.get("td_data")
    code  = store.get("td_code", "COURSE")
    title = store.get("td_title", "Untitled")
    meta  = store.get("td_meta", {})

    if not data:
        return jsonify({"error": "No teaching diary generated. Please generate one first."}), 400

    req  = request.get_json() or {}
    fmt  = req.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            td_build_txt(data, meta, code, title, tmp.name)
        elif fmt == "docx":
            td_build_docx(data, meta, code, title, tmp.name)
        elif fmt == "pdf":
            td_build_pdf(data, meta, code, title, tmp.name)

        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    filename = f"{code}_TeachingDiary{suffix}"
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# â"€â"€ Password reset routes â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    message = None
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        users = load_users()
        match = next(
            (u for u, d in users.items()
             if isinstance(d, dict) and d.get("email", "").lower() == email),
            None,
        )
        if match:
            tokens = _purge_expired(_load_tokens())
            token  = secrets.token_urlsafe(32)
            tokens[token] = {"username": match, "expires": time.time() + 3600}
            _save_tokens(tokens)
            reset_url = url_for("reset_password", token=token, _external=True)
            threading.Thread(target=_send_reset_email, args=(email, reset_url), daemon=True).start()
        # Always show the same message  -  don't reveal whether the email exists
        message = ("If an account with that email exists, we've sent a reset link. "
                   "Check your inbox (and spam folder).")
    return render_template("forgot_password.html", message=message)


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    tokens = _purge_expired(_load_tokens())
    entry  = tokens.get(token)
    if not entry:
        return render_template("reset_password.html",
                               token=None, error="This reset link is invalid or has expired.")
    error = None
    if request.method == "POST":
        password = request.form.get("password", "")
        confirm  = request.form.get("confirm",  "")
        if len(password) < 6:
            error = "Password must be at least 6 characters."
        elif password != confirm:
            error = "Passwords do not match."
        else:
            users = load_users()
            username = entry["username"]
            if username in users:
                users[username]["password"] = generate_password_hash(password)
                _save_users(users)
            del tokens[token]
            _save_tokens(tokens)
            return redirect(url_for("login") + "?reset=1")
    return render_template("reset_password.html", token=token, error=error)


# â"€â"€ Settings â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings():
    username = session["username"]
    email_error = email_ok = None

    if request.method == "POST":
        action = request.form.get("action", "email")

        if action == "email":
            new_email = request.form.get("email", "").strip().lower()
            if not new_email or "@" not in new_email or "." not in new_email.split("@")[-1]:
                email_error = "Please enter a valid email address."
            else:
                users = load_users()
                duplicate = any(
                    u != username and isinstance(d, dict)
                    and d.get("email", "").lower() == new_email
                    for u, d in users.items()
                )
                if duplicate:
                    email_error = "That email is already used by another account."
                else:
                    users[username]["email"] = new_email
                    _save_users(users)
                    email_ok = "Email updated successfully."

    users = load_users()
    user_data     = users.get(username, {})
    current_tokens = user_data.get("tokens", 0)
    current_email  = user_data.get("email", "")
    current_college_type = user_data.get("college_type", "engineering")
    session["tokens"] = current_tokens
    return render_template("settings.html", username=username, tokens=current_tokens,
                           packages=TOPUP_PACKAGES,
                           email=current_email, email_error=email_error, email_ok=email_ok,
                           college_type=current_college_type,
                           college_type_meta=COLLEGE_TYPE_META)


@app.route("/topup", methods=["POST"])
@login_required
def topup():
    pkg_id = request.json.get("package_id", "") if request.is_json else request.form.get("package_id", "")
    pkg = next((p for p in TOPUP_PACKAGES if p["id"] == pkg_id), None)
    if not pkg:
        return jsonify({"error": "Invalid package."}), 400
    username = session["username"]
    users = load_users()
    users[username]["tokens"] = users[username].get("tokens", 0) + pkg["tokens"]
    _save_users(users)
    session["tokens"] = users[username]["tokens"]
    return jsonify({"ok": True, "tokens": users[username]["tokens"], "added": pkg["tokens"]})


@app.route("/api/token_balance")
@login_required
def api_token_balance():
    users = load_users()
    balance = users.get(session["username"], {}).get("tokens", 0)
    session["tokens"] = balance
    return jsonify({"tokens": balance})


@app.route("/api/college_types")
def api_college_types():
    """Return all supported college types with labels."""
    return jsonify([
        {"value": k, "label": v["label"], "accreditation": v["accreditation"]}
        for k, v in COLLEGE_TYPE_META.items()
    ])


@app.route("/api/college_pos")
@login_required
def api_college_pos():
    """Return the standard PO list for the current user's college type."""
    pos = _get_standard_pos()
    return jsonify([
        {"key": p[0], "name": p[1], "desc": p[2]}
        for p in pos
    ])


@app.route("/set_college_type", methods=["POST"])
@login_required
def set_college_type():
    """Update college type for the logged-in user."""
    data         = request.get_json() or {}
    college_type = (data.get("college_type") or "engineering").strip()
    if college_type not in _COLLEGE_CONTEXTS:
        return jsonify({"error": "Unknown college type"}), 400
    users = load_users()
    username = session["username"]
    users[username]["college_type"] = college_type
    _save_users(users)
    session["college_type"] = college_type
    ctx = _COLLEGE_CONTEXTS[college_type]
    return jsonify({"ok": True, "college_type": college_type, "label": ctx["type"],
                    "accreditation": ctx["accreditation"]})


# ── Razorpay payment routes ───────────────────────────────────────────────────

@app.route("/create_order", methods=["POST"])
@login_required
def create_order():
    rz_key_id     = os.environ.get("RAZORPAY_KEY_ID", "").strip()
    rz_key_secret = os.environ.get("RAZORPAY_KEY_SECRET", "").strip()
    if not rz_key_id or not rz_key_secret:
        return jsonify({"error": "Payment gateway not configured. Contact the administrator."}), 503

    data   = request.get_json() or {}
    pkg_id = data.get("package_id", "")
    pkg    = next((p for p in TOPUP_PACKAGES if p["id"] == pkg_id), None)
    if not pkg:
        return jsonify({"error": "Invalid package."}), 400

    try:
        import razorpay
        client = razorpay.Client(auth=(rz_key_id, rz_key_secret))
        order  = client.order.create({
            "amount":   pkg["price"] * 100,   # paise
            "currency": "INR",
            "receipt":  f"{session['username']}_{pkg_id}_{int(time.time())}",
            "notes": {
                "username":   session["username"],
                "package_id": pkg_id,
            },
        })
    except ImportError:
        return jsonify({"error": "razorpay package not installed. Run: pip install razorpay"}), 500
    except Exception as e:
        return jsonify({"error": f"Could not create order: {e}"}), 500

    return jsonify({
        "order_id": order["id"],
        "amount":   order["amount"],
        "currency": order["currency"],
        "key_id":   rz_key_id,
        "package_id": pkg_id,
        "tokens":   pkg["tokens"],
        "label":    pkg["label"],
    })


@app.route("/verify_payment", methods=["POST"])
@login_required
def verify_payment():
    rz_key_id     = os.environ.get("RAZORPAY_KEY_ID", "").strip()
    rz_key_secret = os.environ.get("RAZORPAY_KEY_SECRET", "").strip()
    if not rz_key_id or not rz_key_secret:
        return jsonify({"error": "Payment gateway not configured."}), 503

    data       = request.get_json() or {}
    order_id   = data.get("razorpay_order_id", "")
    payment_id = data.get("razorpay_payment_id", "")
    signature  = data.get("razorpay_signature", "")
    pkg_id     = data.get("package_id", "")

    pkg = next((p for p in TOPUP_PACKAGES if p["id"] == pkg_id), None)
    if not pkg:
        return jsonify({"error": "Invalid package."}), 400

    try:
        import razorpay
        client = razorpay.Client(auth=(rz_key_id, rz_key_secret))
        client.utility.verify_payment_signature({
            "razorpay_order_id":   order_id,
            "razorpay_payment_id": payment_id,
            "razorpay_signature":  signature,
        })
    except Exception:
        return jsonify({"error": "Payment signature verification failed."}), 400

    username = session["username"]
    users    = load_users()
    users[username]["tokens"] = users[username].get("tokens", 0) + pkg["tokens"]
    _save_users(users)
    session["tokens"] = users[username]["tokens"]
    return jsonify({"ok": True, "tokens": users[username]["tokens"], "added": pkg["tokens"]})


@app.route("/razorpay_webhook", methods=["POST"])
def razorpay_webhook():
    webhook_secret = os.environ.get("RAZORPAY_WEBHOOK_SECRET", "").strip()
    if not webhook_secret:
        return jsonify({"status": "webhook secret not configured"}), 200

    sig  = request.headers.get("X-Razorpay-Signature", "")
    body = request.get_data()

    try:
        import razorpay
        rz_key_id     = os.environ.get("RAZORPAY_KEY_ID", "").strip()
        rz_key_secret = os.environ.get("RAZORPAY_KEY_SECRET", "").strip()
        client = razorpay.Client(auth=(rz_key_id, rz_key_secret))
        client.utility.verify_webhook_signature(body.decode(), sig, webhook_secret)
    except Exception:
        return jsonify({"error": "Invalid webhook signature"}), 400

    event = request.get_json(force=True) or {}
    if event.get("event") == "payment.captured":
        notes    = event["payload"]["payment"]["entity"].get("notes", {})
        username = notes.get("username", "")
        pkg_id   = notes.get("package_id", "")
        pkg      = next((p for p in TOPUP_PACKAGES if p["id"] == pkg_id), None)
        if username and pkg:
            users = load_users()
            if username in users:
                users[username]["tokens"] = users[username].get("tokens", 0) + pkg["tokens"]
                _save_users(users)

    return jsonify({"status": "ok"})


# ── SDG Contribution export routes ───────────────────────────────────────────

@app.route("/upload_po_sdg", methods=["POST"])
@login_required
def upload_po_sdg():
    """Parse an uploaded PO attainment file and return [{name, attainment}]."""
    import anthropic as _ant, json as _json

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f   = request.files["file"]
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".txt", ".docx", ".doc", ".pdf"):
        return jsonify({"error": f"Unsupported type '{ext}'. Use TXT, DOCX, or PDF."}), 400

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    try:
        f.save(tmp.name)
        tmp.close()
        raw_txt = extract(tmp.name)
    except Exception as e:
        return jsonify({"error": f"Failed to read file: {e}"}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    if not raw_txt.strip():
        return jsonify({"error": "File appears to be empty or could not be read."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    prompt = (
        "Extract Programme Outcome (PO) names and attainment percentages from the following text.\n"
        "Each PO should have a name (e.g. PO1, PO2) and a numeric attainment value (0-100).\n"
        "Return ONLY a valid JSON array with no explanation:\n"
        '[{"name":"PO1","attainment":75},{"name":"PO2","attainment":80}]\n\n'
        f"Text:\n{raw_txt[:4000]}"
    )
    try:
        client  = _ant.Anthropic(api_key=api_key, timeout=20.0)
        resp    = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=512,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        pos = _json.loads(raw.strip())
    except Exception as e:
        return jsonify({"error": f"Failed to parse PO data: {e}"}), 500

    if not pos:
        return jsonify({"error": "No PO attainment data found in the file."}), 400

    # Normalise
    pos = [{"name": str(p.get("name","")).strip(), "attainment": float(p.get("attainment", 0))}
           for p in pos if p.get("name")]
    return jsonify({"pos": pos})


@app.route("/generate_po_attainments", methods=["POST"])
@tokens_required(1)
def generate_po_attainments():
    """AI estimates PO attainment percentages (PO1-PO12) from the stored syllabus."""
    import anthropic as _ant, json as _json

    body   = request.get_json() or {}
    course = (body.get("course") or "").strip()

    store  = _get_store()
    courses = store.get("courses", {})
    info    = courses.get(course.upper(), {})
    text    = info.get("text", "")[:3000]
    title   = info.get("title", course)

    if not text:
        return jsonify({"error": "Syllabus text not found. Please upload the syllabus first."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    prompt = (
        f"You are an NBA/NAAC accreditation expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({course})\n"
        f"Syllabus excerpt:\n{text}\n\n"
        "Estimate realistic Programme Outcome (PO) attainment percentages for this course. "
        "The 12 standard POs are:\n"
        "PO1: Engineering Knowledge · PO2: Problem Analysis · PO3: Design/Development of Solutions\n"
        "PO4: Conduct Investigations · PO5: Modern Tool Usage · PO6: The Engineer and Society\n"
        "PO7: Environment and Sustainability · PO8: Ethics · PO9: Individual and Team Work\n"
        "PO10: Communication · PO11: Project Management and Finance · PO12: Life-long Learning\n\n"
        "Assign attainment % (0-100) based on how strongly the course content addresses each PO. "
        "Core technical POs directly covered by the syllabus should score higher. "
        "Use realistic values (e.g. 55-85 for relevant POs, 30-55 for partially relevant ones).\n\n"
        "Return ONLY a valid JSON array, no explanation:\n"
        '[{"name":"PO1","attainment":75},{"name":"PO2","attainment":80},...]'
    )

    try:
        client  = _ant.Anthropic(api_key=api_key, timeout=25.0)
        resp    = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=512,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        pos = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    pos = [{"name": str(p.get("name","")).strip(), "attainment": float(p.get("attainment", 0))}
           for p in pos if p.get("name")]
    if not pos:
        return jsonify({"error": "AI returned no PO data."}), 500

    return jsonify({"pos": pos})


@app.route("/generate_po_sdg_weights", methods=["POST"])
@login_required
def generate_po_sdg_weights():
    """AI rates PO1-PO12 contributions to a single SDG (on-demand, no token cost)."""
    import anthropic as _ant, json as _json

    body   = request.get_json() or {}
    course = (body.get("course") or "").strip()
    sdg    = (body.get("sdg") or "").strip()

    if not sdg:
        return jsonify({"error": "SDG label required"}), 400

    store  = _get_store()
    info   = store.get("courses", {}).get(course.upper(), {})
    text   = info.get("text", "")[:2500]
    title  = info.get("title", course)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    prompt = (
        f"You are an NBA/NAAC accreditation expert.\n\n"
        f"Course: {title} ({course})\nSyllabus excerpt:\n{text}\n\n"
        f"Rate how strongly each of PO1–PO12 contributes to {sdg}.\n"
        f"Scale: 3=Strong, 2=Moderate, 1=Low, 0=None.\n\n"
        f"Return ONLY a valid JSON object — no explanation:\n"
        f'{{"PO1":2,"PO2":3,"PO3":2,"PO4":1,"PO5":2,"PO6":1,"PO7":0,"PO8":0,"PO9":1,"PO10":1,"PO11":1,"PO12":1}}'
    )

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=20.0)
        resp   = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=256,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        weights = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    weights = {str(k): float(v) for k, v in weights.items()}
    return jsonify({"weights": weights})


@app.route("/generate_sdg_po_auto", methods=["POST"])
@tokens_required(1)
def generate_sdg_po_auto():
    """AI generates PO attainments + PO→SDG weights from stored syllabus in one call."""
    import anthropic as _ant, json as _json

    body   = request.get_json() or {}
    course = (body.get("course") or "").strip()

    store   = _get_store()
    courses = store.get("courses", {})
    info    = courses.get(course.upper(), {})
    text    = info.get("text", "")[:3500]
    title   = info.get("title", course)

    if not text:
        return jsonify({"error": "Syllabus not found. Upload the syllabus first."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    prompt = (
        f"You are an NBA/NAAC accreditation expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({course})\nSyllabus excerpt:\n{text}\n\n"
        "Based on this syllabus, do TWO things:\n\n"
        "1. Estimate realistic PO attainment percentages (0-100) for PO1-PO12:\n"
        "   PO1: Engineering Knowledge  PO2: Problem Analysis  PO3: Design/Development of Solutions\n"
        "   PO4: Conduct Investigations  PO5: Modern Tool Usage  PO6: The Engineer and Society\n"
        "   PO7: Environment and Sustainability  PO8: Ethics  PO9: Individual and Team Work\n"
        "   PO10: Communication  PO11: Project Management and Finance  PO12: Life-long Learning\n"
        "   Core technical POs directly covered should score 65-85; partially relevant ones 40-65.\n\n"
        "2. Identify the top 3-5 UN SDGs this course contributes to, and rate how strongly each PO\n"
        "   contributes to each selected SDG (3=Strong, 2=Moderate, 1=Low, 0=None).\n\n"
        "Return ONLY a valid JSON object with no explanation or markdown:\n"
        "{\n"
        '  "pos": [{"name":"PO1","attainment":75},{"name":"PO2","attainment":80},...],\n'
        '  "sdgs_selected": ["SDG 4 — Quality Education","SDG 9 — Industry, Innovation and Infrastructure"],\n'
        '  "weights_all": {\n'
        '    "SDG 4 — Quality Education": {"PO1":2,"PO2":1,"PO3":1,"PO4":0,"PO5":1,"PO6":1,"PO7":0,"PO8":1,"PO9":1,"PO10":2,"PO11":1,"PO12":2},\n'
        '    "SDG 9 — Industry, Innovation and Infrastructure": {"PO1":3,"PO2":3,"PO3":3,"PO4":2,"PO5":3,"PO6":1,"PO7":1,"PO8":0,"PO9":1,"PO10":1,"PO11":2,"PO12":1}\n'
        "  }\n"
        "}"
    )

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=35.0)
        resp   = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=1200,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        data = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    pos = [
        {"name": str(p.get("name", "")).strip(), "attainment": float(p.get("attainment", 0))}
        for p in data.get("pos", []) if p.get("name")
    ]
    sdgs_selected = [str(s) for s in data.get("sdgs_selected", []) if s]
    weights_all   = {str(k): {str(pk): float(pv) for pk, pv in v.items()}
                     for k, v in data.get("weights_all", {}).items()}

    if not pos:
        return jsonify({"error": "AI returned no PO data — try again"}), 500
    if not sdgs_selected:
        return jsonify({"error": "AI returned no SDG mapping — try again"}), 500

    return jsonify({"pos": pos, "sdgs_selected": sdgs_selected, "weights_all": weights_all})


@app.route("/generate_sdg_po_all", methods=["POST"])
@tokens_required(1)
def generate_sdg_po_all():
    """AI estimates PO attainments and rates each PO against ALL 17 SDGs in one call."""
    import anthropic as _ant, json as _json

    body   = request.get_json() or {}
    course = (body.get("course") or "").strip()

    store  = _get_store()
    info   = store.get("courses", {}).get(course.upper(), {})
    text   = info.get("text", "")[:3000]
    title  = info.get("title", course)

    if not text:
        return jsonify({"error": "Syllabus not found. Upload the syllabus first."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    # Use short SDG keys to minimise token count in the response
    short_sdgs = [f"SDG{i}" for i in range(1, 18)]
    sdg_key_row = "  ".join(short_sdgs)

    prompt = (
        f"You are an NBA/NAAC OBE expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({course})\nSyllabus excerpt:\n{text}\n\n"
        "Do TWO things:\n\n"
        "1. Estimate realistic PO attainment % (0-100) for PO1-PO12.\n"
        "   Core technical POs directly covered: 65-85; partially relevant: 40-65.\n\n"
        "2. Rate each PO's contribution to ALL 17 UN SDGs using short keys SDG1..SDG17:\n"
        "   SDG1=No Poverty  SDG2=Zero Hunger  SDG3=Good Health  SDG4=Quality Education\n"
        "   SDG5=Gender Equality  SDG6=Clean Water  SDG7=Clean Energy  SDG8=Decent Work\n"
        "   SDG9=Industry & Innovation  SDG10=Reduced Inequalities  SDG11=Sustainable Cities\n"
        "   SDG12=Responsible Consumption  SDG13=Climate Action  SDG14=Life Below Water\n"
        "   SDG15=Life on Land  SDG16=Peace & Justice  SDG17=Partnerships\n\n"
        "   Scale: 3=Strong/direct, 2=Moderate, 1=Low/indirect, 0=No link.\n"
        "   For engineering courses, SDG4 and SDG9 usually score highest.\n\n"
        "Return ONLY valid JSON — no explanation:\n"
        "{\n"
        '  "pos": [{"name":"PO1","attainment":75},...,{"name":"PO12","attainment":50}],\n'
        '  "weights": {\n'
        '    "SDG1":{"PO1":0,"PO2":0,"PO3":0,"PO4":0,"PO5":0,"PO6":1,"PO7":0,"PO8":1,"PO9":1,"PO10":1,"PO11":0,"PO12":1},\n'
        '    "SDG4":{"PO1":2,"PO2":2,"PO3":1,"PO4":1,"PO5":1,"PO6":1,"PO7":0,"PO8":1,"PO9":1,"PO10":2,"PO11":1,"PO12":3},\n'
        '    "SDG9":{"PO1":3,"PO2":3,"PO3":3,"PO4":2,"PO5":3,"PO6":1,"PO7":1,"PO8":0,"PO9":1,"PO10":1,"PO11":2,"PO12":1},\n'
        "    ... all 17 SDGs ...\n"
        "  }\n"
        "}"
    )

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=50.0)
        resp   = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=3500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        data = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    pos = [
        {"name": str(p.get("name","")).strip(), "attainment": float(p.get("attainment", 0))}
        for p in data.get("pos", []) if p.get("name")
    ]
    if not pos:
        return jsonify({"error": "AI returned no PO data — try again"}), 500

    weights_raw = data.get("weights", {})

    # Map short key → full SDG label
    short_to_full = {f"SDG{i}": _ALL_SDGS_PY[i-1] for i in range(1, 18)}

    def _interp(pct):
        if pct >= 85: return "Excellent"
        if pct >= 70: return "Strong"
        if pct >= 50: return "Moderate"
        return "Weak"

    results = []
    for short, full in short_to_full.items():
        wts_for_sdg = weights_raw.get(short, weights_raw.get(full, {}))
        num = den = 0.0
        for po in pos:
            w = float(wts_for_sdg.get(po["name"], 0))
            if w > 0:
                num += po["attainment"] * w
                den += w
        contribution = round(num / den, 2) if den > 0 else 0.0
        results.append({"sdg": full, "contribution": contribution, "interpretation": _interp(contribution)})

    composite = round(sum(r["contribution"] for r in results) / len(results), 2) if results else 0.0

    return jsonify({"pos": pos, "results": results, "composite": composite})


@app.route("/export_sdg_po", methods=["POST"])
@login_required
def export_sdg_po():
    """Export PO → SDG contribution results as TXT, DOCX, or PDF."""
    body = request.get_json() or {}
    data = body.get("data", {})
    fmt  = body.get("fmt", "txt")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    pos      = data.get("pos", [])
    sdgs     = data.get("sdgs", [])
    results  = data.get("results", [])
    composite = data.get("composite", 0)

    def _interp(v):
        if v >= 85: return "Excellent SDG Integration"
        if v >= 70: return "Strong SDG Contribution"
        if v >= 50: return "Moderate Contribution"
        return "Weak SDG Integration"

    def _make_txt():
        lines = ["PO → SDG CONTRIBUTION ANALYSIS", "=" * 40, ""]
        lines.append("Programme Outcomes & Attainments:")
        for po in pos:
            lines.append(f"  {po.get('name','PO')}: {po.get('attainment',0):.2f}%")
        lines.append("")
        lines.append("SDG Contribution Results:")
        lines.append(f"  {'SDG':<12} {'Contribution (%)':>18}  Interpretation")
        lines.append("  " + "-" * 56)
        for r in results:
            lines.append(f"  {r.get('sdg',''):<12} {r.get('contribution',0):>17.2f}%  {r.get('interpretation','')}")
        lines.append("")
        lines.append(f"Composite SDG Index: {composite:.2f}% ({_interp(composite)})")
        return "\n".join(lines)

    def _make_docx(path):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(1)
            sec.left_margin = sec.right_margin = Inches(1)

        h = doc.add_heading("PO → SDG Contribution Analysis", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x31, 0x2E, 0x81)

        doc.add_heading("Programme Outcomes & Attainments", level=2)
        for po in pos:
            doc.add_paragraph(f"{po.get('name','PO')}: {po.get('attainment',0):.2f}%", style="List Bullet")

        doc.add_heading("SDG Contribution Results", level=2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["SDG", "Contribution (%)", "Interpretation"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for r in results:
            row = tbl.add_row().cells
            row[0].text = r.get("sdg", "")
            row[1].text = f"{r.get('contribution',0):.2f}%"
            row[2].text = r.get("interpretation", "")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Composite SDG Index: ").bold = True
        p.add_run(f"{composite:.2f}% ({_interp(composite)})")
        doc.save(path)

    def _make_pdf(path):
        from fpdf import FPDF, XPos, YPos
        def _safe(s):
            return str(s).replace("—", " - ").replace("–", " - ").encode("latin-1", "replace").decode("latin-1")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_margins(15, 15, 15)

        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 8, "PO -> SDG Contribution Analysis", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Programme Outcomes & Attainments", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        for po in pos:
            pdf.cell(0, 5, _safe(f"  {po.get('name','PO')}: {po.get('attainment',0):.2f}%"),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "SDG Contribution Results", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_fill_color(220, 220, 220)
        col_w = [100, 38, 42]   # 180mm total (page 210 - 2×15 margins)
        row_h = 7
        for lbl, w in zip(["SDG", "Contribution (%)", "Interpretation"], col_w):
            pdf.cell(w, row_h, lbl, border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9)
        for r in results:
            sdg_text   = _safe(r.get("sdg", ""))
            contrib    = _safe(f"{r.get('contribution', 0):.2f}%")
            interp     = _safe(r.get("interpretation", ""))
            x, y = pdf.get_x(), pdf.get_y()
            # SDG name may wrap — use multi_cell then match height for other cells
            pdf.multi_cell(col_w[0], 6, sdg_text, border=1)
            new_y    = pdf.get_y()
            cell_h   = new_y - y
            pdf.set_xy(x + col_w[0], y)
            pdf.cell(col_w[1], cell_h, contrib, border=1, align="C")
            pdf.cell(col_w[2], cell_h, interp,  border=1)
            pdf.set_y(new_y)
        pdf.ln(3)

        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, _safe(f"Composite SDG Index: {composite:.2f}% ({_interp(composite)})"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.output(path)

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            with open(tmp.name, "w", encoding="utf-8") as f:
                f.write(_make_txt())
        elif fmt == "docx":
            _make_docx(tmp.name)
        else:
            _make_pdf(tmp.name)
        with open(tmp.name, "rb") as f:
            content = f.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content, mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="SDG_PO_Contribution{suffix}"'},
    )


@app.route("/generate_sdg_co_contributions", methods=["POST"])
@tokens_required(1)
def generate_sdg_co_contributions():
    """AI identifies the most relevant SDG and rates each CO's contribution (0-3)."""
    import anthropic as _ant, json as _json

    body       = request.get_json() or {}
    cos        = body.get("cos", [])   # [{name, statement}]
    course     = (body.get("course") or "").strip()
    target_sdg = (body.get("target_sdg") or "SDG 1 — No Poverty").strip()

    if not cos:
        return jsonify({"error": "No COs provided"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    co_lines = "\n".join(
        f"- {c['name']}: {c.get('statement', '')}"
        for c in cos
    )
    course_ctx = f" for the course \"{course}\"" if course else ""
    prompt = (
        f"You are an expert in Outcome-Based Education and UN Sustainable Development Goals (SDGs).\n\n"
        f"Rate each of the following Course Outcomes{course_ctx} on how strongly each one contributes to "
        f"{target_sdg}. Consider direct AND indirect connections — for example, technology or analytical "
        f"skills can indirectly support any SDG through innovation, capacity building, or efficiency. "
        f"Always assign at least 1 to the COs that are most relevant relative to the others, even if "
        f"the connection is indirect. Use 0 only when there is truly no conceivable link.\n\n"
        f"Course Outcomes:\n{co_lines}\n\n"
        f"Rating scale:\n"
        f"  3 = Strong / direct contribution to {target_sdg}\n"
        f"  2 = Moderate contribution\n"
        f"  1 = Low / indirect contribution\n"
        f"  0 = No conceivable link\n\n"
        f"Return ONLY a valid JSON array with no explanation or markdown:\n"
        f'[{{"co":"CO1","weight":2}},{{"co":"CO2","weight":1}}]'
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=30.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 512,
            messages   = [{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        weights = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    stmt_map = {c["name"]: c.get("statement", "") for c in cos}
    wt_map   = {w.get("co", "").strip(): float(w.get("weight", 0)) for w in weights}

    total_score = sum(wt_map.get(c["name"], 0) for c in cos)
    contributions = []
    for c in cos:
        score = wt_map.get(c["name"], 0)
        pct   = round(score / total_score * 100, 2) if total_score else 0
        contributions.append({
            "co": c["name"], "statement": stmt_map.get(c["name"], ""),
            "score": score, "pct": pct,
        })

    total = round(sum(c["score"] for c in contributions), 2)
    return jsonify({
        "targetSdg": target_sdg, "contributions": contributions, "total": total,
    })


@app.route("/export_sdg_co", methods=["POST"])
@login_required
def export_sdg_co():
    """Export CO → SDG contribution analysis as TXT, DOCX, or PDF."""
    body = request.get_json() or {}
    data = body.get("data", {})
    fmt  = body.get("fmt", "txt")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    target_sdg    = data.get("targetSdg", "Target SDG")
    contributions = data.get("contributions", [])
    total         = data.get("total", 0)

    def _make_txt():
        lines = ["CO -> SDG CONTRIBUTION ANALYSIS", "=" * 34, "",
                 f"Target SDG: {target_sdg}", ""]
        lines.append("CO Contribution Breakdown:")
        lines.append(f"  {'CO':<8} {'Score':>6}  {'Share (%)':>10}  Statement")
        lines.append("  " + "-" * 70)
        for sv in contributions:
            stmt = sv.get("statement", "")
            lines.append(
                f"  {sv.get('co',''):<8} {sv.get('score',0):>6.2f}  {sv.get('pct',0):>9.2f}%  {stmt}"
            )
        lines.append(f"  {'Total':<8} {total:>6.2f}  {'100.00%':>10}")
        lines.append("")
        if contributions:
            top = max(contributions, key=lambda x: x.get("pct", 0))
            lines.append(f"Primary contributor: {top.get('co','')} ({top.get('pct',0):.2f}%)")
        return "\n".join(lines)

    def _make_docx(path):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(1)
            sec.left_margin = sec.right_margin = Inches(1)

        h = doc.add_heading("CO → SDG Contribution Analysis", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x06, 0x4E, 0x5B)
        doc.add_paragraph(f"Target SDG: {target_sdg}")

        doc.add_heading("CO Contribution Breakdown", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for i, lbl in enumerate(["CO", "Statement", "Score (0–3)", "Share (%)"]):
            t.rows[0].cells[i].text = lbl
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for sv in contributions:
            r = t.add_row().cells
            r[0].text = sv.get("co", "")
            r[1].text = sv.get("statement", "")
            r[2].text = f"{sv.get('score',0):.2f}"
            r[3].text = f"{sv.get('pct',0):.2f}%"
        tr = t.add_row().cells
        tr[0].text = "Total"
        tr[1].text = ""
        tr[2].text = f"{total:.2f}"
        tr[3].text = "100%"
        for c in tr:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].bold = True

        if contributions:
            top = max(contributions, key=lambda x: x.get("pct", 0))
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Primary contributor: ").bold = True
            p.add_run(f"{top.get('co','')} — {top.get('pct',0):.2f}%")
        doc.save(path)

    def _make_pdf(path):
        from fpdf import FPDF, XPos, YPos
        def _safe(s): return str(s).encode("latin-1", "replace").decode("latin-1")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_margins(15, 15, 15)

        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, "CO -> SDG Contribution Analysis",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 5, _safe(f"Target SDG: {target_sdg}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        COL_W  = [20, 105, 25, 27]   # CO | Statement | Score | Share(%)
        LINE_H = 5.5
        PAD    = 2

        def _wrap_text(text, max_w):
            """Split text into lines that fit within max_w."""
            words, lines, line = text.split(), [], ""
            for w in words:
                test = (line + " " + w).strip()
                if pdf.get_string_width(test) <= max_w - PAD * 2:
                    line = test
                else:
                    if line:
                        lines.append(line)
                    line = w
            if line:
                lines.append(line)
            return lines or [""]

        def _draw_row(x0, y0, cells):
            """cells = [(text, width, align, bold)]; draws all with equal row height."""
            pdf.set_font("Helvetica", "", 9)
            # calculate row height from the statement column (index 1)
            stmt_lines = _wrap_text(cells[1][0], cells[1][1])
            row_h = max(len(stmt_lines) * LINE_H + PAD * 2, LINE_H + PAD * 2)

            if y0 + row_h > pdf.h - pdf.b_margin:
                pdf.add_page()
                x0 = pdf.get_x()
                y0 = pdf.get_y()

            cx = x0
            for text, w, align, bold in cells:
                pdf.rect(cx, y0, w, row_h)
                pdf.set_font("Helvetica", "B" if bold else "", 9)
                if text == cells[1][0]:          # statement — wrap
                    for j, ln in enumerate(_wrap_text(text, w)):
                        pdf.set_xy(cx + PAD, y0 + PAD + j * LINE_H)
                        pdf.cell(w - PAD * 2, LINE_H, ln, align="L")
                else:
                    pdf.set_xy(cx + PAD, y0 + (row_h - LINE_H) / 2)
                    pdf.cell(w - PAD * 2, LINE_H, text, align=align)
                cx += w

            return y0 + row_h

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "CO Contribution Breakdown", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        # Header
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_fill_color(220, 220, 220)
        for lbl, w in zip(["CO", "Statement", "Score", "Share (%)"], COL_W):
            pdf.cell(w, 7, lbl, border=1, fill=True,
                     align="L" if lbl == "Statement" else "C")
        pdf.ln()

        # Data rows
        x0 = pdf.get_x()
        y0 = pdf.get_y()
        for sv in contributions:
            cells = [
                (_safe(sv.get("co", "")),            COL_W[0], "L", True),
                (_safe(sv.get("statement", "")),      COL_W[1], "L", False),
                (_safe(f"{sv.get('score',0):.2f}"),  COL_W[2], "C", False),
                (_safe(f"{sv.get('pct',0):.2f}%"),   COL_W[3], "C", False),
            ]
            y0 = _draw_row(x0, y0, cells)

        # Total row
        pdf.set_xy(x0, y0)
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_fill_color(235, 235, 235)
        pdf.cell(COL_W[0] + COL_W[1], 7, "Total", border=1, fill=True)
        pdf.cell(COL_W[2], 7, _safe(f"{total:.2f}"), border=1, align="C", fill=True)
        pdf.cell(COL_W[3], 7, "100%", border=1, align="C", fill=True)
        pdf.ln()

        if contributions:
            top = max(contributions, key=lambda x: x.get("pct", 0))
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 5,
                     _safe(f"Primary contributor: {top.get('co','')} ({top.get('pct',0):.2f}%)"),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.output(path)

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            with open(tmp.name, "w", encoding="utf-8") as f:
                f.write(_make_txt())
        elif fmt == "docx":
            _make_docx(tmp.name)
        else:
            _make_pdf(tmp.name)
        with open(tmp.name, "rb") as f:
            content = f.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content, mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="SDG_CO_Contributions{suffix}"'},
    )


_ALL_SDGS_PY = [
    "SDG 1 — No Poverty", "SDG 2 — Zero Hunger",
    "SDG 3 — Good Health and Well-being", "SDG 4 — Quality Education",
    "SDG 5 — Gender Equality", "SDG 6 — Clean Water and Sanitation",
    "SDG 7 — Affordable and Clean Energy", "SDG 8 — Decent Work and Economic Growth",
    "SDG 9 — Industry, Innovation and Infrastructure", "SDG 10 — Reduced Inequalities",
    "SDG 11 — Sustainable Cities and Communities",
    "SDG 12 — Responsible Consumption and Production", "SDG 13 — Climate Action",
    "SDG 14 — Life Below Water", "SDG 15 — Life on Land",
    "SDG 16 — Peace, Justice and Strong Institutions",
    "SDG 17 — Partnerships for the Goals",
]


@app.route("/generate_sdg_co_all", methods=["POST"])
@tokens_required(1)
def generate_sdg_co_all():
    """Rate each CO's contribution to all (or selected) SDGs in one AI call."""
    import anthropic as _ant, json as _json

    body     = request.get_json() or {}
    cos      = body.get("cos", [])         # [{name, statement}]
    course   = (body.get("course") or "").strip()
    sdg_list = body.get("sdg_list") or _ALL_SDGS_PY   # subset or all 17

    if not cos:
        return jsonify({"error": "No COs provided"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    store = _get_store()
    title = store.get("courses", {}).get(course.upper(), {}).get("title", course)
    client = _ant.Anthropic(api_key=api_key, timeout=60.0)

    co_lines  = "\n".join(f"- {c['name']}: {c.get('statement','')}" for c in cos)
    sdg_keys  = [s.split(" — ")[0] for s in sdg_list]   # "SDG 1", "SDG 2", ...
    sdg_block = "\n".join(f'  "{k}": <0-3>' for k in sdg_keys)
    prompt = (
        f"You are an OBE and UN SDG expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({course})\nCourse Outcomes:\n{co_lines}\n\n"
        f"Rate each CO's contribution to each of the following {len(sdg_list)} UN SDGs.\n"
        f"Scale: 3=Strong/direct, 2=Moderate, 1=Low/indirect, 0=No link.\n"
        f"For engineering courses, consider innovation, analytical skills, and capacity-building "
        f"as indirect links (score 1-2). Only use 0 when truly no conceivable connection.\n\n"
        f"Return ONLY a valid JSON object. Keys are CO names, values are objects mapping "
        f"SDG short labels (e.g. 'SDG 1') to integer scores:\n"
        "{{\n" +
        f'  "{cos[0]["name"]}": {{\n{sdg_block}\n  }},\n' +
        (f'  "{cos[1]["name"]}": {{...}},\n' if len(cos) > 1 else '') +
        "  ...\n}}"
    )

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=3000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        ai_matrix = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    # Normalise: map short key ("SDG 1") → score for each CO; also build full-key matrix
    short_to_full = {s.split(" — ")[0]: s for s in sdg_list}
    matrix      = {}   # {co_name: {full_sdg_label: score}}
    co_totals   = {}
    sdg_totals  = {s: 0 for s in sdg_list}

    for c in cos:
        name = c["name"]
        row  = ai_matrix.get(name, {})
        matrix[name] = {}
        co_totals[name] = 0
        for full_sdg in sdg_list:
            short = full_sdg.split(" — ")[0]
            score = int(row.get(short, row.get(full_sdg, 0)))
            score = max(0, min(3, score))
            matrix[name][full_sdg] = score
            co_totals[name]        += score
            sdg_totals[full_sdg]   += score

    top_sdgs = sorted([s for s in sdg_list if sdg_totals[s] > 0],
                      key=lambda s: -sdg_totals[s])

    return jsonify({
        "matrix":     matrix,
        "sdgs":       sdg_list,
        "cos":        [c["name"] for c in cos],
        "co_totals":  co_totals,
        "sdg_totals": sdg_totals,
        "top_sdgs":   top_sdgs,
    })


@app.route("/export_sdg_co_all", methods=["POST"])
@login_required
def export_sdg_co_all():
    """Export CO x SDG full matrix as TXT, DOCX, or PDF."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    body       = request.get_json() or {}
    fmt        = body.get("fmt", "docx")
    data       = body.get("data", {})
    matrix     = data.get("matrix", {})
    sdgs       = data.get("sdgs", [])
    cos        = data.get("cos", [])
    sdg_totals = data.get("sdg_totals", {})
    co_totals  = data.get("co_totals", {})
    course     = (data.get("course") or "COURSE").strip()
    title      = data.get("courseTitle", "")

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400
    if not matrix:
        return jsonify({"error": "No matrix data"}), 400

    def _s(v): return str(v) if v is not None else ""
    short = lambda sdg: sdg.split(" — ")[0]   # "SDG 1 — ..." → "SDG 1"

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {"txt": "text/plain",
              "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              "pdf": "application/pdf"}[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            lines = [f"CO-SDG Contribution Matrix — {title} ({course})", "=" * 56, ""]
            hdr = "CO    " + "  ".join(short(s).ljust(6) for s in sdgs)
            lines.append(hdr)
            lines.append("-" * len(hdr))
            for co in cos:
                row = matrix.get(co, {})
                lines.append(co.ljust(6) + "  ".join(str(row.get(s, 0)).ljust(6) for s in sdgs))
            lines += ["", "Top Contributing SDGs:"]
            top = sorted(sdgs, key=lambda s: -sdg_totals.get(s, 0))[:8]
            for s in top:
                if sdg_totals.get(s, 0) > 0:
                    lines.append(f"  {s}: {sdg_totals[s]}")
            with open(tmp.name, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))

        elif fmt == "docx":
            doc = Document()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Inches(0.7)
                sec.left_margin = sec.right_margin = Inches(0.7)
            h = doc.add_heading("CO-SDG Contribution Matrix", level=0)
            h.runs[0].font.color.rgb = RGBColor(6, 78, 91)
            sub = doc.add_paragraph(f"{title} ({course})")
            sub.runs[0].font.size = Pt(9)
            tbl = doc.add_table(rows=1 + len(cos), cols=1 + len(sdgs))
            tbl.style = "Table Grid"
            tbl.rows[0].cells[0].text = "CO \\ SDG"
            for i, s in enumerate(sdgs):
                tbl.rows[0].cells[i + 1].text = short(s)
            for ri, co in enumerate(cos):
                row = matrix.get(co, {})
                tbl.rows[ri + 1].cells[0].text = co
                for ci, s in enumerate(sdgs):
                    tbl.rows[ri + 1].cells[ci + 1].text = _s(row.get(s, 0))
            for cell in tbl.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            doc.save(tmp.name)

        else:  # pdf
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=12)
            pdf.add_page(orientation="L")   # landscape for wide matrix
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(6, 78, 91)
            pdf.cell(0, 8, f"CO-SDG Contribution Matrix — {title} ({course})",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
            cw  = max(10, min(16, 255 // (1 + len(sdgs))))
            co_w = 20
            pdf.set_font("Helvetica", "B", 7.5)
            pdf.set_fill_color(209, 250, 229)
            pdf.cell(co_w, 6, "CO", border=1, fill=True, align="C")
            for s in sdgs:
                pdf.cell(cw, 6, short(s), border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 7.5)
            for co in cos:
                row = matrix.get(co, {})
                pdf.cell(co_w, 5, co, border=1, align="C")
                for s in sdgs:
                    v = row.get(s, 0)
                    if v == 3:
                        pdf.set_fill_color(198, 224, 180); fill = True
                    elif v == 2:
                        pdf.set_fill_color(255, 242, 204); fill = True
                    elif v == 1:
                        pdf.set_fill_color(221, 238, 247); fill = True
                    else:
                        fill = False
                    pdf.cell(cw, 5, _s(v) if v else "–", border=1, align="C", fill=fill)
                    if fill:
                        pdf.set_fill_color(255, 255, 255)
                pdf.ln()
            pdf.output(tmp.name)

        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(content, mimetype=mime,
                    headers={"Content-Disposition":
                             f'attachment; filename="{course}_CO_SDG_Matrix{suffix}"'})


@app.route("/export_co_attainment", methods=["POST"])
@login_required
def export_co_attainment():
    """Export CO Attainment (Method 2 Tier-I) results as TXT, DOCX, or PDF."""
    body       = request.get_json() or {}
    data       = body.get("data", {})
    fmt        = body.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400

    co_results   = data.get("coResults", [])
    po_results   = data.get("poResults", [])
    n_students   = data.get("studentCount", 0)
    n_questions  = data.get("questionCount", 0)
    course_name  = data.get("courseName", "")
    course_code  = data.get("courseCode", "")
    thr          = data.get("thresholds", {"t1": 70, "t2": 60, "t3": 50})

    def _thr_line():
        return (f"≥{thr['t1']}% → Level 3  |  "
                f"≥{thr['t2']}% → Level 2  |  "
                f"≥{thr['t3']}% → Level 1  |  below → Level 0")

    def _make_txt():
        lines = ["CO ATTAINMENT REPORT  (Method 2 – Tier I)", "=" * 47, ""]
        if course_name:
            lines.append(f"Course : {course_name}" + (f"  ({course_code})" if course_code else ""))
        lines += [f"Students : {n_students}   |   Question columns : {n_questions}",
                  f"Thresholds : {_thr_line()}", ""]
        lines += ["CO ATTAINMENT", "-" * 38,
                  f"  {'CO':<10} {'CO %':>8}  {'Level':>6}"]
        for c in co_results:
            lines.append(f"  {c['co']:<10} {c['pct']:>7.2f}%  {c['level']:>6}")
        if po_results:
            lines += ["", "PO ATTAINMENT", "-" * 38]
            lines.append("  " + "  ".join(f"{p['po']:<8}" for p in po_results))
            lines.append("  " + "  ".join(f"{p['attainment']:<8}" for p in po_results))
        return "\n".join(lines)

    def _make_docx(path):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(1)
            sec.left_margin = sec.right_margin = Inches(1)

        h = doc.add_heading("CO Attainment Report  —  Method 2 Tier-I", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
        if course_name:
            doc.add_paragraph(f"Course: {course_name}" + (f"  ({course_code})" if course_code else ""))
        doc.add_paragraph(f"Students: {n_students}   |   Question columns: {n_questions}")
        doc.add_paragraph(f"Thresholds: {_thr_line()}")

        doc.add_heading("CO Attainment", level=2)
        t1 = doc.add_table(rows=1, cols=3)
        t1.style = "Table Grid"
        for i, lbl in enumerate(["CO", "CO %", "Attainment Level"]):
            cell = t1.rows[0].cells[i]
            cell.text = lbl
            cell.paragraphs[0].runs[0].bold = True
        for c in co_results:
            row = t1.add_row().cells
            row[0].text = c["co"]
            row[1].text = f"{c['pct']:.2f}%"
            row[2].text = str(c["level"])

        if po_results:
            doc.add_paragraph()
            doc.add_heading("PO Attainment", level=2)
            t2 = doc.add_table(rows=2, cols=len(po_results))
            t2.style = "Table Grid"
            for i, p in enumerate(po_results):
                t2.rows[0].cells[i].text = p["po"]
                t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                t2.rows[1].cells[i].text = str(p["attainment"])
        doc.save(path)

    def _make_pdf(path):
        from fpdf import FPDF, XPos, YPos
        def _s(v): return str(v).encode("latin-1", "replace").decode("latin-1")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_margins(15, 15, 15)

        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 9, "CO Attainment Report  -  Method 2 Tier-I",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        if course_name:
            label = f"Course: {course_name}" + (f"  ({course_code})" if course_code else "")
            pdf.cell(0, 6, _s(label), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 6, f"Students: {n_students}   |   Question columns: {n_questions}",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, _s(f"Thresholds: >={thr['t1']}%->L3, >={thr['t2']}%->L2, >={thr['t3']}%->L1, below->L0"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 7, "CO Attainment", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_fill_color(220, 220, 240)
        for lbl, w in [("CO", 55), ("CO %", 55), ("Attainment Level", 70)]:
            pdf.cell(w, 7, lbl, border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 9.5)
        for c in co_results:
            pdf.cell(55, 6, _s(c["co"]), border=1)
            pdf.cell(55, 6, _s(f"{c['pct']:.2f}%"), border=1, align="C")
            pdf.cell(70, 6, _s(str(c["level"])), border=1, align="C")
            pdf.ln()

        if po_results:
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, "PO Attainment", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            col_w = max(18, min(30, 170 // len(po_results)))
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(220, 220, 240)
            for p in po_results:
                pdf.cell(col_w, 7, _s(p["po"]), border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 9)
            for p in po_results:
                pdf.cell(col_w, 6, _s(str(p["attainment"])), border=1, align="C")
            pdf.ln()
        pdf.output(path)

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            with open(tmp.name, "w", encoding="utf-8") as f:
                f.write(_make_txt())
        elif fmt == "docx":
            _make_docx(tmp.name)
        else:
            _make_pdf(tmp.name)
        with open(tmp.name, "rb") as f:
            content = f.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content, mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="CO_Attainment{suffix}"'},
    )


@app.route("/regenerate_selected_cos", methods=["POST"])
@tokens_required(1)
def regenerate_selected_cos():
    """Regenerate specific COs (by number) while keeping the rest unchanged."""
    import anthropic as _ant, json as _json

    body         = request.get_json() or {}
    code         = (body.get("code") or "").strip().upper()
    to_replace   = [int(n) for n in body.get("to_replace", [])]
    existing_cos = body.get("existing_cos", [])   # [{name, statement}]

    if not to_replace:
        return jsonify({"error": "No COs selected for regeneration"}), 400

    store  = _get_store()
    info   = store.get("courses", {}).get(code, {})
    text   = info.get("text", "")[:2500]
    title  = info.get("title", code)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    replace_labels = ", ".join(f"CO{n}" for n in sorted(to_replace))
    keep_cos = [c for c in existing_cos if c["name"] not in {f"CO{n}" for n in to_replace}]
    keep_str = "\n".join(f"- {c['name']}: {c['statement']}" for c in keep_cos) if keep_cos else "None"

    prompt = (
        f"You are an NBA/NAAC accreditation expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({code})\n"
        + (f"Syllabus excerpt:\n{text}\n\n" if text else "\n")
        + f"Existing Course Outcomes to KEEP unchanged:\n{keep_str}\n\n"
        f"Generate NEW, DIFFERENT statements ONLY for: {replace_labels}\n"
        f"Requirements: start with a Bloom's taxonomy action verb, be specific and measurable, "
        f"clearly different from the existing COs listed above.\n\n"
        f"Return ONLY a valid JSON array — one entry per replaced CO — no explanation:\n"
        f'[{{"num": 1, "statement": "..."}}, {{"num": 3, "statement": "..."}}]'
    )

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=30.0)
        resp   = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=800,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        parsed = _json.loads(raw.strip())
        result = []
        for i, n in enumerate(sorted(to_replace)):
            stmt = parsed[i].get("statement", "") if i < len(parsed) else ""
            result.append({"num": n, "statement": stmt})
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI failed: {e}"}), 500

    return jsonify({"cos": result})


@app.route("/export_coatt_cos", methods=["POST"])
@login_required
def export_coatt_cos():
    """Export CO Attainment COs as TXT/DOCX/PDF."""
    body = request.get_json() or {}
    fmt  = body.get("fmt", "docx")
    cos  = body.get("cos", [])
    code = (body.get("code") or "COURSE").strip().upper()

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400
    if not cos:
        return jsonify({"error": "No COs to export"}), 400

    co_text = "\n".join(f"{c['name']}: {c['statement']}" for c in cos)

    store   = _get_store()
    info    = store.get("courses", {}).get(code, {"title": "", "semester": None})
    sem_lbl = f"  [Semester {info['semester']}]" if info.get("semester") else ""
    summary = bloom_level_summary(co_text)

    # Build taxonomy grid from cached co_taxonomy (most reliable)
    from generate_cos import _KDIMS, _LEVELS, _KDIM_NORM
    co_taxonomy = store.get("co_taxonomy", {})
    grid = {kd: {lv: [] for lv in _LEVELS} for kd in _KDIMS}
    if co_taxonomy:
        for c in cos:
            m = re.search(r'\d+', str(c.get('name', '')))
            if not m:
                continue
            co_num = int(m.group())
            entry = co_taxonomy.get(co_num)
            if not entry:
                continue
            kdim = _KDIM_NORM.get(entry['kdim'].lower(), entry['kdim'])
            bloom_lv = entry['bloom']
            if kdim in grid and bloom_lv in grid[kdim]:
                grid[kdim][bloom_lv].append(f'CO{co_num}')
    # Fallback: parse from current_result which has the markdown table
    if not any(grid[kd][lv] for kd in grid for lv in grid[kd]):
        grid = build_taxonomy_grid(store.get("current_result", ""))

    all_output = [
        f"# Course Outcomes\nCourse: {code} — {info.get('title', '')}{sem_lbl}\n",
        (summary + "\n\n" + co_text if summary else co_text) + "\n",
    ]

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            cos_save_txt(all_output, tmp.name, taxonomy_grid=grid)
        elif fmt == "docx":
            cos_save_docx(all_output, tmp.name, taxonomy_grid=grid)
        else:
            cos_save_pdf(all_output, tmp.name, taxonomy_grid=grid)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(
        content, mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="COs_{code}{suffix}"'},
    )


@app.route("/generate_co_attainment_ai", methods=["POST"])
@tokens_required(1)
def generate_co_attainment_ai():
    """AI estimates CO attainment % and PO attainment from CO statements."""
    import anthropic as _ant, json as _json

    body   = request.get_json() or {}
    cos    = body.get("cos", [])       # [{name, statement}, ...]
    course = (body.get("course") or "").strip()
    t1     = float(body.get("t1", 70))
    t2     = float(body.get("t2", 60))
    t3     = float(body.get("t3", 50))

    if not cos:
        return jsonify({"error": "No Course Outcomes provided"}), 400

    store  = _get_store()
    info   = store.get("courses", {}).get(course.upper(), {})
    title  = info.get("title", course)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    co_list_str = "\n".join(f"- {c['name']}: {c['statement']}" for c in cos)

    prompt = (
        f"You are an NBA/NAAC accreditation expert for Indian engineering colleges.\n\n"
        f"Course: {title} ({course})\nCourse Outcomes:\n{co_list_str}\n\n"
        "Based on these COs for a typical engineering course, do TWO things:\n\n"
        "1. Estimate CO attainment % (0-100) for each CO. Well-run courses typically attain 55-80%. "
        "Higher-order Bloom's levels (analyze/evaluate/create) usually attain lower than "
        "knowledge/comprehension COs.\n\n"
        "2. Rate each CO's contribution to PO1-PO12 (3=Strong, 2=Moderate, 1=Low, 0=None).\n"
        "   PO1: Engineering Knowledge  PO2: Problem Analysis  PO3: Design/Development\n"
        "   PO4: Conduct Investigations  PO5: Modern Tool Usage  PO6: Engineer and Society\n"
        "   PO7: Environment and Sustainability  PO8: Ethics  PO9: Individual and Team Work\n"
        "   PO10: Communication  PO11: Project Management  PO12: Life-long Learning\n\n"
        "Return ONLY valid JSON, no explanation or markdown:\n"
        '{"co_attainments":[{"co":"CO1","pct":72.5},...],'
        '"co_po_weights":{"CO1":{"PO1":3,"PO2":2,"PO3":2,"PO4":0,"PO5":1,"PO6":0,"PO7":0,"PO8":0,"PO9":1,"PO10":1,"PO11":1,"PO12":1},...}}'
    )

    try:
        client = _ant.Anthropic(api_key=api_key, timeout=40.0)
        resp   = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        data = _json.loads(raw.strip())
    except _ant.APITimeoutError:
        return jsonify({"error": "AI request timed out — try again"}), 504
    except _ant.APIConnectionError:
        return jsonify({"error": "Could not reach AI service"}), 503
    except Exception as e:
        return jsonify({"error": f"AI analysis failed: {e}"}), 500

    co_atts  = {a["co"]: float(a.get("pct", 60.0)) for a in data.get("co_attainments", [])}
    co_po_w  = data.get("co_po_weights", {})

    def _level(pct):
        if pct >= t1: return 3
        if pct >= t2: return 2
        if pct >= t3: return 1
        return 0

    co_results = []
    for c in cos:
        pct = co_atts.get(c["name"], 60.0)
        co_results.append({"co": c["name"], "pct": round(pct, 2), "level": _level(pct)})

    po_names = [f"PO{i}" for i in range(1, 13)]
    po_results = []
    for po in po_names:
        total_w = 0
        weighted_sum = 0.0
        for c in cos:
            w = int(co_po_w.get(c["name"], {}).get(po, 0))
            if w > 0:
                weighted_sum += _level(co_atts.get(c["name"], 60.0)) * w
                total_w += w
        att = round(weighted_sum / total_w, 2) if total_w > 0 else 0
        po_results.append({"po": po, "attainment": att})

    return jsonify({
        "coResults":     co_results,
        "poResults":     po_results,
        "studentCount":  0,
        "questionCount": len(cos),
        "courseName":    title,
        "courseCode":    course,
        "thresholds":    {"t1": t1, "t2": t2, "t3": t3},
    })


@app.route("/aio_prepare", methods=["POST"])
@login_required
def aio_prepare():
    """Set up session stores so the AIO pipeline can reuse existing SSE generation routes.

    Copies the uploaded course from store["courses"] into lp_courses, td_courses, and
    sets pomap_cos / lp_uploaded_cos / td_uploaded_cos so downstream SSE routes pick up
    the AIO-generated (and possibly user-edited) COs.

    Body: {code: str, cos: [{name: str, statement: str}]}
    """
    body    = request.get_json() or {}
    code    = (body.get("code") or "").strip().upper()
    cos_raw = body.get("cos", [])   # [{name:"CO1", statement:"..."}, ...]

    store   = _get_store()
    courses = store.get("courses", {})
    if code not in courses:
        return jsonify({"error": "Course not found — upload the syllabus first."}), 400

    info = courses[code]

    # Normalise COs and renumber sequentially (1, 2, 3 …)
    cos_norm = []
    orig_to_new = {}
    for i, c in enumerate(cos_raw):
        name     = str(c.get("name", ""))
        m        = re.search(r"\d+", name)
        orig_num = int(m.group()) if m else (i + 1)
        new_num  = i + 1
        orig_to_new[orig_num] = new_num
        stmt = (c.get("statement") or "").strip()
        cos_norm.append({"num": new_num, "statement": stmt})

    # Remap cached CO taxonomy to new numbers
    _old_tax = store.get("co_taxonomy", {})
    if _old_tax and orig_to_new:
        store["co_taxonomy"] = {
            new_n: _old_tax[orig_n]
            for orig_n, new_n in orig_to_new.items()
            if orig_n in _old_tax
        }

    # Copy course into LP / TD stores
    for key in ("lp_courses", "td_courses"):
        if key not in store:
            store[key] = {}
        store[key][code] = info

    # PO-Mapping store
    store["pomap_cos"]          = cos_norm
    store["pomap_course_code"]  = code
    store["pomap_course_title"] = info.get("title", code)
    store["pomap_rows"]         = None
    store["pomap_pos"]          = None   # use all 12 standard POs

    # LP / TD uploaded COs — generate_lp/td use these when use_uploaded_cos=true
    store["lp_uploaded_cos"] = cos_norm if cos_norm else None
    store["td_uploaded_cos"] = cos_norm if cos_norm else None

    # Also update current_result so generate_qbank can pick up these COs if needed.
    # Preserve the Bloom's table rows from the original result so the M1 report
    # RBT matrix can still be populated.
    if cos_norm:
        stmt_lines = [f"CO{c['num']}: {c['statement']}" for c in cos_norm]
        # Rebuild table rows with renumbered CO references
        table_lines = []
        for line in store.get("current_result", "").split("\n"):
            s = line.strip()
            if s.startswith("|") and s.endswith("|"):
                m = re.search(r'CO(\d+)', s)
                if m:
                    orig = int(m.group(1))
                    if orig in orig_to_new:
                        new = orig_to_new[orig]
                        table_lines.append(s.replace(f"CO{orig}", f"CO{new}", 1))
        header = ("| CO | Unit | Knowledge Dimension | Bloom Level |\n"
                  "|----|------|---------------------|-------------|")
        store["current_result"] = (
            "\n".join(stmt_lines) + "\n\n" +
            (header + "\n" + "\n".join(table_lines) if table_lines else "")
        )
    store["current_code"] = code

    return jsonify({"ok": True, "cos": cos_norm})


# ── Module 1 Comprehensive Report ────────────────────────────────────────────

@app.route("/generate_module1_report", methods=["POST"])
@login_required
def generate_module1_report():
    """Generate a single comprehensive DOCX report for all Module-1 deliverables."""
    import anthropic as _ant

    store   = _get_store()
    body    = request.get_json() or {}

    co_text      = store.get("current_result", "")
    qbank_blocks = store.get("qbank_all_blocks", [])
    co_tally     = store.get("qbank_co_tally", {})
    code         = (store.get("qbank_code") or store.get("current_code", "COURSE")).strip().upper()
    courses      = store.get("courses", {})
    info         = courses.get(code, {"title": "Untitled", "semester": None})
    title        = info.get("title", "Untitled")
    semester     = info.get("semester")

    if not co_text.strip():
        return jsonify({"error": "No COs found in session. Run Module 1 generation first."}), 400
    if not qbank_blocks:
        return jsonify({"error": "No Question Bank found in session. Run Module 1 generation first."}), 400

    charged, err = _charge_report("m1", co_text + "\x01" + "\n".join(qbank_blocks))
    if err:
        return err

    # Parse COs and Bloom's data
    from generate_qbank import parse_cos_from_text
    from generate_cos import bloom_level_summary
    from generate_qpaper import parse_qbank as _parse_qbank, is_lab_qbank

    cos           = parse_cos_from_text(co_text)
    bloom_summary = bloom_level_summary(co_text)
    raw_qb_text   = "\n".join(qbank_blocks)

    is_lab    = is_lab_qbank(raw_qb_text)
    lab_tally = store.get("qbank_lab_tally", {}) if is_lab else {}

    # Compute rule-based analytics (lab-aware)
    analytics = _m1r.compute_analytics(cos, co_tally, raw_qb_text,
                                       is_lab=is_lab, lab_tally=lab_tally)

    # Build sample paper
    sample_paper = []
    if not is_lab:
        try:
            qb_data = _parse_qbank(raw_qb_text)
            if qb_data and qb_data.get("units"):
                sample_paper = _m1r.build_sample_paper_text(qb_data, code, title)
                qb_data_for_report = qb_data
            else:
                qb_data_for_report = None
        except Exception:
            qb_data_for_report = None
    else:
        # Lab course: build a sample practical paper from the lab question bank
        try:
            from generate_qpaper import parse_lab_qbank
            lab_data = parse_lab_qbank(raw_qb_text)
            if lab_data and lab_data.get("units"):
                sample_paper = _m1r.build_sample_lab_paper_text(lab_data, code, title)
                qb_data_for_report = lab_data
            else:
                qb_data_for_report = None
        except Exception:
            qb_data_for_report = None

    # AI qualitative analysis (one call)
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    ai_data = {}
    if api_key and cos:
        try:
            client  = _ant.Anthropic(api_key=api_key, timeout=60.0)
            ai_data = _m1r.get_ai_analysis(client, cos, co_tally, bloom_summary,
                                            code, title, analytics)
        except Exception:
            ai_data = {}

    # Build RBT grid from stored co_taxonomy (most reliable — parsed at generation time)
    from generate_cos import _KDIMS as _M1_KDIMS, _LEVELS as _M1_LEVELS, _KDIM_NORM as _M1_KDIM_NORM
    _co_tax = store.get("co_taxonomy", {})
    if _co_tax:
        _m1_grid = {kd: {lv: [] for lv in _M1_LEVELS} for kd in _M1_KDIMS}
        for _co_n, _entry in _co_tax.items():
            _kdim = _M1_KDIM_NORM.get(_entry['kdim'].lower(), _entry['kdim'])
            _bl   = _entry['bloom']
            if _kdim in _m1_grid and _bl in _m1_grid[_kdim]:
                _m1_grid[_kdim][_bl].append(f'CO{_co_n}')
        taxonomy_grid = _m1_grid
    else:
        taxonomy_grid = None

    # Replace the AI's free-form, non-deterministic key metrics with deterministic,
    # data-driven ones spanning the WHOLE report, so the dashboard and the improvement
    # panel stay stable across rebuilds (applying a fix then measurably moves the right
    # metric and the weak list converges instead of churning new names each time).
    _la = _m1r._compute_bloom_analytics(cos, taxonomy_grid)
    _acc = ai_data.setdefault("accreditation_readiness", {})
    _km  = _m1r.build_key_metrics(analytics, la=_la, cos=cos)
    _acc["key_metrics"] = _km
    if not _acc.get("overall_score") and _km:
        _acc["overall_score"] = round(sum(m["score"] for m in _km) / len(_km))
    # Full current weak set; the panel filters this against a frozen baseline (see
    # /m1_improve_start) so applying fixes can only shrink the list, never grow it.
    store["m1_weak_actions_all"] = _m1r.weak_metric_actions(_km)

    fmt = body.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        if charged:
            _refund_tokens(2)
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    report_kwargs = dict(
        cos=cos, co_text=co_text, qb_data=qb_data_for_report,
        co_tally=co_tally, bloom_summary=bloom_summary,
        analytics=analytics, ai=ai_data, sample_paper=sample_paper,
        code=code, title=title, semester=semester, is_lab=is_lab,
        taxonomy_grid=taxonomy_grid,
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "docx":
            _m1r.build_docx(**report_kwargs, output_path=tmp.name)
        elif fmt == "pdf":
            _m1r.build_pdf(**report_kwargs, output_path=tmp.name)
        else:
            _m1r.build_txt(**report_kwargs, output_path=tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        if charged:
            _refund_tokens(2)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    filename = f"{code}_Module1_Report{suffix}"
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


import generate_m3_report as _m3r
import generate_m4_report as _m4r


# ── "Suggest ways to improve" context builders (Modules 3 & 4) ────────────────
# Suggestions-only panels (unlike Module 1, there is no source data to regenerate).
# At report-generation time we distil the analytics + AI summary into a compact text
# brief, store it in the session, and feed it to /module_report_suggest on demand.

def _fmt_metrics(metrics):
    rows = []
    for m in metrics or []:
        if not isinstance(m, dict):
            continue
        name   = str(m.get("metric", "")).strip()
        status = str(m.get("status", "")).strip()
        score  = m.get("score")
        if name:
            tail = f" ({status}{', ' + str(score) + '/100' if score is not None else ''})" if status else ""
            rows.append(f"  - {name}{tail}")
    return "\n".join(rows)


def _m3_suggest_context(analytics, ai_data):
    """Compact improvement brief for the Module 3 (delivery) report."""
    a    = analytics or {}
    ai   = ai_data or {}
    acc  = ai.get("accreditation_readiness", {}) or {}
    obe  = ai.get("obe_compliance", {}) or {}
    miss = [c for c in a.get("cos", []) if f"CO{c.get('num')}" not in set(a.get("cos_in_sessions", []))]
    parts = [
        f"Report type: Module 3 - Course Delivery (lesson plan + teaching diary).",
        f"CO-session coverage: {a.get('lp_co_coverage_pct', '?')}% "
        f"({len(miss)} of {a.get('n_cos', 0)} COs not yet mapped to a session).",
        f"Sessions planned: {a.get('n_sessions', 0)} | Teaching methods used: "
        f"{', '.join(f'{k}({v})' for k, v in (a.get('method_dist') or {}).items()) or 'none recorded'}.",
        f"SDGs covered: {', '.join(a.get('all_sdgs', [])) or 'none'}.",
        f"Readiness: {acc.get('readiness_level', '?')} ({acc.get('overall_score', '?')}/100).",
    ]
    km = _fmt_metrics(acc.get("key_metrics"))
    if km:
        parts.append("Readiness metrics:\n" + km)
    recs = (obe.get("recommendations") or []) + (acc.get("action_items") or [])
    if recs:
        parts.append("Existing recommendations:\n" + "\n".join(f"  - {r}" for r in recs[:6]))
    return "\n".join(parts)


def _m4_suggest_context(analytics, ai_data):
    """Compact improvement brief for the Module 4 (attainment) report."""
    a   = analytics or {}
    ai  = ai_data or {}
    th  = a.get("thresholds", {}) or {}
    t3  = th.get("t3", 50)
    weak_cos = [r.get("co", "") for r in a.get("co_results", []) if (r.get("pct") or 0) < t3]
    below_po = [p.get("po", "") for p in a.get("atr_rows", [])]
    parts = [
        f"Report type: Module 4 - CO/PO Attainment & SDG mapping.",
        f"Mean CO attainment: {a.get('mean_att_pct', '?')}% across {a.get('n_cos', 0)} COs "
        f"(target floor {t3}%).",
        f"CO attainment level spread (L3/L2/L1/L0): "
        f"{'/'.join(str((a.get('level_dist') or {}).get(l, 0)) for l in (3, 2, 1, 0))}.",
        f"COs below the attainment floor: {', '.join(weak_cos) or 'none'}.",
        f"POs not meeting target: {', '.join(below_po) or 'none'}.",
        f"Target SDG: {a.get('target_sdg') or 'not set'} | SDGs covered: "
        f"{', '.join(a.get('sdgs_covered', [])) or 'none'}.",
        f"Readiness score: {ai.get('readiness_score', '?')}/100.",
    ]
    recs = ai.get("recommendations") or []
    if recs:
        parts.append("Existing recommendations:\n" + "\n".join(f"  - {r}" for r in recs[:6]))
    return "\n".join(parts)


_SUGGEST_MODULES = {
    "m3": ("Module 3 (Course Delivery)",
           "course delivery quality, CO-session coverage, teaching-method variety, "
           "Bloom alignment, SDG integration and NBA/NAAC delivery evidence"),
    "m4": ("Module 4 (CO/PO Attainment)",
           "CO and PO attainment levels, attainment gaps, CO-PO mapping strength, "
           "SDG coverage and NBA/NAAC attainment evidence"),
}


@app.route("/module_report_suggest", methods=["POST"])
@tokens_required(1)
def module_report_suggest():
    """Return prioritised, actionable improvement suggestions for a Module 3 or 4 report.
    Suggestions-only: nothing is regenerated, the faculty applies the advice themselves."""
    import anthropic as _ant
    import json as _json

    body   = request.get_json() or {}
    module = (body.get("module") or "").strip().lower()
    if module not in _SUGGEST_MODULES:
        _refund_tokens(1)
        return jsonify({"error": "Unknown module"}), 400

    store = _get_store()
    ctx   = store.get(f"{module}_suggest_ctx")
    if not ctx or not ctx.get("context"):
        _refund_tokens(1)
        return jsonify({"error": f"No report in session. Generate the {_SUGGEST_MODULES[module][0]} report first."}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        _refund_tokens(1)
        return jsonify({"error": "API key not configured"}), 500

    label, focus = _SUGGEST_MODULES[module]
    prompt = (
        "You are an NBA/NAAC accreditation quality advisor for an Indian engineering college.\n"
        f"You are reviewing the {label} report for {ctx.get('title', 'a course')} "
        f"({ctx.get('code', '')}). Focus on {focus}.\n\n"
        "Report brief:\n"
        f"{ctx['context']}\n\n"
        "Identify the most impactful ways the faculty can improve this report's weak areas. "
        "Return 3 to 5 suggestions, hardest-hitting first. For each, give a short area title, "
        "a severity, and one or two sentences of concrete, specific advice the faculty can act "
        "on (no generic filler). If the report already looks strong, return fewer items.\n\n"
        'Return ONLY a JSON object of this exact shape, no markdown, no commentary:\n'
        '{"suggestions": [{"area": "CO-Session Coverage", "severity": "High", '
        '"advice": "Map CO3 and CO5 to specific sessions ..."}]}\n'
        "severity must be one of: High, Medium, Low."
    )

    try:
        client   = _ant.Anthropic(api_key=api_key, timeout=40.0)
        response = client.messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 900,
            messages   = [{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        raw = re.sub(r"```[a-zA-Z]*", "", raw).replace("```", "").strip()
        start, end = raw.find("{"), raw.rfind("}")
        data = _json.loads(raw[start:end + 1]) if (start >= 0 and end > start) else {}
        out  = []
        for s in (data.get("suggestions") or []):
            if not isinstance(s, dict):
                continue
            area   = str(s.get("area", "")).strip()
            advice = str(s.get("advice", "")).strip()
            sev    = str(s.get("severity", "")).strip().title()
            if sev not in ("High", "Medium", "Low"):
                sev = "Medium"
            if area and advice:
                out.append({"area": area, "severity": sev, "advice": advice})
        return jsonify({"suggestions": out})
    except _ant.APITimeoutError:
        _refund_tokens(1)
        return jsonify({"error": "AI request timed out - try again"}), 504
    except _ant.APIConnectionError:
        _refund_tokens(1)
        return jsonify({"error": "Could not reach the AI service"}), 503
    except Exception as e:
        _refund_tokens(1)
        return jsonify({"error": str(e)}), 500


# ── Module 3 Comprehensive Report ────────────────────────────────────────────

@app.route("/generate_module3_report", methods=["POST"])
@login_required
def generate_module3_report():
    """Generate a single comprehensive DOCX/PDF/TXT report for all Module-3 deliverables."""
    import anthropic as _ant

    store    = _get_store()
    body     = request.get_json() or {}
    lp_data  = store.get("lp_data")
    td_data  = store.get("td_data")

    code  = (store.get("lp_code") or store.get("td_code") or "COURSE").strip().upper()
    title = store.get("lp_title") or store.get("td_title") or "Untitled"
    courses = store.get("lp_courses") or store.get("td_courses") or store.get("courses", {})
    info  = courses.get(code, {})
    semester = info.get("semester")

    if not lp_data and not td_data:
        return jsonify({"error": "No Module 3 data in session. Run Module 3 generation first."}), 400

    charged, err = _charge_report(
        "m3",
        json.dumps(lp_data or {}, sort_keys=True, default=str)
        + json.dumps(td_data or {}, sort_keys=True, default=str))
    if err:
        return err

    analytics = _m3r.compute_analytics(lp_data or {}, td_data or {})

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    ai_data = {}
    if api_key and analytics["cos"]:
        try:
            client  = _ant.Anthropic(api_key=api_key, timeout=60.0)
            ai_data = _m3r.get_ai_analysis(client, lp_data or {}, td_data or {},
                                            analytics, code, title)
        except Exception:
            ai_data = {}

    # Brief for the "Suggest ways to improve" panel (see /module_report_suggest)
    store["m3_suggest_ctx"] = {
        "code": code, "title": title,
        "context": _m3_suggest_context(analytics, ai_data),
    }

    fmt = body.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        if charged:
            _refund_tokens(2)
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    report_kwargs = dict(
        lp_data=lp_data or {}, td_data=td_data or {},
        analytics=analytics, ai=ai_data,
        code=code, title=title, semester=semester,
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "docx":
            _m3r.build_docx(**report_kwargs, output_path=tmp.name)
        elif fmt == "pdf":
            _m3r.build_pdf(**report_kwargs, output_path=tmp.name)
        else:
            _m3r.build_txt(**report_kwargs, output_path=tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        if charged:
            _refund_tokens(2)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    filename = f"{code}_Module3_Report{suffix}"
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# -- PO Attainment ------------------------------------------------------------

@app.route("/generate_po_attainment_ai", methods=["POST"])
@tokens_required(1)
def generate_po_attainment_ai():
    """Compute PO attainment via NBA method from CO-PO mapping + CO attainment, with AI ATR."""
    import anthropic as _ant, json as _json

    body            = request.get_json() or {}
    cos             = body.get("cos", [])             # [{name, statement}]
    course          = (body.get("course") or "").strip()
    pomap_rows      = body.get("pomap_rows", [])      # [{co, scores:{PO1:3,...}}]
    co_results      = body.get("co_results", [])      # [{co, pct, level}]
    target_threshold = float(body.get("target_threshold", 60))

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    store = _get_store()
    info  = store.get("courses", {}).get(course.upper(), {})
    title = info.get("title", course)

    po_name_map = {p[0]: p[1] for p in STANDARD_POS}

    client = _ant.Anthropic(api_key=api_key, timeout=45.0)

    # ── Path A: we have pomap + co_results → compute mathematically, AI adds ATR ──
    if pomap_rows and co_results:
        level_map = {r["co"]: int(r.get("level", 0)) for r in co_results}
        po_keys   = list(pomap_rows[0].get("scores", {}).keys()) if pomap_rows else [f"PO{i}" for i in range(1, 13)]

        po_att_raw = {}
        for pk in po_keys:
            total_w      = 0
            weighted_sum = 0.0
            for row in pomap_rows:
                w = int(row.get("scores", {}).get(pk, 0))
                if w > 0:
                    co_level     = level_map.get(row.get("co", ""), 0)
                    weighted_sum += co_level * w
                    total_w      += w
            po_att_raw[pk] = round(weighted_sum / total_w, 3) if total_w else 0.0

        # Express attainment as % (0-3 scale → 0-100 %)
        po_pct_map = {pk: round(v / 3 * 100, 1) for pk, v in po_att_raw.items()}

        # AI: generate ATR for below-threshold POs only
        below = [pk for pk, pct in po_pct_map.items() if pct < target_threshold]
        atr_map = {}
        if below:
            co_lines  = "\n".join(f"  {c['name']}: {c.get('statement','')}" for c in cos[:8])
            po_lines  = "\n".join(f"  {pk}: {po_pct_map[pk]:.1f}% (name: {po_name_map.get(pk,'?')})" for pk in below)
            atr_prompt = (
                f"You are an NBA/NAAC accreditation expert.\n\n"
                f"Course: {title} ({course})\nCourse Outcomes:\n{co_lines}\n\n"
                f"The following Programme Outcomes are below the target threshold of {target_threshold}%:\n{po_lines}\n\n"
                "For each PO listed, write ONE concise Action Taken Report (ATR) sentence "
                "(what the faculty should do to improve attainment of this PO).\n\n"
                "Return ONLY a valid JSON object, no explanation:\n"
                '{"PO3": "Introduce more design projects...", "PO8": "Add ethics case studies..."}'
            )
            try:
                resp = client.messages.create(
                    model="claude-haiku-4-5-20251001", max_tokens=600,
                    messages=[{"role": "user", "content": atr_prompt}],
                )
                raw = resp.content[0].text.strip()
                if raw.startswith("```"):
                    raw = raw.split("```")[1]
                    if raw.startswith("json"):
                        raw = raw[4:]
                atr_map = _json.loads(raw.strip())
            except Exception:
                atr_map = {}

        def _level(pct):
            if pct >= 70: return 3
            if pct >= 60: return 2
            if pct >= 50: return 1
            return 0

        po_attainment = [
            {
                "po":          pk,
                "name":        po_name_map.get(pk, pk),
                "pct":         po_pct_map[pk],
                "level":       _level(po_pct_map[pk]),
                "target_met":  po_pct_map[pk] >= target_threshold,
                "atr":         atr_map.get(pk, "Continue current strategies." if po_pct_map[pk] >= target_threshold else "Review pedagogy and increase practice opportunities."),
            }
            for pk in po_keys
        ]

    # ── Path B: CO-only → AI estimates CO att % + CO-PO weights, then compute ──
    else:
        if not cos:
            return jsonify({"error": "No COs provided"}), 400

        co_lines = "\n".join(f"  {c['name']}: {c.get('statement','')}" for c in cos)
        prompt = (
            f"You are an NBA/NAAC accreditation expert for Indian engineering colleges.\n\n"
            f"Course: {title} ({course})\nCourse Outcomes:\n{co_lines}\n\n"
            "Do THREE things:\n\n"
            "1. Estimate CO attainment levels (0-3) for each CO. "
            "Well-run courses: level 2-3 for most COs.\n\n"
            "2. Rate each CO's contribution to PO1-PO12 (3=Strong, 2=Moderate, 1=Low, 0=None).\n"
            "   PO1:Engineering Knowledge  PO2:Problem Analysis  PO3:Design/Dev\n"
            "   PO4:Investigations  PO5:Modern Tools  PO6:Engineer&Society\n"
            "   PO7:Environment  PO8:Ethics  PO9:Team Work  PO10:Communication\n"
            "   PO11:Project Mgmt  PO12:Life-long Learning\n\n"
            "3. For POs likely below 60% attainment, write a short ATR sentence.\n\n"
            "Return ONLY valid JSON:\n"
            "{\n"
            '  "co_levels": {"CO1":2,"CO2":3,...},\n'
            '  "co_po_weights": {"CO1":{"PO1":3,"PO2":2,...},...},\n'
            '  "atr": {"PO3":"Introduce more design tasks...","PO8":"Add ethics case studies..."}\n'
            "}"
        )
        try:
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001", max_tokens=1500,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = resp.content[0].text.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            ai = _json.loads(raw.strip())
        except _ant.APITimeoutError:
            return jsonify({"error": "AI request timed out — try again"}), 504
        except _ant.APIConnectionError:
            return jsonify({"error": "Could not reach AI service"}), 503
        except Exception as e:
            return jsonify({"error": f"AI analysis failed: {e}"}), 500

        co_levels  = ai.get("co_levels", {})
        co_po_wts  = ai.get("co_po_weights", {})
        atr_map    = ai.get("atr", {})
        po_keys    = [f"PO{i}" for i in range(1, 13)]

        po_att_raw = {}
        for pk in po_keys:
            total_w = 0; weighted_sum = 0.0
            for c in cos:
                w = int(co_po_wts.get(c["name"], {}).get(pk, 0))
                if w > 0:
                    lv = int(co_levels.get(c["name"], 1))
                    weighted_sum += lv * w
                    total_w      += w
            po_att_raw[pk] = round(weighted_sum / total_w, 3) if total_w else 0.0

        po_pct_map = {pk: round(v / 3 * 100, 1) for pk, v in po_att_raw.items()}

        def _level(pct):
            if pct >= 70: return 3
            if pct >= 60: return 2
            if pct >= 50: return 1
            return 0

        po_attainment = [
            {
                "po":         pk,
                "name":       po_name_map.get(pk, pk),
                "pct":        po_pct_map[pk],
                "level":      _level(po_pct_map[pk]),
                "target_met": po_pct_map[pk] >= target_threshold,
                "atr":        atr_map.get(pk, "Continue current strategies." if po_pct_map[pk] >= target_threshold else "Review pedagogy and increase practice opportunities."),
            }
            for pk in po_keys
        ]

    met   = sum(1 for p in po_attainment if p["target_met"])
    mean  = round(sum(p["pct"] for p in po_attainment) / len(po_attainment), 1) if po_attainment else 0
    return jsonify({
        "po_attainment":      po_attainment,
        "summary": {
            "mean_pct":            mean,
            "target_threshold":    target_threshold,
            "pos_meeting_target":  met,
            "pos_below_target":    len(po_attainment) - met,
        },
        "courseName": title,
        "courseCode": course,
    })


@app.route("/export_po_attainment", methods=["POST"])
@login_required
def export_po_attainment():
    """Export PO Attainment table as TXT, DOCX, or PDF."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from docx import Document
    from docx.shared import Pt, RGBColor

    body         = request.get_json() or {}
    fmt          = body.get("fmt", "docx")
    data         = body.get("data", {})
    po_att       = data.get("po_attainment", [])
    summary      = data.get("summary", {})
    course_code  = data.get("courseCode", "COURSE")
    course_name  = data.get("courseName", "")

    if fmt not in ("txt", "docx", "pdf"):
        return jsonify({"error": "Invalid format"}), 400
    if not po_att:
        return jsonify({"error": "No PO attainment data"}), 400

    def _s(v): return str(v) if v is not None else ""

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {"txt": "text/plain",
              "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              "pdf": "application/pdf"}[fmt]

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "txt":
            lines = [f"PO Attainment Report — {course_name} ({course_code})", "=" * 60, ""]
            lines.append(f"{'PO':<6}  {'Name':<36}  {'Att%':>6}  {'Level':>5}  Target")
            lines.append("-" * 64)
            for p in po_att:
                lines.append(f"{p['po']:<6}  {p['name']:<36}  {p['pct']:>5.1f}%  {p['level']:>5}  {'Met' if p['target_met'] else 'Below'}")
            lines.extend(["", f"Mean Attainment : {summary.get('mean_pct', 0):.1f}%",
                          f"POs Meeting Target: {summary.get('pos_meeting_target', 0)} / {len(po_att)}", ""])
            lines.append("Action Taken Report")
            lines.append("-" * 40)
            for p in po_att:
                if not p.get("target_met"):
                    lines.append(f"  {p['po']} ({p['name']}): {p.get('atr', '')}")
            with open(tmp.name, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))

        elif fmt == "docx":
            doc = Document()
            from docx.shared import Inches
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Inches(0.8)
                sec.left_margin = sec.right_margin = Inches(0.9)

            GREEN = RGBColor(5, 150, 105)
            p = doc.add_heading(f"PO Attainment Report", level=0)
            p.runs[0].font.color.rgb = GREEN
            sub = doc.add_paragraph(f"{course_name} ({course_code})")
            sub.runs[0].font.size = Pt(9)

            tbl = doc.add_table(rows=1 + len(po_att), cols=5)
            tbl.style = "Table Grid"
            for i, h in enumerate(["PO", "Name", "Attainment %", "Level", "Target"]):
                tbl.rows[0].cells[i].text = h
                for run in tbl.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            for i, p in enumerate(po_att):
                cells = tbl.rows[i + 1].cells
                cells[0].text = _s(p["po"])
                cells[1].text = _s(p["name"])
                cells[2].text = f"{p['pct']:.1f}%"
                cells[3].text = _s(p["level"])
                cells[4].text = "Met" if p["target_met"] else "Below"

            doc.add_paragraph()
            kv = doc.add_paragraph()
            kv.add_run(f"Mean Attainment: ").bold = True
            kv.add_run(f"{summary.get('mean_pct', 0):.1f}%  |  "
                       f"POs Meeting Target: {summary.get('pos_meeting_target', 0)} / {len(po_att)}")

            atr_rows = [p for p in po_att if not p.get("target_met")]
            if atr_rows:
                doc.add_heading("Action Taken Report", level=2)
                for p in atr_rows:
                    par = doc.add_paragraph(style="List Bullet")
                    par.add_run(f"{p['po']} ({p['name']}): ").bold = True
                    par.add_run(p.get("atr", ""))
            doc.save(tmp.name)

        else:  # pdf
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(5, 150, 105)
            pdf.cell(0, 9, f"PO Attainment Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, f"{course_name} ({course_code})", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)

            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(209, 250, 229)
            for lbl, w in [("PO", 15), ("Name", 75), ("Att%", 22), ("Level", 18), ("Target", 22)]:
                pdf.cell(w, 7, lbl, border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 9)
            for p in po_att:
                pdf.cell(15, 6, _s(p["po"]), border=1)
                pdf.cell(75, 6, _s(p["name"])[:40], border=1)
                pdf.cell(22, 6, f"{p['pct']:.1f}%", border=1, align="C")
                pdf.cell(18, 6, _s(p["level"]), border=1, align="C")
                pdf.cell(22, 6, "Met" if p["target_met"] else "Below", border=1, align="C")
                pdf.ln()
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 6, f"Mean Attainment: {summary.get('mean_pct', 0):.1f}%   "
                           f"POs Meeting Target: {summary.get('pos_meeting_target', 0)} / {len(po_att)}",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            atr_rows = [p for p in po_att if not p.get("target_met")]
            if atr_rows:
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 10)
                pdf.cell(0, 7, "Action Taken Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_font("Helvetica", "", 9)
                for p in atr_rows:
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(30, 5, f"{p['po']}:", border=0)
                    pdf.set_font("Helvetica", "", 9)
                    pdf.multi_cell(0, 5, p.get("atr", ""),
                                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.output(tmp.name)

        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    return Response(content, mimetype=mime,
                    headers={"Content-Disposition": f'attachment; filename="{course_code}_PO_Attainment{suffix}"'})


# -- Module 4 Comprehensive Report --------------------------------------------

@app.route("/generate_module4_report", methods=["POST"])
@login_required
def generate_module4_report():
    """Generate a single comprehensive DOCX/PDF/TXT report for all Module-4 deliverables."""
    import anthropic as _ant

    store  = _get_store()
    body   = request.get_json() or {}

    pomap_rows = body.get("pomap_rows") or store.get("pomap_rows") or []
    coatt_data = body.get("coatt") or {}
    poatt_data = body.get("poatt") or {}
    sdgco_data = body.get("sdgco") or {}
    sdgpo_data = body.get("sdgpo") or {}

    code  = (store.get("pomap_course_code") or body.get("course") or "COURSE").strip().upper()
    title = store.get("pomap_course_title") or body.get("course_title") or "Untitled"
    courses = store.get("courses", {})
    info  = courses.get(code, {})
    semester = info.get("semester")

    if not pomap_rows and not coatt_data and not sdgco_data and not sdgpo_data:
        return jsonify({"error": "No Module 4 data found. Run Module 4 generation first."}), 400

    charged, err = _charge_report(
        "m4",
        json.dumps([pomap_rows, coatt_data, poatt_data, sdgco_data, sdgpo_data],
                   sort_keys=True, default=str))
    if err:
        return err

    analytics = _m4r.compute_analytics(pomap_rows, coatt_data, poatt_data, sdgco_data, sdgpo_data)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    ai_data = {}
    if api_key and (analytics["co_results"] or analytics["pomap_rows"]):
        try:
            client  = _ant.Anthropic(api_key=api_key, timeout=60.0)
            all_data = dict(pomap=pomap_rows, coatt=coatt_data, poatt=poatt_data, sdgco=sdgco_data, sdgpo=sdgpo_data)
            ai_data = _m4r.get_ai_analysis(client, all_data, analytics, code, title)
        except Exception:
            ai_data = {}

    # Brief for the "Suggest ways to improve" panel (see /module_report_suggest)
    store["m4_suggest_ctx"] = {
        "code": code, "title": title,
        "context": _m4_suggest_context(analytics, ai_data),
    }

    fmt = body.get("fmt", "docx")
    if fmt not in ("txt", "docx", "pdf"):
        if charged:
            _refund_tokens(2)
        return jsonify({"error": "Invalid format"}), 400

    suffix = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}[fmt]
    mime   = {
        "txt":  "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf":  "application/pdf",
    }[fmt]

    report_kwargs = dict(
        pomap_data=pomap_rows, coatt_data=coatt_data,
        poatt_data=poatt_data, sdgco_data=sdgco_data, sdgpo_data=sdgpo_data,
        analytics=analytics, ai=ai_data,
        code=code, title=title, semester=semester,
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.close()
    try:
        if fmt == "docx":
            _m4r.build_docx(**report_kwargs, output_path=tmp.name)
        elif fmt == "pdf":
            _m4r.build_pdf(**report_kwargs, output_path=tmp.name)
        else:
            _m4r.build_txt(**report_kwargs, output_path=tmp.name)
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    except Exception as e:
        if charged:
            _refund_tokens(2)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    filename = f"{code}_Module4_Report{suffix}"
    return Response(
        content,
        mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000, threaded=True)

