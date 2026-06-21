import os
import re
import json
import sys
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ── Standard POs (12) ─────────────────────────────────────────────────────────
_STANDARD_POS = [
    "PO1", "PO2", "PO3", "PO4", "PO5", "PO6",
    "PO7", "PO9", "PO10", "PO12",
]

# ── System prompt ─────────────────────────────────────────────────────────────
_SYSTEM = """\
You are an expert in Outcome-Based Education (OBE) and SDG integration for engineering courses.

Generate a complete Course Lesson Plan as a single JSON object. The JSON must be valid,
complete, and follow the exact structure specified. No markdown fences, no explanation —
output ONLY the JSON object.
"""

_PEDAGOGICAL_INNOVATIONS = [
    {"method": "Experiential Learning",     "desc": "Real-world coding exercises"},
    {"method": "AI-Assisted Learning",      "desc": "ChatGPT-assisted debugging"},
    {"method": "Flipped Classroom",         "desc": "Pre-class videos and readings"},
    {"method": "SDG Challenge Activities",  "desc": "Sustainable software tasks"},
    {"method": "Collaborative Learning",    "desc": "GitHub team coding"},
    {"method": "ICT Enabled Learning",      "desc": "Interactive visualization tools"},
]

def _build_assessment_framework(num_cos: int) -> list:
    last = f"CO{num_cos}"
    co_range = f"CO1–{last}"
    mid = f"CO{min(4, num_cos - 1)}" if num_cos >= 3 else "CO2"
    coding_cos = f"CO2, {mid}" if mid != "CO2" else "CO2"
    return [
        {"method": "Internal Test",     "cos": co_range,    "bloom": "Understand–Analyze", "marks": 30,  "evidence": "Answer scripts"},
        {"method": "Coding Assignment", "cos": coding_cos,  "bloom": "Apply",              "marks": 20,  "evidence": "GitHub"},
        {"method": "SDG Mini Project",  "cos": last,        "bloom": "Create",             "marks": 20,  "evidence": "Project report"},
        {"method": "Quiz",              "cos": "CO1, CO2",  "bloom": "Remember",           "marks": 10,  "evidence": "Quiz sheets"},
        {"method": "SEE",               "cos": co_range,    "bloom": "All Levels",         "marks": 100, "evidence": "Answer booklets"},
    ]

_IMPACT_MEASUREMENT = [
    {"indicator": "CO Attainment",      "method": "Direct assessment",   "evidence": "Marks analysis"},
    {"indicator": "SDG Awareness",      "method": "Student survey",       "evidence": "Feedback"},
    {"indicator": "AI Tool Usage",      "method": "Lab activity",         "evidence": "Screenshots"},
    {"indicator": "Innovation Skills",  "method": "Mini projects",        "evidence": "Rubrics"},
    {"indicator": "Industry Readiness", "method": "Internship outcomes",  "evidence": "Placement data"},
]

_OBE_ALIGNMENT = [
    {"framework": "OBE",                 "details": "CO-PO attainment, CEP"},
    {"framework": "SDGs",                "details": "Sustainability integration"},
]


def _build_prompt(course_code: str, course_title: str, course_text: str,
                  num_cos: int, meta: dict,
                  existing_cos: list | None = None) -> str:
    semester = meta.get("semester", "III")
    academic_year = meta.get("academic_year", "2026-2027")
    regulation = meta.get("regulation", "R26")
    program = meta.get("program", "B.Tech – CSE")
    ltp = meta.get("ltp", "3-0-2")
    credits = meta.get("credits", 4)
    course_type = meta.get("course_type", "Theory + Lab")

    if existing_cos:
        co_block = "\n".join(f"CO{c['num']}: {c['statement']}" for c in existing_cos)
        co_section = f"""PRE-DEFINED COURSE OUTCOMES (use EXACTLY as given — do NOT change wording or numbering):
{co_block}"""
        co_requirement = (
            f"- Use the {len(existing_cos)} pre-defined COs above EXACTLY as provided in the \"cos\" array "
            f"(preserve num, statement, and infer bloom/knowledge from the statement)"
        )
    else:
        co_section = ""
        co_requirement = f"- Generate exactly {num_cos} Course Outcomes (CO1 to CO{num_cos})"

    return f"""Generate a complete Course Lesson Plan for the following course.

COURSE METADATA:
- Course Code: {course_code}
- Course Title: {course_title}
- Academic Year: {academic_year}
- Semester: {semester}
- Regulation: {regulation}
- Program: {program}
- Course Type: {course_type}
- Credits: {credits}
- L-T-P: {ltp}

SYLLABUS CONTENT:
{course_text[:5000]}

{co_section}

REQUIREMENTS:
{co_requirement}
- Detect units from the syllabus; generate one entry per unit found (typically 5)
- Each session in the session_plan should map to one topic from one unit
- Generate exactly {num_cos} sessions total (one per CO), ensuring every CO from CO1 to CO{num_cos} appears exactly once in the session_plan
- Generate 3 mini projects
- CO-PO scores use 0/1/2/3 scale (3=Major, 2=Moderate, 1=Minor, 0=None)
- SDG goals: SDG4 (Quality Education), SDG7 (Clean Energy), SDG9 (Industry Innovation), SDG13 (Climate Action)
- Bloom's levels: Remember, Understand, Apply, Analyze, Evaluate, Create
- Knowledge levels: Conceptual, Procedural, Analytical, Strategic

OUTPUT — return ONLY this JSON (no markdown, no explanation):
{{
  "course_overview": "<2-3 sentence description of the course integrating AI, SDGs, OBE>",
  "cos": [
    {{
      "num": 1,
      "statement": "<Measurable CO using Bloom's verb>",
      "bloom": "<Bloom's level>",
      "knowledge": "<Knowledge level>"
    }}
  ],
  "co_po_sdg_mapping": [
    {{
      "co": "CO1",
      "po_scores": {{"PO1": 0, "PO2": 0, "PO3": 0, "PO4": 0, "PO5": 0, "PO6": 0, "PO7": 0, "PO9": 0, "PO10": 0, "PO12": 0}},
      "pso1": 0,
      "sdg": "SDG4"
    }}
  ],
  "graduate_attributes": [
    {{"unit": 1, "problem": "<Real-world problem solved>", "characteristic": "<GA characteristic>"}}
  ],
  "sdg_framework": [
    {{"unit": 1, "sdg_goal": "SDG4", "indicator": "<SDG indicator>", "sustainability": "<focus>", "impact": "<expected impact>"}}
  ],
  "industry_ai": [
    {{"topic": "<topic name>", "industry": "<industry application>", "ai_integration": "<AI/ML use>"}}
  ],
  "units": [
    {{
      "num": 1,
      "title": "<Unit Title>",
      "topics": ["<topic1>", "<topic2>", "<topic3>", "<topic4>"],
      "co": "CO1",
      "pos": "PO1, PO2",
      "ga": "<graduate attribute>",
      "sdg": "SDG4",
      "ai_tool": "<AI tool used>",
      "sustainability": "<sustainability context>",
      "industry": "<industry application>"
    }}
  ],
  "session_plan": [
    {{
      "session": 1,
      "topic": "<topic>",
      "co": "CO1",
      "po": "PO1",
      "ga": "<GA>",
      "sdg": "SDG4",
      "method": "<teaching method>",
      "ai_tool": "<AI tool>",
      "assessment": "<assessment type>",
      "evidence": "<evidence collected>"
    }}
  ],
  "mini_projects": [
    {{
      "title": "<project title>",
      "sdg": "SDG9",
      "ga": "<GA>",
      "innovation": "<innovation aspect>",
      "impact": "<expected societal impact>"
    }}
  ]
}}"""


def _extract_json(raw: str) -> dict | None:
    cleaned = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    start = cleaned.find('{')
    end = cleaned.rfind('}')
    if start != -1 and end > start:
        try:
            return json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError:
            pass
    return None


def generate_lesson_plan(client, course_code: str, course_title: str,
                         course_text: str, meta: dict,
                         num_cos: int = 5,
                         existing_cos: list | None = None) -> dict:
    """Call Claude and return parsed lesson plan dict. Raises on failure."""
    if existing_cos:
        num_cos = len(existing_cos)
    prompt = _build_prompt(course_code, course_title, course_text, num_cos, meta,
                           existing_cos=existing_cos)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        system=_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text
    data = _extract_json(raw)
    if not data:
        raise ValueError(
            f"AI returned invalid JSON. Raw response (first 300 chars): {raw[:300]}"
        )
    return data


def generate_lesson_plan_stream(client, course_code: str, course_title: str,
                                course_text: str, meta: dict,
                                num_cos: int = 5,
                                existing_cos: list | None = None):
    """Yield SSE-style progress dicts. Last item has key 'done' and 'data'.

    Streams the model response so the SSE connection keeps receiving bytes.
    A single blocking call sends nothing for tens of seconds while the large
    JSON is generated, which a hosting proxy treats as an idle connection and
    drops — surfacing client-side as a generic "Connection error". The
    periodic 'heartbeat' items keep the connection alive (the route emits them
    as SSE comments that the browser ignores).
    """
    if existing_cos:
        num_cos = len(existing_cos)
        yield {"progress": f"Using {len(existing_cos)} uploaded COs — generating lesson plan sections…"}
    else:
        yield {"progress": "Generating COs and lesson plan sections via AI…"}

    prompt = _build_prompt(course_code, course_title, course_text, num_cos, meta,
                           existing_cos=existing_cos)
    raw = ""
    try:
        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        ) as stream:
            for chunk in stream.text_stream:
                raw += chunk
                yield {"heartbeat": True}
    except Exception as e:
        yield {"error": str(e)}
        return

    data = _extract_json(raw)
    if not data:
        yield {"error": f"AI returned invalid JSON. Raw response (first 300 chars): {raw[:300]}"}
        return
    yield {"done": True, "data": data}


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(12 if level == 1 else 11)
        run.bold = True
    return p


def _cell_bg(cell, rgb):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r, g, b = rgb
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, size_pt=9, center=False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = str(text)
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.bold = bold


def _make_table(doc, headers, rows_data, header_bg=(180, 198, 231)):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ncols = len(headers)
    nrows = 1 + len(rows_data)
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'

    # Header row
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _cell_bg(hdr[i], header_bg)
        _set_cell_text(hdr[i], h, bold=True, size_pt=9)

    # Data rows
    for ri, row in enumerate(rows_data, 1):
        cells = tbl.rows[ri].cells
        for ci, val in enumerate(row):
            _set_cell_text(cells[ci], val, size_pt=9)

    return tbl


def _add_info_table(doc, meta: dict, course_code: str, course_title: str):
    """Build the header metadata table."""
    rows = [
        ("Academic Year",  meta.get("academic_year", "2026–2027")),
        ("Semester",       meta.get("semester", "III Semester")),
        ("Course Code",    course_code),
        ("Course Title",   course_title),
        ("Course Type",    meta.get("course_type", "Theory + Lab")),
        ("Credits",        str(meta.get("credits", 4))),
        ("L-T-P",          meta.get("ltp", "3-0-2")),
        ("Faculty Name",   meta.get("faculty_name", "Dr. Faculty")),
        ("Regulation",     meta.get("regulation", "R26")),
        ("Program",        meta.get("program", "B.Tech – CSE")),
        ("Prepared By",    "Course Faculty"),
        ("Reviewed By",    "Course Coordinator"),
        ("Approved By",    "HoD"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    from docx.shared import Pt
    for ri, (label, val) in enumerate(rows):
        cells = tbl.rows[ri].cells
        _set_cell_text(cells[0], label, bold=True, size_pt=9)
        _set_cell_text(cells[1], val, size_pt=9)


def _build_co_po_table(doc, mapping_rows: list, cos: list):
    """Build CO–PO–PSO–SDG mapping table."""
    from docx.shared import Pt
    pos = _STANDARD_POS
    headers = ["CO"] + pos + ["PSO1", "SDG"]
    ncols = len(headers)

    co_map = {f"CO{c['num']}": c['statement'] for c in cos}
    nrows = 1 + len(mapping_rows)
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'

    _SCORE_COLORS = {
        3: (198, 224, 180),
        2: (255, 242, 204),
        1: (221, 235, 247),
        0: (255, 255, 255),
    }

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _cell_bg(hdr[i], (180, 198, 231))
        _set_cell_text(hdr[i], h, bold=True, size_pt=8, center=True)

    for ri, row in enumerate(mapping_rows, 1):
        cells = tbl.rows[ri].cells
        _set_cell_text(cells[0], row.get("co", ""), bold=True, size_pt=8)
        po_scores = row.get("po_scores", {})
        for ci, k in enumerate(pos, 1):
            v = po_scores.get(k, 0)
            try:
                v = int(v)
            except (ValueError, TypeError):
                v = 0
            v = max(0, min(3, v))
            display = str(v) if v > 0 else "–"
            _set_cell_text(cells[ci], display, size_pt=8, center=True)
            _cell_bg(cells[ci], _SCORE_COLORS[v])
        # PSO1
        pso = row.get("pso1", 0)
        try:
            pso = int(pso)
        except (ValueError, TypeError):
            pso = 0
        pso = max(0, min(3, pso))
        pso_idx = len(pos) + 1
        _set_cell_text(cells[pso_idx], str(pso) if pso > 0 else "–", size_pt=8, center=True)
        _cell_bg(cells[pso_idx], _SCORE_COLORS[pso])
        # SDG
        sdg_idx = pso_idx + 1
        _set_cell_text(cells[sdg_idx], row.get("sdg", ""), size_pt=8, center=True)

    return tbl


def build_docx(data: dict, meta: dict, course_code: str, course_title: str,
               output_path: str):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1.0)
        sec.top_margin = sec.bottom_margin = Inches(1.0)

    # ── Title ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Course Lesson Plan")
    run.bold = True
    run.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"({course_title})")
    run2.bold = True
    run2.font.size = Pt(12)

    dept = meta.get("department", "Department of Computer Science and Engineering")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(dept)
    run3.bold = True
    run3.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("COURSE DESIGN, DELIVERY, ASSESSMENT & SDG IMPACT FRAMEWORK")
    run4.bold = True
    run4.font.size = Pt(11)

    doc.add_paragraph()
    _add_info_table(doc, meta, course_code, course_title)
    doc.add_paragraph()

    # ── 1. Course Overview ─────────────────────────────────────────────────────
    _heading(doc, "1. COURSE OVERVIEW")
    overview = data.get("course_overview", "")
    doc.add_paragraph(overview)
    doc.add_paragraph()

    # ── 2. Course Outcomes ─────────────────────────────────────────────────────
    _heading(doc, "2. COURSE OUTCOMES (COs)")
    cos = data.get("cos", [])
    _make_table(doc,
        ["CO No.", "Course Outcome", "Bloom's Level", "Knowledge Level"],
        [(f"CO{c['num']}", c['statement'], c.get('bloom',''), c.get('knowledge','')) for c in cos]
    )
    doc.add_paragraph()

    # ── 3. CO–PO–PSO–SDG Mapping ──────────────────────────────────────────────
    _heading(doc, "3. CO–PO–PSO–SDG MAPPING")
    mapping = data.get("co_po_sdg_mapping", [])
    if mapping:
        _build_co_po_table(doc, mapping, cos)
    doc.add_paragraph()

    # ── 4. Graduate Attributes ─────────────────────────────────────────────────
    _heading(doc, "4. Graduate Attributes")
    gas = data.get("graduate_attributes", [])
    _make_table(doc,
        ["Unit", "Attribute Related Problems", "Graduate Attributes Characteristics"],
        [(f"Unit {g['unit']}", g.get('problem',''), g.get('characteristic','')) for g in gas]
    )
    doc.add_paragraph()

    # ── 5. SDG Integration Framework ──────────────────────────────────────────
    _heading(doc, "5. SDG INTEGRATION FRAMEWORK")
    sdg_rows = data.get("sdg_framework", [])
    _make_table(doc,
        ["Unit", "SDG Goal", "SDG Indicator", "Sustainability Focus", "Expected Impact"],
        [(f"Unit {r['unit']}", r.get('sdg_goal',''), r.get('indicator',''),
          r.get('sustainability',''), r.get('impact','')) for r in sdg_rows]
    )
    doc.add_paragraph()

    # ── 6. Industry 4.0 / AI Integration ──────────────────────────────────────
    _heading(doc, "6. INDUSTRY 4.0 / AI INTEGRATION")
    ind_rows = data.get("industry_ai", [])
    _make_table(doc,
        ["Topic", "Industry Relevance", "AI/ML/GenAI Integration"],
        [(r.get('topic',''), r.get('industry',''), r.get('ai_integration','')) for r in ind_rows]
    )
    doc.add_paragraph()

    # ── 7. Pedagogical Innovations ─────────────────────────────────────────────
    _heading(doc, "7. PEDAGOGICAL INNOVATIONS")
    _make_table(doc,
        ["Teaching Method", "Description"],
        [(r["method"], r["desc"]) for r in _PEDAGOGICAL_INNOVATIONS]
    )
    doc.add_paragraph()

    # ── 8. Course Contents ─────────────────────────────────────────────────────
    _heading(doc, "8. COURSE CONTENTS")
    units = data.get("units", [])
    for unit in units:
        _heading(doc, f"UNIT–{unit['num']}: {unit.get('title','')}", level=2)

        # Topics list
        p_topics = doc.add_paragraph()
        p_topics.add_run("Topics").bold = True
        for topic in unit.get("topics", []):
            doc.add_paragraph(topic, style="List Bullet")

        # Integrated Elements table
        p_ie = doc.add_paragraph()
        p_ie.add_run("Integrated Elements").bold = True

        ie_rows = [
            ("CO Addressed",        unit.get("co", "")),
            ("PO Addressed",        unit.get("pos", "")),
            ("GA Addressed",        unit.get("ga", "")),
            ("SDG Addressed",       unit.get("sdg", "")),
            ("AI/Tool Integration", unit.get("ai_tool", "")),
            ("Sustainability Context", unit.get("sustainability", "")),
            ("Industry Application", unit.get("industry", "")),
        ]
        _make_table(doc, ["Component", "Details"], ie_rows)
        doc.add_paragraph()

    # ── 9. Innovative Lesson Delivery Plan ────────────────────────────────────
    _heading(doc, "9. INNOVATIVE LESSON DELIVERY PLAN")
    sessions = data.get("session_plan", [])
    _make_table(doc,
        ["Session", "Topic", "CO", "PO", "GA", "SDG", "Teaching Method", "AI Tool", "Assessment", "Evidence"],
        [(str(s.get('session','')), s.get('topic',''), s.get('co',''),
          s.get('po',''), s.get('ga',''), s.get('sdg',''),
          s.get('method',''), s.get('ai_tool',''),
          s.get('assessment',''), s.get('evidence','')) for s in sessions]
    )
    doc.add_paragraph()

    # ── 10. Mini Projects / SDG Challenge Activities ───────────────────────────
    _heading(doc, "10. MINI PROJECT / SDG CHALLENGE ACTIVITIES")
    projects = data.get("mini_projects", [])
    _make_table(doc,
        ["Project Title", "SDG", "GA", "Innovation Aspect", "Expected Societal Impact"],
        [(p.get('title',''), p.get('sdg',''), p.get('ga',''),
          p.get('innovation',''), p.get('impact','')) for p in projects]
    )
    doc.add_paragraph()

    # ── 11. Course Assessment Framework ───────────────────────────────────────
    _heading(doc, "11. COURSE ASSESSMENT FRAMEWORK")
    _assessment = _build_assessment_framework(len(data.get("cos", [])) or 5)
    _make_table(doc,
        ["Assessment Method", "COs Assessed", "Bloom's Level", "Marks", "Evidence"],
        [(r["method"], r["cos"], r["bloom"], str(r["marks"]), r["evidence"])
         for r in _assessment]
    )
    doc.add_paragraph()

    # ── 12. Course Impact Measurement ─────────────────────────────────────────
    _heading(doc, "12. COURSE IMPACT MEASUREMENT")
    _make_table(doc,
        ["Indicator", "Measurement Method", "Evidence"],
        [(r["indicator"], r["method"], r["evidence"]) for r in _IMPACT_MEASUREMENT]
    )
    doc.add_paragraph()

    # ── 13. OBE & SDG Alignment ───────────────────────────────────────────────
    _heading(doc, "13. OBE & SDG ALIGNMENT")
    _make_table(doc,
        ["Framework", "Alignment Details"],
        [(r["framework"], r["details"]) for r in _OBE_ALIGNMENT]
    )

    doc.save(output_path)


def build_pdf(data: dict, meta: dict, course_code: str, course_title: str,
              output_path: str):
    """Build a PDF version of the lesson plan using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'],
                                  fontSize=14, spaceAfter=4, alignment=TA_CENTER)
    sub_style = ParagraphStyle('SubStyle', parent=styles['Normal'],
                                fontSize=11, spaceAfter=4, alignment=TA_CENTER, fontName='Helvetica-Bold')
    heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'],
                                    fontSize=11, spaceAfter=6, spaceBefore=10,
                                    fontName='Helvetica-Bold')
    sub_heading_style = ParagraphStyle('SubHeadingStyle', parent=styles['Heading3'],
                                        fontSize=10, spaceAfter=4, spaceBefore=6,
                                        fontName='Helvetica-Bold')
    body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'],
                                 fontSize=9, spaceAfter=4)
    cell_style = ParagraphStyle('CellStyle', parent=styles['Normal'],
                                 fontSize=8, leading=10)

    HDR_COLOR = colors.Color(180/255, 198/255, 231/255)
    ALT_COLOR = colors.Color(242/255, 242/255, 242/255)

    def _tbl_style(nrows, header_only=False):
        return TableStyle([
            ('BACKGROUND',  (0, 0), (-1, 0), HDR_COLOR),
            ('TEXTCOLOR',   (0, 0), (-1, 0), colors.black),
            ('FONTNAME',    (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',    (0, 0), (-1, -1), 8),
            ('GRID',        (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN',      (0, 0), (-1, -1), 'TOP'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, ALT_COLOR]),
            ('LEFTPADDING',  (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING',   (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 3),
        ])

    def _wrap(text):
        return Paragraph(str(text), cell_style)

    def _info_table(meta, code, title):
        rows = [
            ("Academic Year",   meta.get("academic_year", "2026–2027")),
            ("Semester",        meta.get("semester", "III Semester")),
            ("Course Code",     code),
            ("Course Title",    title),
            ("Course Type",     meta.get("course_type", "Theory + Lab")),
            ("Credits",         str(meta.get("credits", 4))),
            ("L-T-P",           meta.get("ltp", "3-0-2")),
            ("Faculty Name",    meta.get("faculty_name", "Dr. Faculty")),
            ("Regulation",      meta.get("regulation", "R26")),
            ("Program",         meta.get("program", "B.Tech – CSE")),
            ("Prepared By",     "Course Faculty"),
            ("Reviewed By",     "Course Coordinator"),
            ("Approved By",     "HoD"),
        ]
        data_rows = [[_wrap(k), _wrap(v)] for k, v in rows]
        t = Table(data_rows, colWidths=[2.5 * inch, 4.0 * inch])
        t.setStyle(TableStyle([
            ('GRID',    (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME',(0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE',(0, 0), (-1, -1), 9),
            ('VALIGN',  (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING',  (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING',   (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 3),
        ]))
        return t

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    W = A4[0] - 2 * inch  # usable width
    story = []

    dept = meta.get("department", "Department of Computer Science and Engineering")
    story += [
        Paragraph("Course Lesson Plan", title_style),
        Paragraph(f"({course_title})", sub_style),
        Paragraph(dept, sub_style),
        Paragraph("COURSE DESIGN, DELIVERY, ASSESSMENT &amp; SDG IMPACT FRAMEWORK", sub_style),
        Spacer(1, 10),
        _info_table(meta, course_code, course_title),
        Spacer(1, 10),
    ]

    # 1. Overview
    story.append(Paragraph("1. COURSE OVERVIEW", heading_style))
    story.append(Paragraph(data.get("course_overview", ""), body_style))
    story.append(Spacer(1, 8))

    # 2. COs
    story.append(Paragraph("2. COURSE OUTCOMES (COs)", heading_style))
    cos = data.get("cos", [])
    co_rows = [["CO No.", "Course Outcome", "Bloom's Level", "Knowledge Level"]]
    for c in cos:
        co_rows.append([_wrap(f"CO{c['num']}"), _wrap(c['statement']),
                         _wrap(c.get('bloom','')), _wrap(c.get('knowledge',''))])
    t = Table(co_rows, colWidths=[0.6*inch, 3.5*inch, 1.2*inch, 1.0*inch])
    t.setStyle(_tbl_style(len(co_rows)))
    story += [t, Spacer(1, 8)]

    # 3. CO-PO-SDG Mapping
    story.append(Paragraph("3. CO–PO–PSO–SDG MAPPING", heading_style))
    mapping = data.get("co_po_sdg_mapping", [])
    pos = _STANDARD_POS
    map_headers = ["CO"] + pos + ["PSO1", "SDG"]
    map_rows = [map_headers]
    _SCORE_COLORS_RL = {
        3: colors.Color(198/255, 224/255, 180/255),
        2: colors.Color(255/255, 242/255, 204/255),
        1: colors.Color(221/255, 235/255, 247/255),
        0: colors.white,
    }
    score_style_cmds = []
    for ri, row in enumerate(mapping, 1):
        po_scores = row.get("po_scores", {})
        pso = row.get("pso1", 0)
        try:
            pso = int(pso)
        except (ValueError, TypeError):
            pso = 0
        cells = [row.get("co","")]
        for k in pos:
            v = po_scores.get(k, 0)
            try:
                v = int(v)
            except (ValueError, TypeError):
                v = 0
            v = max(0, min(3, v))
            cells.append(str(v) if v > 0 else "–")
            col_idx = pos.index(k) + 1
            score_style_cmds.append(('BACKGROUND', (col_idx, ri), (col_idx, ri), _SCORE_COLORS_RL[v]))
        cells.append(str(pso) if pso > 0 else "–")
        cells.append(row.get("sdg",""))
        map_rows.append(cells)

    ncols_map = len(map_headers)
    col_w = W / ncols_map
    t_map = Table(map_rows, colWidths=[col_w] * ncols_map)
    base_style = _tbl_style(len(map_rows))
    for cmd in score_style_cmds:
        base_style.add(*cmd)
    t_map.setStyle(base_style)
    story += [t_map, Spacer(1, 8)]

    # 4. Graduate Attributes
    story.append(Paragraph("4. Graduate Attributes", heading_style))
    gas = data.get("graduate_attributes", [])
    ga_rows = [["Unit", "Attribute Related Problems", "Graduate Attributes Characteristics"]]
    for g in gas:
        ga_rows.append([_wrap(f"Unit {g['unit']}"), _wrap(g.get('problem','')),
                         _wrap(g.get('characteristic',''))])
    t = Table(ga_rows, colWidths=[0.6*inch, 3.0*inch, 2.7*inch])
    t.setStyle(_tbl_style(len(ga_rows)))
    story += [t, Spacer(1, 8)]

    # 5. SDG Framework
    story.append(Paragraph("5. SDG INTEGRATION FRAMEWORK", heading_style))
    sdg_rows = data.get("sdg_framework", [])
    sdg_table = [["Unit", "SDG Goal", "SDG Indicator", "Sustainability Focus", "Expected Impact"]]
    for r in sdg_rows:
        sdg_table.append([_wrap(f"Unit {r['unit']}"), _wrap(r.get('sdg_goal','')),
                           _wrap(r.get('indicator','')), _wrap(r.get('sustainability','')),
                           _wrap(r.get('impact',''))])
    cw5 = W / 5
    t = Table(sdg_table, colWidths=[cw5]*5)
    t.setStyle(_tbl_style(len(sdg_table)))
    story += [t, Spacer(1, 8)]

    # 6. Industry AI
    story.append(Paragraph("6. INDUSTRY 4.0 / AI INTEGRATION", heading_style))
    ind_rows = data.get("industry_ai", [])
    ind_table = [["Topic", "Industry Relevance", "AI/ML/GenAI Integration"]]
    for r in ind_rows:
        ind_table.append([_wrap(r.get('topic','')), _wrap(r.get('industry','')),
                           _wrap(r.get('ai_integration',''))])
    cw3 = W / 3
    t = Table(ind_table, colWidths=[cw3]*3)
    t.setStyle(_tbl_style(len(ind_table)))
    story += [t, Spacer(1, 8)]

    # 7. Pedagogical Innovations
    story.append(Paragraph("7. PEDAGOGICAL INNOVATIONS", heading_style))
    ped_table = [["Teaching Method", "Description"]]
    for r in _PEDAGOGICAL_INNOVATIONS:
        ped_table.append([_wrap(r["method"]), _wrap(r["desc"])])
    t = Table(ped_table, colWidths=[2.5*inch, W - 2.5*inch])
    t.setStyle(_tbl_style(len(ped_table)))
    story += [t, Spacer(1, 8)]

    # 8. Course Contents
    story.append(Paragraph("8. COURSE CONTENTS", heading_style))
    units = data.get("units", [])
    for unit in units:
        story.append(Paragraph(f"UNIT–{unit['num']}: {unit.get('title','')}", sub_heading_style))
        story.append(Paragraph("<b>Topics</b>", body_style))
        for topic in unit.get("topics", []):
            story.append(Paragraph(f"• {topic}", body_style))
        story.append(Paragraph("<b>Integrated Elements</b>", body_style))
        ie_rows = [
            ["Component", "Details"],
            [_wrap("CO Addressed"),         _wrap(unit.get("co",""))],
            [_wrap("PO Addressed"),         _wrap(unit.get("pos",""))],
            [_wrap("GA Addressed"),         _wrap(unit.get("ga",""))],
            [_wrap("SDG Addressed"),        _wrap(unit.get("sdg",""))],
            [_wrap("AI/Tool Integration"),  _wrap(unit.get("ai_tool",""))],
            [_wrap("Sustainability Context"),_wrap(unit.get("sustainability",""))],
            [_wrap("Industry Application"), _wrap(unit.get("industry",""))],
        ]
        t = Table(ie_rows, colWidths=[2.0*inch, W - 2.0*inch])
        t.setStyle(_tbl_style(len(ie_rows)))
        story += [t, Spacer(1, 6)]
    story.append(Spacer(1, 4))

    # 9. Session Plan
    story.append(Paragraph("9. INNOVATIVE LESSON DELIVERY PLAN", heading_style))
    sessions = data.get("session_plan", [])
    sp_headers = ["Sess.", "Topic", "CO", "PO", "GA", "SDG", "Method", "AI Tool", "Assess.", "Evidence"]
    sp_rows = [sp_headers]
    for s in sessions:
        sp_rows.append([
            _wrap(str(s.get('session',''))),
            _wrap(s.get('topic','')),
            _wrap(s.get('co','')),
            _wrap(s.get('po','')),
            _wrap(s.get('ga','')),
            _wrap(s.get('sdg','')),
            _wrap(s.get('method','')),
            _wrap(s.get('ai_tool','')),
            _wrap(s.get('assessment','')),
            _wrap(s.get('evidence','')),
        ])
    col_widths = [0.4*inch, 1.0*inch, 0.45*inch, 0.45*inch,
                  0.7*inch, 0.55*inch, 0.9*inch, 0.65*inch, 0.7*inch, 0.65*inch]
    t = Table(sp_rows, colWidths=col_widths)
    t.setStyle(_tbl_style(len(sp_rows)))
    story += [t, Spacer(1, 8)]

    # 10. Mini Projects
    story.append(Paragraph("10. MINI PROJECT / SDG CHALLENGE ACTIVITIES", heading_style))
    projects = data.get("mini_projects", [])
    mp_headers = ["Project Title", "SDG", "GA", "Innovation Aspect", "Expected Societal Impact"]
    mp_rows = [mp_headers]
    for p in projects:
        mp_rows.append([_wrap(p.get('title','')), _wrap(p.get('sdg','')),
                         _wrap(p.get('ga','')), _wrap(p.get('innovation','')),
                         _wrap(p.get('impact',''))])
    cw_mp = W / 5
    t = Table(mp_rows, colWidths=[cw_mp]*5)
    t.setStyle(_tbl_style(len(mp_rows)))
    story += [t, Spacer(1, 8)]

    # 11. Assessment Framework
    story.append(Paragraph("11. COURSE ASSESSMENT FRAMEWORK", heading_style))
    af_headers = ["Assessment Method", "COs Assessed", "Bloom's Level", "Marks", "Evidence"]
    af_rows = [af_headers]
    _assessment = _build_assessment_framework(len(data.get("cos", [])) or 5)
    for r in _assessment:
        af_rows.append([_wrap(r["method"]), _wrap(r["cos"]), _wrap(r["bloom"]),
                         _wrap(str(r["marks"])), _wrap(r["evidence"])])
    cw_af = W / 5
    t = Table(af_rows, colWidths=[cw_af]*5)
    t.setStyle(_tbl_style(len(af_rows)))
    story += [t, Spacer(1, 8)]

    # 12. Impact Measurement
    story.append(Paragraph("12. COURSE IMPACT MEASUREMENT", heading_style))
    im_headers = ["Indicator", "Measurement Method", "Evidence"]
    im_rows = [im_headers]
    for r in _IMPACT_MEASUREMENT:
        im_rows.append([_wrap(r["indicator"]), _wrap(r["method"]), _wrap(r["evidence"])])
    cw_im = W / 3
    t = Table(im_rows, colWidths=[cw_im]*3)
    t.setStyle(_tbl_style(len(im_rows)))
    story += [t, Spacer(1, 8)]

    # 13. OBE & SDG Alignment
    story.append(Paragraph("13. OBE & SDG ALIGNMENT", heading_style))
    nba_headers = ["Framework", "Alignment Details"]
    nba_rows = [nba_headers]
    for r in _OBE_ALIGNMENT:
        nba_rows.append([_wrap(r["framework"]), _wrap(r["details"])])
    t = Table(nba_rows, colWidths=[2.0*inch, W - 2.0*inch])
    t.setStyle(_tbl_style(len(nba_rows)))
    story.append(t)

    doc.build(story)


def build_txt(data: dict, meta: dict, course_code: str, course_title: str,
              output_path: str):
    """Build a plain-text version of the lesson plan."""
    lines = []

    def sec(title):
        lines.append("")
        lines.append("=" * 70)
        lines.append(title)
        lines.append("=" * 70)

    def table_txt(headers, rows, widths=None):
        if widths is None:
            widths = [max(len(str(r[i])) for r in [headers] + rows) + 2
                      for i in range(len(headers))]
        sep = "+" + "+".join("-" * w for w in widths) + "+"
        def fmt_row(row):
            return "|" + "|".join(
                (" " + str(r)).ljust(w) for r, w in zip(row, widths)
            ) + "|"
        lines.append(sep)
        lines.append(fmt_row(headers))
        lines.append(sep)
        for row in rows:
            lines.append(fmt_row(row))
        lines.append(sep)

    lines.append("COURSE LESSON PLAN")
    lines.append(f"Course: {course_code} — {course_title}")
    lines.append(f"Academic Year: {meta.get('academic_year','')}")
    lines.append(f"Semester: {meta.get('semester','')}")
    lines.append(f"Faculty: {meta.get('faculty_name','')}")

    sec("1. COURSE OVERVIEW")
    lines.append(data.get("course_overview", ""))

    sec("2. COURSE OUTCOMES (COs)")
    cos = data.get("cos", [])
    table_txt(
        ["CO No.", "Course Outcome", "Bloom's Level", "Knowledge Level"],
        [(f"CO{c['num']}", c['statement'][:60], c.get('bloom',''), c.get('knowledge','')) for c in cos]
    )

    sec("3. CO–PO–PSO–SDG MAPPING")
    mapping = data.get("co_po_sdg_mapping", [])
    pos = _STANDARD_POS
    table_txt(
        ["CO"] + pos + ["PSO1", "SDG"],
        [([row.get("co","")]
          + [str(row.get("po_scores", {}).get(k, 0)) for k in pos]
          + [str(row.get("pso1", 0)), row.get("sdg","")]) for row in mapping]
    )

    sec("4. GRADUATE ATTRIBUTES")
    gas = data.get("graduate_attributes", [])
    table_txt(
        ["Unit", "Attribute Related Problems", "GA Characteristics"],
        [(f"Unit {g['unit']}", g.get('problem','')[:40], g.get('characteristic','')) for g in gas]
    )

    sec("5. SDG INTEGRATION FRAMEWORK")
    sdg_rows = data.get("sdg_framework", [])
    table_txt(
        ["Unit", "SDG Goal", "Indicator", "Sustainability Focus", "Impact"],
        [(f"Unit {r['unit']}", r.get('sdg_goal',''), r.get('indicator','')[:30],
          r.get('sustainability',''), r.get('impact','')[:30]) for r in sdg_rows]
    )

    sec("6. INDUSTRY 4.0 / AI INTEGRATION")
    ind_rows = data.get("industry_ai", [])
    table_txt(
        ["Topic", "Industry Relevance", "AI/ML/GenAI Integration"],
        [(r.get('topic',''), r.get('industry',''), r.get('ai_integration','')) for r in ind_rows]
    )

    sec("7. PEDAGOGICAL INNOVATIONS")
    table_txt(
        ["Teaching Method", "Description"],
        [(r["method"], r["desc"]) for r in _PEDAGOGICAL_INNOVATIONS]
    )

    sec("8. COURSE CONTENTS")
    units = data.get("units", [])
    for unit in units:
        lines.append(f"\nUNIT–{unit['num']}: {unit.get('title','')}")
        lines.append("Topics: " + " | ".join(unit.get("topics", [])))
        lines.append(f"CO: {unit.get('co','')}  PO: {unit.get('pos','')}  SDG: {unit.get('sdg','')}")
        lines.append(f"GA: {unit.get('ga','')}  AI Tool: {unit.get('ai_tool','')}")
        lines.append(f"Industry: {unit.get('industry','')}")

    sec("9. INNOVATIVE LESSON DELIVERY PLAN")
    sessions = data.get("session_plan", [])
    table_txt(
        ["S#", "Topic", "CO", "PO", "SDG", "Method", "AI Tool", "Assessment", "Evidence"],
        [(str(s.get('session','')), s.get('topic','')[:25], s.get('co',''),
          s.get('po',''), s.get('sdg',''), s.get('method','')[:20],
          s.get('ai_tool','')[:15], s.get('assessment',''), s.get('evidence','')) for s in sessions]
    )

    sec("10. MINI PROJECT / SDG CHALLENGE ACTIVITIES")
    projects = data.get("mini_projects", [])
    table_txt(
        ["Project Title", "SDG", "GA", "Innovation", "Impact"],
        [(p.get('title',''), p.get('sdg',''), p.get('ga',''),
          p.get('innovation',''), p.get('impact','')) for p in projects]
    )

    sec("11. COURSE ASSESSMENT FRAMEWORK")
    _assessment = _build_assessment_framework(len(data.get("cos", [])) or 5)
    table_txt(
        ["Assessment Method", "COs Assessed", "Bloom's Level", "Marks", "Evidence"],
        [(r["method"], r["cos"], r["bloom"], str(r["marks"]), r["evidence"])
         for r in _assessment]
    )

    sec("12. COURSE IMPACT MEASUREMENT")
    table_txt(
        ["Indicator", "Measurement Method", "Evidence"],
        [(r["indicator"], r["method"], r["evidence"]) for r in _IMPACT_MEASUREMENT]
    )

    sec("13. OBE & SDG ALIGNMENT")
    table_txt(
        ["Framework", "Alignment Details"],
        [(r["framework"], r["details"]) for r in _OBE_ALIGNMENT]
    )

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
