import sys
import os
import re
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')


def extract_pdf(path):
    try:
        import fitz
        text = []
        with fitz.open(path) as doc:
            for page in doc:
                t = page.get_text(sort=True)
                if t:
                    text.append(t)
        result = "\n".join(text)
        if result.strip():
            return result
    except Exception:
        pass

    import pdfplumber
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            try:
                t = page.extract_text()
                if t:
                    text.append(t)
            except Exception:
                continue
    return "\n".join(text)


def extract_docx(path):
    from docx import Document
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_txt(path):
    with open(path, encoding='utf-8', errors='replace') as f:
        return f.read()


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    elif ext == ".txt":
        return extract_txt(path)
    else:
        print(f"Error: Unsupported file type '{ext}'")
        sys.exit(1)


ROMAN = {
    'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
    'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10
}

_LTP_RE = re.compile(
    r'L\s*-\s*T\s*-\s*P\s*:?\s*(\d+)\s*-\s*(\d+)\s*-\s*(\d+)',
    re.IGNORECASE
)


def _extract_has_lab(text):
    """Return True if the course text contains an L-T-P line with P > 0."""
    m = _LTP_RE.search(text[:600])
    if m:
        return int(m.group(3)) > 0
    return False

def extract_semester(window_lines):
    """Return semester number (int) from surrounding lines, or None."""
    for line in reversed(window_lines):
        s = line.strip()
        # Format B: "Semester: 3"
        m = re.match(r'Semester\s*:\s*(\d+)', s)
        if m:
            return int(m.group(1))
        # Format A: "III Semester" or "Semester III"
        m = re.match(r'(I{1,3}V?|VI{0,3}|IX|X)\s+Semester', s)
        if m:
            return ROMAN.get(m.group(1))
        m = re.match(r'Semester\s+(I{1,3}V?|VI{0,3}|IX|X)\b', s)
        if m:
            return ROMAN.get(m.group(1))
    return None


def split_into_courses(full_text):
    """Split syllabus text into individual course sections.

    Handles two common formats:

    Format A (VTU / similar):
        <ROMAN> Semester
        COURSE TITLE IN ALL CAPS
        Course Code <CODE>  CIE Marks ...

    Format B (Autonomous colleges):
        Semester: <N>
        Course Title: <title>
        Course Code: <CODE>

    Course codes can be any alphanumeric string that contains at least one
    letter and at least one digit — no year prefix assumed.
    Examples: MAT301, CS3ESCOA, BCS401, 23CS32, BMAT3A, CS-101
    """
    courses = {}

    # Match any alphanumeric code (4-20 chars) that has at least one letter
    # AND at least one digit — prevents false matches on pure words like "MARKS"
    pattern = re.compile(
        r'Course Code\s*:?\s*'           # anchor: "Course Code" with optional colon
        r'((?=[A-Z0-9]*[A-Z])(?=[A-Z0-9]*\d)[A-Z0-9\-]{4,20})'  # code itself
    )
    matches = list(pattern.finditer(full_text))

    if not matches:
        return courses

    lines = full_text.split('\n')
    line_starts = []
    pos = 0
    for line in lines:
        line_starts.append(pos)
        pos += len(line) + 1

    def get_line_num(char_pos):
        for i in range(len(line_starts) - 1, -1, -1):
            if line_starts[i] <= char_pos:
                return i
        return 0

    for idx, match in enumerate(matches):
        code = match.group(1).strip()

        line_num = get_line_num(match.start())
        # Wider window (10 lines back) to catch semester info
        window = lines[max(0, line_num - 10): line_num + 2]

        # Extract semester
        semester = extract_semester(window)

        # Strategy 1: explicit "Course Title:" label (Format B)
        title = None
        for line in reversed(window):
            m = re.match(r'Course Title\s*:\s*(.+)', line.strip())
            if m:
                title = m.group(1).strip()
                break

        # Strategy 2: last ALL-CAPS line before Course Code (Format A)
        if not title:
            caps_lines = [
                l.strip() for l in window[:-1]
                if l.strip().isupper() and len(l.strip()) > 4
            ]
            if caps_lines:
                title = caps_lines[-1]

        # Strategy 3: last non-empty line before Course Code
        if not title:
            prev_lines = [l.strip() for l in window[:-1] if l.strip()]
            title = prev_lines[-1] if prev_lines else code

        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(full_text)
        content = full_text[start:end].strip()

        courses[code] = {
            "title":   title[:100],
            "semester": semester,
            "text":    content,
            "has_lab": _extract_has_lab(content),
        }

    return courses


_KDIMS  = ['Factual', 'Conceptual', 'Procedural', 'Meta-Cognitive']
_KDIM_LABELS = ['FACTUAL\nKNOWLEDGE', 'CONCEPTUAL\nKNOWLEDGE', 'PROCEDURAL\nKNOWLEDGE', 'METACOGNITIVE\nKNOWLEDGE']
_LEVELS = ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']
_LNAMES = ['Remember', 'Understand', 'Apply', 'Analyze', 'Evaluate', 'Create']
_LVERBS = [
    'Recognizing\nRecalling',
    'Interpreting\nExemplifying\nClassifying\nSummarizing\nInferring\nComparing\nExplaining',
    'Executing\nImplementing',
    'Differentiating\nOrganizing\nAttributing',
    'Checking\nCritiquing',
    'Generating\nPlanning\nProducing',
]


_CO_PREFIXES = [
    "The students will be able to",
    "Learners will be able to",
    "Students will demonstrate the ability to",
    "Upon completion, students will be able to",
    "Students will be capable of",
]


def _build_system_prompt(num_cos):
    def _prefix(i):
        return _CO_PREFIXES[(i - 1) % len(_CO_PREFIXES)]

    co_lines   = "\n".join(f"CO{i}: {_prefix(i)} <Behavior: verb> <Condition: context/tool>, <Degree: standard>" for i in range(1, num_cos + 1))
    table_rows = "\n".join(f"| CO{i} | Unit N | <knowledge dimension> | <bloom level> |" for i in range(1, num_cos + 1))
    prefix_list = " | ".join(f'"{p}"' for p in _CO_PREFIXES)
    return f"""You are an expert in Outcome-Based Education (OBE) and NBA accreditation.
Generate exactly {num_cos} Course Outcomes (COs) from the syllabus content provided.

Internal rules (never mention these):
- Apply Revised Bloom's Taxonomy to select appropriate cognitive levels
- Structure each CO using the ABCD framework:
    Audience  : use the exact prefix already shown for each CO in the output format below — do NOT change it
    Behavior  : one measurable action verb from the approved list
    Condition : the real-world context, tool, or constraint under which the skill is applied (e.g. "given an unsorted dataset", "using a relational database", "when presented with a network trace") — never reference the syllabus, unit numbers, or course material as the condition
    Degree    : the standard of performance (e.g. "correctly", "without error", "for at least 3 use cases")
    Combine as: "<Audience prefix> <Behavior> <Condition>, <Degree>"
- Allowed audience prefixes (use only these, exactly as written): {prefix_list}
- Each CO must cover a single competency only
- Distribute COs across all units in the syllabus
- You MUST select action verbs strictly from this list:
    Level 1 - Remember:   Recognize, Recall
    Level 2 - Understand: Interpret, Exemplify, Classify, Summarize, Infer, Compare, Explain
    Level 3 - Apply:      Execute, Implement
    Level 4 - Analyze:    Differentiate, Organize, Attribute
    Level 5 - Evaluate:   Check, Critique
    Level 6 - Create:     Generate, Plan, Produce
- Never use vague verbs like: know, learn, understand deeply, appreciate, study, be aware of
- Assign each CO exactly one Knowledge Dimension:
    Factual        — facts, terminology, specific details
    Conceptual     — classifications, theories, principles, interrelationships
    Procedural     — algorithms, techniques, methods, criteria for performing skills
    Meta-Cognitive — strategic knowledge, self-awareness, cognitive task knowledge

Output format (strictly follow this):
{co_lines}

| CO | Unit | Knowledge Dimension | Bloom Level |
|----|------|---------------------|-------------|
{table_rows}

For the Unit column use EXACTLY the format: Unit 1 | Unit 2 | Unit 3 ... (number only, no unit name)
For Bloom Level use EXACTLY: L1 - Remember | L2 - Understand | L3 - Apply | L4 - Analyze | L5 - Evaluate | L6 - Create
For Knowledge Dimension use EXACTLY one of: Factual | Conceptual | Procedural | Meta-Cognitive

Do not add any explanation or preamble. Output only the COs and table."""


def generate_cos_stream(client, course_code, course_title, course_text, num_cos=5):
    """Generator that yields raw text chunks from the API stream."""
    system_prompt = _build_system_prompt(num_cos)
    user_message = f"Course Code: {course_code}\nCourse Title: {course_title}\n\nSyllabus:\n{course_text[:6000]}"
    max_tokens = 400 + (num_cos * 200)  # enough for all CO statements + full table
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}]
    ) as stream:
        for text in stream.text_stream:
            yield text


def generate_cos_for_course(client, course_code, course_title, course_text, num_cos=5):
    result = ""
    for chunk in generate_cos_stream(client, course_code, course_title, course_text, num_cos):
        print(chunk, end="", flush=True)
        result += chunk
    print("\n")
    return result


_KDIM_NORM = {
    'factual': 'Factual',
    'conceptual': 'Conceptual',
    'procedural': 'Procedural',
    'meta-cognitive': 'Meta-Cognitive',
    'metacognitive': 'Meta-Cognitive',
    'meta cognitive': 'Meta-Cognitive',
}

def build_taxonomy_grid(result_text):
    grid = {kd: {lv: [] for lv in _LEVELS} for kd in _KDIMS}
    for line in result_text.split('\n'):
        s = line.strip()
        if s.startswith('|') and s.endswith('|'):
            cells = [c.strip() for c in s.strip('|').split('|')]
            if len(cells) >= 4 and re.match(r'CO\d+', cells[0]):
                co       = cells[0]
                bloom    = cells[-1]   # always last
                kdim_raw = cells[-2]   # always second-to-last
                m = re.match(r'(L\d+)', bloom)
                if not m:
                    continue
                kdim = _KDIM_NORM.get(kdim_raw.lower(), kdim_raw)
                if kdim in grid:
                    grid[kdim][m.group(1)].append(co)
    return grid


def _taxonomy_txt(grid):
    KD_W = 18; COL_W = 10
    sep = '+' + '-'*(KD_W+2) + ('+' + '-'*(COL_W+2))*6 + '+'
    def row(kd, cells):
        return '| ' + kd.ljust(KD_W) + ' |' + ''.join(' ' + c.ljust(COL_W) + ' |' for c in cells)
    hdr_cells = [f'L{i+1}-{ln[:3]}' for i, ln in enumerate(_LNAMES)]
    lines = ["\nBLOOM'S TAXONOMY MATRIX", sep, row('Knowledge Dim.', hdr_cells), sep]
    for kd in _KDIMS:
        lines.append(row(kd, [', '.join(grid[kd].get(lv, [])) for lv in _LEVELS]))
    lines.append(sep)
    return '\n'.join(lines) + '\n\n'


def _add_taxonomy_docx(doc, grid):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc.add_heading("Bloom's Taxonomy Matrix", level=1)
    tbl = doc.add_table(rows=6, cols=7)  # row 0: names, row 1: verbs, rows 2-5: data
    tbl.style = 'Table Grid'

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _shd(cell, fill_hex):
        tc  = cell._tc
        tcp = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcp.append(shd)

    _LVERBS_FLAT = [v.replace('\n', ', ') for v in _LVERBS]

    # Header row 0 — level numbers and names (bold)
    hdr0 = tbl.rows[0].cells
    p0   = hdr0[0].paragraphs[0]
    r0   = p0.add_run('KNOWLEDGE\nDIMENSION')
    r0.bold = True
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0); p0.paragraph_format.space_after = Pt(0)
    _shd(hdr0[0], 'C8C8C8')
    for i, ln in enumerate(_LNAMES, 1):
        p  = hdr0[i].paragraphs[0]
        r  = p.add_run(f'{i}. {ln.upper()}')
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        _shd(hdr0[i], 'C8C8C8')

    # Header row 1 — verbs (italic, smaller)
    hdr1 = tbl.rows[1].cells
    _shd(hdr1[0], 'C8C8C8')
    for i, verbs in enumerate(_LVERBS_FLAT, 1):
        p  = hdr1[i].paragraphs[0]
        r  = p.add_run(verbs)
        r.italic = True; r.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        _shd(hdr1[i], 'C8C8C8')

    # Data rows (indices 2-5)
    for row_idx, (kd, kd_label) in enumerate(zip(_KDIMS, _KDIM_LABELS), 2):
        row = tbl.rows[row_idx].cells
        pr  = row[0].paragraphs[0]
        rr  = pr.add_run(kd_label.replace('\n', ' ')); rr.bold = True
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col_idx, lv in enumerate(_LEVELS, 1):
            cell_cos = ', '.join(grid[kd].get(lv, []))
            row[col_idx].text = cell_cos
            row[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_taxonomy_pdf(pdf, grid):
    from fpdf import XPos, YPos

    KD_W = 36; LV_W = 22  # 36 + 6×22 = 168mm — fits A4 with margins

    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_fill_color(220, 230, 241)
    pdf.cell(0, 8, "Bloom's Taxonomy Matrix", fill=True,
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # ── Header: two rows so verbs don't overflow ─────────────────────────────
    pdf.set_fill_color(200, 200, 200)
    lm = pdf.l_margin

    # Row 1 — level numbers + names (bold)
    pdf.set_font("Helvetica", "B", 7)
    y0 = pdf.get_y()
    pdf.set_xy(lm, y0)
    pdf.multi_cell(KD_W, 4, 'KNOWLEDGE\nDIMENSION', border='LTR', fill=True, align='C')
    h1 = pdf.get_y() - y0  # always 8 (2 lines x 4mm)
    for i, ln in enumerate(_LNAMES, 1):
        pdf.set_xy(lm + KD_W + (i - 1) * LV_W, y0)
        # cell() with full row height so level name cells match KD cell height
        pdf.cell(LV_W, h1, f'{i}. {ln.upper()}', border='LTR', fill=True, align='C')

    # Row 2 — verbs (smaller, italic): all cells must share the same row height
    pdf.set_font("Helvetica", "I", 5.5)
    y1 = y0 + h1
    _LVERBS_FLAT = [v.replace('\n', ', ') for v in _LVERBS]

    # Pre-compute h2 by simulating word-wrap with actual FPDF font metrics
    def _wrap_lines(text):
        words = text.split(' ')
        lines, cur_w = 1, 0.0
        sp_w = pdf.get_string_width(' ')
        for word in words:
            ww = pdf.get_string_width(word)
            if cur_w == 0:
                cur_w = ww
            elif cur_w + sp_w + ww <= LV_W - 1:
                cur_w += sp_w + ww
            else:
                lines += 1
                cur_w = ww
        return lines

    h2 = max(max(_wrap_lines(v) for v in _LVERBS_FLAT) * 3.2, 3.2)

    # KD empty cell: filled rect + LBR border lines
    pdf.rect(lm, y1, KD_W, h2, style='F')
    pdf.line(lm, y1, lm, y1 + h2)
    pdf.line(lm, y1 + h2, lm + KD_W, y1 + h2)
    pdf.line(lm + KD_W, y1, lm + KD_W, y1 + h2)

    # Verb cells: filled rect at uniform h2, text on top, LBR border lines
    for i, verbs in enumerate(_LVERBS_FLAT, 1):
        x = lm + KD_W + (i - 1) * LV_W
        pdf.rect(x, y1, LV_W, h2, style='F')
        pdf.set_xy(x, y1)
        pdf.multi_cell(LV_W, 3.2, verbs, border=0, fill=False, align='C')
        pdf.line(x, y1, x, y1 + h2)
        pdf.line(x, y1 + h2, x + LV_W, y1 + h2)
        pdf.line(x + LV_W, y1, x + LV_W, y1 + h2)

    pdf.set_xy(lm, y1 + h2)

    # ── Data rows ────────────────────────────────────────────────────────────
    for kd, kd_label in zip(_KDIMS, _KDIM_LABELS):
        pdf.set_x(lm)
        pdf.set_font("Helvetica", "B", 7)
        pdf.cell(KD_W, 8, kd_label.replace('\n', ' '), border=1, align='C')
        pdf.set_font("Helvetica", "B", 8)
        for lv in _LEVELS:
            cos_in_cell = ', '.join(grid[kd].get(lv, []))
            pdf.cell(LV_W, 8, cos_in_cell, border=1, align='C')
        pdf.ln(8)
    pdf.ln(6)


def save_txt(results, output_path, taxonomy_grid=None):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("Course Outcomes\nGenerated from syllabus\n\n")
        if taxonomy_grid:
            f.write(_taxonomy_txt(taxonomy_grid))
        f.write("\n".join(results))


def save_docx(results, output_path, taxonomy_grid=None):
    from docx import Document

    doc = Document()
    doc.add_heading("Course Outcomes", 0)

    if taxonomy_grid:
        _add_taxonomy_docx(doc, taxonomy_grid)

    for block in results:
        lines = block.strip().split('\n')
        if not lines:
            continue

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Section header: ## CODE: TITLE [Semester N]
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)

            # CO statement: CO1: ...
            elif re.match(r'CO\d+:', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True

            # Markdown table header row
            elif line.startswith('| CO |'):
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'CO'
                hdr[1].text = 'Unit'
                hdr[2].text = 'Knowledge Dimension'
                hdr[3].text = 'Bloom Level'
                for cell in hdr:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True

            # Separator row
            elif re.match(r'\|[-| ]+\|', line):
                continue

            # Table data row: | CO1 | Module 1 | Procedural | L3 - Apply |
            elif line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                if len(cells) == 4 and 'table' in dir():
                    row = table.add_row().cells
                    row[0].text = cells[0]
                    row[1].text = cells[1]
                    row[2].text = cells[2]
                    row[3].text = cells[3]

            else:
                doc.add_paragraph(line)

    doc.save(output_path)


def _pdf_safe(text):
    """Replace Unicode punctuation that Helvetica cannot encode with ASCII equivalents."""
    return (text
            .replace('—', '--').replace('–', '-')
            .replace('‘', "'").replace('’', "'")
            .replace('“', '"').replace('”', '"')
            .replace('…', '...').replace(' ', ' ')
            .replace('•', '*').replace('−', '-'))


def save_pdf(results, output_path, taxonomy_grid=None):
    from fpdf import FPDF, XPos, YPos

    # CO table column widths (mm) — 16 + 70 + 40 + 64 = 190mm
    COL_CO   = 16
    COL_UNIT = 70
    COL_KD   = 40
    COL_BL   = 64

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Course Outcomes", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    if taxonomy_grid:
        _add_taxonomy_pdf(pdf, taxonomy_grid)

    for block in results:
        lines = block.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue

            # Section header: ## CODE: TITLE [Semester N]
            if line.startswith('## '):
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_fill_color(220, 230, 241)
                pdf.multi_cell(0, 8, _pdf_safe(line[3:]), fill=True,
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)

            # CO statement: CO1: ...
            elif re.match(r'CO\d+:', line):
                co_label, _, rest = line.partition(':')
                pdf.set_font("Helvetica", "B", 10)
                pdf.multi_cell(0, 7, _pdf_safe(f"{co_label}: {rest.strip()}"),
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Table header row
            elif line.startswith('| CO |'):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_fill_color(200, 200, 200)
                pdf.cell(COL_CO,   6, "CO",                   border=1, fill=True)
                pdf.cell(COL_UNIT, 6, "Unit",                 border=1, fill=True)
                pdf.cell(COL_KD,   6, "Knowledge Dimension",  border=1, fill=True)
                pdf.cell(COL_BL,   6, "Bloom Level",          border=1, fill=True,
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Separator row — skip
            elif re.match(r'\|[-| ]+\|', line):
                continue

            # Table data row
            elif line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                if len(cells) == 4:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.cell(COL_CO,   6, _pdf_safe(cells[0]),      border=1)
                    pdf.cell(COL_UNIT, 6, _pdf_safe(cells[1][:45]), border=1)
                    pdf.cell(COL_KD,   6, _pdf_safe(cells[2]),      border=1)
                    pdf.cell(COL_BL,   6, _pdf_safe(cells[3]),      border=1,
                             new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Generic text (e.g. Bloom's Level Distribution summary)
            else:
                pdf.set_font("Helvetica", "I", 9)
                pdf.multi_cell(0, 6, _pdf_safe(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(output_path)


def _to_sem_num(token):
    """Return integer semester if token looks like a semester number, else None."""
    try:
        n = int(token)
        return n if 1 <= n <= 10 else None
    except ValueError:
        return ROMAN.get(token.strip().upper())


def parse_filter_input(raw):
    """Parse a combined filter string into (pairs, sem_only).

    Formats accepted:
      "Artificial Intelligence,6,Mini Project,5"  -> pairs=[("Artificial Intelligence",6),("Mini Project",5)]
      "5,6"                                        -> sem_only=[5,6]
      ""                                           -> pairs=[], sem_only=[]
    """
    if not raw:
        return [], []

    tokens = [t.strip() for t in raw.split(',') if t.strip()]

    # If every token is a semester number → semester-only mode
    sem_nums = [_to_sem_num(t) for t in tokens]
    if all(s is not None for s in sem_nums):
        return [], sem_nums

    # Alternating CourseName,Semester pairs
    pairs = []
    i = 0
    while i < len(tokens):
        name = tokens[i]
        if i + 1 < len(tokens):
            sem = _to_sem_num(tokens[i + 1])
            if sem is not None:
                pairs.append((name, sem))
                i += 2
                continue
        pairs.append((name, None))
        i += 1

    return pairs, []


def bloom_level_summary(result_text):
    """Return a one-line distribution string, e.g. 'L2 (Understand): 2 | L3 (Apply): 3'."""
    counts = {}
    for line in result_text.split('\n'):
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) == 4 and re.match(r'CO\d+', cells[0]):
                m = re.match(r'(L\d+)', cells[3])
                if m:
                    lv = m.group(1)
                    counts[lv] = counts.get(lv, 0) + 1
    if not counts:
        return ""
    labels = dict(zip(_LEVELS, _LNAMES))
    parts = [f"{lv} ({labels.get(lv, lv)}): {counts[lv]}" for lv in sorted(counts)]
    return "Bloom's Level Distribution - " + " | ".join(parts)


def parse_co_selection(raw, total):
    """Parse '1,3,5', '1-5', 'all', or a mix into a sorted list within [1, total]."""
    raw = raw.strip().lower()
    if not raw or raw == 'all':
        return list(range(1, total + 1))
    selected = set()
    for part in raw.split(','):
        part = part.strip()
        if '-' in part:
            start, _, end = part.partition('-')
            try:
                selected.update(range(int(start.strip()), int(end.strip()) + 1))
            except ValueError:
                pass
        else:
            try:
                selected.add(int(part))
            except ValueError:
                pass
    return sorted(n for n in selected if 1 <= n <= total)


def filter_cos(result_text, selected):
    """Return result_text keeping only the chosen COs, renumbered 1..N."""
    sel_set  = set(selected)
    new_num  = {orig: i for i, orig in enumerate(sorted(sel_set), 1)}

    co_stmts   = []   # (orig_num, rest)
    table_rows = []   # (orig_num, unit, kdim, level)

    for line in result_text.split('\n'):
        s = line.strip()
        m = re.match(r'CO(\d+):\s*(.*)', s)
        if m:
            n = int(m.group(1))
            if n in sel_set:
                co_stmts.append((n, m.group(2)))
            continue
        if re.match(r'\|[-| ]+\|', s):
            continue
        if s.startswith('|') and s.endswith('|'):
            cells = [c.strip() for c in s.strip('|').split('|')]
            if len(cells) >= 4 and re.match(r'CO(\d+)', cells[0]):
                m2 = re.match(r'CO(\d+)', cells[0])
                if m2:
                    n = int(m2.group(1))
                    if n in sel_set:
                        unit  = ', '.join(cells[1:len(cells)-2])
                        kdim  = cells[-2]
                        bloom = cells[-1]
                        table_rows.append((n, unit, kdim, bloom))

    out = [f"CO{new_num[o]}: {rest}" for o, rest in sorted(co_stmts)]
    out.append("")
    out.append("| CO | Unit | Knowledge Dimension | Bloom Level |")
    out.append("|----|------|---------------------|-------------|")
    for o, unit, kdim, level in sorted(table_rows):
        out.append(f"| CO{new_num[o]} | {unit} | {kdim} | {level} |")
    return '\n'.join(out)


def generate_all(syllabus_text, base_path, fmt, settings):
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"], timeout=120.0, max_retries=3)

    filter_pairs = settings["filter_pairs"]   # list of (name, sem_or_None)
    sem_only     = settings["sem_only"]       # list of int
    num_cos      = settings["num_cos"]

    courses = split_into_courses(syllabus_text)

    if not courses:
        print("No course codes detected. Generating COs for the full document...\n")
        generate_cos_for_course(client, "COURSE", "Syllabus", syllabus_text, num_cos)
        return

    # --- Filter courses ---
    if filter_pairs:
        def matches_any_pair(code, info):
            for name, sem in filter_pairs:
                needle = name.lower()
                name_match = needle in code.lower() or needle in info["title"].lower()
                sem_match  = sem is None or info["semester"] == sem
                if name_match and sem_match:
                    return True
            return False

        courses = {code: info for code, info in courses.items() if matches_any_pair(code, info)}
        if not courses:
            desc = ', '.join(f"{n} (sem {s})" if s else n for n, s in filter_pairs)
            print(f"No courses found matching: {desc}. Exiting.")
            return

    elif sem_only:
        sem_set = set(sem_only)
        courses = {code: info for code, info in courses.items() if info["semester"] in sem_set}
        if not courses:
            print(f"No courses found for semester(s): {', '.join(str(s) for s in sem_only)}. Exiting.")
            return

    print(f"\nGenerating COs for {len(courses)} course(s)...\n")
    print("=" * 60)

    all_output = []

    # Document header
    header_lines = ["# Course Outcomes\nGenerated from syllabus"]
    if filter_pairs:
        desc = ', '.join(f"{n} (sem {s})" if s else n for n, s in filter_pairs)
        header_lines.append(f"Filter: {desc}")
    elif sem_only:
        header_lines.append(f"Filter — Semesters: {', '.join(str(s) for s in sem_only)}")
    header_lines.append(f"COs per course: {num_cos}")
    header_lines.append("")
    all_output.append("\n".join(header_lines))

    for i, (code, info) in enumerate(courses.items(), 1):
        sem_label = f"  [Semester {info['semester']}]" if info['semester'] else ""
        progress = f"\n[{i}/{len(courses)}] {code}: {info['title']}{sem_label}"
        print(progress)
        print("-" * 60)
        all_output.append(f"## {code}: {info['title']}{sem_label}\n")

        result = generate_cos_for_course(client, code, info['title'], info['text'], num_cos)
        summary = bloom_level_summary(result)
        all_output.append((summary + "\n\n" + result if summary else result) + "\n")

    output_path = f"{base_path}_COs.{fmt}"
    if fmt == "txt":
        save_txt(all_output, output_path)
    elif fmt == "docx":
        save_docx(all_output, output_path)
    elif fmt == "pdf":
        save_pdf(all_output, output_path)

    print(f"\nSaved to: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_cos.py <file.pdf|file.docx|file.txt>")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"Error: File not found: {path}")
        sys.exit(1)

    print(f"\nExtracting text from: {path}")
    syllabus_text = extract(path)
    if not syllabus_text.strip():
        print("Error: No text could be extracted.")
        sys.exit(1)

    courses = split_into_courses(syllabus_text)
    if not courses:
        print("No course codes detected in the file.")
        sys.exit(1)

    # --- List available courses ---
    print("\n" + "=" * 60)
    print("  Available Courses")
    print("=" * 60)
    for code, info in courses.items():
        sem = f"  [Sem {info['semester']}]" if info['semester'] else ""
        print(f"  {code:<20} {info['title'][:45]}{sem}")
    print("=" * 60)

    # --- Select course ---
    raw_code = input("\nEnter Course Code: ").strip().upper()
    if raw_code in courses:
        matched = raw_code
    else:
        hits = [c for c in courses if raw_code in c]
        if len(hits) == 1:
            matched = hits[0]
        elif len(hits) > 1:
            print(f"Multiple matches: {', '.join(hits)}. Please be more specific.")
            sys.exit(1)
        else:
            print(f"Course code '{raw_code}' not found.")
            sys.exit(1)

    info = courses[matched]
    sem_label = f"  [Semester {info['semester']}]" if info['semester'] else ""
    print(f"\nSelected: {matched} — {info['title']}{sem_label}")

    # --- Generate 15 COs ---
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"], timeout=120.0, max_retries=3)

    NUM_COS = 15
    print(f"\nGenerating {NUM_COS} COs for {matched}...")
    print("=" * 60)
    result = generate_cos_for_course(client, matched, info['title'], info['text'], num_cos=NUM_COS)

    # --- CO selection ---
    print("=" * 60)
    print("Enter CO numbers to keep (e.g.  1,3,5  |  1-5  |  1,3-7,10  |  all) [default: all]:")
    raw_sel  = input("Selection: ").strip()
    selected = parse_co_selection(raw_sel, NUM_COS)

    if not selected:
        print("No valid COs selected. Exiting.")
        sys.exit(0)

    print(f"\nKeeping {len(selected)} CO(s): {', '.join(f'CO{n}' for n in selected)}")
    filtered = filter_cos(result, selected)
    summary  = bloom_level_summary(filtered)
    grid     = build_taxonomy_grid(filtered)

    # --- Output format ---
    print("\nOutput format:")
    print("  1. txt")
    print("  2. docx  (Word)")
    print("  3. pdf")
    fmt_input = input("Enter 1, 2, or 3 [default: 1]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_input, "txt")
    if fmt_input and fmt_input not in ("1", "2", "3"):
        print("Invalid choice — using txt.")

    # --- Build and save output ---
    header       = f"# Course Outcomes\nGenerated from syllabus\nCOs exported: {len(selected)}\n"
    course_block = f"## {matched}: {info['title']}{sem_label}\n"
    body         = (summary + "\n\n" + filtered if summary else filtered) + "\n"
    all_output   = [header, course_block, body]

    base_path   = os.path.splitext(path)[0]
    output_path = f"{base_path}_{matched}_COs.{fmt}"

    if fmt == "txt":
        save_txt(all_output, output_path, taxonomy_grid=grid)
    elif fmt == "docx":
        save_docx(all_output, output_path, taxonomy_grid=grid)
    elif fmt == "pdf":
        save_pdf(all_output, output_path, taxonomy_grid=grid)

    print(f"\nSaved to: {output_path}")


if __name__ == "__main__":
    main()
