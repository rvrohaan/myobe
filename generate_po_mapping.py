import os
import re
import json
import sys
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ── PO definitions (NBA standard, 12 POs) ─────────────────────────────────────

POS = [
    ("PO1",  "Engineering Knowledge",
     "Apply knowledge of mathematics, science, engineering fundamentals, and specialization "
     "to solve complex engineering problems. "
     "1=Understands concepts; 2=Applies formulas/principles; 3=Solves complex problems independently."),

    ("PO2",  "Problem Analysis",
     "Identify, formulate, review research literature, and analyze complex engineering problems "
     "reaching substantiated conclusions using first principles. "
     "1=Identifies/states problem; 2=Analyzes and formulates; 3=Models, validates, reaches conclusions."),

    ("PO3",  "Design/Development of Solutions",
     "Design solutions for complex engineering problems and design system components or processes "
     "meeting specified needs with consideration for health, safety, and environment. "
     "1=Understands design process; 2=Develops partial solution; 3=Complete, validated design."),

    ("PO4",  "Conduct Investigations",
     "Use research-based knowledge including design of experiments, analysis and interpretation of data, "
     "and synthesis of information to provide valid conclusions. "
     "1=Observes/records experiments; 2=Conducts and analyses experiments; 3=Designs experiments and interprets results."),

    ("PO5",  "Modern Tool Usage",
     "Create, select, and apply appropriate techniques, resources, and modern engineering/IT tools "
     "including prediction and modelling, with understanding of limitations. "
     "1=Uses given tools; 2=Applies and adapts tools; 3=Selects, creates, and validates tools."),

    ("PO6",  "The Engineer and Society",
     "Apply reasoning to assess societal, health, safety, legal and cultural issues and responsibilities "
     "relevant to professional engineering practice. "
     "1=Awareness of societal issues; 2=Analyzes impacts; 3=Evaluates and proposes responsible solutions."),

    ("PO7",  "Environment and Sustainability",
     "Understand the impact of engineering solutions in societal and environmental contexts; "
     "demonstrate knowledge of and need for sustainable development. "
     "1=Understands impact; 2=Analyzes impact; 3=Designs sustainable solutions."),

    ("PO8",  "Ethics",
     "Apply ethical principles and commit to professional ethics, responsibilities, and norms of "
     "engineering practice. "
     "1=Awareness of ethics; 2=Applies ethical reasoning; 3=Evaluates and advocates ethical practice."),

    ("PO9",  "Individual and Team Work",
     "Function effectively as an individual, and as a member or leader in diverse and "
     "multidisciplinary teams. "
     "1=Participates in team activities; 2=Contributes and collaborates; 3=Leads and manages teams."),

    ("PO10", "Communication",
     "Communicate effectively on complex engineering activities — write reports, design documentation, "
     "make presentations, and give/receive clear instructions. "
     "1=Aware of communication requirements; 2=Communicates adequately; 3=Communicates proficiently in multiple modes."),

    ("PO11", "Project Management and Finance",
     "Demonstrate knowledge of engineering and management principles; apply to manage projects "
     "and in multidisciplinary environments. "
     "1=Aware of project management concepts; 2=Applies management principles; 3=Plans and manages projects."),

    ("PO12", "Life-long Learning",
     "Recognize the need for, and have the ability to engage in independent and life-long learning "
     "in the broadest context of technological change. "
     "1=Aware of need for continuous learning; 2=Identifies learning needs and pursues them; 3=Demonstrates self-directed, reflective learning."),
]

PO_KEYS = [p[0] for p in POS]   # ["PO1", ..., "PO12"]


# ── CO / course parsing ───────────────────────────────────────────────────────

def parse_cos_for_mapping(text: str) -> list[dict]:
    """Extract CO statements from text produced by generate_cos.py.

    Returns list of {"num": int, "statement": str}.
    Skips table rows (lines starting with |) to avoid grabbing CO column headers.
    """
    cos = []
    for line in text.split('\n'):
        s = line.strip()
        if s.startswith('|'):
            continue
        m = re.match(r'CO(\d+):\s*(.+)', s)
        if m:
            cos.append({"num": int(m.group(1)), "statement": m.group(2).strip()})
    return cos


def parse_course_header(text: str) -> tuple[str, str]:
    """Extract (course_code, course_title) from a CO file's header line.

    Handles two formats produced by generate_cos.py:
      TXT / raw:  ## 23CS5AEAST: Automated Software Testing  [Semester 5]
      DOCX text:  23CS5AEAST: Automated Software Testing  [Semester 5]
    Returns ("COURSE", "Untitled") if not found.
    """
    for line in text.split('\n'):
        s = line.strip().lstrip('#').strip()
        # Match CODE: Title [optional semester tag]
        m = re.match(
            r'((?=[A-Z0-9]*[A-Z])(?=[A-Z0-9]*\d)[A-Z0-9\-]{4,20})\s*:\s*(.+?)(?:\s*\[Semester\s*\d+\])?\s*$',
            s
        )
        if m:
            code  = m.group(1).strip()
            title = m.group(2).strip()
            # Skip generic section headers that look like codes (e.g. "CO1: ...")
            if re.match(r'CO\d+', code):
                continue
            return code, title
    return "COURSE", "Untitled"


# ── Prompt builder ────────────────────────────────────────────────────────────

_MAPPING_SYSTEM = """\
You are an expert NBA/OBE evaluator tasked with generating a CO-PO mapping table.

SCORING RUBRIC (strictly follow):
  3 – Major   : CO directly and thoroughly addresses this PO. The topic is fully introduced,
                developed, and reinforced throughout the course. Learner can SOLVE / DESIGN /
                EVALUATE independently.
  2 – Moderate: CO partially addresses this PO. Topic is introduced and developed with support.
                Learner can ANALYZE / APPLY with guidance.
  1 – Minor   : CO has a minor or indirect connection to this PO. Topic is mentioned for
                awareness. Learner can IDENTIFY / EXPLAIN / OBSERVE.
  0 – None    : CO has no meaningful relationship to this PO.

DECISION LOGIC (one-step):
  For each CO–PO pair ask:
    (a) Does any PO element described below appear in the CO statement?
    (b) What depth of learning does the CO require?
  The answer directly yields 0 / 1 / 2 / 3. No guessing; no padding non-zero values.

NBA EXPECTATION: A typical 5–6 CO course maps to 5–8 POs. Scores of 2 or 3 should be earned,
not given by default. Most COs will map to only 2–4 POs significantly.

OUTPUT FORMAT — return ONLY a JSON object, no explanation, no markdown fences:
{
  "mapping": [
    { "co": "CO1", "scores": { "<PO_KEY>": 0|1|2|3, ... } },
    { "co": "CO2", "scores": { ... } },
    ...
  ]
}
Use EXACTLY the PO keys listed in the user message. Include ALL COs from the input.
"""


def _build_user_message(cos: list[dict], course_code: str, course_title: str,
                        pos=None) -> str:
    _pos = pos if pos is not None else POS
    po_block = "\n".join(
        f"{key} – {name}: {desc}"
        for key, name, desc in _pos
    )
    co_block  = "\n".join(f"CO{co['num']}: {co['statement']}" for co in cos)
    key_list  = ", ".join(f'"{k}"' for k, _, _ in _pos)
    return (
        f"Course: {course_code} — {course_title}\n\n"
        f"PROGRAMME OUTCOMES (with scoring guidance):\n{po_block}\n\n"
        f"COURSE OUTCOMES TO MAP:\n{co_block}\n\n"
        f"PO keys to use in your JSON (in this order): {key_list}\n\n"
        "Generate the CO-PO mapping table as specified."
    )


# ── AI calls ──────────────────────────────────────────────────────────────────

def generate_mapping_stream(client, cos: list[dict], course_code: str, course_title: str,
                            pos=None):
    """Yield raw text chunks from the API stream."""
    _pos = pos if pos is not None else POS
    user_msg = _build_user_message(cos, course_code, course_title, _pos)
    # ~20 tokens per PO per CO, plus overhead
    max_tokens = max(1500, 300 + len(cos) * (100 + len(_pos) * 20))
    with client.messages.stream(
        model="claude-haiku-4-5-20251001",
        max_tokens=max_tokens,
        system=_MAPPING_SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def generate_mapping(client, cos: list[dict], course_code: str, course_title: str,
                     pos=None) -> str:
    """Non-streaming: return the full response text."""
    result = ""
    for chunk in generate_mapping_stream(client, cos, course_code, course_title, pos=pos):
        result += chunk
    return result


# ── Response parsing ──────────────────────────────────────────────────────────

def _extract_json(raw: str) -> dict | None:
    """Try several strategies to extract a JSON object from model output."""
    # Strategy 1: strip fences, parse directly
    cleaned = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Strategy 2: find the outermost { ... } span
    start = cleaned.find('{')
    end   = cleaned.rfind('}')
    if start != -1 and end > start:
        try:
            return json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError:
            pass

    return None


def parse_mapping_response(raw: str, cos: list[dict], po_keys=None) -> list[dict]:
    """Parse the AI JSON response into a list of {co, statement, scores} dicts.

    Falls back to zeros for any CO or PO that is missing from the response.
    po_keys: list of PO key strings to expect; defaults to the standard 12.
    """
    _po_keys   = po_keys if po_keys is not None else PO_KEYS
    co_map     = {co["num"]: co["statement"] for co in cos}
    zero_scores = {k: 0 for k in _po_keys}

    data = _extract_json(raw)
    rows = data.get("mapping", []) if data else []

    parsed = []
    seen = set()
    for row in rows:
        label = str(row.get("co", ""))
        m = re.match(r'CO(\d+)', label)
        if not m:
            continue
        num = int(m.group(1))
        scores_raw = row.get("scores", {})
        scores = {}
        for k in _po_keys:
            try:
                v = int(scores_raw.get(k, 0))
                scores[k] = max(0, min(3, v))
            except (ValueError, TypeError):
                scores[k] = 0
        parsed.append({
            "co":        f"CO{num}",
            "num":       num,
            "statement": co_map.get(num, ""),
            "scores":    scores,
        })
        seen.add(num)

    # Ensure every input CO appears (fill missing with zeros)
    for co in cos:
        if co["num"] not in seen:
            parsed.append({
                "co":        f"CO{co['num']}",
                "num":       co["num"],
                "statement": co["statement"],
                "scores":    dict(zero_scores),
            })

    parsed.sort(key=lambda r: r["num"])
    return parsed


# ── Summary helpers ───────────────────────────────────────────────────────────

def mapping_summary(rows: list[dict]) -> str:
    """Return a one-line summary: how many POs each mapping level covers."""
    counts = {0: 0, 1: 0, 2: 0, 3: 0}
    for row in rows:
        for v in row["scores"].values():
            counts[v] = counts.get(v, 0) + 1
    parts = []
    if counts[3]: parts.append(f"Major (3): {counts[3]}")
    if counts[2]: parts.append(f"Moderate (2): {counts[2]}")
    if counts[1]: parts.append(f"Minor (1): {counts[1]}")
    parts.append(f"None (0): {counts[0]}")
    return "CO–PO Mapping Distribution — " + " | ".join(parts)


def active_pos(rows: list[dict], all_pos=None) -> list[str]:
    """Return PO keys that have at least one non-zero score across all COs.

    When all_pos is None, derives the full key list from the rows themselves
    (preserving standard order then custom keys alphabetically).
    """
    if all_pos is None:
        seen_keys: list[str] = []
        seen_set: set[str]   = set()
        for r in rows:
            for k in r["scores"]:
                if k not in seen_set:
                    seen_set.add(k)
                    seen_keys.append(k)
        std_order = {k: i for i, k in enumerate(PO_KEYS)}
        all_pos = sorted(seen_keys, key=lambda k: (std_order.get(k, 999), k))
    if not all_pos:
        all_pos = PO_KEYS
    active = [k for k in all_pos if any(r["scores"].get(k, 0) > 0 for r in rows)]
    return active if active else all_pos


# ── Save: TXT ─────────────────────────────────────────────────────────────────

def save_txt(rows: list[dict], course_code: str, course_title: str, output_path: str, title: str = "CO-PO Mapping Table"):
    pos = active_pos(rows)
    col = 8
    co_w = 6

    sep = "+" + "-" * (co_w + 2) + ("+" + "-" * (col + 2)) * len(pos) + "+"
    def row_line(cells):
        return "| " + cells[0].ljust(co_w) + " |" + "".join(
            " " + str(c).center(col) + " |" for c in cells[1:]
        )

    lines = [
        title.upper(),
        f"Course: {course_code} — {course_title}",
        "",
        sep,
        row_line(["CO \\ PO"] + pos),
        sep,
    ]
    for r in rows:
        cells = [r["co"]] + [r["scores"].get(k, 0) for k in pos]
        lines.append(row_line(cells))
    lines.append(sep)
    lines.append("")
    lines.append(mapping_summary(rows))
    lines.append("")
    lines.append("PO Descriptions:")
    for key, name, _ in POS:
        if key in pos:
            lines.append(f"  {key}: {name}")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))


# ── Save: DOCX ────────────────────────────────────────────────────────────────

_SCORE_COLORS = {
    3: (198, 224, 180),   # green
    2: (255, 242, 204),   # yellow
    1: (221, 235, 247),   # blue-light
    0: (255, 255, 255),   # white
}


def save_docx(rows: list[dict], course_code: str, course_title: str, output_path: str, title: str = "CO–PO Mapping Table"):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, rgb):
        r, g, b = rgb
        hex_color = f"{r:02X}{g:02X}{b:02X}"
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    pos = active_pos(rows)
    doc = Document()
    for sec in doc.sections:
        sec.orientation = 1  # landscape
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        sec.left_margin = sec.right_margin = Inches(0.7)
        sec.top_margin = sec.bottom_margin = Inches(0.7)

    doc.add_heading(title, 0)
    p = doc.add_paragraph(f"Course: {course_code} — {course_title}")
    p.runs[0].font.size = Pt(11)

    ncols = 1 + len(pos)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    # Header row
    hdr = tbl.rows[0].cells
    hdr[0].text = "CO \\ PO"
    for i, k in enumerate(pos, 1):
        hdr[i].text = k
    for cell in hdr:
        _set_cell_bg(cell, (180, 198, 231))
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if para.runs:
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(9)

    # Data rows
    for ri, r in enumerate(rows, 1):
        cells = tbl.rows[ri].cells
        cells[0].text = r["co"]
        cells[0].paragraphs[0].runs[0].bold = True if cells[0].paragraphs[0].runs else None
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        for ci, k in enumerate(pos, 1):
            v = r["scores"].get(k, 0)
            cells[ci].text = str(v) if v else "–"
            _set_cell_bg(cells[ci], _SCORE_COLORS[v])
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if para.runs:
                    para.runs[0].font.size = Pt(9)
                    if v == 3:
                        para.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(mapping_summary(rows)).runs[0].font.size = Pt(9) if doc.paragraphs[-1].runs else None

    legend_para = doc.add_paragraph()
    legend_para.add_run("Legend: ").bold = True
    legend_para.add_run("3 = Major  |  2 = Moderate  |  1 = Minor  |  0 / – = None")
    if legend_para.runs:
        for run in legend_para.runs:
            run.font.size = Pt(9)

    doc.save(output_path)


# ── Save: PDF ─────────────────────────────────────────────────────────────────

def save_pdf(rows: list[dict], course_code: str, course_title: str, output_path: str, title: str = "CO-PO Mapping Table"):
    from fpdf import FPDF, XPos, YPos

    def _safe(s):
        s = s.replace('—', '-').replace('–', '-').replace('’', "'")
        return s.encode('latin-1', errors='replace').decode('latin-1')

    pos = active_pos(rows)

    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_margins(12, 12, 12)

    page_w = pdf.w - 2 * pdf.l_margin   # usable width in landscape A4 ~270 mm

    # Title
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 9, _safe(title), align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, _safe(f"Course: {course_code} — {course_title}"), align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Column widths
    co_w = 18
    po_w = (page_w - co_w) / len(pos)
    ROW_H = 7

    # Header
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(180, 198, 231)
    pdf.cell(co_w, ROW_H, "CO \\ PO", border=1, fill=True, align="C")
    for k in pos:
        pdf.cell(po_w, ROW_H, k, border=1, fill=True, align="C")
    pdf.ln(ROW_H)

    # Fill colours
    _fill = {
        3: (198, 224, 180),
        2: (255, 242, 204),
        1: (221, 235, 247),
        0: (255, 255, 255),
    }

    # Data rows
    for r in rows:
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(co_w, ROW_H, r["co"], border=1, fill=True, align="C")
        for k in pos:
            v = r["scores"].get(k, 0)
            pdf.set_fill_color(*_fill[v])
            pdf.set_font("Helvetica", "B" if v == 3 else "", 8)
            pdf.cell(po_w, ROW_H, str(v) if v else "-", border=1, fill=True, align="C")
        pdf.ln(ROW_H)

    pdf.ln(5)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 6, _safe(mapping_summary(rows)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 6, "Legend:  3 = Major  |  2 = Moderate  |  1 = Minor  |  0 / - = None",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(output_path)


# ── CLI entry point ───────────────────────────────────────────────────────────

def main():
    import anthropic
    import tempfile

    if len(sys.argv) < 2:
        print("Usage: python generate_po_mapping.py <co_file.txt|co_file.pdf|co_file.docx>")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"Error: File not found: {path}")
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    # Read CO text
    if ext == ".pdf":
        from generate_qbank import load_cos_from_pdf
        from generate_cos import extract_pdf
        cos     = load_cos_from_pdf(path)
        co_text = extract_pdf(path)   # full text for header detection
    elif ext in (".docx", ".doc"):
        from generate_cos import extract_docx
        co_text = extract_docx(path)
        cos = parse_cos_for_mapping(co_text)
    elif ext == ".txt":
        from generate_cos import extract_txt
        co_text = extract_txt(path)
        cos = parse_cos_for_mapping(co_text)
    else:
        print(f"Unsupported file type '{ext}'")
        sys.exit(1)

    if not cos:
        print("No COs found in the file.")
        sys.exit(1)

    course_code, course_title = parse_course_header(co_text)
    print(f"\nCourse : {course_code} — {course_title}")
    print(f"COs found: {len(cos)}")
    for co in cos:
        print(f"  CO{co['num']}: {co['statement'][:80]}{'...' if len(co['statement']) > 80 else ''}")

    print("\nGenerating CO-PO Mapping…")
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"], timeout=60.0)

    raw = ""
    for chunk in generate_mapping_stream(client, cos, course_code, course_title):
        print(chunk, end="", flush=True)
        raw += chunk
    print("\n")

    rows = parse_mapping_response(raw, cos)

    print("\n" + "=" * 60)
    print("Output format:")
    print("  1. txt")
    print("  2. docx  (Word)")
    print("  3. pdf")
    fmt_input = input("Enter 1, 2, or 3 [default: 2]: ").strip()
    fmt = {"1": "txt", "2": "docx", "3": "pdf"}.get(fmt_input, "docx")

    base = os.path.splitext(path)[0]
    output_path = f"{base}_CO_PO_Mapping.{fmt}"

    if fmt == "txt":
        save_txt(rows, course_code, course_title, output_path)
    elif fmt == "docx":
        save_docx(rows, course_code, course_title, output_path)
    elif fmt == "pdf":
        save_pdf(rows, course_code, course_title, output_path)

    print(f"\nSaved to: {output_path}")
    print(mapping_summary(rows))


if __name__ == "__main__":
    main()
